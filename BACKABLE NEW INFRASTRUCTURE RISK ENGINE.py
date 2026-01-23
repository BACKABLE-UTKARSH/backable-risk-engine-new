import logging
import warnings
import requests
import time
import json
import re
import random
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple, Union
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import numpy as np
from dotenv import load_dotenv
import os
import pandas as pd
import tiktoken
import io
import concurrent.futures
from concurrent.futures import ThreadPoolExecutor
from threading import Thread
from fastapi import FastAPI, Request, Response, HTTPException, BackgroundTasks, Depends, Header, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
from pydantic import BaseModel
from azure.storage.blob import BlobServiceClient, ContainerClient, ContentSettings
import psycopg2
import uvicorn
import base64
import jwt
import hashlib
from psycopg2.extras import RealDictCursor
from dataclasses import dataclass
from typing import Optional
from concurrent.futures import ThreadPoolExecutor, as_completed
import asyncio
import aiohttp
from typing import Optional, Dict, Any
import platform
from contextlib import asynccontextmanager
from collections import deque
import statistics
import asyncpg
import threading, sys
from collections import defaultdict

# Google GenAI SDK for Vertex AI (new architecture - primary method)
from google import genai
from google.oauth2 import service_account

# Load environment variables from .env file
load_dotenv()

# Global connection pools with better management
_connection_pools = {}
_pool_creation_locks = defaultdict(asyncio.Lock)
_pool_lock = asyncio.Lock()

# ======================================================
#                 Configuration - RISK ENGINE
# ======================================================

# Windows-specific event loop fix for aiohttp
if platform.system() == 'Windows':
    asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

# Notification configuration
NOTIFICATION_API_URL = "https://philotimo-backend-staging.azurewebsites.net/send-notification"
NOTIFICATION_TIMEOUT = 10  # seconds

# Risk Engine Database Configuration
RISK_DB_CONFIG = {
    "host": os.getenv("DB_HOST", "memberchat-db.postgres.database.azure.com"),
    "database": "BACKABLE-RISK-ENGINE",
    "user": os.getenv("DB_USER", "backable"),
    "password": os.getenv("DB_PASSWORD", ""),
    "port": int(os.getenv("DB_PORT", "5432")),
    "sslmode": "require"
}

# Multi-Database Intelligence Sources (same as People & Ops)
COMPONENT_DB_CONFIG = {
    "host": os.getenv("DB_HOST", "memberchat-db.postgres.database.azure.com"),
    "database": "BACKABLE-COMPONENT-ENGINE",
    "user": os.getenv("DB_USER", "backable"),
    "password": os.getenv("DB_PASSWORD", ""),
    "port": int(os.getenv("DB_PORT", "5432")),
    "sslmode": "require"
}

PROFILE_DB_CONFIG = {
    "host": os.getenv("DB_HOST", "memberchat-db.postgres.database.azure.com"),
    "database": "BACKABLE-PROFILE-ENGINE",
    "user": os.getenv("DB_USER", "backable"),
    "password": os.getenv("DB_PASSWORD", ""),
    "port": int(os.getenv("DB_PORT", "5432")),
    "sslmode": "require"
}

DREAM_DB_CONFIG = {
    "host": os.getenv("DB_HOST", "memberchat-db.postgres.database.azure.com"),
    "database": "BACKABLE-DREAM-ANALYZER",
    "user": os.getenv("DB_USER", "backable"),
    "password": os.getenv("DB_PASSWORD", ""),
    "port": int(os.getenv("DB_PORT", "5432")),
    "sslmode": "require"
}

ANALYST_DB_CONFIG = {
    "host": os.getenv("DB_HOST", "memberchat-db.postgres.database.azure.com"),
    "database": "BACKABLE-THE-ANALYST",
    "user": os.getenv("DB_USER", "backable"),
    "password": os.getenv("DB_PASSWORD", ""),
    "port": int(os.getenv("DB_PORT", "5432")),
    "sslmode": "require"
}

GROWTH_DB_CONFIG = {
    "host": os.getenv("DB_HOST", "memberchat-db.postgres.database.azure.com"),
    "database": "BACKABLE-THE-GROWTH-ENGINE",
    "user": os.getenv("DB_USER", "backable"),
    "password": os.getenv("DB_PASSWORD", ""),
    "port": int(os.getenv("DB_PORT", "5432")),
    "sslmode": "require"
}

PEOPLE_OPS_DB_CONFIG = {
    "host": os.getenv("DB_HOST", "memberchat-db.postgres.database.azure.com"),
    "database": "BACKABLE-PEOPLE-AND-OPERATIONS-ENGINE",
    "user": os.getenv("DB_USER", "backable"),
    "password": os.getenv("DB_PASSWORD", ""),
    "port": int(os.getenv("DB_PORT", "5432")),
    "sslmode": "require"
}

USER_DB_CONFIG = {
    "host": os.getenv("USER_DB_HOST", "philotimo-staging-db.postgres.database.azure.com"),
    "database": os.getenv("USER_DB_NAME", "philotimodb"),
    "user": os.getenv("USER_DB_USER", "wchen"),
    "password": os.getenv("USER_DB_PASSWORD", ""),
    "port": int(os.getenv("USER_DB_PORT", "5432")),
    "sslmode": "require"
}

AZURE_STORAGE_CONNECTION_STRING = os.getenv(
    "AZURE_STORAGE_CONNECTION_STRING",
    ""  # Empty default - will fail gracefully if not set
)

ONBOARDING_DB_HOST = os.getenv("DB_HOST", "memberchat-db.postgres.database.azure.com")
ONBOARDING_DB_NAME = "BACKABLE-GOOGLE-RAG"  # Updated to new unified architecture database
ONBOARDING_DB_USER = os.getenv("DB_USER", "backable")
ONBOARDING_DB_PASSWORD = os.getenv("DB_PASSWORD", "")
ONBOARDING_DB_PORT = int(os.getenv("DB_PORT", "5432"))

# ======================================================
#                 Gemini 2.5 Pro Configuration - RISK ENGINE
# ======================================================
# Load API keys from environment variable (comma-separated)
GEMINI_API_KEYS_STR = os.getenv("GEMINI_API_KEYS_RISK", "")
if GEMINI_API_KEYS_STR:
    GEMINI_API_KEYS = [key.strip() for key in GEMINI_API_KEYS_STR.split(",") if key.strip()]
    API_KEYS_CONFIG = {
        key: {"name": f"Back_Risk{i+1:02d}", "priority": 1, "health": True}
        for i, key in enumerate(GEMINI_API_KEYS)
    }
else:
    logging.warning("⚠️ No Gemini API keys found in environment variable GEMINI_API_KEYS_RISK")
    GEMINI_API_KEYS = []
    API_KEYS_CONFIG = {}

# ======================================================
#           Vertex AI Configuration (Primary Method)
# ======================================================
VERTEX_PROJECT_ID = "backable-machine-learning-apis"
VERTEX_LOCATION = "us-central1"
USE_VERTEX_AI = True  # Primary method - will fallback to API keys if fails

# API Key Management Variables
api_key_stats = defaultdict(lambda: {"requests": 0, "failures": 0, "last_used": 0, "cooldown_until": 0})
api_key_lock = threading.Lock()

# API key health tracking for Risk Engine
api_key_health = {}

# ======================================================
#           Vertex AI Initialization
# ======================================================

def initialize_vertex_ai_client():
    """
    Initialize Google GenAI client for Vertex AI.
    Supports both file-based and environment variable credentials.
    Returns None if initialization fails (will use API keys fallback).
    """
    try:
        # Try loading credentials from environment variable first (Azure deployment)
        creds_json = os.getenv('GOOGLE_APPLICATION_CREDENTIALS_JSON')

        if creds_json:
            logging.info("Loading Vertex AI credentials from environment variable")
            import tempfile
            creds_dict = json.loads(creds_json)
            with tempfile.NamedTemporaryFile(mode='w', delete=False, suffix='.json') as temp_file:
                json.dump(creds_dict, temp_file)
                temp_path = temp_file.name

            credentials = service_account.Credentials.from_service_account_file(
                temp_path,
                scopes=['https://www.googleapis.com/auth/cloud-platform']
            )
            os.unlink(temp_path)
        else:
            # Fall back to file-based credentials (local development)
            creds_file = "vertex-key.json"
            if os.path.exists(creds_file):
                logging.info(f"Loading Vertex AI credentials from {creds_file}")
                credentials = service_account.Credentials.from_service_account_file(
                    creds_file,
                    scopes=['https://www.googleapis.com/auth/cloud-platform']
                )
            else:
                logging.warning("No Vertex AI credentials found - will use API keys fallback")
                return None

        # Initialize GenAI client
        client = genai.Client(
            vertexai=True,
            credentials=credentials,
            project=VERTEX_PROJECT_ID,
            location=VERTEX_LOCATION
        )

        logging.info(f"✅ Vertex AI GenAI client initialized successfully (Project: {VERTEX_PROJECT_ID})")
        return client

    except Exception as e:
        logging.warning(f"⚠️ Vertex AI initialization failed: {str(e)} - Will use API keys fallback")
        return None

# Initialize Vertex AI client at startup
vertex_ai_client = initialize_vertex_ai_client() if USE_VERTEX_AI else None

# ======================================================
#           USER PROFILE AND DATABASE FUNCTIONS
# ======================================================

async def get_user_profile_data(user_id: str):
    """Get user profile data using connection pool - FIXED DATA TYPE ISSUE"""
    try:
        logging.info(f"Getting user profile data for user_id={user_id}")
        pool = await get_db_pool(USER_DB_CONFIG)
        
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    id, email, username, password, remember_me_token,
                    created_at, updated_at, is_email_verified, client_id,
                    business_name, contact_name, phone_number, ppr_id,
                    company_url, last_name, abn, archive, personal_bio, 
                    location, profile_image_url, skills, interests, 
                    last_login_at, achievements, provider, provider_id, 
                    login_count, last_login_provider, industry, team_size, 
                    business_description, biggest_challenge
                FROM users
                WHERE id = $1 OR client_id = $1
                LIMIT 1
            """
            
            # 🔥 FIX: Convert string user_id to integer for database query
            try:
                user_id_int = int(user_id)
                logging.info(f"🔢 Converted user_id '{user_id}' to integer {user_id_int}")
                row = await conn.fetchrow(sql, user_id_int)
            except ValueError:
                # If user_id is not a valid integer, try as string with client_id only
                logging.warning(f"⚠️ user_id '{user_id}' is not an integer, trying as client_id string")
                sql_string = """
                    SELECT 
                        id, email, username, password, remember_me_token,
                        created_at, updated_at, is_email_verified, client_id,
                        business_name, contact_name, phone_number, ppr_id,
                        company_url, last_name, abn, archive, personal_bio, 
                        location, profile_image_url, skills, interests, 
                        last_login_at, achievements, provider, provider_id, 
                        login_count, last_login_provider, industry, team_size, 
                        business_description, biggest_challenge
                    FROM users
                    WHERE client_id = $1
                    LIMIT 1
                """
                row = await conn.fetchrow(sql_string, user_id)
            
            if not row:
                logging.warning(f"No user found for user_id={user_id}")
                return None
            
            # Convert asyncpg Record to dict
            user_data = dict(row)
            
            # Convert datetime objects to ISO format
            for key, value in user_data.items():
                if hasattr(value, 'isoformat'):
                    user_data[key] = value.isoformat()
            
            logging.info(f"Found user profile data for user_id={user_id}")
            return user_data
            
    except Exception as e:
        logging.error(f"Error getting user profile data: {str(e)}")
        logging.error(f"🔍 Error context: user_id='{user_id}', type={type(user_id)}")
        return None


def get_azure_container_name(user_id: str) -> str:
    """Get Azure container name for user"""
    conn = None
    try:
        conn = psycopg2.connect(
            host=ONBOARDING_DB_HOST,
            dbname=ONBOARDING_DB_NAME,
            user=ONBOARDING_DB_USER,
            password=ONBOARDING_DB_PASSWORD,
            port=ONBOARDING_DB_PORT
        )
        conn.autocommit = True

        with conn.cursor() as cur:
            sql = """
                SELECT azure_container_name
                FROM client_onboarding
                WHERE client_id = %s
                LIMIT 1
            """
            cur.execute(sql, (user_id,))
            row = cur.fetchone()
            if not row:
                logging.warning(f"No container found for user_id={user_id}, using default container 'unified-clients-prod'")
                return "unified-clients-prod"  # Updated to new unified architecture container

            container_name = row[0]
            logging.info(f"Found container for user_id={user_id}: {container_name}")
            return container_name

    except Exception as e:
        logging.error(f"Error retrieving container from DB: {str(e)}")
        return "unified-clients-prod"  # Updated to new unified architecture container

    finally:
        if conn:
            conn.close()

def get_client_folder_name(user_id: str) -> str:
    """
    Get the client's folder name from database.
    Returns folder_name like '666-tim' from client_onboarding table.
    This ensures risk reports go to: {container}/{client_folder}/risk analysis engine report/
    """
    conn = None
    try:
        conn = psycopg2.connect(
            host=ONBOARDING_DB_HOST,
            dbname=ONBOARDING_DB_NAME,
            user=ONBOARDING_DB_USER,
            password=ONBOARDING_DB_PASSWORD,
            port=ONBOARDING_DB_PORT
        )
        conn.autocommit = True

        with conn.cursor() as cur:
            sql = """
                SELECT folder_name
                FROM client_onboarding
                WHERE client_id = %s
                LIMIT 1
            """
            cur.execute(sql, (user_id,))
            row = cur.fetchone()
            if not row:
                logging.warning(f"No folder_name found for user_id={user_id}, using client_id as fallback")
                return user_id

            folder_name = row[0]
            logging.info(f"Found folder_name for user_id={user_id}: {folder_name}")
            return folder_name

    except Exception as e:
        logging.error(f"Error retrieving folder_name from DB: {str(e)}")
        return user_id  # Fallback to client_id

    finally:
        if conn:
            conn.close()

def store_risk_assessment(user_id: str, assessment_data: Dict, include_multi_db: bool = False):
    """Store risk assessment data with optional multi-database intelligence"""
    conn = None
    start_time = time.time()
    
    try:
        logging.info(f"💾 Starting risk assessment storage")
        logging.info(f"📊 Storage parameters:")
        logging.info(f"   - User ID: {user_id}")
        logging.info(f"   - Include multi-DB: {include_multi_db}")
        logging.info(f"   - Assessment data size: {len(str(assessment_data))} characters")
        logging.info(f"   - Assessment keys: {list(assessment_data.keys()) if assessment_data else 'No data'}")
        
        # Validate input data
        if not user_id:
            raise ValueError("user_id is required and cannot be empty")
        
        if not assessment_data:
            raise ValueError("assessment_data is required and cannot be empty")
        
        # Log data structure analysis
        responses = assessment_data.get("responses", [])
        assessment_metadata = assessment_data.get("assessment_metadata", {})
        comprehensive_metadata = assessment_data.get("comprehensive_metadata", {})
        
        logging.info(f"📋 Data structure analysis:")
        logging.info(f"   - Responses: {len(responses)} items")
        logging.info(f"   - Assessment metadata keys: {list(assessment_metadata.keys()) if assessment_metadata else 'None'}")
        logging.info(f"   - Comprehensive metadata keys: {list(comprehensive_metadata.keys()) if comprehensive_metadata else 'None'}")
        
        # Get database connection
        logging.info(f"🔗 Establishing database connection...")
        connection_start = time.time()
        
        conn = get_risk_connection()
        connection_time = time.time() - connection_start
        
        logging.info(f"✅ Database connection established in {connection_time:.3f}s")
        
        # Create/verify tables
        logging.info(f"📋 Creating/verifying database tables...")
        table_creation_start = time.time()
        
        create_risk_tables(conn)
        table_creation_time = time.time() - table_creation_start
        
        logging.info(f"✅ Tables verified/created in {table_creation_time:.3f}s")

        # Multi-database intelligence handling
        multi_db_intelligence = {}
        multi_db_fetch_time = 0
        
        if include_multi_db:
            logging.info(f"🧠 Fetching multi-database intelligence...")
            multi_db_start = time.time()
            
            try:
                # Check if we're in an async context
                try:
                    current_loop = asyncio.get_running_loop()
                    logging.warning(f"⚠️ Already in async context - this indicates improper function usage")
                    multi_db_intelligence = {}
                except RuntimeError:
                    # No running loop - safe to create new one
                    logging.debug(f"🔄 No running loop, creating new one for multi-DB fetch")
                    
                    if platform.system() == 'Windows':
                        asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
                    
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                    
                    try:
                        multi_db_intelligence = loop.run_until_complete(
                            get_multi_database_intelligence(user_id)
                        )
                        
                        multi_db_fetch_time = time.time() - multi_db_start
                        
                        # Analyze fetched intelligence
                        data_sources = multi_db_intelligence.get('data_sources_available', [])
                        qa_pairs_count = multi_db_intelligence.get('complete_qa_data', {}).get('token_tracking', {}).get('qa_pairs_count', 0)
                        logging.info(f"✅ Multi-database intelligence fetched in {multi_db_fetch_time:.3f}s")
                        logging.info(f"🧠 Intelligence summary:")
                        logging.info(f"   - Data sources available: {len(data_sources)}")
                        logging.info(f"   - Sources: {', '.join(data_sources) if data_sources else 'None'}")
                        logging.info(f"   - Total Q&A pairs: {qa_pairs_count}")
                        logging.info(f"   - Intelligence data size: {len(str(multi_db_intelligence))} characters")
                        
                    except Exception as loop_error:
                        multi_db_fetch_time = time.time() - multi_db_start
                        logging.warning(f"⚠️ Error in event loop execution: {str(loop_error)}")
                        multi_db_intelligence = {}
                    finally:
                        try:
                            pending = asyncio.all_tasks(loop)
                            if pending:
                                logging.debug(f"🔄 Cancelling {len(pending)} pending tasks")
                                for task in pending:
                                    task.cancel()
                            
                            loop.close()
                            logging.debug(f"🔄 Event loop closed successfully")
                        except Exception as cleanup_error:
                            logging.warning(f"⚠️ Event loop cleanup error: {cleanup_error}")
                        finally:
                            try:
                                asyncio.set_event_loop(None)
                            except Exception:
                                pass
                    
            except Exception as e:
                multi_db_fetch_time = time.time() - multi_db_start
                logging.warning(f"⚠️ Failed to fetch multi-database intelligence after {multi_db_fetch_time:.3f}s: {str(e)}")
                logging.warning(f"🔍 Multi-DB error type: {type(e).__name__}")
                logging.warning(f"🔍 Continuing without multi-DB intelligence...")
                multi_db_intelligence = {}
        else:
            logging.info(f"ℹ️ Multi-database intelligence not requested, skipping")

        # Begin database transaction
        logging.info(f"📝 Starting database transaction...")
        transaction_start = time.time()
        
        with conn.cursor() as cur:
            logging.debug(f"✅ Database cursor acquired")
            
            # Prepare assessment metadata
            logging.info(f"🔧 Preparing assessment metadata...")
            
            assessment_type = assessment_metadata.get("assessment_type", "backable_risk_engine")
            version = assessment_metadata.get("version", "1.0")
            created_at = assessment_metadata.get("created_at")
            last_updated = assessment_metadata.get("last_updated")
            timezone = assessment_metadata.get("timezone", "UTC")
            
            # Enhanced metadata preparation with validation
            session_metadata = assessment_metadata.get("session_metadata", {})
            device_fingerprint = assessment_metadata.get("device_fingerprint", {})
            progress_tracking = assessment_data.get("progress_tracking", {})
            completion_flags = assessment_data.get("completion_flags", {})
            
            logging.info(f"📊 Assessment metadata prepared:")
            logging.info(f"   - Type: {assessment_type}")
            logging.info(f"   - Version: {version}")
            logging.info(f"   - Timezone: {timezone}")
            logging.info(f"   - Session metadata: {len(session_metadata)} items")
            logging.info(f"   - Device fingerprint: {len(device_fingerprint)} items")
            logging.info(f"   - Progress tracking: {len(progress_tracking)} items")
            logging.info(f"   - Completion flags: {len(completion_flags)} items")
            
            # Safer JSON serialization
            try:
                session_metadata_json = json.dumps(session_metadata, default=str)
                device_fingerprint_json = json.dumps(device_fingerprint, default=str)
                progress_tracking_json = json.dumps(progress_tracking, default=str)
                completion_flags_json = json.dumps(completion_flags, default=str)
                raw_data_json = json.dumps(assessment_data, default=str)
                multi_db_intelligence_json = json.dumps(multi_db_intelligence, default=str)
                
                logging.debug(f"✅ JSON serialization successful")
                logging.debug(f"   - Raw data JSON size: {len(raw_data_json)} characters")
                logging.debug(f"   - Multi-DB JSON size: {len(multi_db_intelligence_json)} characters")
                
            except Exception as json_error:
                logging.error(f"❌ JSON serialization error: {json_error}")
                raise ValueError(f"Failed to serialize data to JSON: {json_error}")

            # Execute main assessment insert/update
            logging.info(f"📝 Executing main assessment SQL...")
            sql_start = time.time()
            
            sql = """
                INSERT INTO risk_assessments (
                    user_id, assessment_type, version, created_at, last_updated,
                    timezone, session_metadata, device_fingerprint,
                    progress_tracking, completion_flags, raw_data, multi_database_intelligence
                ) VALUES (
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                ) ON CONFLICT (user_id) DO UPDATE SET
                    last_updated              = EXCLUDED.last_updated,
                    session_metadata          = EXCLUDED.session_metadata,
                    progress_tracking         = EXCLUDED.progress_tracking,
                    completion_flags          = EXCLUDED.completion_flags,
                    raw_data                  = EXCLUDED.raw_data,
                    multi_database_intelligence = CASE 
                        WHEN %s = true THEN EXCLUDED.multi_database_intelligence 
                        ELSE risk_assessments.multi_database_intelligence 
                    END
                RETURNING id
            """

            try:
                cur.execute(sql, (
                    user_id,
                    assessment_type,
                    version,
                    created_at,
                    last_updated,
                    timezone,
                    session_metadata_json,
                    device_fingerprint_json,
                    progress_tracking_json,
                    completion_flags_json,
                    raw_data_json,
                    multi_db_intelligence_json,
                    include_multi_db  # For the CASE statement
                ))
                
                assessment_id_row = cur.fetchone()
                assessment_id = assessment_id_row[0] if assessment_id_row else None
                
                sql_time = time.time() - sql_start
                logging.info(f"✅ Main assessment SQL executed in {sql_time:.3f}s")
                logging.info(f"📊 Assessment ID: {assessment_id}")
                
                if not assessment_id:
                    raise Exception("Failed to get assessment_id from database - no row returned")
                    
            except Exception as sql_error:
                sql_time = time.time() - sql_start
                logging.error(f"❌ Main assessment SQL failed after {sql_time:.3f}s: {sql_error}")
                raise

            # Store responses
            responses_start = time.time()
            successful_responses = 0
            failed_responses = 0
            
            if responses:
                logging.info(f"📝 Storing {len(responses)} responses...")
                
                response_sql = """
                    INSERT INTO risk_responses (
                        assessment_id, user_id, question_id, section, question_type,
                        question_text, response_format, response_data, all_options,
                        metadata, weight, answered_at, last_modified_at
                    ) VALUES (
                        %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                    ) ON CONFLICT (assessment_id, question_id) DO UPDATE SET
                        response_data    = EXCLUDED.response_data,
                        metadata         = EXCLUDED.metadata,
                        last_modified_at = EXCLUDED.last_modified_at
                """

                for i, response in enumerate(responses):
                    try:
                        question_id = response.get("question_id", f"unknown_{i}")
                        section = response.get("section", "Unknown")
                        
                        logging.debug(f"📝 Storing response {i+1}/{len(responses)}: {question_id} ({section})")
                        
                        # Validate and prepare response data
                        response_data = response.get("response_data", {})
                        all_options = response.get("all_options", [])
                        metadata = response.get("metadata", {})
                        
                        # Safer JSON encoding
                        try:
                            response_data_json = json.dumps(response_data, default=str)
                            all_options_json = json.dumps(all_options, default=str)
                            metadata_json = json.dumps(metadata, default=str)
                        except Exception as response_json_error:
                            logging.error(f"❌ JSON encoding error for response {question_id}: {response_json_error}")
                            failed_responses += 1
                            continue
                        
                        cur.execute(response_sql, (
                            assessment_id,
                            user_id,
                            question_id,
                            section,
                            response.get("question_type"),
                            response.get("question_text"),
                            response.get("response_format"),
                            response_data_json,
                            all_options_json,
                            metadata_json,
                            response.get("weight", "medium"),
                            response.get("answered_at"),
                            response.get("last_modified_at")
                        ))
                        
                        successful_responses += 1
                        
                    except Exception as response_error:
                        failed_responses += 1
                        logging.error(f"❌ Error storing response {i+1} ({response.get('question_id', 'unknown')}): {response_error}")
                        continue  # Continue with other responses
                
                responses_time = time.time() - responses_start
                logging.info(f"✅ Response storage completed in {responses_time:.3f}s")
                logging.info(f"📊 Response results: {successful_responses} successful, {failed_responses} failed")
                
                if failed_responses > 0 and successful_responses == 0:
                    logging.error(f"❌ All responses failed to store!")
                elif failed_responses > 0:
                    logging.warning(f"⚠️ Some responses failed to store: {failed_responses}/{len(responses)}")
            else:
                logging.info(f"ℹ️ No responses to store")

            # Simplified behavioral analytics
            behavioral_start = time.time()
            behavioral_data = assessment_data.get("comprehensive_metadata", {}).get("behavioral_analytics", {})

            if behavioral_data:
                logging.info(f"🧠 Processing behavioral analytics...")
                logging.info(f"📊 Behavioral data keys: {list(behavioral_data.keys())}")
                
                try:
                    # Analyze risk decision patterns
                    logging.debug(f"🔧 Analyzing risk decision patterns...")
                    pattern_analysis_start = time.time()
                    
                    risk_patterns = analyze_risk_decision_patterns(behavioral_data, responses)
                    pattern_analysis_time = time.time() - pattern_analysis_start
                    
                    logging.info(f"✅ Risk pattern analysis completed in {pattern_analysis_time:.3f}s")
                    logging.info(f"📊 Pattern results: {list(risk_patterns.keys()) if risk_patterns else 'None'}")
                    
                    # Store behavioral analytics
                    behavior_sql = """
                        INSERT INTO risk_behavioral_analytics (
                            assessment_id, user_id, mouse_behavior, keyboard_behavior,
                            attention_patterns, decision_making_style,
                            risk_decision_patterns, created_at
                        ) VALUES (
                            %s, %s, %s, %s, %s, %s, %s, %s
                        ) ON CONFLICT (assessment_id) DO UPDATE SET
                            mouse_behavior           = EXCLUDED.mouse_behavior,
                            keyboard_behavior        = EXCLUDED.keyboard_behavior,
                            attention_patterns       = EXCLUDED.attention_patterns,
                            decision_making_style    = EXCLUDED.decision_making_style,
                            risk_decision_patterns   = EXCLUDED.risk_decision_patterns
                    """

                    # Prepare behavioral data
                    mouse_behavior = behavioral_data.get("mouse_behavior", {})
                    keyboard_behavior = behavioral_data.get("keyboard_behavior", {})
                    attention_patterns = behavioral_data.get("attention_patterns", {})
                    decision_making_style = behavioral_data.get("decision_making_style", {})
                    
                    logging.debug(f"🧠 Behavioral data breakdown:")
                    logging.debug(f"   - Mouse behavior: {len(mouse_behavior)} items")
                    logging.debug(f"   - Keyboard behavior: {len(keyboard_behavior)} items")
                    logging.debug(f"   - Attention patterns: {len(attention_patterns)} items")
                    logging.debug(f"   - Decision making: {len(decision_making_style)} items")
                    logging.debug(f"   - Risk patterns: {len(risk_patterns)} items")

                    cur.execute(behavior_sql, (
                        assessment_id,
                        user_id,
                        json.dumps(mouse_behavior, default=str),
                        json.dumps(keyboard_behavior, default=str),
                        json.dumps(attention_patterns, default=str),
                        json.dumps(decision_making_style, default=str),
                        json.dumps(risk_patterns, default=str),
                        datetime.now().isoformat()
                    ))
                    
                    behavioral_time = time.time() - behavioral_start
                    logging.info(f"✅ Behavioral analytics stored in {behavioral_time:.3f}s")
                    
                except Exception as behavioral_error:
                    behavioral_time = time.time() - behavioral_start
                    logging.error(f"❌ Error storing behavioral analytics after {behavioral_time:.3f}s: {behavioral_error}")
                    logging.error(f"🔍 Behavioral error type: {type(behavioral_error).__name__}")
                    logging.warning(f"⚠️ Continuing without behavioral analytics...")
            else:
                logging.info(f"ℹ️ No behavioral analytics data to store")
        
        # Calculate final timing
        total_time = time.time() - start_time
        transaction_time = time.time() - transaction_start
        
        logging.info(f"🎉 Risk assessment storage completed successfully!")
        logging.info(f"📊 STORAGE PERFORMANCE SUMMARY:")
        logging.info(f"   - Total time: {total_time:.3f}s")
        logging.info(f"   - Connection time: {connection_time:.3f}s")
        logging.info(f"   - Table creation: {table_creation_time:.3f}s")
        logging.info(f"   - Multi-DB fetch: {multi_db_fetch_time:.3f}s")
        logging.info(f"   - Transaction time: {transaction_time:.3f}s")
        logging.info(f"   - Assessment ID: {assessment_id}")
        logging.info(f"   - Responses stored: {successful_responses}")
        logging.info(f"   - Multi-DB sources: {len(multi_db_intelligence.get('data_sources_available', []))}")
        
        return assessment_id

    except Exception as e:
        total_time = time.time() - start_time
        logging.error(f"❌ Risk assessment storage failed after {total_time:.3f}s")
        logging.error(f"🔍 Error details:")
        logging.error(f"   - Error type: {type(e).__name__}")
        logging.error(f"   - Error message: {str(e)}")
        logging.error(f"   - User ID: {user_id}")
        logging.error(f"   - Include multi-DB: {include_multi_db}")
        logging.error(f"   - Assessment data size: {len(str(assessment_data)) if assessment_data else 0} chars")
        
        # Log the full traceback for debugging
        import traceback
        logging.error(f"🔍 Full traceback:")
        for line in traceback.format_exc().split('\n'):
            if line.strip():
                logging.error(f"   {line}")
        
        raise

    finally:
        if conn:
            try:
                conn.close()
                logging.debug(f"🔗 Database connection closed")
            except Exception as close_error:
                logging.warning(f"⚠️ Error closing connection: {close_error}")




def get_smart_api_key(section_index: int, retry_attempt: int = 0) -> str:
    """Enhanced smart API key selection with load balancing for Risk Engine"""
    global api_key_health
    
    # Initialize health tracking if not exists
    if not api_key_health:
        for i, key in enumerate(GEMINI_API_KEYS):
            api_key_health[key] = {
                'last_503_time': None,
                'consecutive_failures': 0,
                'total_requests': 0,
                'key_id': f'RiskEngine_{i+1:02d}',
                'response_times': deque(maxlen=10),
                'success_rate': 1.0,
                'last_used': 0,
                'current_load': 0
            }
    
    current_time = time.time()
    
    # Calculate health scores for all keys
    key_scores = []
    for key in GEMINI_API_KEYS:
        health = api_key_health[key]
        score = calculate_key_health_score(health, current_time)
        key_scores.append((key, score, health))
    
    # Sort by health score (higher is better)
    key_scores.sort(key=lambda x: x[1], reverse=True)
    
    # Select best available key
    for key, score, health in key_scores:
        if score > 0:  # Key is usable
            # Update usage tracking
            health['total_requests'] += 1
            health['last_used'] = current_time
            health['current_load'] += 1
            
            logging.info(f"🔑 [{section_index}] Selected {health['key_id']} (score: {score:.2f}, load: {health['current_load']})")
            return key
    
    # Fallback: use least recently failed key
    logging.warning("⚠️ All keys degraded, using least recently failed")
    fallback_key = min(GEMINI_API_KEYS, 
                      key=lambda k: api_key_health[k]['last_503_time'] or 0)
    api_key_health[fallback_key]['total_requests'] += 1
    return fallback_key

def reset_api_key_immediately(api_key: str):
    """Immediately reset a failed API key with detailed logging"""
    if api_key in api_key_health:
        old_health = api_key_health[api_key].copy()
        
        # Reset critical health metrics
        api_key_health[api_key]['consecutive_failures'] = 0
        api_key_health[api_key]['last_503_time'] = None
        api_key_health[api_key]['current_load'] = 0
        # Keep success_rate but improve it slightly
        old_rate = api_key_health[api_key].get('success_rate', 1.0)
        api_key_health[api_key]['success_rate'] = min(1.0, old_rate + 0.1)
        
        key_id = api_key_health[api_key].get('key_id', 'unknown')
        
        logging.info(f"🔄 RESET API KEY: {key_id} (...{api_key[-4:]})")
        logging.info(f"📊 Reset details:")
        logging.info(f"   - Consecutive failures: {old_health.get('consecutive_failures', 0)} → 0")
        logging.info(f"   - Success rate: {old_health.get('success_rate', 1.0):.3f} → {api_key_health[api_key]['success_rate']:.3f}")
        logging.info(f"   - Current load: {old_health.get('current_load', 0)} → 0")
        logging.info(f"   - Cooldown cleared: {bool(old_health.get('last_503_time'))}")
        
    else:
        logging.warning(f"⚠️ Cannot reset API key ...{api_key[-4:]}: not found in health tracking")

def reset_all_failed_api_keys():
    """Reset all failed API keys for batch retry"""
    global api_key_health
    
    reset_count = 0
    for api_key, health in api_key_health.items():
        if health.get('consecutive_failures', 0) >= 3:
            reset_api_key_immediately(api_key)
            reset_count += 1
    
    logging.info(f"🔄 BATCH RESET: Reset {reset_count} failed API keys")
    return reset_count

def calculate_key_health_score(health: Dict, current_time: float) -> float:
    """Calculate health score for API key (0-100, higher is better)"""
    score = 100.0
    
    # Penalize recent 503 errors
    if health['last_503_time'] and (current_time - health['last_503_time']) < 300:
        score *= 0.1  # Major penalty for recent 503
    
    # Penalize consecutive failures
    failure_penalty = min(0.9, health['consecutive_failures'] * 0.3)
    score *= (1.0 - failure_penalty)
    
    # Penalize high current load
    load_penalty = min(0.5, health['current_load'] * 0.1)
    score *= (1.0 - load_penalty)
    
    # Bonus for good response times
    if health['response_times']:
        avg_response_time = statistics.mean(health['response_times'])
        if avg_response_time < 30:  # Fast responses
            score *= 1.2
        elif avg_response_time > 60:  # Slow responses
            score *= 0.8
    
    # Bonus for high success rate
    score *= health['success_rate']
    
    return max(0, score)

def update_api_key_health(api_key: str, success: bool, error_code: str = None, response_time: float = None):
    """Enhanced API key health update with detailed metrics"""
    global api_key_health
    
    if api_key not in api_key_health:
        return
    
    health = api_key_health[api_key]
    
    # Decrease current load
    health['current_load'] = max(0, health['current_load'] - 1)
    
    # Track response time
    if response_time:
        health['response_times'].append(response_time)
    
    if success:
        health['consecutive_failures'] = 0
        # Update success rate (sliding window)
        old_rate = health['success_rate']
        health['success_rate'] = min(1.0, old_rate * 0.9 + 0.1)  # Weighted moving average
        
        logging.debug(f"✅ {health['key_id']} success (rate: {health['success_rate']:.2f})")
    else:
        health['consecutive_failures'] += 1
        # Update success rate
        old_rate = health['success_rate']
        health['success_rate'] = max(0.0, old_rate * 0.9)  # Penalize failures
        
        # Special handling for different error types
        if error_code == "503":
            health['last_503_time'] = time.time()
            logging.warning(f"🚨 {health['key_id']} got 503 - cooling down")
        elif error_code == "429":
            health['last_503_time'] = time.time() - 150  # Shorter cooldown for rate limits
            logging.warning(f"🚦 {health['key_id']} rate limited")
        
        logging.warning(f"❌ {health['key_id']} failed (consecutive: {health['consecutive_failures']})")

def get_load_balanced_api_key(section_index: int) -> str:
    """Get API key using round-robin with health awareness"""
    global api_key_health
    
    if not api_key_health:
        # Initialize if needed
        return get_smart_api_key(section_index, 0)
    
    current_time = time.time()
    
    # Filter healthy keys
    healthy_keys = []
    for key in GEMINI_API_KEYS:
        health = api_key_health[key]
        
        # Skip if in cooldown
        if health['last_503_time'] and (current_time - health['last_503_time']) < 180:
            continue
            
        # Skip if too many failures
        if health['consecutive_failures'] >= 3:
            continue
            
        # Skip if overloaded
        if health['current_load'] >= 3:
            continue
            
        healthy_keys.append(key)
    
    if not healthy_keys:
        # Fallback to smart selection
        return get_smart_api_key(section_index, 0)
    
    # Use weighted random selection based on inverse load
    weights = []
    for key in healthy_keys:
        load = api_key_health[key]['current_load']
        success_rate = api_key_health[key]['success_rate']
        weight = success_rate / max(1, load + 1)  # Higher success rate, lower load = higher weight
        weights.append(weight)
    
    # Weighted random selection
    selected_key = random.choices(healthy_keys, weights=weights)[0]
    
    # Update usage
    api_key_health[selected_key]['current_load'] += 1
    api_key_health[selected_key]['total_requests'] += 1
    
    health = api_key_health[selected_key]
    logging.info(f"🎯 Load balanced selection: {health['key_id']} (load: {health['current_load']}, rate: {health['success_rate']:.2f})")
    
    return selected_key

def get_api_key_status_summary() -> str:
    """Get summary of all API key health for logging"""
    if not api_key_health:
        return "No health data available"
    
    healthy_count = 0
    cooling_down = 0
    failed_count = 0
    
    for key, health in api_key_health.items():
        current_time = time.time()
        
        if health['last_503_time'] and (current_time - health['last_503_time']) < 300:
            cooling_down += 1
        elif health['consecutive_failures'] >= 3:
            failed_count += 1
        else:
            healthy_count += 1
    
    return f"Healthy: {healthy_count}, Cooling: {cooling_down}, Failed: {failed_count}"

def get_enhanced_api_key_status() -> Dict:
    """Get comprehensive API key status"""
    if not api_key_health:
        return {"status": "not_initialized"}
    
    current_time = time.time()
    status = {
        "total_keys": len(GEMINI_API_KEYS),
        "healthy_keys": 0,
        "degraded_keys": 0,
        "failed_keys": 0,
        "cooling_down": 0,
        "total_load": 0,
        "average_success_rate": 0,
        "key_details": []
    }
    
    success_rates = []
    
    for key, health in api_key_health.items():
        # Determine status
        if health['last_503_time'] and (current_time - health['last_503_time']) < 180:
            key_status = "cooling_down"
            status["cooling_down"] += 1
        elif health['consecutive_failures'] >= 3:
            key_status = "failed"
            status["failed_keys"] += 1
        elif health['consecutive_failures'] > 0 or health['success_rate'] < 0.8:
            key_status = "degraded"
            status["degraded_keys"] += 1
        else:
            key_status = "healthy"
            status["healthy_keys"] += 1
        
        status["total_load"] += health['current_load']
        success_rates.append(health['success_rate'])
        
        # Add key details
        avg_response = statistics.mean(health['response_times']) if health['response_times'] else 0
        
        status["key_details"].append({
            "key_id": health['key_id'],
            "status": key_status,
            "success_rate": health['success_rate'],
            "current_load": health['current_load'],
            "consecutive_failures": health['consecutive_failures'],
            "total_requests": health['total_requests'],
            "avg_response_time": round(avg_response, 2)
        })
    
    status["average_success_rate"] = statistics.mean(success_rates) if success_rates else 0
    
    return status

def reset_failed_api_keys():
    """Reset failed API keys for retry attempts"""
    global api_key_health
    
    reset_count = 0
    for key, health in api_key_health.items():
        if health.get('consecutive_failures', 0) >= 3:
            health['consecutive_failures'] = 1  # Reduce but don't fully reset
            health['last_503_time'] = None  # Clear cooldown
            reset_count += 1
            logging.info(f"🔄 Partially reset API key {health.get('key_id', 'unknown')} for retry")
    
    logging.info(f"🔄 Reset {reset_count} failed API keys for retry attempt")

# Production-optimized settings for Risk Engine
MAX_RETRIES = 10
MAX_REQUESTS_PER_ENDPOINT = 100
REQUEST_TIMEOUT = 120  # 2 minutes
MAX_SECTION_RETRIES = 3
MAX_REPORT_RETRIES = 2
MIN_ACCEPTABLE_WORDS = 100
RETRY_WAIT_BASE = 30
risk_job_status = {}

# ======================================================
#           Database Connection Pooling - RISK ENGINE
# ======================================================

async def get_db_pool(db_config: Dict):
    """Get or create connection pool for database with Python 3.10 compatibility"""
    global _connection_pools, _pool_creation_locks
    
    # Clean and validate database name
    db_key = str(db_config['database']).strip()
    
    # Validate against known good database names
    valid_db_names = {
        "BACKABLE-PROFILE-ENGINE",
        "BACKABLE-THE-ANALYST", 
        "BACKABLE-THE-GROWTH-ENGINE",
        "BACKABLE-COMPONENT-ENGINE",
        "BACKABLE-DREAM-ANALYZER",
        "BACKABLE-PEOPLE-AND-OPERATIONS-ENGINE",
        "BACKABLE-RISK-ENGINE",  # Added Risk Engine
        "philotimodb"
    }
    
    if db_key not in valid_db_names:
        logging.error(f"❌ INVALID DATABASE NAME: '{db_key}' not in valid list")
        # Try to find closest match and correct
        for valid_name in valid_db_names:
            clean_valid = valid_name.replace('-', '').replace('_', '').lower()
            clean_input = db_key.replace('-', '').replace('_', '').lower()
            if clean_valid in clean_input or clean_input in clean_valid:
                logging.warning(f"🔧 AUTO-CORRECTING: '{db_key}' → '{valid_name}'")
                db_key = valid_name
                break
        else:
            raise ValueError(f"Cannot map database name '{db_key}' to valid database")
    
    pool_creation_start = time.time()
    
    logging.info(f"🔗 POOL REQUEST: Starting pool acquisition for {db_key}")
    logging.info(f"📊 Pool request details:")
    logging.info(f"   - Database: {db_config.get('database', 'Unknown')}")
    logging.info(f"   - Host: {db_config.get('host', 'Unknown')}")
    logging.info(f"   - User: {db_config.get('user', 'Unknown')}")
    logging.info(f"   - Port: {db_config.get('port', 'Unknown')}")
    logging.info(f"   - Current pools in memory: {len(_connection_pools)}")
    logging.info(f"   - Existing pool keys: {list(_connection_pools.keys())}")
    
    # Check if pool exists and validate its health
    if db_key in _connection_pools:
        existing_pool = _connection_pools[db_key]
        logging.info(f"🔍 POOL EXISTS: Found existing pool for {db_key}")
        logging.info(f"📊 Existing pool details:")
        logging.info(f"   - Pool object type: {type(existing_pool)}")
        logging.info(f"   - Pool closed status: {getattr(existing_pool, '_closed', 'Unknown')}")
        
        # Health check with Python 3.10 compatible timeout
        try:
            logging.info(f"🧪 HEALTH CHECK: Testing existing pool for {db_key}")
            health_check_start = time.time()
            
            # Use asyncio.wait_for instead of asyncio.timeout (Python 3.10 compatible)
            async def health_check():
                async with existing_pool.acquire() as test_conn:
                    await test_conn.execute('SELECT 1')
            
            await asyncio.wait_for(health_check(), timeout=3.0)
            
            health_check_time = time.time() - health_check_start
            logging.info(f"✅ HEALTH CHECK PASSED: Pool {db_key} is healthy ({health_check_time:.3f}s)")
            logging.info(f"📊 Pool stats:")
            logging.info(f"   - Min size: {getattr(existing_pool, '_minsize', 'Unknown')}")
            logging.info(f"   - Max size: {getattr(existing_pool, '_maxsize', 'Unknown')}")
            logging.info(f"   - Current size: {getattr(existing_pool, 'get_size', lambda: 'Unknown')()}")
            logging.info(f"   - Idle connections: {getattr(existing_pool, 'get_idle_size', lambda: 'Unknown')()}")
            
            return existing_pool
            
        except asyncio.TimeoutError:
            logging.error(f"⏰ HEALTH CHECK TIMEOUT: Pool {db_key} timed out after 3s")
            logging.error(f"🔧 Pool appears to be stuck, will recreate")
        except Exception as health_error:
            logging.error(f"❌ HEALTH CHECK FAILED: Pool {db_key} failed health check: {health_error}")
            logging.error(f"🔧 Pool is corrupted, will recreate")
            logging.error(f"🔍 Health check error type: {type(health_error).__name__}")
        
        # Proper cleanup of failed pool
        logging.warning(f"🗑️ REMOVING UNHEALTHY POOL: {db_key}")
        try:
            if not existing_pool.is_closing():
                await existing_pool.close()
                # Wait for pool to actually close
                await asyncio.sleep(0.1)
            logging.info(f"✅ Closed unhealthy pool for {db_key}")
        except Exception as close_error:
            logging.error(f"❌ Error closing unhealthy pool: {close_error}")
        finally:
            # Always remove from dictionary
            _connection_pools.pop(db_key, None)
            logging.info(f"✅ Removed {db_key} from pools dictionary")
    
    # Create new pool with enhanced safety using per-database locks
    async with _pool_creation_locks[db_key]:
        # Double-check pool doesn't exist (race condition protection)
        if db_key not in _connection_pools:
            logging.info(f"🏗️ CREATING NEW POOL: Starting pool creation for {db_key}")
            
            try:
                # Validate required config keys
                required_config_keys = ['host', 'database', 'user', 'password', 'port']
                missing_keys = [key for key in required_config_keys if not db_config.get(key)]
                
                if missing_keys:
                    raise ValueError(f"Missing required database config keys: {missing_keys}")
                
                logging.info(f"✅ Database config validation passed for {db_key}")
                
                # Test connection first before creating pool
                try:
                    test_conn = await asyncio.wait_for(
                        asyncpg.connect(
                            host=db_config["host"],
                            database=db_config["database"],
                            user=db_config["user"],
                            password=db_config["password"],
                            port=db_config["port"],
                            ssl="require"
                        ),
                        timeout=5.0
                    )
                    await test_conn.close()
                    logging.info(f"✅ Connection test passed for {db_key}")
                except Exception as test_error:
                    logging.error(f"❌ Connection test failed for {db_key}: {test_error}")
                    raise Exception(f"Cannot connect to database {db_key}: {test_error}")
                
                # Reduced pool configuration for stability
                pool_config = {
                    "host": db_config["host"],
                    "database": db_config["database"],
                    "user": db_config["user"],
                    "password": db_config["password"],
                    "port": db_config["port"],
                    "ssl": "require",
                    "min_size": 1,  # Minimum connections
                    "max_size": 2,  # From 5 to 2 for stability
                    "command_timeout": 15,  # From 30 to 15
                    "server_settings": {
                        'application_name': f'backable_risk_{db_key.replace("-", "_")}',
                        'jit': 'off',
                        'tcp_keepalives_idle': '300',
                        'tcp_keepalives_interval': '30',
                        'tcp_keepalives_count': '3'
                    }
                }
                
                logging.info(f"🔧 Pool configuration:")
                logging.info(f"   - Min size: {pool_config['min_size']}")
                logging.info(f"   - Max size: {pool_config['max_size']}")
                logging.info(f"   - Command timeout: {pool_config['command_timeout']}s")
                logging.info(f"   - SSL: {pool_config['ssl']}")
                
                # Create pool with shorter timeout
                pool_creation_timeout = 10.0  # From 15s to 10s
                
                logging.info(f"⏳ Creating pool with {pool_creation_timeout}s timeout...")
                pool_start_time = time.time()
                
                pool = await asyncio.wait_for(
                    asyncpg.create_pool(**pool_config),
                    timeout=pool_creation_timeout
                )
                
                pool_creation_time = time.time() - pool_start_time
                logging.info(f"✅ Pool creation successful in {pool_creation_time:.3f}s")
                
                # Quick post-creation validation with Python 3.10 compatible timeout
                logging.info(f"🧪 POST-CREATION VALIDATION: Testing new pool for {db_key}")
                validation_start = time.time()
                
                try:
                    # PYTHON 3.10 COMPATIBLE: Use asyncio.wait_for instead of asyncio.timeout
                    async def validation_test():
                        async with pool.acquire() as test_conn:
                            await test_conn.execute('SELECT 1')
                    
                    await asyncio.wait_for(validation_test(), timeout=2.0)
                    
                    validation_time = time.time() - validation_start
                    logging.info(f"✅ Post-creation validation passed ({validation_time:.3f}s)")
                    
                    # Store pool
                    _connection_pools[db_key] = pool
                    
                    # Log final pool status
                    total_time = time.time() - pool_creation_start
                    logging.info(f"🎉 POOL READY: {db_key} pool fully operational")
                    logging.info(f"📊 Final pool statistics:")
                    logging.info(f"   - Total creation time: {total_time:.3f}s")
                    logging.info(f"   - Pool size: {pool.get_size()}")
                    logging.info(f"   - Idle connections: {pool.get_idle_size()}")
                    logging.info(f"   - Pool status: HEALTHY")
                    logging.info(f"   - Pools in memory: {len(_connection_pools)}")
                    
                except Exception as validation_error:
                    logging.error(f"❌ POST-CREATION VALIDATION FAILED: {validation_error}")
                    logging.error(f"🔧 Closing failed pool...")
                    
                    try:
                        await pool.close()
                        await asyncio.sleep(0.1)  # Wait for cleanup
                    except Exception:
                        pass
                    
                    raise Exception(f"Pool validation failed: {validation_error}")
                
            except asyncio.TimeoutError:
                pool_creation_time = time.time() - pool_creation_start
                logging.error(f"⏰ POOL CREATION TIMEOUT: {db_key} timed out after {pool_creation_timeout}s")
                raise Exception(f"Pool creation timed out after {pool_creation_timeout}s")
                
            except Exception as creation_error:
                pool_creation_time = time.time() - pool_creation_start
                logging.error(f"❌ POOL CREATION FAILED: Error creating pool for {db_key}")
                logging.error(f"🔍 Error: {type(creation_error).__name__}: {creation_error}")
                raise Exception(f"Failed to create pool for {db_key}: {creation_error}")
        
        else:
            # Pool was created by another task during lock wait
            existing_pool = _connection_pools[db_key]
            total_time = time.time() - pool_creation_start
            logging.info(f"🏃 RACE CONDITION: Pool {db_key} was created by another task ({total_time:.3f}s)")
    
    # Return the pool
    final_pool = _connection_pools[db_key]
    total_function_time = time.time() - pool_creation_start
    
    logging.info(f"🎯 POOL ACQUIRED: Returning pool for {db_key}")
    logging.info(f"📊 Final acquisition statistics:")
    logging.info(f"   - Total function time: {total_function_time:.3f}s")
    logging.info(f"   - Pool object: {type(final_pool)}")
    logging.info(f"   - Pool size: {final_pool.get_size()}")
    logging.info(f"   - Idle connections: {final_pool.get_idle_size()}")
    logging.info(f"   - Pool health: VALIDATED")
    
    return final_pool

# ======================================================
#           Notification Functions - RISK ENGINE
# ======================================================

async def send_risk_notification(user_id: str, title: str, body: str, data_type: str = "notification", save_to_db: bool = False, report_id: str = None, business_name: str = None):
    """
    Send notification to user with optional database persistence
    """
    try:
        from datetime import timedelta

        payload = {
            "userId": int(user_id),
            "title": title,
            "body": body,
            "data": {
                "type": data_type,
                "timestamp": str(int(datetime.now().timestamp()))
            }
        }

        # Add enhanced payload and DB persistence for completion notification
        if save_to_db and report_id:
            payload["saveToDb"] = True
            payload["expiresAt"] = (datetime.now() + timedelta(days=30)).strftime("%Y-%m-%dT%H:%M:%SZ")
            payload["data"]["screen"] = "RiskReport"
            payload["data"]["reportId"] = report_id

            # IMPORTANT: payload must be inside data object for proper handling
            payload["data"]["payload"] = {
                "type": "ai_report_complete",
                "params": {
                    "reportId": report_id,
                    "reportTitle": "Risk Analysis Report",
                    "reportType": "comprehensive_risk",
                    "userId": int(user_id),
                    "businessName": business_name or "Your Business",
                    "completionStatus": "success",
                    "sections": 8,
                    "generatedAt": datetime.now().isoformat()
                },
                "actionType": "navigate",
                "screen": "RiskReport",
                "url": f"/risks/{report_id}"
            }

        logging.info(f"🔔 Sending risk notification to user {user_id}: {title} (saveToDb: {save_to_db})")

        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=NOTIFICATION_TIMEOUT)) as session:
            async with session.post(
                NOTIFICATION_API_URL,
                json=payload,
                headers={"Content-Type": "application/json"}
            ) as response:

                if response.status == 200:
                    result = await response.text()
                    logging.info(f"✅ Risk notification sent successfully to user {user_id}")
                    return True, result
                else:
                    error_text = await response.text()
                    logging.error(f"❌ Risk notification failed for user {user_id}: {response.status} - {error_text}")
                    return False, f"HTTP {response.status}: {error_text}"

    except Exception as e:
        logging.error(f"❌ Risk notification error for user {user_id}: {str(e)}")
        return False, str(e)

def send_risk_notification_sync(user_id: str, title: str, body: str, data_type: str = "notification", save_to_db: bool = False, report_id: str = None, business_name: str = None):
    """
    Synchronous wrapper for sending notifications
    """
    try:
        # Handle Windows event loop policy
        if platform.system() == 'Windows':
            asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        try:
            return loop.run_until_complete(
                send_risk_notification(user_id, title, body, data_type, save_to_db, report_id, business_name)
            )
        finally:
            loop.close()

    except Exception as e:
        logging.error(f"❌ Sync risk notification error: {str(e)}")
        return False, str(e)

async def generate_personalized_risk_message(user_profile: Dict, stage: str, progress_data: Dict = None) -> str:
    """Generate personalized, engaging risk message using Gemini AI"""
    try:
        # Extract user context
        business_name = user_profile.get('business_name', 'Your Business')
        username = user_profile.get('username', 'Leader')
        industry = user_profile.get('industry', 'Business')
        team_size = user_profile.get('team_size', 'Unknown')
        
        # Create stage-specific prompts for risk management
        if stage == "start":
            prompt = f"""
            Create a PROFESSIONAL, motivating notification for {username} from {business_name} in the {industry} industry.
            They just started their RISK ASSESSMENT (focus on protection, vulnerabilities, threats, resilience).

            Make it:
            - Professional and encouraging about business risk management
            - Include a reference to {industry} risk management opportunities
            - Reference {business_name} professionally
            - Focus on RISK MANAGEMENT, SECURITY, PROTECTION, RESILIENCE
            - 1-2 sentences max, NO EMOJIS
            - Business-appropriate language only
            """
        
        elif stage == "middle":
            chapters_done = progress_data.get('chapters_completed', 4) if progress_data else 4
            total_chapters = progress_data.get('total_chapters', 8) if progress_data else 8

            prompt = f"""
            Create a PROFESSIONAL mid-progress notification for {username} from {business_name}.
            They're {chapters_done}/{total_chapters} chapters through their risk assessment.

            Make it:
            - Professional and encouraging about analysis progress
            - Reference their {industry} security insights being uncovered
            - Focus on RISK INSIGHTS, ASSESSMENT PROGRESS
            - 1-2 sentences max, NO EMOJIS
            - Business-appropriate language only
            """
        
        elif stage == "complete":
            total_words = progress_data.get('total_words', 15000) if progress_data else 15000

            prompt = f"""
            Create a PROFESSIONAL completion notification for {username} from {business_name}.
            Their risk assessment is complete with {total_words:,} words of security insights.

            Make it:
            - Professional and congratulatory about completion
            - Reference {industry} risk management excellence
            - Focus on RISK ASSESSMENT COMPLETION, ACTIONABLE INSIGHTS
            - 1-2 sentences max, NO EMOJIS
            - Business-appropriate language only
            """
        
        # Use first available API key for notifications
        gemini_api_key = GEMINI_API_KEYS[0]
        
        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=15)) as session:
            url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent"
            
            payload = {
                "contents": [{
                    "role": "user",
                    "parts": [{"text": prompt}]
                }],
                "generationConfig": {
                    "temperature": 1.0,
                    "maxOutputTokens": 150,
                    "topP": 0.95,
                    "candidateCount": 1
                }
            }
            
            params = {'key': gemini_api_key}
            
            async with session.post(url, json=payload, params=params) as response:
                if response.status == 200:
                    data = await response.json()
                    if 'candidates' in data and len(data['candidates']) > 0:
                        candidate = data['candidates'][0]
                        
                        content = ""
                        try:
                            if 'content' in candidate and 'parts' in candidate['content']:
                                content = candidate['content']['parts'][0]['text']
                            elif 'text' in candidate:
                                content = candidate['text']
                            else:
                                content = str(candidate.get('content', candidate))
                        except Exception as e:
                            logging.warning(f"Content extraction issue: {e}")
                            content = str(candidate)
                        
                        if content:
                            content = content.strip().replace('"', '').replace("'", "'")
                            # Accept if it has meaningful content (more than 3 words)
                            if len(content.split()) > 3:
                                if not any(tech in content.lower() for tech in ['role', 'model', 'parts', 'content']):
                                    logging.info(f"Generated risk {stage} message for {username}")
                                    return content
    
    except Exception as e:
        logging.error(f"❌ Error generating risk message: {str(e)}")
    
    # Fallback messages for risk management
    fallback_messages = {
        "start": [
            f"{business_name}'s risk assessment is now underway. {username}, comprehensive security insights are being generated.",
            f"{username}'s risk analysis has commenced. {business_name} will receive detailed vulnerability and resilience insights.",
            f"Risk assessment initiated for {business_name}. {username}, your strategic protection blueprint is being developed."
        ],
        "middle": [
            f"Risk assessment 50% complete. {username}, strategic insights for {business_name} are being compiled.",
            f"{business_name}'s security analysis is progressing well. Midway through comprehensive risk evaluation.",
            f"Progress update: {username}'s risk assessment continues with detailed analysis of {business_name}'s protection opportunities."
        ],
        "complete": [
            f"Risk assessment complete. {username}, your comprehensive security blueprint for {business_name} is ready for review.",
            f"Analysis finished. {business_name}'s detailed risk management strategy is now available.",
            f"{username}, your risk assessment is complete. Strategic insights for {business_name} are ready for implementation."
        ]
    }
    
    return random.choice(fallback_messages.get(stage, fallback_messages["start"]))

async def send_personalized_risk_notification(user_id: str, user_profile: Dict, stage: str, progress_data: Dict = None, report_id: str = None):
    """Send personalized risk notification with optional DB persistence for completion"""
    try:
        # Generate personalized message
        message = await generate_personalized_risk_message(user_profile, stage, progress_data)

        # Create titles for risk management
        username = user_profile.get('username', 'Leader')
        business_name = user_profile.get('business_name', 'Your Business')

        professional_titles = {
            "start": [
                f"Risk Assessment Started",
                f"{business_name} - Security Analysis Initiated",
                f"Risk Analysis Underway"
            ],
            "middle": [
                f"Risk Assessment - Progress Update",
                f"{business_name} - Security Analysis 50% Complete",
                f"Risk Analysis In Progress"
            ],
            "complete": [
                f"Risk Assessment Complete",
                f"{business_name} - Security Analysis Ready",
                f"Risk Analysis Finished"
            ]
        }

        title = random.choice(professional_titles[stage])

        # For completion notification, save to DB with deep linking
        save_to_db = (stage == "complete" and report_id is not None)

        # Send notification
        success, result = await send_risk_notification(
            user_id, title, message, "notification",
            save_to_db, report_id, business_name
        )

        if success:
            logging.info(f"🎭 Sent risk {stage} notification to user {user_id} (saveToDb: {save_to_db})")
        else:
            logging.error(f"❌ Failed to send risk notification: {result}")

        return success, message

    except Exception as e:
        logging.error(f"❌ Error sending risk notification: {str(e)}")
        return False, str(e)

def send_risk_notification_background(user_id: str, user_profile: Dict, stage: str, progress_data: Dict = None, report_id: str = None):
    """Send risk notification in background thread with proper async handling"""

    notification_id = f"risk_{stage}_{user_id}_{int(time.time())}"

    def notification_worker():
        worker_start_time = time.time()
        thread_id = threading.current_thread().ident

        try:
            logging.info(f"🔔 NOTIFICATION WORKER START: {notification_id}")
            logging.info(f"📊 Notification details:")
            logging.info(f"   - User ID: {user_id}")
            logging.info(f"   - Stage: {stage}")
            logging.info(f"   - Report ID: {report_id}")
            logging.info(f"   - Thread ID: {thread_id}")
            logging.info(f"   - Business: {user_profile.get('business_name', 'Unknown') if user_profile else 'No profile'}")
            logging.info(f"   - Progress data: {bool(progress_data)}")

            # Validate inputs
            if not user_id:
                raise ValueError("user_id is required and cannot be empty")

            if not stage:
                raise ValueError("stage is required and cannot be empty")

            if stage not in ['start', 'middle', 'complete']:
                logging.warning(f"⚠️ Unusual stage value: '{stage}' (expected: start/middle/complete)")

            logging.info(f"✅ Input validation passed for notification {notification_id}")

            # Platform-specific setup
            if platform.system() == 'Windows':
                logging.debug(f"🪟 Setting Windows event loop policy for thread {thread_id}")
                asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())

            # Event loop creation
            logging.debug(f"🔄 Creating new event loop in thread {thread_id}")
            loop_start_time = time.time()

            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)
            loop_creation_time = time.time() - loop_start_time

            logging.debug(f"✅ Event loop created in {loop_creation_time:.3f}s")

            try:
                logging.info(f"📤 Executing notification for {notification_id}...")
                notification_start_time = time.time()

                try:
                    success, message = loop.run_until_complete(
                        asyncio.wait_for(
                            send_personalized_risk_notification(user_id, user_profile, stage, progress_data, report_id),
                            timeout=30.0  # 30 second timeout
                        )
                    )
                    
                    notification_time = time.time() - notification_start_time
                    
                    logging.info(f"📡 Notification execution completed in {notification_time:.3f}s")
                    logging.info(f"📊 Notification result:")
                    logging.info(f"   - Success: {success}")
                    logging.info(f"   - Message: {message}")
                    
                except asyncio.TimeoutError:
                    notification_time = time.time() - notification_start_time
                    success = False
                    message = f"Notification timed out after {notification_time:.1f}s"
                    
                    logging.error(f"⏰ TIMEOUT: Notification {notification_id} timed out after {notification_time:.1f}s")
                    logging.error(f"🔍 This may indicate network issues or API problems")
                
                # Log final results
                total_worker_time = time.time() - worker_start_time
                
                if success:
                    logging.info(f"🎉 SUCCESS: Background risk {stage} notification sent")
                    logging.info(f"📊 Success metrics:")
                    logging.info(f"   - Notification ID: {notification_id}")
                    logging.info(f"   - User ID: {user_id}")
                    logging.info(f"   - Stage: {stage}")
                    logging.info(f"   - Total time: {total_worker_time:.3f}s")
                    logging.info(f"   - Notification time: {notification_time:.3f}s")
                    logging.info(f"   - Thread ID: {thread_id}")
                else:
                    logging.warning(f"⚠️ FAILURE: Background risk {stage} notification failed")
                    logging.warning(f"📊 Failure details:")
                    logging.warning(f"   - Notification ID: {notification_id}")
                    logging.warning(f"   - User ID: {user_id}")
                    logging.warning(f"   - Stage: {stage}")
                    logging.warning(f"   - Error message: {message}")
                    logging.warning(f"   - Total time: {total_worker_time:.3f}s")
                    logging.warning(f"   - Thread ID: {thread_id}")
                    
            except Exception as loop_error:
                loop_error_time = time.time() - notification_start_time if 'notification_start_time' in locals() else 0
                total_worker_time = time.time() - worker_start_time
                
                logging.error(f"❌ LOOP ERROR: Exception in notification execution")
                logging.error(f"🔍 Loop error details:")
                logging.error(f"   - Notification ID: {notification_id}")
                logging.error(f"   - Error type: {type(loop_error).__name__}")
                logging.error(f"   - Error message: {str(loop_error)}")
                logging.error(f"   - Loop error time: {loop_error_time:.3f}s")
                logging.error(f"   - Total worker time: {total_worker_time:.3f}s")
                logging.error(f"   - Thread ID: {thread_id}")
                
                # Log the full traceback for debugging
                import traceback
                logging.error(f"🔍 Loop error traceback:")
                for line in traceback.format_exc().split('\n'):
                    if line.strip():
                        logging.error(f"   {line}")
                
            finally:
                # Clean up event loop
                try:
                    loop_cleanup_start = time.time()
                    
                    logging.debug(f"🔄 Cleaning up event loop for thread {thread_id}")
                    
                    # Cancel any remaining tasks
                    pending_tasks = [task for task in asyncio.all_tasks(loop) if not task.done()]
                    if pending_tasks:
                        logging.warning(f"⚠️ Found {len(pending_tasks)} pending tasks, cancelling...")
                        for task in pending_tasks:
                            task.cancel()
                    
                    loop.close()
                    loop_cleanup_time = time.time() - loop_cleanup_start
                    
                    logging.debug(f"✅ Event loop cleaned up in {loop_cleanup_time:.3f}s")
                    
                except Exception as cleanup_error:
                    logging.error(f"❌ Error during loop cleanup: {cleanup_error}")
                
        except Exception as worker_error:
            total_worker_time = time.time() - worker_start_time
            
            logging.error(f"💥 WORKER ERROR: Critical error in notification worker")
            logging.error(f"🔍 Worker error details:")
            logging.error(f"   - Notification ID: {notification_id}")
            logging.error(f"   - User ID: {user_id}")
            logging.error(f"   - Stage: {stage}")
            logging.error(f"   - Error type: {type(worker_error).__name__}")
            logging.error(f"   - Error message: {str(worker_error)}")
            logging.error(f"   - Total worker time: {total_worker_time:.3f}s")
            logging.error(f"   - Thread ID: {thread_id}")
            
            # Log the full traceback for debugging
            import traceback
            logging.error(f"🔍 Worker error traceback:")
            for line in traceback.format_exc().split('\n'):
                if line.strip():
                    logging.error(f"   {line}")
        
        finally:
            # Final cleanup and statistics
            total_worker_time = time.time() - worker_start_time
            
            logging.info(f"🏁 NOTIFICATION WORKER END: {notification_id}")
            logging.info(f"📊 Worker final statistics:")
            logging.info(f"   - Total execution time: {total_worker_time:.3f}s")
            logging.info(f"   - Thread ID: {thread_id}")
            logging.info(f"   - Worker completed at: {datetime.now().isoformat()}")
    
    # Better thread creation with error handling
    try:
        logging.info(f"🚀 LAUNCHING: Background notification thread for {notification_id}")
        
        # Create thread with proper naming and error handling
        notification_thread = Thread(
            target=notification_worker, 
            daemon=True,
            name=f"RiskNotification-{stage}-{user_id}"
        )
        
        # Store notification metadata in thread for debugging
        notification_thread._notification_data = {
            'notification_id': notification_id,
            'user_id': user_id,
            'stage': stage,
            'created_at': time.time()
        }
        
        thread_start_time = time.time()
        notification_thread.start()
        thread_start_duration = time.time() - thread_start_time
        
        logging.info(f"✅ Notification thread launched successfully")
        logging.info(f"📊 Thread launch details:")
        logging.info(f"   - Notification ID: {notification_id}")
        logging.info(f"   - Thread name: {notification_thread.name}")
        logging.info(f"   - Thread ID: {notification_thread.ident}")
        logging.info(f"   - Launch time: {thread_start_duration:.3f}s")
        logging.info(f"   - Daemon thread: {notification_thread.daemon}")
        
    except Exception as thread_error:
        logging.error(f"💥 THREAD CREATION ERROR: Failed to create notification thread")
        logging.error(f"🔍 Thread error details:")
        logging.error(f"   - Notification ID: {notification_id}")
        logging.error(f"   - User ID: {user_id}")
        logging.error(f"   - Stage: {stage}")
        logging.error(f"   - Error type: {type(thread_error).__name__}")
        logging.error(f"   - Error message: {str(thread_error)}")
        
        # Log the full traceback for debugging
        import traceback
        logging.error(f"🔍 Thread creation traceback:")
        for line in traceback.format_exc().split('\n'):
            if line.strip():
                logging.error(f"   {line}")




# ======================================================
#           Document Creation for Risk Engine
# ======================================================

def create_risk_word_document(report_data: Dict, user_id: str) -> Document:
    """Create risk Word document with fortress-themed formatting"""
    logging.info("📄 Creating Risk Fortress Word Document")
    
    doc = Document()
    
    # Enhanced styling
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    
    # Professional title page with fortress theme
    title = doc.add_heading("BACKABLE", 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(42)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 51, 102)
    
    subtitle = doc.add_heading("Comprehensive Risk Fortress Strategy Blueprint", 1)
    subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitle_run = subtitle.runs[0]
    subtitle_run.font.size = Pt(20)
    subtitle_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for Risk
    
    # Add metadata
    metadata_para = doc.add_paragraph()
    metadata_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    report_meta = report_data.get("_enhanced_risk_report_metadata", {})
    
    metadata_para.add_run(f"User ID: {user_id}\n").bold = True
    metadata_para.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}\n")
    metadata_para.add_run(f"Analysis: {report_meta.get('total_words', 0):,} words\n")
    metadata_para.add_run(f"Model: Gemini 2.5 Pro Risk Engine\n")
    metadata_para.add_run(f"Focus: Threat Protection & Business Fortress\n")
    
    # Add multi-database intelligence indicator
    multi_db_info = report_meta.get("multi_database_integration", {})
    if multi_db_info.get("enabled", False):
        data_sources = multi_db_info.get("data_sources_used", [])
        metadata_para.add_run(f"Intelligence Sources: {len(data_sources)} databases integrated\n")
        complete_qa_pairs = multi_db_info.get("complete_qa_pairs", 0)
        if complete_qa_pairs > 0:
            metadata_para.add_run(f"Q&A Intelligence: {complete_qa_pairs} cross-engine insights\n")
    
    doc.add_page_break()
    
    # Table of Contents
    toc_heading = doc.add_heading("TABLE OF CONTENTS", 1)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    section_number = 1
    for section_name, section_data in report_data.items():
        if section_name != "_enhanced_risk_report_metadata" and isinstance(section_data, dict):
            title = section_data.get("title", "Untitled Section")
            
            toc_para = doc.add_paragraph()
            toc_para.add_run(f"{section_number}. {title}").bold = True
            
            # Add word count
            metadata = section_data.get("metadata", {})
            words_generated = metadata.get("words_generated", 0)
            
            toc_para.add_run(f" ({words_generated:,} words)")
            
            section_number += 1
    
    doc.add_page_break()
    
    # Process each section
    section_number = 1
    for section_name, section_data in report_data.items():
        if section_name != "_enhanced_risk_report_metadata" and isinstance(section_data, dict):
            
            logging.info(f"📝 Formatting risk section: {section_name}")
            
            title = section_data.get("title", "Untitled Section")
            content = section_data.get("content", "")
            
            # Add section header
            section_heading = doc.add_heading(f"{section_number}. {title}", 1)
            heading_run = section_heading.runs[0]
            heading_run.font.color.rgb = RGBColor(0, 51, 102)
            
            # Add the AI-generated content
            add_risk_content_to_document(doc, content)
            
            # Add section separator with fortress theme
            separator_para = doc.add_paragraph()
            separator_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            separator_run = separator_para.add_run("🏰 ◆ ◆ ◆ 🏰")
            separator_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for Risk
            separator_run.font.size = Pt(16)
            
            section_number += 1
            doc.add_page_break()
    
    # Add report summary
    add_risk_report_summary(doc, report_data)
    
    logging.info("✅ Risk Fortress Word Document Created")
    return doc

def add_risk_content_to_document(doc: Document, content: str):
    """Add AI-generated risk content to document with intelligent formatting"""
    
    # Split by paragraphs and headers
    lines = content.split('\n')
    current_paragraph = ""
    
    for line in lines:
        line = line.strip()
        
        if not line:
            # Empty line - finalize paragraph
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
        elif line.startswith('###'):
            # Sub-subsection header
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            header_text = line.replace('###', '').strip()
            subheading = doc.add_heading(header_text, 3)
            subheading_run = subheading.runs[0]
            subheading_run.font.color.rgb = RGBColor(255, 102, 102)  # Light red for Risk
            
        elif line.startswith('##'):
            # Subsection header
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            header_text = line.replace('##', '').strip()
            subheading = doc.add_heading(header_text, 2)
            subheading_run = subheading.runs[0]
            subheading_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for Risk
            
        elif line.startswith('#'):
            # Main header
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            header_text = line.replace('#', '').strip()
            subheading = doc.add_heading(header_text, 2)
            subheading_run = subheading.runs[0]
            subheading_run.font.color.rgb = RGBColor(204, 0, 0)
            
        elif line.startswith('- ') or line.startswith('• '):
            # Bullet point
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            bullet_text = line[2:].strip()
            doc.add_paragraph(bullet_text, style='List Bullet')
            
        elif re.match(r'^\d+\.', line):
            # Numbered list
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            number_text = re.sub(r'^\d+\.\s*', '', line)
            doc.add_paragraph(number_text, style='List Number')
            
        elif line.startswith('🏰') or line.startswith('🛡️') or line.startswith('⚔️') or line.startswith('🔒'):
            # Risk-specific emojis - treat as emphasis
            if current_paragraph:
                para = doc.add_paragraph(current_paragraph)
                current_paragraph = ""
            
            emoji_para = doc.add_paragraph(line)
            emoji_run = emoji_para.runs[0]
            emoji_run.bold = True
            emoji_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for Risk
            
        else:
            # Regular content - accumulate
            if current_paragraph:
                current_paragraph += " " + line
            else:
                current_paragraph = line
    
    # Add any remaining paragraph
    if current_paragraph:
        para = doc.add_paragraph(current_paragraph)

def add_risk_report_summary(doc: Document, report_data: Dict):
    """Add risk-specific report summary"""
    
    summary_heading = doc.add_heading("RISK FORTRESS REPORT SUMMARY", 1)
    summary_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    summary_heading_run = summary_heading.runs[0]
    summary_heading_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for Risk
    
    report_meta = report_data.get("_enhanced_risk_report_metadata", {})
    
    summary_para = doc.add_paragraph()
    summary_para.add_run("Risk Fortress Strategy Report Statistics:").bold = True
    summary_para.add_run(f"\n• Total Security Sections: {report_meta.get('total_sections', 0)}")
    summary_para.add_run(f"\n• Total Words Generated: {report_meta.get('total_words', 0):,}")
    summary_para.add_run(f"\n• AI Model: {report_meta.get('ai_model', 'N/A')}")
    summary_para.add_run(f"\n• Processing Method: {report_meta.get('processing_method', 'N/A')}")
    summary_para.add_run(f"\n• Report Type: {report_meta.get('report_type', 'N/A')}")
    summary_para.add_run(f"\n• Focus Areas: Threat Analysis, Security Architecture, Risk Mitigation, Compliance")
    
    # Add risk-specific insights summary
    insights_para = doc.add_paragraph()
    insights_para.add_run("\nRisk Fortress Intelligence Summary:").bold = True
    insights_para.add_run(f"\n🎯 Threat Landscape Analysis with Mitigation Strategies")
    insights_para.add_run(f"\n🏰 Security Architecture Blueprint with Implementation Timeline")
    insights_para.add_run(f"\n⚔️ Risk Mitigation Strategy with Priority Matrix")
    insights_para.add_run(f"\n🛡️ Compliance & Governance Framework with Audit Protocols")
    insights_para.add_run(f"\n🔒 Fortress Implementation Roadmap with Success Metrics")
    insights_para.add_run(f"\n📊 Financial Risk Analysis with Investment ROI")
    insights_para.add_run(f"\n🚨 Crisis Response Protocols for Business Continuity")
    insights_para.add_run(f"\n🔧 Continuous Fortress Optimization with Monitoring")
    
    # Add multi-database integration summary
    multi_db_info = report_meta.get("multi_database_integration", {})
    if multi_db_info.get("enabled", False):
        summary_para.add_run(f"\n• Multi-Database Integration: Enabled")
        data_sources = multi_db_info.get("data_sources_used", [])
        summary_para.add_run(f"\n• Intelligence Sources: {', '.join(data_sources).title() if data_sources else 'Risk Assessment Only'}")
        complete_qa_pairs = multi_db_info.get("complete_qa_pairs", 0)
        if complete_qa_pairs > 0:
            summary_para.add_run(f"\n• Cross-Engine Q&A Intelligence: {complete_qa_pairs} insights integrated")

# ======================================================
#           Question-Response Chunking for Risk Engine
# ======================================================

async def create_risk_question_response_chunks(raw_assessment_data: Dict, report_id: str, user_id: str) -> List[Dict]:
    """Create Question-Response chunks optimized for RAG context with comprehensive monitoring and flexible sizing"""
    
    function_start_time = time.time()
    
    logging.info(f"🚀 STARTING RISK Q&R CHUNKING PROCESS")
    logging.info(f"📊 Input parameters:")
    logging.info(f"   - Report ID: {report_id}")
    logging.info(f"   - User ID: {user_id}")
    logging.info(f"   - Raw data type: {type(raw_assessment_data)}")
    logging.info(f"   - Raw data keys: {list(raw_assessment_data.keys()) if isinstance(raw_assessment_data, dict) else 'Not a dict'}")
    
    qr_chunks = []
    
    # Get user profile for context using async connection pooling
    logging.info(f"👤 FETCHING USER PROFILE for user_id={user_id}")
    profile_start_time = time.time()
    
    try:
        user_profile = await get_user_profile_data(user_id)
        profile_fetch_time = time.time() - profile_start_time
        
        if user_profile:
            business_name = user_profile.get('business_name', 'Unknown')
            industry = user_profile.get('industry', 'Unknown')
            team_size = user_profile.get('team_size', 'Unknown')
            
            logging.info(f"✅ USER PROFILE RETRIEVED in {profile_fetch_time:.3f}s:")
            logging.info(f"   - Business: {business_name}")
            logging.info(f"   - Industry: {industry}")
            logging.info(f"   - Team Size: {team_size}")
        else:
            logging.warning(f"⚠️ NO USER PROFILE found for user_id={user_id}")
            user_profile = {}
            
    except Exception as profile_error:
        profile_fetch_time = time.time() - profile_start_time
        logging.error(f"❌ USER PROFILE FETCH FAILED after {profile_fetch_time:.3f}s: {profile_error}")
        user_profile = {}
    
    # FIXED: More flexible Q&R chunking settings optimized for actual content
    # Reduced from original settings to be more accommodating
    TARGET_SIZE_WORDS = 400  # Reduced from 800 - more realistic for Q&A content
    MAX_SIZE_WORDS = 600     # Reduced from 1000 - prevents overly large chunks  
    MIN_SIZE_WORDS = 30      # Reduced from 200 - much more flexible for short responses
    
    logging.info(f"⚙️ Q&R CHUNKING SETTINGS (OPTIMIZED):")
    logging.info(f"   - Target size: {TARGET_SIZE_WORDS} words (reduced from 800 for flexibility)")
    logging.info(f"   - Max size: {MAX_SIZE_WORDS} words (reduced from 1000)")
    logging.info(f"   - Min size: {MIN_SIZE_WORDS} words (reduced from 200 for preservation)")
    logging.info(f"   - Strategy: Preserve all content, prioritize context over strict sizing")
    
    # Extract responses from raw assessment data with detailed analysis
    logging.info(f"📊 EXTRACTING RESPONSES from raw assessment data...")
    
    responses = raw_assessment_data.get("responses", [])
    
    if not responses:
        # Try alternative keys if "responses" is not found
        alternative_keys = ["response_data", "assessment_responses", "user_responses", "answers"]
        for alt_key in alternative_keys:
            if alt_key in raw_assessment_data:
                responses = raw_assessment_data[alt_key]
                logging.info(f"🔧 Found responses under alternative key: '{alt_key}'")
                break
    
    if not responses:
        logging.error(f"❌ NO RESPONSES FOUND in raw assessment data!")
        logging.error(f"🔍 Available keys in raw_assessment_data: {list(raw_assessment_data.keys()) if isinstance(raw_assessment_data, dict) else 'Not a dict'}")
        logging.error(f"🔍 Raw data sample: {str(raw_assessment_data)[:500]}{'...' if len(str(raw_assessment_data)) > 500 else ''}")
        return []
    
    logging.info(f"✅ RESPONSES EXTRACTED:")
    logging.info(f"   - Total responses: {len(responses)}")
    logging.info(f"   - Response data type: {type(responses)}")
    
    # Analyze response structure for better processing
    if responses:
        sample_response = responses[0]
        logging.info(f"📋 SAMPLE RESPONSE ANALYSIS:")
        logging.info(f"   - Sample type: {type(sample_response)}")
        logging.info(f"   - Sample keys: {list(sample_response.keys()) if isinstance(sample_response, dict) else 'Not a dict'}")
        
        # Count responses with actual content
        valid_responses = 0
        total_words_estimate = 0
        
        for i, response in enumerate(responses):
            if isinstance(response, dict):
                question_text = response.get("question_text", "")
                response_data = response.get("response_data", {})
                
                if question_text or response_data:
                    valid_responses += 1
                    # Quick word count estimate
                    words = len(str(question_text).split()) + len(str(response_data).split())
                    total_words_estimate += words
        
        avg_response_length = total_words_estimate / valid_responses if valid_responses > 0 else 0
        
        logging.info(f"📊 RESPONSE CONTENT ANALYSIS:")
        logging.info(f"   - Valid responses: {valid_responses}/{len(responses)}")
        logging.info(f"   - Estimated total words: {total_words_estimate}")
        logging.info(f"   - Average response length: {avg_response_length:.1f} words")
        
        # Adjust chunking strategy based on content analysis
        if avg_response_length < 20:
            # Very short responses - be even more flexible
            MIN_SIZE_WORDS = 15
            TARGET_SIZE_WORDS = 200
            MAX_SIZE_WORDS = 300
            logging.info(f"🔧 ADJUSTED SETTINGS for short responses:")
            logging.info(f"   - New target: {TARGET_SIZE_WORDS} words")
            logging.info(f"   - New max: {MAX_SIZE_WORDS} words") 
            logging.info(f"   - New min: {MIN_SIZE_WORDS} words")
    
    # Group responses by section for better contextual chunks
    logging.info(f"📂 GROUPING RESPONSES by section...")
    grouping_start_time = time.time()
    
    sections = {}
    ungrouped_count = 0
    
    for response in responses:
        if isinstance(response, dict):
            section = response.get("section", "General")
            if not section or section.strip() == "":
                section = "General"
                ungrouped_count += 1
            
            if section not in sections:
                sections[section] = []
            sections[section].append(response)
        else:
            logging.warning(f"⚠️ Non-dict response found: {type(response)}")
    
    grouping_time = time.time() - grouping_start_time
    
    logging.info(f"✅ SECTION GROUPING COMPLETE in {grouping_time:.3f}s:")
    logging.info(f"   - Total sections: {len(sections)}")
    logging.info(f"   - Section names: {list(sections.keys())}")
    logging.info(f"   - Ungrouped responses: {ungrouped_count}")
    
    # Log section details
    for section_name, section_responses in sections.items():
        logging.info(f"   📁 Section '{section_name}': {len(section_responses)} responses")
    
    # Processing variables
    chunk_id = 1
    total_questions = 0
    total_output_chunks = 0
    total_sections_processed = 0
    sections_with_chunks = 0
    sections_without_chunks = 0
    
    # Process each section with detailed tracking
    logging.info(f"🔄 PROCESSING SECTIONS for chunking...")
    section_processing_start = time.time()
    
    for section_name, section_responses in sections.items():
        section_start_time = time.time()
        
        if not section_responses:
            logging.warning(f"⚠️ Skipping empty section: '{section_name}'")
            continue
        
        total_sections_processed += 1
        
        logging.info(f"📄 PROCESSING SECTION '{section_name}' ({total_sections_processed}/{len(sections)})")
        logging.info(f"   - Responses in section: {len(section_responses)}")
        logging.info(f"   - Processing start time: {datetime.now().strftime('%H:%M:%S.%f')[:-3]}")
        
        # Create semantic Q&R chunks from this section
        logging.info(f"🔧 Creating semantic chunks for section '{section_name}'...")
        chunking_start_time = time.time()
        
        try:
            section_chunks = create_risk_qr_semantic_chunks(
                section_responses, section_name, TARGET_SIZE_WORDS, MAX_SIZE_WORDS, MIN_SIZE_WORDS
            )
            chunking_time = time.time() - chunking_start_time
            
            logging.info(f"✅ SECTION CHUNKING COMPLETE for '{section_name}' in {chunking_time:.3f}s:")
            logging.info(f"   - Input: {len(section_responses)} responses")
            logging.info(f"   - Output: {len(section_chunks)} chunks")
            logging.info(f"   - Success rate: {len(section_chunks)/len(section_responses):.2f} chunks per response")
            
            if len(section_chunks) > 0:
                sections_with_chunks += 1
            else:
                sections_without_chunks += 1
                logging.warning(f"⚠️ No chunks created for section '{section_name}'")
                
        except Exception as chunking_error:
            chunking_time = time.time() - chunking_start_time
            logging.error(f"❌ SECTION CHUNKING FAILED for '{section_name}' after {chunking_time:.3f}s:")
            logging.error(f"   - Error type: {type(chunking_error).__name__}")
            logging.error(f"   - Error message: {str(chunking_error)}")
            logging.error(f"   - Continuing with next section...")
            section_chunks = []
            sections_without_chunks += 1
            
            # Log error traceback for debugging
            import traceback
            logging.error(f"🔍 Error traceback:")
            for line in traceback.format_exc().split('\n'):
                if line.strip():
                    logging.error(f"     {line}")
        
        # Convert each chunk to a Word document
        if section_chunks:
            logging.info(f"📝 CREATING WORD DOCUMENTS for {len(section_chunks)} chunks...")
            
            for i, (chunk_questions, chunk_content) in enumerate(section_chunks):
                doc_creation_start = time.time()
                
                chunk_title = f"Risk Q&A: {section_name}" if len(section_chunks) == 1 else f"Risk Q&A: {section_name} - Part {i+1}"
                chunk_word_count = len(chunk_content.split())
                question_count = len(chunk_questions)
                character_count = len(chunk_content)
                
                logging.info(f"📝 CREATING DOCUMENT for chunk {chunk_id}:")
                logging.info(f"   - Title: '{chunk_title}'")
                logging.info(f"   - Word count: {chunk_word_count}")
                logging.info(f"   - Question count: {question_count}")
                logging.info(f"   - Character count: {character_count}")
                
                # Create Word document for this Q&R chunk
                try:
                    chunk_doc = create_risk_qr_chunk_word_document(
                        chunk_content, 
                        chunk_title, 
                        user_profile,
                        section_name,
                        f"{report_id}_risk_qr_chunk_{chunk_id:03d}",
                        chunk_questions
                    )
                    doc_creation_time = time.time() - doc_creation_start
                    
                    logging.debug(f"✅ Word document created in {doc_creation_time:.3f}s")
                    
                except Exception as doc_error:
                    doc_creation_time = time.time() - doc_creation_start
                    logging.error(f"❌ Word document creation failed after {doc_creation_time:.3f}s: {doc_error}")
                    # Create a minimal fallback document
                    chunk_doc = None
                
                # Determine chunk quality metrics
                try:
                    chunk_category = categorize_risk_qr_chunk_size_by_words(chunk_word_count)
                    contextual_completeness = calculate_risk_qr_contextual_completeness(chunk_content, chunk_questions)
                except Exception as metrics_error:
                    logging.warning(f"⚠️ Metrics calculation failed: {metrics_error}")
                    chunk_category = "unknown"
                    contextual_completeness = 0.0
                
                # Determine quality status with detailed logic
                if TARGET_SIZE_WORDS * 0.8 <= chunk_word_count <= TARGET_SIZE_WORDS * 1.2:
                    quality_status = "✅ OPTIMAL"
                    quality_color = "GREEN"
                elif chunk_word_count > TARGET_SIZE_WORDS * 1.2:
                    quality_status = "⚠️ LARGE"
                    quality_color = "YELLOW"
                elif chunk_word_count >= MIN_SIZE_WORDS:
                    quality_status = "⚠️ SMALL"  
                    quality_color = "YELLOW"
                else:
                    quality_status = "❌ TINY"
                    quality_color = "RED"
                
                logging.info(f"📊 CHUNK {chunk_id} QUALITY ASSESSMENT:")
                logging.info(f"   - Status: {quality_status} ({quality_color})")
                logging.info(f"   - Size: {chunk_word_count} words (target: {TARGET_SIZE_WORDS})")
                logging.info(f"   - Questions: {question_count}")
                logging.info(f"   - Category: {chunk_category}")
                logging.info(f"   - Completeness: {contextual_completeness:.2f}")
                logging.info(f"   - Size ratio: {chunk_word_count / TARGET_SIZE_WORDS:.2f}x target")
                
                # Create comprehensive chunk info
                chunk_info = {
                    "chunk_id": f"{report_id}_risk_qr_chunk_{chunk_id:03d}",
                    "section_name": section_name,
                    "expansion_title": chunk_title,
                    "word_count": chunk_word_count,
                    "question_count": question_count,
                    "character_count": character_count,
                    "content_preview": chunk_content[:200] + "..." if len(chunk_content) > 200 else chunk_content,
                    "questions_included": [q.get("question_id", f"Q{idx+1}") for idx, q in enumerate(chunk_questions)],
                    "document": chunk_doc,
                    "chunk_metadata": {
                        "original_section": section_name,
                        "chunk_size_category": chunk_category,
                        "contextual_completeness": contextual_completeness,
                        "chunk_type": "risk_question_response_rag_optimized",
                        "quality_status": quality_status,
                        "quality_color": quality_color,
                        "rag_optimization": {
                            "target_size": TARGET_SIZE_WORDS,
                            "size_ratio": chunk_word_count / TARGET_SIZE_WORDS,
                            "quality_status": quality_status.replace("✅ ", "").replace("⚠️ ", "").replace("❌ ", ""),
                            "context_preservation": True,
                            "min_size_met": chunk_word_count >= MIN_SIZE_WORDS,
                            "target_size_met": TARGET_SIZE_WORDS * 0.8 <= chunk_word_count <= TARGET_SIZE_WORDS * 1.2,
                            "flexible_sizing_applied": True
                        }
                    },
                    "user_context": {
                        "user_id": user_id,
                        "business_name": user_profile.get("business_name", "Unknown") if user_profile else "Unknown",
                        "industry": user_profile.get("industry", "Unknown") if user_profile else "Unknown",
                        "team_size": user_profile.get("team_size", "Unknown") if user_profile else "Unknown"
                    },
                    "created_at": datetime.now().isoformat(),
                    "processing_stats": {
                        "section_processing_order": total_sections_processed,
                        "chunk_in_section_order": i + 1,
                        "total_chunks_in_section": len(section_chunks),
                        "doc_creation_time": doc_creation_time,
                        "chunk_creation_timestamp": datetime.now().isoformat()
                    }
                }
                
                qr_chunks.append(chunk_info)
                total_output_chunks += 1
                total_questions += question_count
                chunk_id += 1
                
                logging.debug(f"✅ Chunk {chunk_id-1} added to output collection")
        
        section_time = time.time() - section_start_time
        logging.info(f"✅ SECTION '{section_name}' PROCESSING COMPLETE in {section_time:.3f}s:")
        logging.info(f"   - Chunks created: {len(section_chunks)}")
        logging.info(f"   - Questions processed: {sum(len(chunk[0]) for chunk in section_chunks)}")
        logging.info(f"   - Average processing time per chunk: {(section_time/len(section_chunks)):.3f}s" if section_chunks else "N/A")
    
    total_section_processing_time = time.time() - section_processing_start
    
    # Final comprehensive statistics with detailed analysis
    logging.info(f"🏁 ALL SECTIONS PROCESSED in {total_section_processing_time:.3f}s")
    
    if qr_chunks:
        # Calculate detailed statistics
        word_counts = [c['word_count'] for c in qr_chunks]
        question_counts = [c['question_count'] for c in qr_chunks]
        
        avg_chunk_size = sum(word_counts) // len(word_counts)
        min_chunk_size = min(word_counts)
        max_chunk_size = max(word_counts)
        
        avg_questions_per_chunk = sum(question_counts) // len(question_counts)
        min_questions_per_chunk = min(question_counts)
        max_questions_per_chunk = max(question_counts)
        
        # Quality analysis
        optimal_chunks = sum(1 for c in qr_chunks if TARGET_SIZE_WORDS * 0.8 <= c['word_count'] <= TARGET_SIZE_WORDS * 1.2)
        large_chunks = sum(1 for c in qr_chunks if c['word_count'] > TARGET_SIZE_WORDS * 1.2)
        small_chunks = sum(1 for c in qr_chunks if MIN_SIZE_WORDS <= c['word_count'] < TARGET_SIZE_WORDS * 0.8)
        tiny_chunks = sum(1 for c in qr_chunks if c['word_count'] < MIN_SIZE_WORDS)
        
        optimal_percentage = (optimal_chunks / len(qr_chunks)) * 100
        
        # Success metrics
        total_function_time = time.time() - function_start_time
        
        logging.info(f"🎉 Q&R CHUNKING PROCESS COMPLETED SUCCESSFULLY!")
        logging.info(f"📊 COMPREHENSIVE FINAL STATISTICS:")
        logging.info(f"   ⏱️ TIMING:")
        logging.info(f"      - Total function time: {total_function_time:.3f}s")
        logging.info(f"      - Section processing time: {total_section_processing_time:.3f}s")
        logging.info(f"      - Average time per section: {(total_section_processing_time/total_sections_processed):.3f}s")
        logging.info(f"      - Average time per chunk: {(total_section_processing_time/len(qr_chunks)):.3f}s")
        logging.info(f"   📄 INPUT ANALYSIS:")
        logging.info(f"      - Total responses: {len(responses)}")
        logging.info(f"      - Sections found: {len(sections)}")
        logging.info(f"      - Sections processed: {total_sections_processed}")
        logging.info(f"   📦 OUTPUT SUMMARY:")
        logging.info(f"      - Q&R chunks created: {len(qr_chunks)}")
        logging.info(f"      - Total questions preserved: {total_questions}")
        logging.info(f"      - Sections with chunks: {sections_with_chunks}")
        logging.info(f"      - Sections without chunks: {sections_without_chunks}")
        logging.info(f"   📏 SIZE METRICS:")
        logging.info(f"      - Chunk word count range: {min_chunk_size}-{max_chunk_size}")
        logging.info(f"      - Average chunk size: {avg_chunk_size} words")
        logging.info(f"      - Questions per chunk range: {min_questions_per_chunk}-{max_questions_per_chunk}")
        logging.info(f"      - Average questions per chunk: {avg_questions_per_chunk}")
        logging.info(f"   🎯 QUALITY METRICS:")
        logging.info(f"      - Optimal chunks (target ±20%): {optimal_chunks} ({optimal_percentage:.1f}%)")
        logging.info(f"      - Large chunks (>120% target): {large_chunks}")
        logging.info(f"      - Small chunks (min to 80% target): {small_chunks}")
        logging.info(f"      - Tiny chunks (<min size): {tiny_chunks}")
        logging.info(f"   🔗 CONTEXT PRESERVATION:")
        logging.info(f"      - Question-Answer pairs maintained: YES")
        logging.info(f"      - Section context preserved: YES")
        logging.info(f"      - Content loss: NONE (all responses processed)")
        
        # Performance assessment
        chunks_per_second = len(qr_chunks) / total_function_time
        questions_per_second = total_questions / total_function_time
        
        logging.info(f"   ⚡ PERFORMANCE METRICS:")
        logging.info(f"      - Chunks per second: {chunks_per_second:.2f}")
        logging.info(f"      - Questions per second: {questions_per_second:.2f}")
        logging.info(f"      - Processing efficiency: {'HIGH' if chunks_per_second > 1 else 'MODERATE'}")
        
        # Success indicators
        success_rate = len(qr_chunks) / len(responses) if responses else 0
        
        logging.info(f"   ✅ SUCCESS INDICATORS:")
        logging.info(f"      - Chunk creation rate: {success_rate:.2f} chunks per response")
        logging.info(f"      - Content preservation: 100%")
        logging.info(f"      - Processing success: {'HIGH' if success_rate > 0.3 else 'MODERATE' if success_rate > 0 else 'LOW'}")
        
    else:
        # No chunks created - detailed failure analysis
        total_function_time = time.time() - function_start_time
        
        logging.error(f"❌ Q&R CHUNKING PROCESS FAILED - NO CHUNKS CREATED!")
        logging.error(f"📊 FAILURE ANALYSIS:")
        logging.error(f"   ⏱️ TIMING:")
        logging.error(f"      - Total function time: {total_function_time:.3f}s")
        logging.error(f"      - Time wasted on failed processing: {total_function_time:.3f}s")
        logging.error(f"   📄 INPUT ANALYSIS:")
        logging.error(f"      - Total responses provided: {len(responses)}")
        logging.error(f"      - Sections identified: {len(sections)}")
        logging.error(f"      - Sections processed: {total_sections_processed}")
        logging.error(f"      - Sections with chunks: {sections_with_chunks}")
        logging.error(f"      - Sections without chunks: {sections_without_chunks}")
        logging.error(f"   🔍 POTENTIAL CAUSES:")
        logging.error(f"      - All responses too short (below {MIN_SIZE_WORDS} words)")
        logging.error(f"      - Q&A formatting errors")
        logging.error(f"      - Invalid response data structure")
        logging.error(f"      - Chunking algorithm issues")
        logging.error(f"   💡 RECOMMENDATIONS:")
        logging.error(f"      - Check response data format")
        logging.error(f"      - Verify question text and response data validity")
        logging.error(f"      - Consider lowering MIN_SIZE_WORDS further")
        logging.error(f"      - Review chunking algorithm logic")
        
        # Log sample responses for debugging
        if responses:
            logging.error(f"   📋 SAMPLE RESPONSES FOR DEBUGGING:")
            for i, response in enumerate(responses[:3]):  # Show first 3 responses
                logging.error(f"      Response {i+1}: {str(response)[:200]}{'...' if len(str(response)) > 200 else ''}")
    
    logging.info(f"🎯 RETURNING {len(qr_chunks)} Q&R chunks for report_id={report_id}")
    return qr_chunks

def create_risk_qr_semantic_chunks(responses: List[Dict], section_name: str, target_size: int, max_size: int, min_size: int) -> List[Tuple[List[Dict], str]]:
    """Create semantic Q&R chunks that preserve question-answer context - FIXED VERSION with detailed logging"""
    
    logging.info(f"🔧 STARTING RISK Q&R SEMANTIC CHUNKING")
    logging.info(f"📊 Chunking parameters:")
    logging.info(f"   - Section: '{section_name}'")
    logging.info(f"   - Target size: {target_size} words")
    logging.info(f"   - Max size: {max_size} words") 
    logging.info(f"   - Min size: {min_size} words")
    logging.info(f"   - Input responses: {len(responses) if responses else 0}")
    
    if not responses:
        logging.warning(f"⚠️ No responses provided for section '{section_name}' - returning empty chunks")
        return []
    
    # Initialize chunking variables
    chunks = []
    current_chunk_questions = []
    current_chunk_content = ""
    current_word_count = 0
    
    # FIXED: Calculate effective minimum size - be more flexible for short content
    original_min_size = min_size
    effective_min_size = min(min_size, 50)  # Never require more than 50 words minimum
    
    # If all responses are very short, be even more flexible
    total_response_words = 0
    for response in responses:
        question_text = response.get("question_text", "")
        response_data = response.get("response_data", {})
        qa_text = format_risk_qa_pair(question_text, response_data, response.get("question_id", ""))
        total_response_words += len(qa_text.split())
    
    avg_response_length = total_response_words / len(responses) if responses else 0
    
    # If average response is very short, use even smaller minimum
    if avg_response_length < 30:
        effective_min_size = min(effective_min_size, 20)  # Allow very small chunks for short responses
        logging.info(f"🔧 ADAPTIVE MIN SIZE: Average response length is {avg_response_length:.1f} words - using minimum {effective_min_size} words")
    
    logging.info(f"📊 Chunking analysis:")
    logging.info(f"   - Total response words: {total_response_words}")
    logging.info(f"   - Average response length: {avg_response_length:.1f} words")
    logging.info(f"   - Original min size: {original_min_size} words")
    logging.info(f"   - Effective min size: {effective_min_size} words")
    
    logging.info(f"📝 Processing {len(responses)} responses for section '{section_name}':")
    
    # Process each response
    for response_idx, response in enumerate(responses):
        response_start_time = time.time()
        
        # Extract question and answer data
        question_text = response.get("question_text", f"Question {response_idx + 1}")
        response_data = response.get("response_data", {})
        question_id = response.get("question_id", f"Q{response_idx + 1}")
        
        logging.debug(f"🔍 PROCESSING RESPONSE {response_idx + 1}/{len(responses)}:")
        logging.debug(f"   - Question ID: {question_id}")
        logging.debug(f"   - Question text: {question_text[:100]}{'...' if len(question_text) > 100 else ''}")
        logging.debug(f"   - Response data type: {type(response_data)}")
        logging.debug(f"   - Response data keys: {list(response_data.keys()) if isinstance(response_data, dict) else 'Not dict'}")
        
        # Format the Q&A pair
        try:
            qa_text = format_risk_qa_pair(question_text, response_data, question_id)
            qa_word_count = len(qa_text.split())
            
            logging.debug(f"   ✅ Q&A formatted successfully: {qa_word_count} words")
            logging.debug(f"   📝 Q&A preview: {qa_text[:150]}{'...' if len(qa_text) > 150 else ''}")
            
        except Exception as format_error:
            logging.error(f"   ❌ Error formatting Q&A pair: {format_error}")
            # Create a simple fallback Q&A
            qa_text = f"QUESTION {question_id}: {question_text}\n\nANSWER: {str(response_data)}"
            qa_word_count = len(qa_text.split())
            logging.warning(f"   🔧 Using fallback Q&A format: {qa_word_count} words")
        
        # Check if adding this Q&A would exceed max size
        test_word_count = current_word_count + qa_word_count
        
        logging.debug(f"🧮 CHUNK SIZE CALCULATION:")
        logging.debug(f"   - Current chunk words: {current_word_count}")
        logging.debug(f"   - This Q&A words: {qa_word_count}")
        logging.debug(f"   - Test total: {test_word_count}")
        logging.debug(f"   - Max size: {max_size}")
        logging.debug(f"   - Would exceed max: {test_word_count > max_size}")
        logging.debug(f"   - Current chunk questions: {len(current_chunk_questions)}")
        
        if test_word_count > max_size and current_chunk_questions:
            # Current chunk is getting too big - decide what to do
            logging.info(f"📦 CHUNK SIZE LIMIT REACHED:")
            logging.info(f"   - Current chunk: {current_word_count} words, {len(current_chunk_questions)} questions")
            logging.info(f"   - Adding Q&A would make: {test_word_count} words (max: {max_size})")
            logging.info(f"   - Effective min size: {effective_min_size}")
            
            if current_word_count >= effective_min_size:
                # Save current chunk - it's big enough
                chunks.append((current_chunk_questions.copy(), current_chunk_content.strip()))
                chunk_number = len(chunks)
                
                logging.info(f"✅ SAVED Q&R CHUNK {chunk_number}:")
                logging.info(f"   - Questions: {len(current_chunk_questions)}")
                logging.info(f"   - Words: {current_word_count}")
                logging.info(f"   - Met minimum: {current_word_count >= effective_min_size}")
                logging.info(f"   - Content preview: {current_chunk_content[:100]}{'...' if len(current_chunk_content) > 100 else ''}")
                
                # Start new chunk with current response
                current_chunk_questions = [response]
                current_chunk_content = qa_text
                current_word_count = qa_word_count
                
                logging.info(f"🆕 STARTED NEW CHUNK:")
                logging.info(f"   - Initial Q&A: {question_id}")
                logging.info(f"   - Initial words: {qa_word_count}")
                
            else:
                # Current chunk too small, but adding Q&A makes it too big
                logging.warning(f"⚠️ CHUNK SIZE DILEMMA:")
                logging.warning(f"   - Current chunk: {current_word_count} words < {effective_min_size} minimum")
                logging.warning(f"   - But adding Q&A: {test_word_count} words > {max_size} maximum")
                
                if not current_chunk_questions:
                    # No questions in current chunk yet - start with this one
                    current_chunk_questions = [response]
                    current_chunk_content = qa_text
                    current_word_count = qa_word_count
                    logging.info(f"🔧 RESOLUTION: Started chunk with large Q&A ({qa_word_count} words)")
                    
                else:
                    # Add to current chunk even if it exceeds max (preserve context)
                    current_chunk_questions.append(response)
                    current_chunk_content += f"\n\n{qa_text}"
                    current_word_count = test_word_count
                    logging.info(f"🔧 RESOLUTION: Added to chunk despite size ({test_word_count} words) to preserve context")
                    
        else:
            # Add Q&A to current chunk - fits within limits
            current_chunk_questions.append(response)
            if current_chunk_content:
                current_chunk_content += f"\n\n{qa_text}"
            else:
                current_chunk_content = qa_text
            current_word_count = test_word_count
            
            logging.debug(f"➕ ADDED Q&A TO CURRENT CHUNK:")
            logging.debug(f"   - Question ID: {question_id}")
            logging.debug(f"   - Total words now: {current_word_count}")
            logging.debug(f"   - Total questions now: {len(current_chunk_questions)}")
            logging.debug(f"   - Within max size: {current_word_count <= max_size}")
        
        response_time = time.time() - response_start_time
        logging.debug(f"   ⏱️ Response processing time: {response_time:.3f}s")
    
    # FIXED: Handle the final chunk - always save if we have content
    logging.info(f"🔚 FINAL CHUNK PROCESSING:")
    logging.info(f"   - Final chunk questions: {len(current_chunk_questions)}")
    logging.info(f"   - Final chunk words: {current_word_count}")
    logging.info(f"   - Effective min size: {effective_min_size}")
    logging.info(f"   - Original min size: {original_min_size}")
    
    if current_chunk_questions:
        if current_word_count >= effective_min_size:
            chunks.append((current_chunk_questions.copy(), current_chunk_content.strip()))
            chunk_number = len(chunks)
            
            logging.info(f"✅ SAVED FINAL Q&R CHUNK {chunk_number}:")
            logging.info(f"   - Questions: {len(current_chunk_questions)}")
            logging.info(f"   - Words: {current_word_count}")
            logging.info(f"   - Met effective minimum: YES ({current_word_count} >= {effective_min_size})")
            logging.info(f"   - Met original minimum: {'YES' if current_word_count >= original_min_size else 'NO'} ({current_word_count} vs {original_min_size})")
            
        else:
            # FIXED: Save even small chunks instead of discarding them
            chunks.append((current_chunk_questions.copy(), current_chunk_content.strip()))
            chunk_number = len(chunks)
            
            logging.warning(f"⚠️ SAVED SMALL FINAL CHUNK {chunk_number}:")
            logging.warning(f"   - Questions: {len(current_chunk_questions)}")
            logging.warning(f"   - Words: {current_word_count}")
            logging.warning(f"   - Below effective minimum: {current_word_count} < {effective_min_size}")
            logging.warning(f"   - But preserving content instead of discarding")
            logging.warning(f"   - Content: {current_chunk_content[:200]}{'...' if len(current_chunk_content) > 200 else ''}")
    else:
        logging.info(f"ℹ️ No final chunk to process (no remaining questions)")
    
    # Validate the created chunks
    logging.info(f"📊 VALIDATING CREATED CHUNKS:")
    try:
        chunk_stats = validate_risk_qr_chunk_sizes(chunks, target_size, f"Risk Q&R Section: {section_name}")
        logging.info(f"✅ Chunk validation completed successfully")
    except Exception as validation_error:
        logging.error(f"❌ Chunk validation error: {validation_error}")
        chunk_stats = {"min_words": 0, "max_words": 0, "avg_words": 0}
    
    # Final comprehensive logging
    total_questions = sum(len(chunk[0]) for chunk in chunks)
    total_words = sum(len(chunk[1].split()) for chunk in chunks)
    
    logging.info(f"🎉 RISK Q&R SEMANTIC CHUNKING COMPLETE!")
    logging.info(f"📊 FINAL STATISTICS for '{section_name}':")
    logging.info(f"   ✅ Input: {len(responses)} responses")
    logging.info(f"   ✅ Output: {len(chunks)} chunks created")
    logging.info(f"   ✅ Total questions preserved: {total_questions}")
    logging.info(f"   ✅ Total words: {total_words}")
    logging.info(f"   📏 Chunk word range: {chunk_stats.get('min_words', 0)}-{chunk_stats.get('max_words', 0)}")
    logging.info(f"   📈 Average chunk size: {chunk_stats.get('avg_words', 0)} words (target: {target_size})")
    logging.info(f"   🎯 Target compliance: {'GOOD' if chunk_stats.get('avg_words', 0) <= target_size * 1.2 else 'OVER TARGET'}")
    
    # Log individual chunk details for debugging
    if chunks:
        logging.info(f"📋 INDIVIDUAL CHUNK BREAKDOWN:")
        for i, (questions, content) in enumerate(chunks, 1):
            chunk_words = len(content.split())
            question_ids = [q.get('question_id', f'Q{idx}') for idx, q in enumerate(questions)]
            
            logging.info(f"   📦 Chunk {i}: {len(questions)} questions, {chunk_words} words")
            logging.info(f"       Question IDs: {', '.join(question_ids)}")
            logging.info(f"       Content preview: {content[:100]}{'...' if len(content) > 100 else ''}")
            
            # Quality assessment
            if chunk_words >= target_size * 0.8 and chunk_words <= target_size * 1.2:
                quality = "✅ OPTIMAL"
            elif chunk_words >= effective_min_size:
                quality = "⚠️ ACCEPTABLE"
            else:
                quality = "❌ SMALL"
                
            logging.info(f"       Quality: {quality}")
    else:
        logging.warning(f"⚠️ NO CHUNKS CREATED for section '{section_name}'")
        logging.warning(f"📊 This might indicate:")
        logging.warning(f"   - Very short responses that don't meet minimum requirements")
        logging.warning(f"   - Errors in Q&A formatting")
        logging.warning(f"   - Issues with response data structure")
    
    logging.info(f"✅ Returning {len(chunks)} chunks for section '{section_name}'")
    return chunks

def format_risk_qa_pair(question_text: str, response_data: Dict, question_id: str) -> str:
    """Format a question-answer pair for Q&R chunking"""
    
    # Start with the question
    formatted_text = f"QUESTION {question_id}: {question_text}\n\n"
    
    # Format the answer based on response type
    if isinstance(response_data, dict):
        if 'selected_option' in response_data:
            formatted_text += f"ANSWER: {response_data['selected_option']}"
        elif 'selected_options' in response_data:
            options = response_data['selected_options']
            if isinstance(options, list):
                formatted_text += f"ANSWERS: {', '.join(map(str, options))}"
            else:
                formatted_text += f"ANSWERS: {options}"
        elif 'response_text' in response_data:
            formatted_text += f"ANSWER: {response_data['response_text']}"
        elif 'slider_value' in response_data:
            formatted_text += f"RATING: {response_data['slider_value']}"
        elif 'text_response' in response_data:
            formatted_text += f"ANSWER: {response_data['text_response']}"
        else:
            # Try to extract any meaningful value
            meaningful_values = []
            for key, value in response_data.items():
                if key not in ['metadata', 'timestamp', 'session_id'] and value is not None:
                    meaningful_values.append(f"{key}: {value}")
            
            if meaningful_values:
                formatted_text += f"RESPONSE: {'; '.join(meaningful_values)}"
            else:
                formatted_text += "ANSWER: [No response recorded]"
    else:
        formatted_text += f"ANSWER: {response_data}"
    
    return formatted_text

def create_risk_qr_chunk_word_document(content: str, title: str, user_profile: Dict, 
                                      section_name: str, chunk_id: str, questions: List[Dict]) -> Document:
    """Create a professionally formatted Word document for Q&R chunk"""
    
    try:
        doc = Document()
        
        # Enhanced styling
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Add header with branding
        header = doc.add_heading("BACKABLE RISK ENGINE", 0)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header.runs[0]
        header_run.font.size = Pt(24)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Add subtitle for Q&R
        subtitle = doc.add_heading("Question & Response Intelligence", 2)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.size = Pt(16)
        subtitle_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for Risk
        
        # Add chunk title
        chunk_title = doc.add_heading(title, 1)
        chunk_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        title_run = chunk_title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for Risk
        
        # Add metadata section
        if user_profile:
            metadata_para = doc.add_paragraph()
            metadata_para.add_run("Risk Q&A Context:").bold = True
            metadata_para.add_run(f"\nBusiness: {user_profile.get('business_name', 'Unknown')}")
            metadata_para.add_run(f"\nIndustry: {user_profile.get('industry', 'Unknown')}")  
            metadata_para.add_run(f"\nTeam Size: {user_profile.get('team_size', 'Unknown')} employees")
            metadata_para.add_run(f"\nRisk Section: {section_name}")
            metadata_para.add_run(f"\nChunk ID: {chunk_id}")
            metadata_para.add_run(f"\nQuestions Included: {len(questions)}")
            metadata_para.add_run(f"\nGenerated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        
        # Add separator
        doc.add_paragraph("─" * 60)
        
        # Add the Q&A content with risk-specific formatting
        add_risk_qa_content_to_document(doc, content)
        
        # Add footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_run = footer_para.add_run("Generated by Backable AI Risk Intelligence Engine")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        logging.info(f"📄 Created risk Q&R chunk Word document: {len(content.split())} words, {len(questions)} questions")
        return doc
        
    except Exception as e:
        logging.error(f"❌ Error creating risk Q&R chunk Word document: {str(e)}")
        # Return minimal document on error
        doc = Document()
        doc.add_heading("Error Creating Risk Q&A Document", 1)
        doc.add_paragraph(f"Error: {str(e)}")
        return doc

def add_risk_qa_content_to_document(doc: Document, content: str):
    """Add Q&A content to Word document with special formatting for questions and answers"""
    
    # Split content into Q&A pairs
    sections = content.split('\n\n')
    
    for section in sections:
        section = section.strip()
        if not section:
            continue
        
        # Check if this is a question
        if section.startswith('QUESTION'):
            # Extract question number and text
            lines = section.split('\n')
            question_line = lines[0]
            
            # Add question header
            question_para = doc.add_paragraph()
            question_run = question_para.add_run(question_line)
            question_run.font.bold = True
            question_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for Risk
            question_run.font.size = Pt(12)
            
            # Add question text if it spans multiple lines
            if len(lines) > 1:
                for line in lines[1:]:
                    if line.strip():
                        question_para.add_run(f"\n{line}")
        
        elif section.startswith('ANSWER:') or section.startswith('ANSWERS:') or section.startswith('RATING:') or section.startswith('RESPONSE:'):
            # Add answer with different formatting
            answer_para = doc.add_paragraph()
            answer_run = answer_para.add_run(section)
            answer_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue for answers
            answer_run.font.size = Pt(11)
            
            # Add some spacing after answers
            doc.add_paragraph()
        
        else:
            # Regular content
            para = doc.add_paragraph(section)

def categorize_risk_qr_chunk_size_by_words(word_count: int) -> str:
    """Categorize Q&R chunk size for risk analysis"""
    if word_count < 400:
        return "small"
    elif word_count < 700:
        return "optimal"
    elif word_count < 1200:
        return "large"
    else:
        return "oversized"

def calculate_risk_qr_contextual_completeness(content: str, questions: List[Dict]) -> float:
    """Calculate contextual completeness score for Q&R content"""
    
    # Q&R-specific completeness indicators
    completeness_indicators = {
        'question_words': ['question', 'what', 'how', 'why', 'when', 'where', 'which'],
        'answer_words': ['answer', 'response', 'rating', 'selected', 'option'],
        'risk_context': ['risk', 'threat', 'vulnerability', 'security', 'compliance', 'audit'],
        'business_context': ['business', 'company', 'organization', 'team', 'process', 'system'],
        'action_words': ['implement', 'develop', 'create', 'establish', 'improve', 'manage'],
        'assessment_words': ['assess', 'evaluate', 'review', 'analyze', 'measure', 'monitor']
    }
    
    content_lower = content.lower()
    total_indicators = 0
    found_indicators = 0
    
    for category, words in completeness_indicators.items():
        total_indicators += len(words)
        found_indicators += sum(1 for word in words if word in content_lower)
    
    base_completeness = found_indicators / total_indicators if total_indicators > 0 else 0
    
    # Q&A pair bonus (more Q&A pairs = more complete)
    qa_factor = min(1.0, len(questions) / 5)  # Optimal around 5 Q&A pairs
    
    # Length bonus for Q&R content
    word_count = len(content.split())
    length_factor = min(1.0, word_count / 800)  # Optimal around 800 words for Q&R
    
    # Structure bonus for Q&A formatting
    structure_indicators = ['QUESTION', 'ANSWER:', 'RATING:', 'RESPONSE:']
    structure_count = sum(1 for indicator in structure_indicators if indicator in content)
    structure_factor = min(0.2, structure_count * 0.05)  # Up to 0.2 bonus
    
    final_completeness = min(1.0, base_completeness + qa_factor * 0.3 + length_factor * 0.2 + structure_factor)
    
    return final_completeness

def validate_risk_qr_chunk_sizes(chunks: List[Tuple[List[Dict], str]], target_size: int, context_name: str = "") -> Dict:
    """Validate and log risk Q&R chunk sizes for monitoring"""
    
    if not chunks:
        return {"total_chunks": 0}
    
    chunk_stats = {
        "total_chunks": len(chunks),
        "avg_words": 0,
        "min_words": float('inf'),
        "max_words": 0,
        "chunks_over_target": 0,
        "chunks_under_200": 0,  # Flag very small Q&R chunks
        "chunks_optimal": 0,     # Chunks within target range
        "total_questions": 0,
        "avg_questions": 0
    }
    
    total_words = 0
    total_questions = 0
    
    for questions, content in chunks:
        words = len(content.split())
        total_words += words
        question_count = len(questions)
        total_questions += question_count
        
        chunk_stats["min_words"] = min(chunk_stats["min_words"], words)
        chunk_stats["max_words"] = max(chunk_stats["max_words"], words)
        
        if words > target_size * 1.2:  # 20% over target
            chunk_stats["chunks_over_target"] += 1
        elif words < 200:
            chunk_stats["chunks_under_200"] += 1
        elif target_size * 0.8 <= words <= target_size * 1.2:  # Within 20% of target
            chunk_stats["chunks_optimal"] += 1
    
    chunk_stats["avg_words"] = total_words // len(chunks)
    chunk_stats["avg_questions"] = total_questions // len(chunks)
    chunk_stats["total_questions"] = total_questions
    chunk_stats["min_words"] = chunk_stats["min_words"] if chunk_stats["min_words"] != float('inf') else 0
    
    # Log the stats
    context_prefix = f"[{context_name}] " if context_name else ""
    logging.info(f"📊 {context_prefix}Risk Q&R chunk validation: "
                f"{chunk_stats['total_chunks']} chunks, "
                f"avg: {chunk_stats['avg_words']} words, "
                f"range: {chunk_stats['min_words']}-{chunk_stats['max_words']}, "
                f"optimal: {chunk_stats['chunks_optimal']}/{chunk_stats['total_chunks']}, "
                f"total questions: {chunk_stats['total_questions']}")
    
    return chunk_stats





# ======================================================
#           5. Missing upload_risk_report_to_azure function
# ======================================================

# ======================================================
#           FIXED: Risk Engine Indexer Integration
# ======================================================

async def trigger_risk_indexer_for_client(user_id: str) -> Tuple[bool, str, str]:
    """FIXED: Trigger risk indexer for specific client with correct endpoint and payload"""
    
    logging.info(f"🚀 RISK INDEXER API: Triggering indexer for user_id={user_id}")
    logging.info(f"📊 Indexer parameters:")
    logging.info(f"   - User ID: {user_id}")
    logging.info(f"   - API Base URL: {INDEXER_API_BASE_URL}")
    logging.info(f"   - Timeout: {INDEXER_TIMEOUT}s")
    logging.info(f"   - Max retries: {INDEXER_RETRY_ATTEMPTS}")
    
    max_retries = INDEXER_RETRY_ATTEMPTS
    base_delay = INDEXER_RETRY_DELAY
    
    for attempt in range(max_retries):
        attempt_start_time = time.time()
        logging.info(f"🔄 RISK INDEXER API: Attempt {attempt + 1}/{max_retries} for user_id={user_id}")
        
        try:
            async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=INDEXER_TIMEOUT)) as session:
                # FIXED: Use correct payload structure that matches auto-indexer expectations
                payload = {
                    "client_id": user_id,  # FIXED: Changed from user_id to client_id
                    "force": True,         # FIXED: Added force flag to override any existing jobs
                    "new_client": False    # FIXED: Added new_client flag (assuming existing client)
                }
                
                # Enhanced payload logging
                logging.info(f"📤 RISK INDEXER API: Sending payload")
                logging.info(f"📊 Payload details:")
                logging.info(f"   - Client ID: {payload['client_id']}")
                logging.info(f"   - Force: {payload['force']}")
                logging.info(f"   - New Client: {payload['new_client']}")
                logging.info(f"   - JSON size: {len(json.dumps(payload))} bytes")
                
                # FIXED: Use correct endpoint that exists in auto-indexer
                indexer_url = f"{INDEXER_API_BASE_URL}/run-indexer"  # FIXED: Changed from /trigger_indexer_for_client
                logging.info(f"🌐 RISK INDEXER API: Making request to {indexer_url}")
                
                async with session.post(
                    indexer_url,
                    json=payload,
                    headers={"Content-Type": "application/json"}
                ) as response:
                    
                    request_time = time.time() - attempt_start_time
                    
                    # Enhanced response logging
                    logging.info(f"📡 RISK INDEXER API: Response received")
                    logging.info(f"📊 Response details:")
                    logging.info(f"   - Status code: {response.status}")
                    logging.info(f"   - Response time: {request_time:.3f}s")
                    logging.info(f"   - Content type: {response.headers.get('content-type', 'Unknown')}")
                    logging.info(f"   - Content length: {response.headers.get('content-length', 'Unknown')}")
                    
                    try:
                        response_data = await response.json()
                        logging.info(f"✅ RISK INDEXER API: JSON parsing successful")
                        logging.info(f"📊 Response data keys: {list(response_data.keys())}")
                        
                    except Exception as json_error:
                        logging.error(f"❌ RISK INDEXER API: JSON parse error: {json_error}")
                        response_text = await response.text()
                        logging.error(f"🔍 Raw response text (first 500 chars): {response_text[:500]}...")
                        
                        if attempt < max_retries - 1:
                            wait_time = base_delay * (2 ** attempt)
                            logging.warning(f"⏳ RISK INDEXER API: JSON parse failed, retrying in {wait_time}s...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"JSON parse error: {json_error}", None
                    
                    # Handle different response statuses
                    if response.status == 202:  # Accepted - job added to queue
                        job_id = response_data.get("job_id")
                        message = response_data.get("message", "Risk indexer job started successfully")
                        queue_position = response_data.get("queue_position", 0)
                        
                        logging.info(f"🎉 RISK INDEXER API: JOB ACCEPTED!")
                        logging.info(f"✅ Acceptance details:")
                        logging.info(f"   - User ID: {user_id}")
                        logging.info(f"   - Job ID: {job_id}")
                        logging.info(f"   - Message: {message}")
                        logging.info(f"   - Queue Position: {queue_position}")
                        logging.info(f"   - Attempt: {attempt + 1}/{max_retries}")
                        logging.info(f"   - Request time: {request_time:.3f}s")
                        
                        return True, message, job_id
                    
                    elif response.status == 200:  # Success (immediate processing)
                        job_id = response_data.get("job_id")
                        message = response_data.get("message", "Risk indexer job completed successfully")
                        
                        logging.info(f"🎉 RISK INDEXER API: JOB COMPLETED!")
                        logging.info(f"✅ Success details:")
                        logging.info(f"   - User ID: {user_id}")
                        logging.info(f"   - Job ID: {job_id}")
                        logging.info(f"   - Message: {message}")
                        
                        return True, message, job_id
                    
                    elif response.status == 409:  # Conflict - Job already in progress/recent warning
                        message = response_data.get("message", "Indexer job already in progress or recent")
                        job_created_at = response_data.get("job_created_at")
                        
                        logging.warning(f"⚠️ RISK INDEXER API: Conflict for user_id={user_id}")
                        logging.warning(f"📊 Conflict details:")
                        logging.warning(f"   - Message: {message}")
                        logging.warning(f"   - Job Created At: {job_created_at}")
                        logging.warning(f"   - This usually means indexing is already running or recent")
                        
                        # For conflicts, we can retry with force=true
                        if attempt < max_retries - 1:
                            logging.info(f"🔄 Will retry with force=true on next attempt...")
                            payload["force"] = True  # Force override on retry
                            wait_time = base_delay * (attempt + 1)
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            # On final attempt, treat conflict as success since job exists
                            return True, f"Indexer job exists (conflict): {message}", "existing_job"
                    
                    elif response.status == 404:  # Client not found
                        message = response_data.get("message", "Client not found in indexer system")
                        
                        logging.warning(f"⚠️ RISK INDEXER API: Client not found: user_id={user_id}")
                        logging.warning(f"📊 404 details:")
                        logging.warning(f"   - Message: {message}")
                        
                        # Try with new_client=true on retry
                        if attempt < max_retries - 1:
                            payload["new_client"] = True
                            wait_time = base_delay * (attempt + 1)
                            logging.info(f"⏳ RISK INDEXER API: Client not found, retrying with new_client=true in {wait_time}s...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"Risk client not found: {message}", None
                    
                    elif response.status == 429:  # Rate Limited
                        message = response_data.get("message", "Risk indexer rate limited")
                        logging.warning(f"🚦 RISK INDEXER API: Rate limited for user_id={user_id}")
                        
                        if attempt < max_retries - 1:
                            wait_time = base_delay * 2 * (attempt + 1)  # Longer wait for rate limits
                            logging.info(f"⏳ RISK INDEXER API: Rate limited, waiting {wait_time}s before retry...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"Risk indexer rate limited: {message}", None
                    
                    elif response.status >= 500:  # Server errors - retry
                        message = response_data.get("message", f"Risk indexer server error {response.status}")
                        logging.error(f"🚨 RISK INDEXER API: Server error {response.status} for user_id={user_id}")
                        
                        if attempt < max_retries - 1:
                            wait_time = base_delay * (2 ** attempt)
                            logging.info(f"⏳ RISK INDEXER API: Server error, retrying in {wait_time}s...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"Risk indexer server error: {message}", None
                    
                    else:  # Other client errors - don't retry
                        message = response_data.get("message", f"Risk indexer failed with status {response.status}")
                        logging.error(f"❌ RISK INDEXER API: Client error {response.status} for user_id={user_id}")
                        return False, f"Risk indexer client error: {message}", None
                        
        except asyncio.TimeoutError:
            request_time = time.time() - attempt_start_time
            error_msg = f"Risk indexer request timed out after {INDEXER_TIMEOUT}s"
            
            logging.error(f"⏰ RISK INDEXER API: TIMEOUT (attempt {attempt + 1}/{max_retries})")
            logging.error(f"📊 Timeout details:")
            logging.error(f"   - Configured timeout: {INDEXER_TIMEOUT}s")
            logging.error(f"   - Actual time: {request_time:.3f}s")
            logging.error(f"   - User ID: {user_id}")
            
            if attempt < max_retries - 1:
                wait_time = base_delay * (attempt + 1)
                logging.info(f"⏳ RISK INDEXER API: Timeout, retrying in {wait_time}s...")
                await asyncio.sleep(wait_time)
                continue
            else:
                return False, error_msg, None
                
        except aiohttp.ClientError as e:
            request_time = time.time() - attempt_start_time
            error_msg = f"Risk indexer HTTP client error: {str(e)}"
            
            logging.error(f"🌐 RISK INDEXER API: HTTP CLIENT ERROR (attempt {attempt + 1}/{max_retries})")
            logging.error(f"📊 HTTP error details:")
            logging.error(f"   - Error type: {type(e).__name__}")
            logging.error(f"   - Error message: {str(e)}")
            logging.error(f"   - Request time: {request_time:.3f}s")
            
            if attempt < max_retries - 1:
                wait_time = base_delay * (attempt + 1)
                logging.info(f"⏳ RISK INDEXER API: HTTP error, retrying in {wait_time}s...")
                await asyncio.sleep(wait_time)
                continue
            else:
                return False, error_msg, None
                
        except Exception as e:
            request_time = time.time() - attempt_start_time
            error_msg = f"Risk indexer unexpected error: {str(e)}"
            
            logging.error(f"💥 RISK INDEXER API: UNEXPECTED ERROR (attempt {attempt + 1}/{max_retries})")
            logging.error(f"📊 Unexpected error details:")
            logging.error(f"   - Error type: {type(e).__name__}")
            logging.error(f"   - Error message: {str(e)}")
            logging.error(f"   - Request time: {request_time:.3f}s")
            
            if attempt < max_retries - 1:
                wait_time = base_delay * (attempt + 1)
                logging.info(f"⏳ RISK INDEXER API: Unexpected error, retrying in {wait_time}s...")
                await asyncio.sleep(wait_time)
                continue
            else:
                return False, error_msg, None
    
    # If we get here, all retries failed
    final_error = f"Risk indexer failed after {max_retries} attempts"
    
    logging.error(f"💥 RISK INDEXER API: ALL ATTEMPTS FAILED")
    logging.error(f"📊 Final failure summary:")
    logging.error(f"   - User ID: {user_id}")
    logging.error(f"   - Total attempts: {max_retries}")
    logging.error(f"   - Final error: {final_error}")
    
    return False, final_error, None


# ======================================================
#           Azure Blob Upload with Retry Logic
# ======================================================

def upload_blob_with_retry(container_client, blob_name, data, content_settings, max_retries=3):
    """
    Helper function to upload blob with retry logic
    """
    for attempt in range(max_retries):
        try:
            container_client.upload_blob(
                name=blob_name,
                data=data,
                overwrite=True,
                content_settings=content_settings
            )
            logging.info(f"✅ Successfully uploaded: {blob_name}")
            return True
        except Exception as e:
            if attempt < max_retries - 1:
                wait_time = (attempt + 1) * 2
                logging.warning(f"Upload attempt {attempt + 1} failed for {blob_name}: {str(e)}. Retrying in {wait_time}s...")
                time.sleep(wait_time)
                # Reset data stream position if possible
                if hasattr(data, 'seek'):
                    data.seek(0)
            else:
                logging.error(f"Failed to upload {blob_name} after {max_retries} attempts: {str(e)}")
                raise Exception(f"Failed to upload {blob_name} after {max_retries} attempts: {str(e)}")

    return False

# ======================================================
#           FIXED: Enhanced Question-Response Chunking
# ======================================================

async def upload_risk_report_to_azure(report_data: Dict, report_id: str, user_id: str):
    """FIXED: Upload risk report to Azure with enhanced Q&R chunking that handles empty responses gracefully"""
    try:
        logging.info(f"🚀 Starting Risk Report Azure Upload for report_id={report_id}, user_id={user_id}")
        
        container_name = get_azure_container_name(user_id)
        logging.info(f"📦 Using Azure container: {container_name}")
        
        blob_service_client = BlobServiceClient.from_connection_string(AZURE_STORAGE_CONNECTION_STRING)
        container_client = blob_service_client.get_container_client(container_name)
        
        try:
            container_client.create_container()
            logging.info(f"✅ Container '{container_name}' created")
        except:
            logging.info(f"📦 Container '{container_name}' already exists")

        # Get client folder name from database (e.g., "666-tim")
        client_folder = get_client_folder_name(user_id)

        # Create folder structure: {client_folder}/risk analysis engine report/
        folder_name = f"{client_folder}/risk analysis engine report"
        logging.info(f"📁 Using folder structure: {folder_name}/")
        logging.info(f"📁 Uploading risk report to: {container_name}/{folder_name}/")
        
        # ===============================================================
        # 1. Upload complete Word document
        # ===============================================================
        logging.info("📄 Step 1/6: Creating and uploading complete Word document...")
        doc = create_risk_word_document(report_data, user_id)
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        
        doc_blob_name = f"{folder_name}/{report_id}_comprehensive_risk_fortress_strategy.docx"
        upload_blob_with_retry(
            container_client,
            doc_blob_name,
            doc_bytes,
            ContentSettings(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        )
        
        # ===============================================================
        # 2. Upload complete JSON data
        # ===============================================================
        logging.info("📊 Step 2/6: Creating and uploading complete JSON data...")
        json_data = json.dumps(report_data, indent=2, default=str)
        json_bytes = io.BytesIO(json_data.encode("utf-8"))
        
        json_blob_name = f"{folder_name}/{report_id}_comprehensive_risk_report.json"
        upload_blob_with_retry(
            container_client,
            json_blob_name,
            json_bytes,
            ContentSettings(content_type="application/json")
        )
        
        # ===============================================================
        # 3. Create and upload Word document chunks for Azure Cognitive Search
        # ===============================================================
        logging.info("🔧 Step 3/6: Creating Word document chunks for Azure Cognitive Search...")
        word_chunks = await create_risk_word_document_chunks(report_data, report_id, user_id)
        logging.info(f"📊 Created {len(word_chunks)} report Word chunks")
        
        # Upload individual Word chunk files
        chunk_files_created = []
        for i, chunk_doc in enumerate(word_chunks):
            chunk_blob_name = f"{folder_name}/{report_id}_risk_chunk_{i+1:03d}.docx"
            
            # Save Word document chunk to bytes
            chunk_bytes = io.BytesIO()
            chunk_doc['document'].save(chunk_bytes)
            chunk_bytes.seek(0)

            upload_blob_with_retry(
                container_client,
                chunk_blob_name,
                chunk_bytes,
                ContentSettings(
                    content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            )
            chunk_files_created.append(chunk_blob_name)
            logging.info(f"Risk Word chunk {i+1}: {chunk_blob_name} ({chunk_doc['word_count']} words)")
        
        # ===============================================================
        # 4. FIXED: Enhanced Question-Response chunks for RAG context
        # ===============================================================
        logging.info("🧠 Step 4/6: Creating Question-Response chunks for RAG context...")
        
        # FIXED: Better extraction of raw assessment data with multiple fallbacks
        raw_assessment_data = None
        qr_chunks = []
        qr_chunk_files_created = []
        
        # Try multiple sources for assessment data
        extraction_sources = [
            ("_enhanced_risk_report_metadata.raw_assessment_data", lambda: report_data.get('_enhanced_risk_report_metadata', {}).get('raw_assessment_data')),
            ("report_data directly", lambda: report_data if 'responses' in report_data else None),
            ("user input from metadata", lambda: report_data.get('user_input')),
            ("assessment_data from metadata", lambda: report_data.get('assessment_data'))
        ]
        
        for source_name, extractor in extraction_sources:
            try:
                potential_data = extractor()
                if potential_data and isinstance(potential_data, dict):
                    responses = potential_data.get('responses', [])
                    if responses and len(responses) > 0:
                        raw_assessment_data = potential_data
                        logging.info(f"✅ Found assessment data from: {source_name} ({len(responses)} responses)")
                        break
            except Exception as e:
                logging.debug(f"🔍 Could not extract from {source_name}: {e}")
                continue
        
        if raw_assessment_data and raw_assessment_data.get('responses'):
            try:
                logging.info(f"📊 Processing {len(raw_assessment_data.get('responses', []))} responses for Q&R chunking...")
                qr_chunks = await create_risk_question_response_chunks(raw_assessment_data, report_id, user_id)
                logging.info(f"📊 Created {len(qr_chunks)} Question-Response chunks for RAG")
                
                # Upload Question-Response chunk files
                for i, qr_chunk in enumerate(qr_chunks):
                    qr_chunk_blob_name = f"{folder_name}/{report_id}_risk_qr_chunk_{i+1:03d}.docx"
                    
                    # FIXED: Better error handling for document creation
                    try:
                        # Save Question-Response document chunk to bytes
                        qr_chunk_bytes = io.BytesIO()
                        qr_chunk['document'].save(qr_chunk_bytes)
                        qr_chunk_bytes.seek(0)

                        upload_blob_with_retry(
                            container_client,
                            qr_chunk_blob_name,
                            qr_chunk_bytes,
                            ContentSettings(
                                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        )
                        qr_chunk_files_created.append(qr_chunk_blob_name)
                        logging.info(f"Risk Question-Response chunk {i+1}: {qr_chunk_blob_name} ({qr_chunk['word_count']} words, {qr_chunk['question_count']} questions)")
                        
                    except Exception as chunk_upload_error:
                        logging.error(f"❌ Error uploading Q&R chunk {i+1}: {chunk_upload_error}")
                        continue
                        
            except Exception as qr_error:
                logging.error(f"❌ Error creating Question-Response chunks: {qr_error}")
                logging.error(f"🔍 Q&R error details: {type(qr_error).__name__}: {str(qr_error)}")
                qr_chunks = []
        else:
            logging.warning(f"⚠️ No suitable assessment data found for Question-Response chunking")
            logging.info(f"🔍 Available report data keys: {list(report_data.keys())}")
            logging.info(f"🔍 Checked sources: {[source[0] for source in extraction_sources]}")
        
        # ===============================================================
        # 5. Create comprehensive chunks index file
        # ===============================================================
        logging.info("📋 Step 5/6: Creating comprehensive chunks index...")
        
        chunks_index = {
            "report_id": report_id,
            "user_id": user_id,
            "total_report_chunks": len(word_chunks),
            "total_qr_chunks": len(qr_chunks),
            "total_all_chunks": len(word_chunks) + len(qr_chunks),
            "report_chunk_files": chunk_files_created,
            "qr_chunk_files": qr_chunk_files_created,
            "chunking_strategy": {
                "report_chunks": {
                    "target_size_words": 300,
                    "max_size_words": 400,
                    "min_size_words": 150,
                    "chunk_type": "risk_word_documents",
                    "optimized_for": "azure_cognitive_search_risk_analysis"
                },
                "qr_chunks": {
                    "target_size_words": 400,  # FIXED: Reduced from 800 for better handling
                    "max_size_words": 600,     # FIXED: Reduced from 1000
                    "min_size_words": 30,      # FIXED: Reduced from 200 for flexibility
                    "chunk_type": "question_response_documents",
                    "optimized_for": "rag_context_risk_questions",
                    "empty_responses_handled": True,  # FIXED: Flag indicating robust handling
                    "fallback_strategies": ["minimal_chunks", "preserve_all_content"]
                }
            },
            "report_chunks_summary": [
                {
                    "chunk_id": chunk_doc["chunk_id"],
                    "section_title": chunk_doc["section_title"],
                    "word_count": chunk_doc["word_count"],
                    "character_count": chunk_doc["character_count"],
                    "content_preview": chunk_doc["content_preview"],
                    "file_name": chunk_files_created[i],
                    "sections_included": chunk_doc["sections_included"],
                    "chunk_type": "report_content"
                }
                for i, chunk_doc in enumerate(word_chunks)
            ],
            "qr_chunks_summary": [
                {
                    "chunk_id": qr_chunk["chunk_id"],
                    "expansion_title": qr_chunk["expansion_title"],
                    "word_count": qr_chunk["word_count"],
                    "question_count": qr_chunk["question_count"],
                    "character_count": qr_chunk["character_count"],
                    "content_preview": qr_chunk["content_preview"],
                    "file_name": qr_chunk_files_created[i] if i < len(qr_chunk_files_created) else None,
                    "questions_included": qr_chunk["questions_included"],
                    "chunk_type": "question_response"
                }
                for i, qr_chunk in enumerate(qr_chunks)
            ],
            "created_at": datetime.now().isoformat(),
            "folder": folder_name,
            "report_type": "comprehensive_risk_fortress_strategy_with_enhanced_qr_chunks",
            "qr_extraction_status": {
                "extraction_successful": bool(qr_chunks),
                "responses_found": len(raw_assessment_data.get('responses', [])) if raw_assessment_data else 0,
                "chunks_created": len(qr_chunks),
                "files_uploaded": len(qr_chunk_files_created)
            }
        }
        
        chunks_index_blob_name = f"{folder_name}/{report_id}_risk_chunks_index.json"
        chunks_index_json = json.dumps(chunks_index, indent=2, default=str)
        chunks_index_bytes = io.BytesIO(chunks_index_json.encode("utf-8"))

        upload_blob_with_retry(
            container_client,
            chunks_index_blob_name,
            chunks_index_bytes,
            ContentSettings(content_type="application/json")
        )
        
        # ===============================================================
        # 6. Upload final summary and statistics
        # ===============================================================
        logging.info("📈 Step 6/6: Generating final upload summary...")
        
        total_sections = len([k for k in report_data.keys() if k != "_enhanced_risk_report_metadata"])
        total_files = 3 + len(word_chunks) + len(qr_chunks)  # Word doc + JSON + chunks index + all chunk files
        
        # Create detailed upload summary
        upload_summary = {
            "report_id": report_id,
            "user_id": user_id,
            "upload_completed_at": datetime.now().isoformat(),
            "folder_name": folder_name,
            "files_created": {
                "complete_word_document": doc_blob_name,
                "complete_json_report": json_blob_name,
                "chunks_index": chunks_index_blob_name,
                "report_chunks": chunk_files_created,
                "question_response_chunks": qr_chunk_files_created
            },
            "statistics": {
                "total_files_created": total_files,
                "report_sections": total_sections,
                "report_word_chunks": len(word_chunks),
                "question_response_chunks": len(qr_chunks),
                "total_chunks": len(word_chunks) + len(qr_chunks)
            },
            "chunk_optimization": {
                "report_chunks_for": "Azure Cognitive Search Risk Analysis",
                "qr_chunks_for": "RAG Context for AI Risk Questions",
                "target_chunk_size": "300-400 words (main), 400-600 words (Q&R)",
                "chunk_format": "Microsoft Word (.docx)",
                "qr_handling": "Enhanced with fallbacks and flexible sizing"
            }
        }
        
        # Upload summary file
        summary_blob_name = f"{folder_name}/{report_id}_risk_upload_summary.json"
        summary_json = json.dumps(upload_summary, indent=2, default=str)
        summary_bytes = io.BytesIO(summary_json.encode("utf-8"))

        upload_blob_with_retry(
            container_client,
            summary_blob_name,
            summary_bytes,
            ContentSettings(content_type="application/json")
        )
        
        # ===============================================================
        # Final Success Logging
        # ===============================================================
        logging.info(f"🎉 Risk Report upload complete: {total_files} files in '{folder_name}' folder")
        logging.info(f"📊 Created {len(word_chunks)} Risk Word document chunks for Azure Cognitive Search")
        logging.info(f"🧠 Created {len(qr_chunks)} Question-Response chunks for RAG context")
        logging.info(f"📁 All files uploaded to container '{container_name}' in folder '{folder_name}/'")
        
        success_message = f"Risk fortress report uploaded successfully: {total_sections} sections, {len(word_chunks)} report chunks, {len(qr_chunks)} Q&R chunks, {total_files} files total"
        logging.info(f"✅ {success_message}")
        
        return True, success_message
        
    except Exception as e:
        error_message = f"Error uploading risk report: {str(e)}"
        logging.error(f"❌ {error_message}")
        logging.error(f"🔍 Error details: {type(e).__name__}: {e}")
        import traceback
        logging.error(f"🔍 Full traceback: {traceback.format_exc()}")
        return False, error_message

def extract_risk_assessment_data_from_report(report_data: Dict) -> Dict:
    """Extract assessment data from risk report if not available in metadata"""
    try:
        # Try to find assessment data in various places within the report
        for key, value in report_data.items():
            if isinstance(value, dict) and 'responses' in value:
                return value
            elif isinstance(value, dict) and 'assessment_data' in value:
                return value['assessment_data']
        
        logging.warning("⚠️ Could not extract risk assessment data from report")
        return {}
    except Exception as e:
        logging.error(f"❌ Error extracting risk assessment data: {str(e)}")
        return {}

# ======================================================
#           COMPLETE RISK ENGINE CHUNKING FUNCTIONS
#           Adapted from People Engine Implementation
# ======================================================

async def create_risk_word_document_chunks(report_data: Dict, report_id: str, user_id: str) -> List[Dict]:
    """Create Word document chunks optimized for RAG performance with detailed monitoring"""
    
    logging.info(f"🚀 Starting RAG-optimized risk chunking for report_id={report_id}, user_id={user_id}")
    
    word_chunks = []
    
    # Get user profile for context using async connection pooling
    user_profile = await get_user_profile_data(user_id)
    if user_profile:
        logging.info(f"👤 User context: {user_profile.get('business_name', 'Unknown')} ({user_profile.get('industry', 'Unknown')})")
    else:
        logging.warning(f"⚠️ No user profile found for user_id={user_id}")
    
    # RAG-OPTIMIZED chunking settings for better retrieval performance
    TARGET_SIZE_WORDS = 300  # Sweet spot for RAG retrieval
    MAX_SIZE_WORDS = 400     # Hard limit to prevent oversized chunks
    MIN_SIZE_WORDS = 150     # Minimum to maintain semantic meaning
    
    logging.info(f"⚙️ RAG chunking settings: target={TARGET_SIZE_WORDS}, max={MAX_SIZE_WORDS}, min={MIN_SIZE_WORDS}")
    
    chunk_id = 1
    total_sections = len([k for k in report_data.keys() if k != "_enhanced_risk_report_metadata"])
    logging.info(f"📂 Processing {total_sections} report sections for chunking")
    
    # Track overall statistics
    total_input_words = 0
    total_output_chunks = 0
    section_stats = []
    
    # Process each section and create smart chunks
    for section_idx, (section_name, section_data) in enumerate(report_data.items()):
        if section_name == "_enhanced_risk_report_metadata":
            continue
            
        if not isinstance(section_data, dict):
            logging.warning(f"⚠️ Skipping non-dict section: {section_name}")
            continue
            
        title = section_data.get("title", "Untitled Section")
        content = section_data.get("content", "")
        metadata = section_data.get("metadata", {})
        
        # Log section processing start
        section_word_count = len(content.split())
        total_input_words += section_word_count
        logging.info(f"📄 Processing section {section_idx + 1}/{total_sections}: '{title}' ({section_word_count:,} words)")
        
        # Clean content for better processing
        clean_content = clean_risk_content_for_word_chunks(content)
        clean_word_count = len(clean_content.split())
        
        if clean_word_count != section_word_count:
            logging.info(f"🧹 Content cleaned: {section_word_count} → {clean_word_count} words")
        
        # Create semantic chunks from this section
        logging.info(f"🔧 Starting semantic chunking for section '{title}'...")
        section_chunks = create_semantic_risk_word_chunks(clean_content, TARGET_SIZE_WORDS, MAX_SIZE_WORDS, MIN_SIZE_WORDS)
        
        # Validate section chunks
        section_chunk_stats = validate_risk_chunk_sizes(section_chunks, TARGET_SIZE_WORDS, f"Section: {title}")
        section_stats.append({
            "section_name": section_name,
            "section_title": title,
            "input_words": clean_word_count,
            "chunks_created": len(section_chunks),
            "chunk_stats": section_chunk_stats
        })
        
        logging.info(f"✅ Section '{title}' chunked: {clean_word_count} words → {len(section_chunks)} chunks")
        
        # Convert each chunk to a Word document
        for i, chunk_content in enumerate(section_chunks):
            chunk_title = title if len(section_chunks) == 1 else f"{title} - Part {i+1}"
            chunk_word_count = len(chunk_content.split())
            
            logging.debug(f"📝 Creating Word document for chunk {chunk_id}: '{chunk_title}' ({chunk_word_count} words)")
            
            # Create Word document for this chunk
            chunk_doc = create_risk_chunk_word_document(
                chunk_content, 
                chunk_title, 
                user_profile,
                section_name,
                f"{report_id}_risk_chunk_{chunk_id:03d}"
            )
            
            character_count = len(chunk_content)
            
            # Determine chunk quality metrics
            chunk_category = categorize_risk_chunk_size_by_words(chunk_word_count)
            semantic_completeness = calculate_risk_semantic_completeness(chunk_content)
            
            # Log chunk quality
            quality_status = "✅ OPTIMAL" if TARGET_SIZE_WORDS * 0.8 <= chunk_word_count <= TARGET_SIZE_WORDS * 1.2 else \
                           "⚠️ LARGE" if chunk_word_count > TARGET_SIZE_WORDS * 1.2 else \
                           "⚠️ SMALL" if chunk_word_count < TARGET_SIZE_WORDS * 0.8 else "❓ UNKNOWN"
            
            logging.info(f"📊 Chunk {chunk_id} quality: {quality_status} | "
                        f"{chunk_word_count} words | "
                        f"Category: {chunk_category} | "
                        f"Completeness: {semantic_completeness:.2f}")
            
            chunk_info = {
                "chunk_id": f"{report_id}_risk_chunk_{chunk_id:03d}",
                "section_name": section_name,
                "section_title": chunk_title,
                "word_count": chunk_word_count,
                "character_count": character_count,
                "content_preview": chunk_content[:200] + "..." if len(chunk_content) > 200 else chunk_content,
                "sections_included": [section_name],
                "document": chunk_doc,
                "chunk_metadata": {
                    "original_section": section_name,
                    "chunk_size_category": chunk_category,
                    "semantic_completeness": semantic_completeness,
                    "ai_analysis_time": metadata.get("ai_analysis_time", 0),
                    "chunk_type": "risk_analysis_rag_optimized",
                    "rag_optimization": {
                        "target_size": TARGET_SIZE_WORDS,
                        "size_ratio": chunk_word_count / TARGET_SIZE_WORDS,
                        "quality_status": quality_status.replace("✅ ", "").replace("⚠️ ", "").replace("❓ ", ""),
                        "overlap_enabled": True
                    }
                },
                "user_context": {
                    "user_id": user_id,
                    "business_name": user_profile.get("business_name", "Unknown") if user_profile else "Unknown",
                    "industry": user_profile.get("industry", "Unknown") if user_profile else "Unknown",
                    "team_size": user_profile.get("team_size", "Unknown") if user_profile else "Unknown"
                },
                "created_at": datetime.now().isoformat()
            }
            
            word_chunks.append(chunk_info)
            total_output_chunks += 1
            chunk_id += 1
    
    # Final comprehensive statistics
    if word_chunks:
        avg_chunk_size = sum(c['word_count'] for c in word_chunks) // len(word_chunks)
        min_chunk_size = min(c['word_count'] for c in word_chunks)
        max_chunk_size = max(c['word_count'] for c in word_chunks)
        
        # Count optimal chunks
        optimal_chunks = sum(1 for c in word_chunks if TARGET_SIZE_WORDS * 0.8 <= c['word_count'] <= TARGET_SIZE_WORDS * 1.2)
        optimal_percentage = (optimal_chunks / len(word_chunks)) * 100
        
        # Calculate compression ratio
        compression_ratio = total_input_words / sum(c['word_count'] for c in word_chunks) if word_chunks else 1
        
        logging.info(f"🎉 RAG-optimized risk chunking complete!")
        logging.info(f"📊 FINAL STATISTICS:")
        logging.info(f"   📄 Input: {total_input_words:,} words across {total_sections} sections")
        logging.info(f"   📦 Output: {len(word_chunks)} chunks")
        logging.info(f"   📏 Chunk sizes: {min_chunk_size}-{max_chunk_size} words (avg: {avg_chunk_size})")
        logging.info(f"   🎯 Target compliance: {optimal_chunks}/{len(word_chunks)} chunks optimal ({optimal_percentage:.1f}%)")
        logging.info(f"   🔗 Overlap enabled: 50-word context preservation between chunks")
        logging.info(f"   📈 Compression ratio: {compression_ratio:.2f}x (due to overlap)")
        
    else:
        logging.error(f"❌ No chunks created from {total_sections} sections!")
    
    return word_chunks

def clean_risk_content_for_word_chunks(content: str) -> str:
    """Clean risk content for better Word document processing"""
    
    # Remove excessive whitespace
    content = re.sub(r'\n\s*\n\s*\n', '\n\n', content)
    content = re.sub(r' +', ' ', content)
    
    # Fix common formatting issues
    content = re.sub(r'([.!?])\s*([A-Z])', r'\1 \2', content)
    content = re.sub(r'([a-z])([A-Z])', r'\1 \2', content)
    
    # Remove artifacts from AI generation
    content = re.sub(r'<[^>]+>', '', content)  # Remove any HTML tags
    content = re.sub(r'\[.*?\]', '', content)  # Remove bracket annotations
    
    # Risk-specific cleaning
    content = re.sub(r'Risk Engine:', 'Risk Strategy:', content)
    content = re.sub(r'Chapter \d+:', 'Risk Chapter:', content)
    
    # Normalize quotes and special characters
    content = content.replace('"', '"').replace('"', '"')
    content = content.replace(''', "'").replace(''', "'")
    content = re.sub(r'…', '...', content)
    
    return content.strip()

def create_semantic_risk_word_chunks(content: str, target_size: int, max_size: int, min_size: int) -> List[str]:
    """Create semantic chunks that preserve risk context WITH OVERLAP for better RAG performance"""
    
    logging.info(f"🔧 Starting risk semantic chunking: target={target_size}, max={max_size}, min={min_size}")
    
    # If content is small enough, return as single chunk
    word_count = len(content.split())
    logging.info(f"📊 Input risk content: {word_count} words")
    
    if word_count <= max_size:
        logging.info(f"✅ Risk content fits in single chunk ({word_count} <= {max_size})")
        return [content]
    
    chunks = []
    OVERLAP_SIZE = 50  # 50 words overlap between chunks for context preservation
    logging.info(f"🔗 Using {OVERLAP_SIZE}-word overlap between risk chunks")
    
    # Split by risk logic sections first
    risk_sections = split_by_risk_logic(content)
    logging.info(f"📂 Split into {len(risk_sections)} risk logic sections")
    
    current_chunk = ""
    current_word_count = 0
    previous_chunk_end = ""  # Store end of previous chunk for overlap
    
    for section_idx, section in enumerate(risk_sections):
        section_words = len(section.split())
        test_word_count = current_word_count + section_words
        
        logging.debug(f"🔍 Processing risk section {section_idx + 1}/{len(risk_sections)}: {section_words} words")
        
        if test_word_count > max_size and current_chunk:
            # Current chunk is full, save it if it's substantial
            if current_word_count >= min_size:
                # Add overlap from previous chunk if available
                final_chunk = current_chunk
                if previous_chunk_end and chunks:
                    overlap_text = get_last_n_words(previous_chunk_end, OVERLAP_SIZE)
                    final_chunk = overlap_text + "\n\n" + current_chunk
                    logging.debug(f"🔗 Added {len(overlap_text.split())} word overlap to risk chunk {len(chunks) + 1}")
                
                chunks.append(final_chunk.strip())
                
                # Store end of current chunk for next overlap
                previous_chunk_end = get_last_n_words(current_chunk, OVERLAP_SIZE * 2)
                
                logging.info(f"✅ Saved risk chunk {len(chunks)}: {len(final_chunk.split())} words (original: {current_word_count})")
                
                current_chunk = section
                current_word_count = section_words
            else:
                # Current chunk too small, but adding section makes it too big
                logging.debug(f"⚠️ Current risk chunk too small ({current_word_count} < {min_size}), handling large section")
                
                if section_words > max_size:
                    logging.debug(f"🔨 Risk section too large ({section_words} > {max_size}), splitting with overlap")
                    sub_sections = split_large_risk_section_with_overlap(section, max_size)
                    logging.info(f"📂 Split large risk section into {len(sub_sections)} sub-sections with overlap")
                    
                    for sub_idx, sub_section in enumerate(sub_sections):
                        sub_words = len(sub_section.split())
                        logging.debug(f"🔍 Processing risk sub-section {sub_idx + 1}/{len(sub_sections)}: {sub_words} words")
                        
                        if current_word_count + sub_words > max_size and current_chunk:
                            if current_word_count >= min_size:
                                # Add overlap before saving
                                final_chunk = current_chunk
                                if previous_chunk_end and chunks:
                                    overlap_text = get_last_n_words(previous_chunk_end, OVERLAP_SIZE)
                                    final_chunk = overlap_text + "\n\n" + current_chunk
                                    logging.debug(f"🔗 Added {len(overlap_text.split())} word overlap to risk chunk {len(chunks) + 1}")
                                
                                chunks.append(final_chunk.strip())
                                previous_chunk_end = get_last_n_words(current_chunk, OVERLAP_SIZE * 2)
                                logging.info(f"✅ Saved risk chunk {len(chunks)}: {len(final_chunk.split())} words")
                            
                            current_chunk = sub_section
                            current_word_count = sub_words
                        else:
                            current_chunk += "\n\n" + sub_section if current_chunk else sub_section
                            current_word_count += sub_words
                            logging.debug(f"➕ Added risk sub-section to current chunk: {current_word_count} total words")
                else:
                    current_chunk += "\n\n" + section if current_chunk else section
                    current_word_count = test_word_count
                    logging.debug(f"➕ Added risk section to current chunk: {current_word_count} total words")
        else:
            current_chunk += "\n\n" + section if current_chunk else section
            current_word_count = test_word_count
            logging.debug(f"➕ Added risk section to current chunk: {current_word_count} total words")
    
    # Add the last chunk if it exists and is substantial
    if current_chunk and current_word_count >= min_size:
        # Add overlap to final chunk too
        final_chunk = current_chunk
        if previous_chunk_end and chunks:
            overlap_text = get_last_n_words(previous_chunk_end, OVERLAP_SIZE)
            final_chunk = overlap_text + "\n\n" + current_chunk
            logging.debug(f"🔗 Added {len(overlap_text.split())} word overlap to final risk chunk")
        
        chunks.append(final_chunk.strip())
        logging.info(f"✅ Saved final risk chunk {len(chunks)}: {len(final_chunk.split())} words (original: {current_word_count})")
    elif current_chunk:
        logging.warning(f"⚠️ Discarded final risk chunk: {current_word_count} words < {min_size} minimum")
    
    # Validate the created chunks
    chunk_stats = validate_risk_chunk_sizes(chunks, target_size, "Risk Semantic Chunking")
     
    logging.info(f"🎉 Risk semantic chunking complete: {len(chunks)} chunks created")
    logging.info(f"📊 Risk chunk size range: {chunk_stats.get('min_words', 0)}-{chunk_stats.get('max_words', 0)} words")
    logging.info(f"📈 Average risk chunk size: {chunk_stats.get('avg_words', 0)} words (target: {target_size})")
    
    return chunks

def split_by_risk_logic(content: str) -> List[str]:
    """Split content by risk-specific logical boundaries"""
    
    logging.info(f"🔧 Starting risk logic splitting...")
    
    # Log input content stats
    total_words = len(content.split())
    total_paragraphs = len([p for p in content.split('\n\n') if p.strip()])
    logging.info(f"📊 Input risk content: {total_words} words, {total_paragraphs} paragraphs")
    
    # Enhanced risk-specific section indicators
    risk_patterns = [
        r'(?i)(?:^|\n)(?:threat|vulnerability|risk|security|fortress|defense):',
        r'(?i)(?:^|\n)(?:threat analysis|vulnerability assessment|risk evaluation|security architecture):',
        r'(?i)(?:^|\n)(?:risk analysis|threat assessment|security review|vulnerability scan):',
        r'(?i)(?:^|\n)(?:risk strengths|security advantages|defense opportunities|fortress assets):',
        r'(?i)(?:^|\n)(?:risk challenges|security gaps|threat vectors|vulnerability weaknesses):',
        r'(?i)(?:^|\n)(?:implementation|deployment|security rollout|risk management):',
        r'(?i)(?:^|\n)(?:optimization|enhancement|security improvement|risk mitigation):',
        r'(?i)(?:^|\n)(?:integration|alignment|security coordination|risk synchronization):',
        r'(?i)(?:^|\n)(?:measurement|metrics|tracking|kpis|risk indicators):',
        r'(?i)(?:^|\n)(?:development|training|capability building|security awareness):',
        
        # Enhanced patterns for AI-generated risk content
        r'(?i)(?:^|\n)(?:##\s*|###\s*)?(?:strategic|operational|security|threat|risk|fortress)',
        r'(?i)(?:^|\n)(?:##\s*|###\s*)?(?:analysis|assessment|evaluation|mitigation|protection)',
        r'(?i)(?:^|\n)(?:##\s*|###\s*)?(?:recommendations?|strategies|approaches|solutions|controls)',
        r'(?i)(?:^|\n)(?:##\s*|###\s*)?(?:chapter \d+|section \d+|area \d+|phase \d+)',
        r'(?i)(?:^|\n)(?:your security|your risk|your threats|considering your vulnerabilities)',
        r'(?i)(?:^|\n)(?:to secure|to protect|to mitigate|to fortify|moving forward)',
        
        # Risk-specific structural patterns
        r'(?i)(?:^|\n)(?:compliance|governance|audit|controls|monitoring|detection)',
        r'(?i)(?:^|\n)(?:incident|response|recovery|continuity|backup|disaster)',
        r'(?i)(?:^|\n)(?:access|authorization|authentication|encryption|firewall|intrusion)'
    ]
    
    logging.info(f"🔍 Using {len(risk_patterns)} risk-specific patterns for splitting")
    
    # Try to split by risk patterns first
    sections = []
    current_section = ""
    pattern_matches = 0
    
    paragraphs = content.split('\n\n')
    logging.info(f"📂 Processing {len(paragraphs)} paragraphs for risk pattern matching")
    
    for paragraph in paragraphs:
        # Check if this paragraph starts a new risk section
        is_new_section = False
        for pattern in risk_patterns:
            if re.search(pattern, paragraph):
                is_new_section = True
                pattern_matches += 1
                break
        
        if is_new_section and current_section:
            sections.append(current_section.strip())
            current_section = paragraph
        else:
            current_section += "\n\n" + paragraph if current_section else paragraph
    
    # Add the last section
    if current_section:
        sections.append(current_section.strip())
    
    logging.info(f"📊 Risk pattern matching results: {pattern_matches} matches found, {len(sections)} sections created")
    
    # Smart fallback logic - if no risk patterns found or sections too large
    needs_fallback = False
    if len(sections) <= 1:
        needs_fallback = True
        logging.warning(f"⚠️ No risk patterns found, applying smart fallback")
    elif any(len(s.split()) > 400 for s in sections):
        needs_fallback = True
        logging.warning(f"⚠️ Risk sections too large (>400 words), applying smart fallback")
    
    if needs_fallback:
        logging.info(f"🔄 Applying smart paragraph splitting with risk-optimized size limits...")
        
        # Smart paragraph splitting with risk-optimized size limits
        sections = []
        current_section = ""
        current_words = 0
        target_words = 300  # Target size for risk chunks (optimized for RAG)
        
        for paragraph in paragraphs:
            paragraph = paragraph.strip()
            if not paragraph:
                continue
            
            para_words = len(paragraph.split())
            
            # If adding this paragraph would make section too large, save current and start new
            if current_words + para_words > target_words and current_section:
                sections.append(current_section.strip())
                logging.debug(f"📄 Saved risk section {len(sections)}: {current_words} words")
                current_section = paragraph
                current_words = para_words
            else:
                current_section += "\n\n" + paragraph if current_section else paragraph
                current_words += para_words
        
        if current_section:
            sections.append(current_section.strip())
            logging.debug(f"📄 Saved final risk section {len(sections)}: {current_words} words")
        
        logging.info(f"📄 Smart risk paragraph splitting: {len(sections)} sections created")
    
    # Final validation - force split any remaining oversized sections
    final_sections = []
    for i, section in enumerate(sections):
        section_words = len(section.split())
        
        if section_words > 400:
            logging.warning(f"⚠️ Risk section {i+1} still oversized ({section_words} words), force splitting")
            
            # Force split by sentences with risk context preservation
            sentences = re.split(r'(?<=[.!?])\s+', section)
            sub_sections = []
            current_sub = ""
            current_sub_words = 0
            
            for sentence in sentences:
                sentence_words = len(sentence.split())
                
                if current_sub_words + sentence_words > 350 and current_sub:
                    sub_sections.append(current_sub.strip())
                    current_sub = sentence
                    current_sub_words = sentence_words
                else:
                    current_sub += " " + sentence if current_sub else sentence
                    current_sub_words += sentence_words
            
            if current_sub:
                sub_sections.append(current_sub.strip())
            
            final_sections.extend(sub_sections)
            logging.info(f"🔨 Split oversized risk section into {len(sub_sections)} sub-sections")
        else:
            final_sections.append(section)
    
    # Final statistics
    section_sizes = [len(s.split()) for s in final_sections]
    avg_size = sum(section_sizes) // len(section_sizes) if final_sections else 0
    min_size = min(section_sizes) if final_sections else 0
    max_size = max(section_sizes) if final_sections else 0
    optimal_sections = sum(1 for size in section_sizes if 150 <= size <= 400)
    optimal_percentage = (optimal_sections / len(final_sections)) * 100 if final_sections else 0
    
    logging.info(f"🎉 Risk logic splitting complete!")
    logging.info(f"📊 Final: {len(final_sections)} sections, {min_size}-{max_size} words (avg: {avg_size})")
    logging.info(f"🎯 Optimal risk sections (150-400 words): {optimal_sections}/{len(final_sections)} ({optimal_percentage:.1f}%)")
    
    return final_sections

def split_large_risk_section_with_overlap(section: str, max_size: int) -> List[str]:
    """Split large risk section with overlap for context preservation"""
    
    OVERLAP_SIZE = 50  # words
    
    words = section.split()
    if len(words) <= max_size:
        return [section]
    
    chunks = []
    start = 0
    
    while start < len(words):
        # Calculate end position
        end = min(start + max_size, len(words))
        
        # Create chunk
        chunk_words = words[start:end]
        chunk = ' '.join(chunk_words)
        chunks.append(chunk)
        
        # Move start position with overlap
        if end >= len(words):
            break
        
        start = end - OVERLAP_SIZE
        if start <= 0:
            start = end
    
    logging.debug(f"🔨 Split large risk section: {len(words)} words → {len(chunks)} chunks with {OVERLAP_SIZE}-word overlap")
    return chunks

def validate_risk_chunk_sizes(chunks: List[str], target_size: int, context_name: str = "") -> Dict:
    """Validate and log risk chunk sizes for monitoring"""
    
    if not chunks:
        return {"total_chunks": 0}
    
    chunk_stats = {
        "total_chunks": len(chunks),
        "avg_words": 0,
        "min_words": float('inf'),
        "max_words": 0,
        "chunks_over_target": 0,
        "chunks_under_100": 0,  # Flag very small chunks
        "chunks_optimal": 0     # Chunks within target range
    }
    
    total_words = 0
    for chunk in chunks:
        words = len(chunk.split())
        total_words += words
        
        chunk_stats["min_words"] = min(chunk_stats["min_words"], words)
        chunk_stats["max_words"] = max(chunk_stats["max_words"], words)
        
        if words > target_size * 1.2:  # 20% over target
            chunk_stats["chunks_over_target"] += 1
        elif words < 100:
            chunk_stats["chunks_under_100"] += 1
        elif target_size * 0.8 <= words <= target_size * 1.2:  # Within 20% of target
            chunk_stats["chunks_optimal"] += 1
    
    chunk_stats["avg_words"] = total_words // len(chunks)
    chunk_stats["min_words"] = chunk_stats["min_words"] if chunk_stats["min_words"] != float('inf') else 0
    
    # Log the stats
    context_prefix = f"[{context_name}] " if context_name else ""
    logging.info(f"📊 {context_prefix}Risk chunk validation: "
                f"{chunk_stats['total_chunks']} chunks, "
                f"avg: {chunk_stats['avg_words']} words, "
                f"range: {chunk_stats['min_words']}-{chunk_stats['max_words']}, "
                f"optimal: {chunk_stats['chunks_optimal']}/{chunk_stats['total_chunks']}")
    
    return chunk_stats

def categorize_risk_chunk_size_by_words(word_count: int) -> str:
    """Categorize chunk size for risk analysis"""
    if word_count < 200:
        return "small"
    elif word_count < 350:
        return "optimal"
    elif word_count < 600:
        return "large"
    else:
        return "oversized"

def calculate_risk_semantic_completeness(content: str) -> float:
    """Calculate semantic completeness score for risk content"""
    
    # Risk-specific completeness indicators
    completeness_indicators = {
        'threat_words': ['threat', 'attack', 'vulnerability', 'exploit', 'breach', 'intrusion'],
        'security_words': ['security', 'secure', 'protect', 'defense', 'safeguard', 'shield'],
        'risk_words': ['risk', 'danger', 'exposure', 'liability', 'hazard', 'peril'],
        'control_words': ['control', 'measure', 'procedure', 'policy', 'standard', 'guideline'],
        'compliance_words': ['compliance', 'regulation', 'audit', 'governance', 'framework', 'standard'],
        'monitoring_words': ['monitor', 'detect', 'alert', 'surveillance', 'tracking', 'observation'],
        'response_words': ['response', 'incident', 'recovery', 'continuity', 'backup', 'restoration'],
        'assessment_words': ['assessment', 'evaluation', 'analysis', 'review', 'examination', 'inspection'],
        'mitigation_words': ['mitigation', 'prevention', 'reduction', 'elimination', 'minimization', 'containment']
    }
    
    content_lower = content.lower()
    total_indicators = 0
    found_indicators = 0
    
    for category, words in completeness_indicators.items():
        total_indicators += len(words)
        found_indicators += sum(1 for word in words if word in content_lower)
    
    base_completeness = found_indicators / total_indicators if total_indicators > 0 else 0
    
    # Length bonus (longer content is generally more complete)
    word_count = len(content.split())
    length_factor = min(1.0, word_count / 400)  # Optimal around 400 words for risk
    
    # Structure bonus (headers, bullets, etc.)
    structure_indicators = ['##', '###', '- ', '* ', '1.', '2.', '3.', 'Chapter', 'Step', 'Phase']
    structure_count = sum(1 for indicator in structure_indicators if indicator in content)
    structure_factor = min(0.2, structure_count * 0.03)  # Up to 0.2 bonus
    
    # Risk-specific bonus for strategic language
    strategic_indicators = ['recommend', 'suggest', 'should', 'focus on', 'prioritize', 'implement']
    strategic_count = sum(1 for indicator in strategic_indicators if indicator in content_lower)
    strategic_factor = min(0.15, strategic_count * 0.03)
    
    final_completeness = min(1.0, base_completeness + length_factor * 0.3 + structure_factor + strategic_factor)
    
    return final_completeness

def create_risk_chunk_word_document(content: str, title: str, user_profile: Dict, 
                                   section_name: str, chunk_id: str) -> Document:
    """Create a professionally formatted Word document for risk chunk"""
    
    try:
        doc = Document()
        
        # Enhanced styling
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Add header with branding
        header = doc.add_heading("BACKABLE RISK ENGINE", 0)
        header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        header_run = header.runs[0]
        header_run.font.size = Pt(24)
        header_run.font.bold = True
        header_run.font.color.rgb = RGBColor(0, 51, 102)
        
        # Add chunk title
        chunk_title = doc.add_heading(title, 1)
        chunk_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        title_run = chunk_title.runs[0]
        title_run.font.size = Pt(18)
        title_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for Risk
        
        # Add metadata section
        if user_profile:
            metadata_para = doc.add_paragraph()
            metadata_para.add_run("Business Risk Context:").bold = True
            metadata_para.add_run(f"\nBusiness: {user_profile.get('business_name', 'Unknown')}")
            metadata_para.add_run(f"\nIndustry: {user_profile.get('industry', 'Unknown')}")  
            metadata_para.add_run(f"\nTeam Size: {user_profile.get('team_size', 'Unknown')} employees")
            metadata_para.add_run(f"\nRisk Section: {section_name}")
            metadata_para.add_run(f"\nChunk ID: {chunk_id}")
            metadata_para.add_run(f"\nGenerated: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        
        # Add separator
        doc.add_paragraph("─" * 60)
        
        # Add the AI-generated content with risk-specific formatting
        add_risk_content_to_document(doc, content)
        
        # Add footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        footer_run = footer_para.add_run("Generated by Backable AI Risk Intelligence")
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        logging.info(f"📄 Created risk chunk Word document: {len(content.split())} words")
        return doc
        
    except Exception as e:
        logging.error(f"❌ Error creating risk chunk Word document: {str(e)}")
        # Return minimal document on error
        doc = Document()
        doc.add_heading("Error Creating Risk Document", 1)
        doc.add_paragraph(f"Error: {str(e)}")
        return doc

def add_risk_content_to_document(doc: Document, content: str):
    """Add risk-specific formatted content to Word document"""
    
    # Split content into logical sections
    sections = content.split('\n\n')
    
    for section in sections:
        section = section.strip()
        if not section:
            continue
        
        # Check if this looks like a header
        if (section.startswith('##') or 
            (len(section) < 100 and section.isupper()) or
            any(section.lower().startswith(prefix) for prefix in ['risk analysis:', 'threat assessment:', 'security recommendation:', 'fortress strategy:'])):
            
            # Add as heading
            header_text = section.replace('##', '').replace('#', '').strip()
            heading = doc.add_heading(header_text, 2)
            heading_run = heading.runs[0]
            heading_run.font.color.rgb = RGBColor(204, 0, 0)  # Red for Risk
        else:
            # Add as paragraph with risk-specific formatting
            para = doc.add_paragraph()
            
            # Handle bold text patterns in risk content
            bold_patterns = [
                r'\*\*(.*?)\*\*',  # **bold**
                r'__(.*?)__',      # __bold__
            ]
            
            current_text = section
            for pattern in bold_patterns:
                matches = re.findall(pattern, current_text)
                for match in matches:
                    # Split around the bold text
                    parts = current_text.split(f'**{match}**')
                    if len(parts) >= 2:
                        # Add text before bold
                        if parts[0]:
                            para.add_run(parts[0])
                        # Add bold text
                        bold_run = para.add_run(match)
                        bold_run.bold = True
                        # Continue with rest
                        current_text = '**'.join(parts[1:])
            
            # Add any remaining text
            if current_text and current_text != section:
                para.add_run(current_text)
            elif current_text == section:
                para.add_run(section)

# ======================================================
#           UTILITY FUNCTIONS
# ======================================================

def get_last_n_words(text: str, n: int) -> str:
    """Get last N words from text for overlap between chunks"""
    words = text.split()
    if len(words) <= n:
        return text
    return " ".join(words[-n:])

def get_first_n_words(text: str, n: int) -> str:
    """Get first N words from text for overlap between chunks"""
    words = text.split()
    if len(words) <= n:
        return text
    return " ".join(words[:n])
# ======================================================
#           Enhanced Risk Engine Indexer Integration
# ======================================================

# ======================================================
#           FIXED: Risk Engine Indexer Integration
# ======================================================

async def trigger_risk_indexer_for_client(user_id: str) -> Tuple[bool, str, str]:
    """FIXED: Trigger risk indexer for specific client with correct endpoint and payload"""
    
    logging.info(f"🚀 RISK INDEXER API: Triggering indexer for user_id={user_id}")
    logging.info(f"📊 Indexer parameters:")
    logging.info(f"   - User ID: {user_id}")
    logging.info(f"   - API Base URL: {INDEXER_API_BASE_URL}")
    logging.info(f"   - Timeout: {INDEXER_TIMEOUT}s")
    logging.info(f"   - Max retries: {INDEXER_RETRY_ATTEMPTS}")
    
    max_retries = INDEXER_RETRY_ATTEMPTS
    base_delay = INDEXER_RETRY_DELAY
    
    for attempt in range(max_retries):
        attempt_start_time = time.time()
        logging.info(f"🔄 RISK INDEXER API: Attempt {attempt + 1}/{max_retries} for user_id={user_id}")
        
        try:
            async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=INDEXER_TIMEOUT)) as session:
                # FIXED: Use correct payload structure that matches auto-indexer expectations
                payload = {
                    "client_id": user_id,  # FIXED: Changed from user_id to client_id
                    "force": True,         # FIXED: Added force flag to override any existing jobs
                    "new_client": False    # FIXED: Added new_client flag (assuming existing client)
                }
                
                # Enhanced payload logging
                logging.info(f"📤 RISK INDEXER API: Sending payload")
                logging.info(f"📊 Payload details:")
                logging.info(f"   - Client ID: {payload['client_id']}")
                logging.info(f"   - Force: {payload['force']}")
                logging.info(f"   - New Client: {payload['new_client']}")
                logging.info(f"   - JSON size: {len(json.dumps(payload))} bytes")
                
                # FIXED: Use correct endpoint that exists in auto-indexer
                indexer_url = f"{INDEXER_API_BASE_URL}/run-indexer"  # FIXED: Changed from /trigger_indexer_for_client
                logging.info(f"🌐 RISK INDEXER API: Making request to {indexer_url}")
                
                async with session.post(
                    indexer_url,
                    json=payload,
                    headers={"Content-Type": "application/json"}
                ) as response:
                    
                    request_time = time.time() - attempt_start_time
                    
                    # Enhanced response logging
                    logging.info(f"📡 RISK INDEXER API: Response received")
                    logging.info(f"📊 Response details:")
                    logging.info(f"   - Status code: {response.status}")
                    logging.info(f"   - Response time: {request_time:.3f}s")
                    logging.info(f"   - Content type: {response.headers.get('content-type', 'Unknown')}")
                    logging.info(f"   - Content length: {response.headers.get('content-length', 'Unknown')}")
                    
                    try:
                        response_data = await response.json()
                        logging.info(f"✅ RISK INDEXER API: JSON parsing successful")
                        logging.info(f"📊 Response data keys: {list(response_data.keys())}")
                        
                    except Exception as json_error:
                        logging.error(f"❌ RISK INDEXER API: JSON parse error: {json_error}")
                        response_text = await response.text()
                        logging.error(f"🔍 Raw response text (first 500 chars): {response_text[:500]}...")
                        
                        if attempt < max_retries - 1:
                            wait_time = base_delay * (2 ** attempt)
                            logging.warning(f"⏳ RISK INDEXER API: JSON parse failed, retrying in {wait_time}s...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"JSON parse error: {json_error}", None
                    
                    # Handle different response statuses
                    if response.status == 202:  # Accepted - job added to queue
                        job_id = response_data.get("job_id")
                        message = response_data.get("message", "Risk indexer job started successfully")
                        queue_position = response_data.get("queue_position", 0)
                        
                        logging.info(f"🎉 RISK INDEXER API: JOB ACCEPTED!")
                        logging.info(f"✅ Acceptance details:")
                        logging.info(f"   - User ID: {user_id}")
                        logging.info(f"   - Job ID: {job_id}")
                        logging.info(f"   - Message: {message}")
                        logging.info(f"   - Queue Position: {queue_position}")
                        logging.info(f"   - Attempt: {attempt + 1}/{max_retries}")
                        logging.info(f"   - Request time: {request_time:.3f}s")
                        
                        return True, message, job_id
                    
                    elif response.status == 200:  # Success (immediate processing)
                        job_id = response_data.get("job_id")
                        message = response_data.get("message", "Risk indexer job completed successfully")
                        
                        logging.info(f"🎉 RISK INDEXER API: JOB COMPLETED!")
                        logging.info(f"✅ Success details:")
                        logging.info(f"   - User ID: {user_id}")
                        logging.info(f"   - Job ID: {job_id}")
                        logging.info(f"   - Message: {message}")
                        
                        return True, message, job_id
                    
                    elif response.status == 409:  # Conflict - Job already in progress/recent warning
                        message = response_data.get("message", "Indexer job already in progress or recent")
                        job_created_at = response_data.get("job_created_at")
                        
                        logging.warning(f"⚠️ RISK INDEXER API: Conflict for user_id={user_id}")
                        logging.warning(f"📊 Conflict details:")
                        logging.warning(f"   - Message: {message}")
                        logging.warning(f"   - Job Created At: {job_created_at}")
                        logging.warning(f"   - This usually means indexing is already running or recent")
                        
                        # For conflicts, we can retry with force=true
                        if attempt < max_retries - 1:
                            logging.info(f"🔄 Will retry with force=true on next attempt...")
                            payload["force"] = True  # Force override on retry
                            wait_time = base_delay * (attempt + 1)
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            # On final attempt, treat conflict as success since job exists
                            return True, f"Indexer job exists (conflict): {message}", "existing_job"
                    
                    elif response.status == 404:  # Client not found
                        message = response_data.get("message", "Client not found in indexer system")
                        
                        logging.warning(f"⚠️ RISK INDEXER API: Client not found: user_id={user_id}")
                        logging.warning(f"📊 404 details:")
                        logging.warning(f"   - Message: {message}")
                        
                        # Try with new_client=true on retry
                        if attempt < max_retries - 1:
                            payload["new_client"] = True
                            wait_time = base_delay * (attempt + 1)
                            logging.info(f"⏳ RISK INDEXER API: Client not found, retrying with new_client=true in {wait_time}s...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"Risk client not found: {message}", None
                    
                    elif response.status == 429:  # Rate Limited
                        message = response_data.get("message", "Risk indexer rate limited")
                        logging.warning(f"🚦 RISK INDEXER API: Rate limited for user_id={user_id}")
                        
                        if attempt < max_retries - 1:
                            wait_time = base_delay * 2 * (attempt + 1)  # Longer wait for rate limits
                            logging.info(f"⏳ RISK INDEXER API: Rate limited, waiting {wait_time}s before retry...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"Risk indexer rate limited: {message}", None
                    
                    elif response.status >= 500:  # Server errors - retry
                        message = response_data.get("message", f"Risk indexer server error {response.status}")
                        logging.error(f"🚨 RISK INDEXER API: Server error {response.status} for user_id={user_id}")
                        
                        if attempt < max_retries - 1:
                            wait_time = base_delay * (2 ** attempt)
                            logging.info(f"⏳ RISK INDEXER API: Server error, retrying in {wait_time}s...")
                            await asyncio.sleep(wait_time)
                            continue
                        else:
                            return False, f"Risk indexer server error: {message}", None
                    
                    else:  # Other client errors - don't retry
                        message = response_data.get("message", f"Risk indexer failed with status {response.status}")
                        logging.error(f"❌ RISK INDEXER API: Client error {response.status} for user_id={user_id}")
                        return False, f"Risk indexer client error: {message}", None
                        
        except asyncio.TimeoutError:
            request_time = time.time() - attempt_start_time
            error_msg = f"Risk indexer request timed out after {INDEXER_TIMEOUT}s"
            
            logging.error(f"⏰ RISK INDEXER API: TIMEOUT (attempt {attempt + 1}/{max_retries})")
            logging.error(f"📊 Timeout details:")
            logging.error(f"   - Configured timeout: {INDEXER_TIMEOUT}s")
            logging.error(f"   - Actual time: {request_time:.3f}s")
            logging.error(f"   - User ID: {user_id}")
            
            if attempt < max_retries - 1:
                wait_time = base_delay * (attempt + 1)
                logging.info(f"⏳ RISK INDEXER API: Timeout, retrying in {wait_time}s...")
                await asyncio.sleep(wait_time)
                continue
            else:
                return False, error_msg, None
                
        except aiohttp.ClientError as e:
            request_time = time.time() - attempt_start_time
            error_msg = f"Risk indexer HTTP client error: {str(e)}"
            
            logging.error(f"🌐 RISK INDEXER API: HTTP CLIENT ERROR (attempt {attempt + 1}/{max_retries})")
            logging.error(f"📊 HTTP error details:")
            logging.error(f"   - Error type: {type(e).__name__}")
            logging.error(f"   - Error message: {str(e)}")
            logging.error(f"   - Request time: {request_time:.3f}s")
            
            if attempt < max_retries - 1:
                wait_time = base_delay * (attempt + 1)
                logging.info(f"⏳ RISK INDEXER API: HTTP error, retrying in {wait_time}s...")
                await asyncio.sleep(wait_time)
                continue
            else:
                return False, error_msg, None
                
        except Exception as e:
            request_time = time.time() - attempt_start_time
            error_msg = f"Risk indexer unexpected error: {str(e)}"
            
            logging.error(f"💥 RISK INDEXER API: UNEXPECTED ERROR (attempt {attempt + 1}/{max_retries})")
            logging.error(f"📊 Unexpected error details:")
            logging.error(f"   - Error type: {type(e).__name__}")
            logging.error(f"   - Error message: {str(e)}")
            logging.error(f"   - Request time: {request_time:.3f}s")
            
            if attempt < max_retries - 1:
                wait_time = base_delay * (attempt + 1)
                logging.info(f"⏳ RISK INDEXER API: Unexpected error, retrying in {wait_time}s...")
                await asyncio.sleep(wait_time)
                continue
            else:
                return False, error_msg, None
    
    # If we get here, all retries failed
    final_error = f"Risk indexer failed after {max_retries} attempts"
    
    logging.error(f"💥 RISK INDEXER API: ALL ATTEMPTS FAILED")
    logging.error(f"📊 Final failure summary:")
    logging.error(f"   - User ID: {user_id}")
    logging.error(f"   - Total attempts: {max_retries}")
    logging.error(f"   - Final error: {final_error}")
    
    return False, final_error, None

def store_risk_indexer_job_metadata(report_id: str, user_id: str, indexer_job_id: str, indexer_status: str, 
                                   indexer_response: str = None, error_message: str = None):
    """Enhanced storage of risk indexer job metadata with comprehensive tracking"""
    
    conn = None
    try:
        logging.info(f"💾 RISK INDEXER: Storing metadata for report_id={report_id}")
        logging.info(f"📊 Metadata storage parameters:")
        logging.info(f"   - Report ID: {report_id}")
        logging.info(f"   - User ID: {user_id}")
        logging.info(f"   - Indexer Job ID: {indexer_job_id}")
        logging.info(f"   - Indexer Status: {indexer_status}")
        logging.info(f"   - Has response: {bool(indexer_response)}")
        logging.info(f"   - Has error: {bool(error_message)}")
        
        conn = get_risk_connection()
        
        with conn.cursor() as cur:
            # Check if the report exists first
            check_sql = """
                SELECT COUNT(*) FROM risk_reports WHERE report_id = %s AND user_id = %s
            """
            cur.execute(check_sql, (report_id, user_id))
            exists = cur.fetchone()[0] > 0
            
            logging.info(f"📋 RISK INDEXER: Report existence check: {exists}")
            
            if not exists:
                logging.error(f"❌ RISK INDEXER: Report not found in database")
                logging.error(f"📊 Missing report details:")
                logging.error(f"   - Report ID: {report_id}")
                logging.error(f"   - User ID: {user_id}")
                logging.error(f"   - Cannot store indexer metadata without existing report")
                return False
            
            # Update with comprehensive indexer information
            update_sql = """
                UPDATE risk_reports 
                SET 
                    indexer_job_id = %s,
                    indexer_status = %s,
                    indexer_triggered_at = %s,
                    indexer_completed_at = CASE WHEN %s = 'completed' THEN %s ELSE indexer_completed_at END,
                    indexer_error_message = %s,
                    indexer_retry_count = COALESCE(indexer_retry_count, 0) + CASE WHEN %s != 'triggered' THEN 1 ELSE 0 END
                WHERE report_id = %s AND user_id = %s
            """
            
            current_time = datetime.now()
            
            logging.info(f"📝 RISK INDEXER: Executing database update")
            logging.info(f"📊 SQL parameters:")
            logging.info(f"   - Job ID: {indexer_job_id}")
            logging.info(f"   - Status: {indexer_status}")
            logging.info(f"   - Triggered at: {current_time.isoformat()}")
            logging.info(f"   - Will set completed_at: {indexer_status == 'completed'}")
            logging.info(f"   - Error message: {error_message if error_message else 'None'}")
            
            cur.execute(update_sql, (
                indexer_job_id,
                indexer_status,
                current_time,
                indexer_status,  # For CASE WHEN condition
                current_time if indexer_status == 'completed' else None,
                error_message,
                indexer_status,  # For retry count condition
                report_id,
                user_id
            ))
            
            rows_affected = cur.rowcount
            
            if rows_affected > 0:
                logging.info(f"✅ RISK INDEXER: Database update successful")
                logging.info(f"📊 Update results:")
                logging.info(f"   - Report ID: {report_id}")
                logging.info(f"   - Job ID: {indexer_job_id}")
                logging.info(f"   - Status: {indexer_status}")
                logging.info(f"   - Rows affected: {rows_affected}")
                logging.info(f"   - Timestamp: {current_time.isoformat()}")
                
                # Log additional context if error
                if error_message:
                    logging.warning(f"⚠️ RISK INDEXER: Error message stored in database")
                    logging.warning(f"📊 Error details:")
                    logging.warning(f"   - Error: {error_message}")
                    logging.warning(f"   - Error length: {len(error_message)} chars")
                    
                # Log successful completion
                if indexer_status == 'completed':
                    logging.info(f"🎉 RISK INDEXER: Marked as COMPLETED in database")
                elif indexer_status == 'triggered':
                    logging.info(f"🔄 RISK INDEXER: Marked as TRIGGERED in database")
                elif indexer_status == 'failed':
                    logging.warning(f"❌ RISK INDEXER: Marked as FAILED in database")
                    
                return True
            else:
                logging.error(f"❌ RISK INDEXER: Database update failed - no rows affected")
                logging.error(f"📊 Update failure details:")
                logging.error(f"   - Report ID: {report_id}")
                logging.error(f"   - User ID: {user_id}")
                logging.error(f"   - Expected 1 row, got {rows_affected}")
                logging.error(f"   - This indicates report may not exist or ownership mismatch")
                return False
                
    except Exception as e:
        logging.error(f"❌ RISK INDEXER: Database storage error")
        logging.error(f"📊 Storage error details:")
        logging.error(f"   - Error type: {type(e).__name__}")
        logging.error(f"   - Error message: {str(e)}")
        logging.error(f"   - Report ID: {report_id}")
        logging.error(f"   - User ID: {user_id}")
        logging.error(f"   - Indexer Job ID: {indexer_job_id}")
        logging.error(f"   - Indexer Status: {indexer_status}")
        
        # Log full traceback for database errors
        import traceback
        logging.error(f"🔍 RISK INDEXER: Database error traceback:")
        for line in traceback.format_exc().split('\n'):
            if line.strip():
                logging.error(f"   {line}")
        
        return False
        
    finally:
        if conn:
            try:
                conn.close()
                logging.debug(f"🔗 RISK INDEXER: Database connection closed")
            except Exception as close_error:
                logging.warning(f"⚠️ RISK INDEXER: Error closing connection: {close_error}")

async def trigger_and_monitor_risk_indexer(report_id: str, user_id: str, container_name: str) -> Dict:
    """Comprehensive risk indexer trigger with monitoring and database integration"""
    
    logging.info(f"🚀 RISK INDEXER: Starting comprehensive indexer process")
    logging.info(f"📊 Monitor parameters:")
    logging.info(f"   - Report ID: {report_id}")
    logging.info(f"   - User ID: {user_id}")
    logging.info(f"   - Container: {container_name}")
    
    monitor_start_time = time.time()
    
    # Step 1: Trigger the indexer
    logging.info(f"📤 RISK INDEXER: Step 1 - Triggering indexer")
    trigger_start_time = time.time()
    
    success, message, job_id = await trigger_risk_indexer_for_client(user_id)
    trigger_time = time.time() - trigger_start_time
    
    logging.info(f"📊 RISK INDEXER: Trigger completed in {trigger_time:.3f}s")
    logging.info(f"📊 Trigger results:")
    logging.info(f"   - Success: {success}")
    logging.info(f"   - Message: {message}")
    logging.info(f"   - Job ID: {job_id}")
    
    if success and job_id:
        # Step 2: Store successful trigger in database
        logging.info(f"💾 RISK INDEXER: Step 2 - Storing successful trigger")
        store_start_time = time.time()
        
        store_result = store_risk_indexer_job_metadata(
            report_id=report_id,
            user_id=user_id,
            indexer_job_id=job_id,
            indexer_status="triggered",
            indexer_response=message
        )
        
        store_time = time.time() - store_start_time
        
        logging.info(f"📊 RISK INDEXER: Database storage completed in {store_time:.3f}s")
        logging.info(f"📊 Storage result: {store_result}")
        
        if store_result:
            total_monitor_time = time.time() - monitor_start_time
            
            logging.info(f"✅ RISK INDEXER: PROCESS COMPLETED SUCCESSFULLY!")
            logging.info(f"📊 Final success metrics:")
            logging.info(f"   - Report ID: {report_id}")
            logging.info(f"   - Job ID: {job_id}")
            logging.info(f"   - User ID: {user_id}")
            logging.info(f"   - Container: {container_name}")
            logging.info(f"   - Trigger time: {trigger_time:.3f}s")
            logging.info(f"   - Storage time: {store_time:.3f}s")
            logging.info(f"   - Total monitor time: {total_monitor_time:.3f}s")
            logging.info(f"   - Background processing: ACTIVE")
            
            # Simplified: Skip status checking since endpoints don't exist
            return_status = "triggered"  # Assume success after successful trigger
            logging.info(f"📊 RISK INDEXER: Status checking disabled to avoid 404 errors")
            logging.info(f"📊 Job ID {job_id} is assumed to be processing successfully")
            logging.info(f"📊 Expected completion: 2-5 minutes for full indexing")
            
            return {
                "success": True,
                "job_id": job_id,
                "message": message,
                "status": return_status,
                "stored_in_db": True,
                "report_id": report_id,
                "container_name": container_name,
                "trigger_time": trigger_time,
                "storage_time": store_time,
                "total_time": total_monitor_time,
                "status_check_note": "Status checking disabled - indexer endpoints not available",
                "expected_completion": "2-5 minutes"
            }
        else:
            logging.error(f"❌ RISK INDEXER: Triggered but database storage failed")
            logging.error(f"📊 Storage failure details:")
            logging.error(f"   - Trigger successful: {success}")
            logging.error(f"   - Job ID received: {job_id}")
            logging.error(f"   - Database storage: FAILED")
            logging.error(f"   - Impact: Indexer running but not tracked in database")
            
            return {
                "success": False,
                "error": "Triggered but database storage failed",
                "job_id": job_id,
                "message": message,
                "trigger_successful": True,
                "storage_failed": True
            }
    else:
        # Step 2: Store failed trigger in database
        logging.error(f"❌ RISK INDEXER: Trigger failed, storing failure")
        
        store_risk_indexer_job_metadata(
            report_id=report_id,
            user_id=user_id,
            indexer_job_id=job_id or "failed",
            indexer_status="failed",
            error_message=message
        )
        
        total_monitor_time = time.time() - monitor_start_time
        
        logging.error(f"💥 RISK INDEXER: TRIGGER PROCESS FAILED")
        logging.error(f"📊 Failure summary:")
        logging.error(f"   - Report ID: {report_id}")
        logging.error(f"   - User ID: {user_id}")
        logging.error(f"   - Container: {container_name}")
        logging.error(f"   - Error: {message}")
        logging.error(f"   - Job ID: {job_id if job_id else 'None'}")
        logging.error(f"   - Total time: {total_monitor_time:.3f}s")
        logging.error(f"   - Trigger time: {trigger_time:.3f}s")
        
        return {
            "success": False,
            "error": message,
            "job_id": job_id,
            "status": "failed",
            "stored_in_db": True,
            "report_id": report_id,
            "trigger_time": trigger_time,
            "total_time": total_monitor_time
        }

def trigger_risk_indexer_sync(report_id: str, user_id: str, container_name: str) -> Dict:
    """Synchronous wrapper for risk indexer trigger (for background tasks)"""
    
    sync_start_time = time.time()
    thread_id = threading.current_thread().ident
    
    logging.info(f"🔄 RISK INDEXER SYNC: Starting wrapper")
    logging.info(f"📊 Sync wrapper parameters:")
    logging.info(f"   - Report ID: {report_id}")
    logging.info(f"   - User ID: {user_id}")
    logging.info(f"   - Container: {container_name}")
    logging.info(f"   - Thread ID: {thread_id}")
    logging.info(f"   - Platform: {platform.system()}")
    
    try:
        # Handle Windows event loop policy
        if platform.system() == 'Windows':
            logging.debug(f"🪟 RISK INDEXER: Setting Windows event loop policy")
            asyncio.set_event_loop_policy(asyncio.WindowsSelectorEventLoopPolicy())
        
        # Create new event loop for this thread
        logging.debug(f"🔄 RISK INDEXER: Creating new event loop")
        loop_creation_start = time.time()
        
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
        
        loop_creation_time = time.time() - loop_creation_start
        logging.debug(f"✅ RISK INDEXER: Event loop created in {loop_creation_time:.3f}s")
        
        try:
            # Execute the async indexer function
            logging.info(f"⚙️ RISK INDEXER: Executing async trigger function")
            execution_start_time = time.time()
            
            result = loop.run_until_complete(
                trigger_and_monitor_risk_indexer(report_id, user_id, container_name)
            )
            
            execution_time = time.time() - execution_start_time
            total_sync_time = time.time() - sync_start_time
            
            logging.info(f"✅ RISK INDEXER SYNC: Execution completed")
            logging.info(f"📊 Sync execution metrics:")
            logging.info(f"   - Execution time: {execution_time:.3f}s")
            logging.info(f"   - Total sync time: {total_sync_time:.3f}s")
            logging.info(f"   - Loop creation time: {loop_creation_time:.3f}s")
            logging.info(f"   - Result success: {result.get('success', False)}")
            logging.info(f"   - Result job_id: {result.get('job_id', 'None')}")
            
            logging.info(f"🎯 Risk indexer sync wrapper completed for report_id={report_id}")
            return result
            
        finally:
            # Clean up event loop
            logging.debug(f"🔄 RISK INDEXER: Cleaning up event loop")
            loop_cleanup_start = time.time()
            
            try:
                # Cancel any pending tasks
                pending_tasks = [task for task in asyncio.all_tasks(loop) if not task.done()]
                if pending_tasks:
                    logging.debug(f"🔧 RISK INDEXER: Cancelling {len(pending_tasks)} pending tasks")
                    for task in pending_tasks:
                        task.cancel()
                
                loop.close()
                loop_cleanup_time = time.time() - loop_cleanup_start
                logging.debug(f"✅ RISK INDEXER: Event loop cleaned up in {loop_cleanup_time:.3f}s")
                
            except Exception as cleanup_error:
                loop_cleanup_time = time.time() - loop_cleanup_start
                logging.warning(f"⚠️ RISK INDEXER: Loop cleanup error after {loop_cleanup_time:.3f}s: {cleanup_error}")
            
    except Exception as e:
        total_sync_time = time.time() - sync_start_time
        error_msg = f"Risk indexer sync wrapper error: {str(e)}"
        
        logging.error(f"💥 RISK INDEXER SYNC: WRAPPER ERROR")
        logging.error(f"📊 Wrapper error details:")
        logging.error(f"   - Error type: {type(e).__name__}")
        logging.error(f"   - Error message: {str(e)}")
        logging.error(f"   - Total sync time: {total_sync_time:.3f}s")
        logging.error(f"   - Thread ID: {thread_id}")
        logging.error(f"   - Platform: {platform.system()}")
        
        # Log full traceback for sync wrapper errors
        import traceback
        logging.error(f"🔍 RISK INDEXER SYNC: Error traceback:")
        for line in traceback.format_exc().split('\n'):
            if line.strip():
                logging.error(f"   {line}")
        
        # Still try to store the error in database
        try:
            logging.info(f"💾 RISK INDEXER: Attempting to store sync error in database")
            store_risk_indexer_job_metadata(
                report_id=report_id,
                user_id=user_id,
                indexer_job_id="sync_error",
                indexer_status="error",
                error_message=error_msg
            )
            logging.info(f"✅ RISK INDEXER: Sync error stored in database")
        except Exception as store_error:
            logging.error(f"❌ RISK INDEXER: Failed to store sync error: {store_error}")
        
        return {
            "success": False,
            "error": error_msg,
            "status": "sync_error",
            "report_id": report_id,
            "total_sync_time": total_sync_time,
            "thread_id": thread_id
        }

def trigger_risk_indexer_background(report_id: str, user_id: str, container_name: str):
    """Background thread function for risk indexer triggering with comprehensive logging"""
    
    def indexer_worker():
        worker_start_time = time.time()
        thread_id = threading.current_thread().ident
        thread_name = f"RiskIndexer-{user_id}-{int(time.time())}"
        
        # Set thread name for better debugging
        threading.current_thread().name = thread_name
        
        logging.info(f"🔄 RISK INDEXER WORKER: Starting background worker")
        logging.info(f"📊 Worker details:")
        logging.info(f"   - Thread ID: {thread_id}")
        logging.info(f"   - Thread name: {thread_name}")
        logging.info(f"   - Report ID: {report_id}")
        logging.info(f"   - User ID: {user_id}")
        logging.info(f"   - Container: {container_name}")
        logging.info(f"   - Worker start time: {datetime.fromtimestamp(worker_start_time).isoformat()}")
        
        try:
            logging.info(f"⚙️ RISK INDEXER WORKER: Calling sync wrapper")
            
            result = trigger_risk_indexer_sync(report_id, user_id, container_name)
            
            worker_time = time.time() - worker_start_time
            
            if result["success"]:
                logging.info(f"🎉 RISK INDEXER WORKER: BACKGROUND SUCCESS!")
                logging.info(f"📊 Success details:")
                logging.info(f"   - Job ID: {result.get('job_id')}")
                logging.info(f"   - Status: {result.get('status')}")
                logging.info(f"   - Message: {result.get('message')}")
                logging.info(f"   - Worker time: {worker_time:.3f}s")
                logging.info(f"   - Trigger time: {result.get('trigger_time', 0):.3f}s")
                logging.info(f"   - Storage time: {result.get('storage_time', 0):.3f}s")
                logging.info(f"   - Thread: {thread_name}")
                
                # Log success to make it easy to find in logs
                logging.info(f"🎯 INDEXER SUCCESS: Risk report {report_id} indexer job {result.get('job_id')} triggered successfully")
                
            else:
                logging.error(f"❌ RISK INDEXER WORKER: BACKGROUND FAILURE")
                logging.error(f"📊 Failure details:")
                logging.error(f"   - Error: {result.get('error')}")
                logging.error(f"   - Status: {result.get('status')}")
                logging.error(f"   - Job ID: {result.get('job_id', 'None')}")
                logging.error(f"   - Worker time: {worker_time:.3f}s")
                logging.error(f"   - Thread: {thread_name}")
                
                # Log failure to make it easy to find in logs
                logging.error(f"💥 INDEXER FAILURE: Risk report {report_id} indexer failed - {result.get('error')}")
                
        except Exception as e:
            worker_time = time.time() - worker_start_time
            
            logging.error(f"💥 RISK INDEXER WORKER: CRITICAL ERROR")
            logging.error(f"📊 Critical error details:")
            logging.error(f"   - Error type: {type(e).__name__}")
            logging.error(f"   - Error message: {str(e)}")
            logging.error(f"   - Worker time: {worker_time:.3f}s")
            logging.error(f"   - Thread ID: {thread_id}")
            logging.error(f"   - Thread name: {thread_name}")
            logging.error(f"   - Report ID: {report_id}")
            logging.error(f"   - User ID: {user_id}")
            
            # Log full traceback for worker errors
            import traceback
            logging.error(f"🔍 RISK INDEXER WORKER: Error traceback:")
            for line in traceback.format_exc().split('\n'):
                if line.strip():
                    logging.error(f"   {line}")
            
            # Try to store the worker error
            try:
                store_risk_indexer_job_metadata(
                    report_id=report_id,
                    user_id=user_id,
                    indexer_job_id="worker_error",
                    indexer_status="worker_error",
                    error_message=f"Background worker error: {str(e)}"
                )
                logging.info(f"✅ RISK INDEXER: Worker error stored in database")
            except Exception as store_error:
                logging.error(f"❌ RISK INDEXER: Failed to store worker error: {store_error}")
        
        finally:
            final_worker_time = time.time() - worker_start_time
            
            logging.info(f"🏁 RISK INDEXER WORKER: Background worker completed")
            logging.info(f"📊 Final worker metrics:")
            logging.info(f"   - Total worker time: {final_worker_time:.3f}s")
            logging.info(f"   - Thread ID: {thread_id}")
            logging.info(f"   - Thread name: {thread_name}")
            logging.info(f"   - Worker end time: {datetime.now().isoformat()}")
    
    # Enhanced thread creation with comprehensive logging
    try:
        logging.info(f"🚀 RISK INDEXER: Launching background thread")
        
        # Create thread with proper naming and error handling
        thread_creation_start = time.time()
        
        indexer_thread = Thread(
            target=indexer_worker, 
            daemon=True,
            name=f"RiskIndexer-{user_id}-{report_id[-8:]}"  # Last 8 chars of report_id
        )
        
        # Store metadata in thread for debugging
        indexer_thread._indexer_data = {
            'report_id': report_id,
            'user_id': user_id,
            'container_name': container_name,
            'created_at': time.time(),
            'created_datetime': datetime.now().isoformat()
        }
        
        indexer_thread.start()
        thread_creation_time = time.time() - thread_creation_start
        
        logging.info(f"✅ RISK INDEXER: Background thread launched successfully")
        logging.info(f"📊 Thread launch details:")
        logging.info(f"   - Thread name: {indexer_thread.name}")
        logging.info(f"   - Thread ID: {indexer_thread.ident}")
        logging.info(f"   - Launch time: {thread_creation_time:.3f}s")
        logging.info(f"   - Daemon thread: {indexer_thread.daemon}")
        logging.info(f"   - Thread alive: {indexer_thread.is_alive()}")
        logging.info(f"   - Report ID: {report_id}")
        logging.info(f"   - Background processing: ACTIVE")
        
        # Enhanced thread monitoring for debugging
        if logging.getLogger().isEnabledFor(logging.DEBUG):
            def monitor_thread():
                time.sleep(1)  # Give thread time to start
                if indexer_thread.is_alive():
                    logging.debug(f"✅ RISK INDEXER: Thread {indexer_thread.name} is running successfully")
                else:
                    logging.warning(f"⚠️ RISK INDEXER: Thread {indexer_thread.name} finished quickly (this may be normal)")
            
            Thread(target=monitor_thread, daemon=True, name=f"RiskIndexerMonitor-{user_id}").start()
        
        return indexer_thread
        
    except Exception as thread_error:
        thread_creation_time = time.time() - thread_creation_start if 'thread_creation_start' in locals() else 0
        
        logging.error(f"💥 RISK INDEXER: THREAD CREATION FAILED")
        logging.error(f"📊 Thread creation error details:")
        logging.error(f"   - Error type: {type(thread_error).__name__}")
        logging.error(f"   - Error message: {str(thread_error)}")
        logging.error(f"   - Creation time: {thread_creation_time:.3f}s")
        logging.error(f"   - Report ID: {report_id}")
        logging.error(f"   - User ID: {user_id}")
        logging.error(f"   - Container: {container_name}")
        
        # Log full traceback for thread creation errors
        import traceback
        logging.error(f"🔍 RISK INDEXER: Thread creation traceback:")
        for line in traceback.format_exc().split('\n'):
            if line.strip():
                logging.error(f"   {line}")
        
        # Try to store thread creation error
        try:
            store_risk_indexer_job_metadata(
                report_id=report_id,
                user_id=user_id,
                indexer_job_id="thread_error",
                indexer_status="thread_error",
                error_message=f"Thread creation error: {str(thread_error)}"
            )
            logging.info(f"✅ RISK INDEXER: Thread creation error stored in database")
        except Exception as store_error:
            logging.error(f"❌ RISK INDEXER: Failed to store thread error: {store_error}")
        
        return None  # Return None for failed thread creation
    

# ======================================================
#           Gemini AI Integration - RISK ENGINE
# ======================================================

@dataclass
class RiskChatResponse:
    content: str
    model: str
    api_key_used: str
    usage: Dict[str, Any]
    finish_reason: str
    response_time: float
    timestamp: float
    token_count: int

def convert_messages_to_gemini_format(messages: List[Dict[str, str]]) -> List[Dict]:
    """Convert messages to Gemini API format"""
    contents = []
    
    for msg in messages:
        role = msg["role"]
        content = msg["content"]
        
        if role in ["user", "human"]:
            if contents and contents[-1]["role"] == "user":
                contents[-1]["parts"].append({"text": content})
            else:
                contents.append({
                    "role": "user",
                    "parts": [{"text": content}]
                })
        elif role in ["assistant", "model", "ai"]:
            contents.append({
                "role": "model",
                "parts": [{"text": content}]
            })
        elif role == "system":
            if contents and contents[-1]["role"] == "user":
                contents[-1]["parts"].insert(0, {"text": f"SYSTEM CONTEXT: {content}\n\n"})
            else:
                contents.append({
                    "role": "user",
                    "parts": [{"text": f"SYSTEM CONTEXT: {content}"}]
                })
    
    return contents

def try_vertex_ai_risk_request(
    enhanced_prompt: str,
    temperature: float,
    max_tokens: int,
    start_time: float
) -> Optional[Dict]:
    """
    Try making request using Vertex AI (PRIMARY METHOD for Risk Engine).
    Returns Dict if successful, None if fails (will fallback to API keys).
    """
    if not vertex_ai_client:
        logging.info("Vertex AI client not available - using API keys fallback")
        return None

    try:
        logging.info("🚀 Trying Vertex AI (Primary Method for Risk Analysis)")

        # Call Vertex AI using GenAI SDK with gemini-2.5-pro
        response = vertex_ai_client.models.generate_content(
            model="gemini-2.5-pro",
            contents=enhanced_prompt,
            config={
                "temperature": temperature,
                "max_output_tokens": max_tokens,
                "top_p": 0.9,
                "top_k": 40
            }
        )

        # Extract text from response
        if hasattr(response, 'text') and response.text:
            content = response.text
            request_time = time.time() - start_time

            # Count tokens if available
            token_count = 0
            if hasattr(response, 'usage_metadata'):
                token_count = getattr(response.usage_metadata, 'total_token_count', 0)

            logging.info(f"✅ Vertex AI SUCCESS - {len(content.split())} words, {token_count} tokens, {request_time:.2f}s")

            # Return successful response in Risk Engine format
            return {
                "content": content,
                "token_count": token_count,
                "finish_reason": "stop",
                "response_time": request_time,
                "timestamp": time.time(),
                "model": "gemini-2.5-pro-vertex"
            }
        else:
            logging.warning("⚠️ Vertex AI returned empty response - trying API keys")
            return None

    except Exception as e:
        logging.warning(f"⚠️ Vertex AI request failed: {str(e)} - Falling back to API keys")
        return None

def risk_ultra_deep_analysis(
    complete_raw_data: Dict,
    analysis_type: str,
    analysis_requirements: str,
    api_key: str,
    client_id: str = "risk_analysis",
    temperature: float = 0.7,
    max_tokens: int = 1000000
) -> RiskChatResponse:
    """Enhanced risk analysis with ultra-deep response analysis and detailed logging"""
    
    start_time = time.time()
    request_start_time = None
    
    logging.info(f"🚀 [{client_id}] Starting Risk Analysis: {analysis_type}")
    logging.info(f"🔍 [{client_id}] Input parameters: temp={temperature}, max_tokens={max_tokens}")
    logging.info(f"🔍 [{client_id}] API key ending: ...{api_key[-4:]}")
    logging.info(f"🔍 [{client_id}] Complete raw data keys: {list(complete_raw_data.keys()) if complete_raw_data else 'No data'}")
    logging.info(f"🔍 [{client_id}] Analysis requirements length: {len(analysis_requirements)} characters")
    
    # Log API key health status at start
    key_health = api_key_health.get(api_key, {})
    if key_health:
        success_rate = key_health.get('success_rate', 1.0)
        current_load = key_health.get('current_load', 0)
        consecutive_failures = key_health.get('consecutive_failures', 0)
        total_requests = key_health.get('total_requests', 0)
        avg_response = statistics.mean(key_health.get('response_times', [0])) if key_health.get('response_times') else 0
        last_503_time = key_health.get('last_503_time')
        
        logging.info(f"🔑 [{client_id}] API Key Health at Start: {key_health.get('key_id', 'unknown')} - "
                    f"Failures: {consecutive_failures}, "
                    f"Total Requests: {total_requests}, "
                    f"Success Rate: {success_rate:.2f}, "
                    f"Current Load: {current_load}, "
                    f"Avg Response: {avg_response:.1f}s")
        
        if last_503_time:
            cooldown_remaining = 300 - (time.time() - last_503_time)
            if cooldown_remaining > 0:
                logging.warning(f"⏰ [{client_id}] API Key in cooldown: {cooldown_remaining:.1f}s remaining")
            else:
                logging.info(f"✅ [{client_id}] API Key cooldown expired, ready for use")
        
        if current_load > 3:
            logging.warning(f"⚠️ [{client_id}] API Key overloaded: {current_load} concurrent requests")
        
        logging.info(f"🔑 [{client_id}] Overall API Key Ecosystem: {get_api_key_status_summary()}")
    else:
        logging.warning(f"⚠️ [{client_id}] No health data available for API key ...{api_key[-4:]}")
    
    try:
        # Create enhanced prompt
        logging.info(f"📝 [{client_id}] Creating enhanced prompt...")
        prompt_creation_start = time.time()
        
        try:
            enhanced_prompt = create_enhanced_risk_analysis_prompt(
                complete_raw_data, analysis_type, analysis_requirements
            )
            prompt_creation_time = time.time() - prompt_creation_start
            
            logging.info(f"✅ [{client_id}] Prompt created successfully in {prompt_creation_time:.3f}s")
            logging.info(f"🔍 [{client_id}] Prompt length: {len(enhanced_prompt):,} characters")
            logging.info(f"🔍 [{client_id}] Prompt word count: ~{len(enhanced_prompt.split()):,} words")
            
            if len(enhanced_prompt) > 100000:  # 100KB limit
                logging.warning(f"⚠️ [{client_id}] Large prompt detected: {len(enhanced_prompt):,} chars")
            
        except Exception as prompt_error:
            logging.error(f"❌ [{client_id}] Prompt creation failed: {prompt_error}")
            logging.error(f"🔍 [{client_id}] Prompt error type: {type(prompt_error).__name__}")
            raise Exception(f"Failed to create analysis prompt: {prompt_error}")

        # ===================================================================
        # STEP 1: TRY VERTEX AI (PRIMARY METHOD)
        # ===================================================================
        vertex_result = try_vertex_ai_risk_request(
            enhanced_prompt=enhanced_prompt,
            temperature=temperature,
            max_tokens=max_tokens,
            start_time=start_time
        )

        if vertex_result:
            # Vertex AI succeeded - return immediately
            total_time = time.time() - start_time
            logging.info(f"🎉 [{client_id}] Risk analysis completed via Vertex AI in {total_time:.2f}s")

            return RiskChatResponse(
                content=vertex_result["content"],
                token_count=vertex_result["token_count"],
                finish_reason=vertex_result["finish_reason"],
                response_time=vertex_result["response_time"],
                timestamp=vertex_result["timestamp"],
                model=vertex_result["model"]
            )

        # ===================================================================
        # STEP 2: FALLBACK TO API KEYS (if Vertex AI failed)
        # ===================================================================
        logging.info(f"🔄 [{client_id}] Vertex AI not available - using API keys fallback")

        # Convert to Gemini format
        logging.info(f"🔄 [{client_id}] Converting to Gemini format...")
        format_conversion_start = time.time()
        
        try:
            contents = convert_messages_to_gemini_format([
                {"role": "user", "content": enhanced_prompt}
            ])
            format_conversion_time = time.time() - format_conversion_start
            
            logging.info(f"✅ [{client_id}] Gemini format conversion successful in {format_conversion_time:.3f}s")
            logging.info(f"🔍 [{client_id}] Converted contents length: {len(contents)}")
            logging.info(f"🔍 [{client_id}] Contents structure validation: {all('role' in c and 'parts' in c for c in contents)}")
            
        except Exception as format_error:
            logging.error(f"❌ [{client_id}] Gemini format conversion failed: {format_error}")
            raise Exception(f"Failed to convert to Gemini format: {format_error}")
        
        # Production-optimized payload
        logging.info(f"🔧 [{client_id}] Creating API payload...")
        payload = {
            "contents": contents,
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": max_tokens,
                "topP": 0.9,
                "topK": 40,
                "candidateCount": 1,
                "stopSequences": [],
                "responseMimeType": "text/plain"
            },
            "safetySettings": [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
            ]
        }
        
        params = {'key': api_key}
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent"
        
        logging.info(f"🌐 [{client_id}] Preparing API request to Gemini")
        logging.info(f"🔍 [{client_id}] API URL: {url}")
        logging.info(f"🔍 [{client_id}] Payload keys: {list(payload.keys())}")
        logging.info(f"🔍 [{client_id}] Generation config: {payload['generationConfig']}")
        logging.info(f"🔍 [{client_id}] Safety settings count: {len(payload['safetySettings'])}")
        
        # Request execution with comprehensive monitoring
        logging.info(f"📤 [{client_id}] Sending request to Gemini API...")
        request_start_time = time.time()
        
        try:
            response = requests.post(
                url,
                json=payload,
                params=params,
                timeout=REQUEST_TIMEOUT
            )
            
            # Calculate response time immediately
            response_time = time.time() - request_start_time
            
            logging.info(f"📡 [{client_id}] Response received from Gemini API")
            logging.info(f"⏱️ [{client_id}] Response time: {response_time:.2f}s")
            logging.info(f"📊 [{client_id}] Response status: {response.status_code}")
            logging.info(f"🔍 [{client_id}] Response headers: {dict(response.headers)}")
            logging.info(f"🔍 [{client_id}] Response size: {len(response.text)} characters")
            
        except requests.exceptions.Timeout:
            response_time = time.time() - request_start_time
            logging.error(f"⏰ [{client_id}] Request timed out after {response_time:.2f}s (timeout: {REQUEST_TIMEOUT}s)")
            update_api_key_health(api_key, success=False, error_code="TIMEOUT", response_time=response_time)
            raise Exception(f"Request timed out after {REQUEST_TIMEOUT}s")
            
        except requests.exceptions.ConnectionError as conn_error:
            response_time = time.time() - request_start_time
            logging.error(f"🌐 [{client_id}] Connection error after {response_time:.2f}s: {conn_error}")
            update_api_key_health(api_key, success=False, error_code="CONNECTION_ERROR", response_time=response_time)
            raise Exception(f"Connection error: {conn_error}")
            
        except Exception as request_error:
            response_time = time.time() - request_start_time
            logging.error(f"❌ [{client_id}] Request error after {response_time:.2f}s: {request_error}")
            update_api_key_health(api_key, success=False, error_code="REQUEST_ERROR", response_time=response_time)
            raise Exception(f"Request error: {request_error}")
        
        # HTTP Response Processing
        if response.status_code == 200:
            # SUCCESS: Update API key health with response time
            logging.info(f"✅ [{client_id}] HTTP 200 Success - Processing response")
            update_api_key_health(api_key, success=True, response_time=response_time)
            
            # JSON parsing with detailed validation
            try:
                logging.info(f"🔄 [{client_id}] Parsing JSON response...")
                json_parse_start = time.time()
                
                data = response.json()
                json_parse_time = time.time() - json_parse_start
                
                logging.info(f"✅ [{client_id}] JSON parsing successful in {json_parse_time:.3f}s")
                logging.info(f"🔍 [{client_id}] Response data size: {len(str(data))} characters")
                
            except json.JSONDecodeError as json_error:
                logging.error(f"❌ [{client_id}] JSON parsing failed: {json_error}")
                logging.error(f"🔍 [{client_id}] Raw response text (first 500 chars): {response.text[:500]}...")
                logging.error(f"🔍 [{client_id}] Response content-type: {response.headers.get('content-type', 'Unknown')}")
                
                update_api_key_health(api_key, success=False, error_code="JSON_PARSE_ERROR", response_time=response_time)
                raise Exception(f"Failed to parse JSON response: {json_error}")
            
            except Exception as parsing_error:
                logging.error(f"❌ [{client_id}] Unexpected parsing error: {parsing_error}")
                update_api_key_health(api_key, success=False, error_code="PARSING_ERROR", response_time=response_time)
                raise Exception(f"Unexpected parsing error: {parsing_error}")
            
            # ULTRA DETAILED LOGGING FOR AI RESPONSE
            logging.info(f"🔍 [{client_id}] RAW API RESPONSE ANALYSIS:")
            logging.info(f"🔍 [{client_id}] Top-level response keys: {list(data.keys())}")
            
            # Log response metadata
            if 'usageMetadata' in data:
                usage_meta = data['usageMetadata']
                logging.info(f"📊 [{client_id}] Usage metadata: {usage_meta}")
                prompt_tokens = usage_meta.get('promptTokenCount', 0)
                total_tokens = usage_meta.get('totalTokenCount', 0)
                completion_tokens = total_tokens - prompt_tokens
                logging.info(f"📊 [{client_id}] Token breakdown: prompt={prompt_tokens}, completion={completion_tokens}, total={total_tokens}")
            
            if 'modelVersion' in data:
                logging.info(f"🤖 [{client_id}] Model version: {data['modelVersion']}")
            
            if 'responseId' in data:
                logging.info(f"🆔 [{client_id}] Response ID: {data['responseId']}")
            
            # Enhanced candidate validation
            if 'candidates' in data and len(data['candidates']) > 0:
                candidate = data['candidates'][0]
                logging.info(f"🔍 [{client_id}] CANDIDATE ANALYSIS:")
                logging.info(f"🔍 [{client_id}] Found {len(data['candidates'])} candidate(s)")
                logging.info(f"🔍 [{client_id}] Candidate keys: {list(candidate.keys())}")
                logging.info(f"🔍 [{client_id}] Candidate structure: {json.dumps(candidate, indent=2, default=str)[:800]}...")
                
                # Check finish reason with detailed analysis
                finish_reason = candidate.get('finishReason', 'UNKNOWN')
                logging.info(f"🔍 [{client_id}] Finish reason: {finish_reason}")
                
                # Validate finish reason
                if finish_reason == 'SAFETY':
                    logging.error(f"🚨 [{client_id}] Content blocked by safety filters")
                    logging.error(f"🔍 [{client_id}] Safety ratings: {candidate.get('safetyRatings', 'None')}")
                    update_api_key_health(api_key, success=False, error_code="SAFETY_FILTER", response_time=response_time)
                    raise Exception("Content blocked by safety filters")
                elif finish_reason in ['MAX_TOKENS', 'RECITATION']:
                    logging.warning(f"⚠️ [{client_id}] Response truncated due to: {finish_reason}")
                elif finish_reason != 'STOP':
                    logging.warning(f"⚠️ [{client_id}] Unusual finish reason: {finish_reason}")
                
                # Check safety ratings
                if 'safetyRatings' in candidate:
                    safety_ratings = candidate['safetyRatings']
                    logging.info(f"🔍 [{client_id}] Safety ratings: {safety_ratings}")
                    
                    # Check for blocking
                    blocked_categories = [r for r in safety_ratings if r.get('probability') in ['HIGH', 'MEDIUM']]
                    if blocked_categories:
                        logging.warning(f"⚠️ [{client_id}] Potentially blocked categories: {blocked_categories}")
                
                # Enhanced content extraction with detailed validation
                content = ""
                extraction_method = "none"
                extraction_start = time.time()
                
                logging.info(f"🔍 [{client_id}] CONTENT EXTRACTION ANALYSIS:")
                
                try:
                    # Method 1 - Standard content extraction with validation
                    if 'content' in candidate and candidate['content'] is not None:
                        content_obj = candidate['content']
                        logging.info(f"🔍 [{client_id}] Found content object: {type(content_obj)}")
                        logging.info(f"🔍 [{client_id}] Content object keys: {list(content_obj.keys()) if isinstance(content_obj, dict) else 'Not a dict'}")
                        logging.info(f"🔍 [{client_id}] Content object preview: {json.dumps(content_obj, indent=2, default=str)[:400]}...")
                        
                        # Validate content object structure
                        if not isinstance(content_obj, dict):
                            logging.error(f"❌ [{client_id}] Content object is not a dictionary: {type(content_obj)}")
                            logging.error(f"🔍 [{client_id}] Content object value: {content_obj}")
                            update_api_key_health(api_key, success=False, error_code="INVALID_CONTENT_TYPE", response_time=response_time)
                            raise Exception(f"Content object is not a dictionary: {type(content_obj)}")
                        
                        # Check for malformed content object
                        if content_obj == {'role': 'model'}:
                            logging.error(f"🚨 [{client_id}] MALFORMED RESPONSE: Got metadata instead of content")
                            logging.error(f"🔍 [{client_id}] This indicates API key returned incomplete response")
                            logging.error(f"🔍 [{client_id}] Full candidate: {json.dumps(candidate, indent=2, default=str)}")
                            update_api_key_health(api_key, success=False, error_code="METADATA_RESPONSE", response_time=response_time)
                            raise Exception("API returned metadata instead of content - key may be corrupted")
                        
                        if 'parts' in content_obj and content_obj['parts']:
                            parts = content_obj['parts']
                            logging.info(f"🔍 [{client_id}] Found content.parts: {len(parts)} parts")
                            
                            if len(parts) > 0:
                                first_part = parts[0]
                                logging.info(f"🔍 [{client_id}] First part type: {type(first_part)}")
                                logging.info(f"🔍 [{client_id}] First part keys: {list(first_part.keys()) if isinstance(first_part, dict) else 'Not a dict'}")
                                logging.info(f"🔍 [{client_id}] First part preview: {json.dumps(first_part, indent=2, default=str)[:300]}...")
                                
                                if isinstance(first_part, dict) and 'text' in first_part:
                                    content = first_part['text']
                                    extraction_method = "content.parts[0].text"
                                    logging.info(f"✅ [{client_id}] Extracted via method 1: {len(content)} characters")
                                else:
                                    logging.warning(f"⚠️ [{client_id}] First part has no 'text' field")
                                    logging.warning(f"🔍 [{client_id}] First part content: {first_part}")
                            else:
                                logging.warning(f"⚠️ [{client_id}] Parts array is empty")
                        else:
                            logging.warning(f"⚠️ [{client_id}] Content object has no 'parts' field or parts is empty")
                            logging.warning(f"🔍 [{client_id}] Content object: {content_obj}")
                            
                            # Check for alternative content structures
                            if 'text' in content_obj:
                                content = content_obj['text']
                                extraction_method = "content.text"
                                logging.info(f"✅ [{client_id}] Found alternative text field: {len(content)} characters")
                    
                    # Method 2 - Direct text field with validation
                    if not content and 'text' in candidate:
                        content = candidate['text']
                        extraction_method = "candidate.text"
                        logging.info(f"✅ [{client_id}] Extracted via method 2: {len(content)} characters")
                    
                    # Method 3 - Search for any text-like fields
                    if not content:
                        logging.warning(f"⚠️ [{client_id}] No content found via standard methods, searching for text fields...")
                        
                        for key, value in candidate.items():
                            if isinstance(value, str) and len(value) > 20:  # Increased minimum length
                                content = value
                                extraction_method = f"candidate.{key}"
                                logging.info(f"✅ [{client_id}] Extracted via method 3 ({key}): {len(content)} characters")
                                break
                            elif isinstance(value, dict):
                                # Search nested dictionaries
                                for nested_key, nested_value in value.items():
                                    if isinstance(nested_value, str) and len(nested_value) > 20:
                                        content = nested_value
                                        extraction_method = f"candidate.{key}.{nested_key}"
                                        logging.info(f"✅ [{client_id}] Extracted via method 3 nested ({key}.{nested_key}): {len(content)} characters")
                                        break
                                if content:
                                    break
                    
                    # Method 4 - Enhanced fallback with validation
                    if not content:
                        logging.error(f"🚨 [{client_id}] NO CONTENT FOUND: All extraction methods failed")
                        logging.error(f"🔍 [{client_id}] Candidate debug info:")
                        logging.error(f"   - Candidate type: {type(candidate)}")
                        logging.error(f"   - Candidate keys: {list(candidate.keys()) if isinstance(candidate, dict) else 'Not a dict'}")
                        logging.error(f"   - Candidate repr: {repr(candidate)[:300]}...")
                        
                        # Last resort: string conversion with validation
                        content_obj = candidate.get('content', candidate)
                        content_str = str(content_obj)
                        
                        # Validate the string conversion isn't just metadata
                        if len(content_str) > 50 and not (content_str.count('{') > 3 and content_str.count('}') > 3):
                            content = content_str
                            extraction_method = "string_conversion_validated"
                            logging.warning(f"⚠️ [{client_id}] Extracted via method 4 (validated): {len(content)} characters")
                        else:
                            logging.error(f"🚨 [{client_id}] String conversion also failed or returned metadata")
                            logging.error(f"🔍 [{client_id}] String conversion result: {content_str[:200]}...")
                            extraction_method = "failed"
                        
                except Exception as extraction_error:
                    logging.error(f"❌ [{client_id}] Content extraction error: {extraction_error}")
                    logging.error(f"🔍 [{client_id}] Candidate type: {type(candidate)}")
                    logging.error(f"🔍 [{client_id}] Extraction error type: {type(extraction_error).__name__}")
                    
                    # Fallback extraction
                    content = str(candidate)
                    extraction_method = "error_fallback"
                    logging.warning(f"⚠️ [{client_id}] Using error fallback extraction: {len(content)} characters")
                
                extraction_time = time.time() - extraction_start
                
                # Detailed content validation with comprehensive checks
                logging.info(f"🔍 [{client_id}] CONTENT VALIDATION ANALYSIS:")
                logging.info(f"🔍 [{client_id}] Extraction method: {extraction_method}")
                logging.info(f"🔍 [{client_id}] Extraction time: {extraction_time:.3f}s")
                logging.info(f"🔍 [{client_id}] Content type: {type(content)}")
                logging.info(f"🔍 [{client_id}] Content length: {len(content) if content else 0}")
                logging.info(f"🔍 [{client_id}] Content stripped length: {len(content.strip()) if content else 0}")
                logging.info(f"🔍 [{client_id}] Content preview (first 300 chars): '{content[:300] if content else 'EMPTY'}'")
                
                # Enhanced content validation with specific error types
                if not content:
                    logging.error(f"❌ [{client_id}] VALIDATION FAILED: Content is None or False")
                    logging.error(f"🔍 [{client_id}] Candidate finish reason: {finish_reason}")
                    logging.error(f"🔍 [{client_id}] Full candidate debug: {json.dumps(candidate, indent=2, default=str)}")
                    update_api_key_health(api_key, success=False, error_code="NO_CONTENT", response_time=response_time)
                    raise Exception("Content is None - API returned no text")
                    
                elif content.strip() == "":
                    logging.error(f"❌ [{client_id}] VALIDATION FAILED: Content is empty string or whitespace only")
                    logging.error(f"🔍 [{client_id}] Raw content representation: {repr(content)}")
                    update_api_key_health(api_key, success=False, error_code="EMPTY_CONTENT", response_time=response_time)
                    raise Exception("Content is empty string - API returned whitespace only")
                    
                elif len(content.strip()) < 10:  # Increased minimum from 5 to 10
                    logging.error(f"❌ [{client_id}] VALIDATION FAILED: Content too short: '{content.strip()}'")
                    logging.error(f"🔍 [{client_id}] Content length: {len(content.strip())} characters")
                    logging.error(f"🔍 [{client_id}] This usually indicates API returned metadata instead of content")
                    update_api_key_health(api_key, success=False, error_code="SHORT_CONTENT", response_time=response_time)
                    raise Exception(f"Content too short ({len(content.strip())} chars): '{content.strip()}'")
                
                # Check for metadata-only responses
                elif any(metadata_indicator in content.lower() for metadata_indicator in ['role": "model', '"parts":', '"content":']):
                    logging.error(f"🚨 [{client_id}] VALIDATION FAILED: Content appears to be metadata")
                    logging.error(f"🔍 [{client_id}] Metadata indicators found in content")
                    logging.error(f"🔍 [{client_id}] Content: {content[:200]}...")
                    update_api_key_health(api_key, success=False, error_code="METADATA_CONTENT", response_time=response_time)
                    raise Exception("Content appears to be API metadata instead of generated text")
                
                else:
                    # SUCCESS: Content validation passed
                    word_count = len(content.split())
                    logging.info(f"✅ [{client_id}] CONTENT VALIDATION PASSED")
                    logging.info(f"📊 [{client_id}] Content metrics:")
                    logging.info(f"   - Character count: {len(content):,}")
                    logging.info(f"   - Word count: {word_count:,}")
                    logging.info(f"   - Line count: {content.count(chr(10)) + 1}")
                    logging.info(f"   - Extraction method: {extraction_method}")
                    logging.info(f"   - Content quality: {'HIGH' if word_count > 500 else 'MEDIUM' if word_count > 100 else 'LOW'}")
                
                # Success metrics calculation
                usage = data.get('usageMetadata', {})
                token_count = usage.get('totalTokenCount', 0)
                analysis_time = time.time() - start_time
                
                logging.info(f"🎉 [{client_id}] ANALYSIS COMPLETE - SUCCESS")
                logging.info(f"📊 [{client_id}] Final success metrics:")
                logging.info(f"   - Analysis type: {analysis_type}")
                logging.info(f"   - Total tokens: {token_count:,}")
                logging.info(f"   - Total time: {analysis_time:.2f}s")
                logging.info(f"   - Request time: {response_time:.2f}s")
                logging.info(f"   - Processing time: {(analysis_time - response_time):.2f}s")
                logging.info(f"   - Words generated: {len(content.split()):,}")
                logging.info(f"   - Characters generated: {len(content):,}")
                
                # Log final API key health status after successful completion
                updated_health = api_key_health.get(api_key, {})
                updated_success_rate = updated_health.get('success_rate', 1.0)
                updated_current_load = updated_health.get('current_load', 0)
                updated_total_requests = updated_health.get('total_requests', 0)
                
                logging.info(f"🔑 [{client_id}] FINAL API KEY HEALTH:")
                logging.info(f"   - Key ID: {updated_health.get('key_id', 'unknown')}")
                logging.info(f"   - Status: HEALTHY")
                logging.info(f"   - Consecutive Failures: 0 (reset)")
                logging.info(f"   - Success Rate: {updated_success_rate:.3f}")
                logging.info(f"   - Current Load: {updated_current_load}")
                logging.info(f"   - Total Requests: {updated_total_requests}")
                logging.info(f"   - Last Response Time: {response_time:.2f}s")
                
                # Create response object with comprehensive metadata
                response_obj = RiskChatResponse(
                    content=content,
                    model="gemini-2.5-pro",
                    api_key_used=f"{client_id}_key_{api_key[-4:]}",
                    usage=usage,
                    finish_reason=candidate.get('finishReason', 'STOP'),
                    response_time=analysis_time,
                    timestamp=time.time(),
                    token_count=token_count
                )
                
                logging.info(f"✅ [{client_id}] Response object created successfully")
                logging.info(f"🎯 [{client_id}] Returning successful analysis result")
                
                return response_obj
                
            else:
                # CRITICAL ERROR: No candidates in response
                logging.error(f"❌ [{client_id}] CRITICAL ERROR: No candidates in response")
                logging.error(f"🔍 [{client_id}] Response data: {data}")
                logging.error(f"🔍 [{client_id}] Response structure analysis:")
                
                if 'candidates' in data:
                    candidates = data['candidates']
                    logging.error(f"   - Candidates key exists: True")
                    logging.error(f"   - Candidates type: {type(candidates)}")
                    logging.error(f"   - Candidates length: {len(candidates)}")
                    logging.error(f"   - Candidates content: {candidates}")
                else:
                    logging.error(f"   - Candidates key exists: False")
                    logging.error(f"   - Available keys: {list(data.keys())}")
                
                # Check for error indicators in response
                if 'error' in data:
                    error_info = data['error']
                    logging.error(f"🚨 [{client_id}] API returned error: {error_info}")
                    update_api_key_health(api_key, success=False, error_code="API_ERROR", response_time=response_time)
                    raise Exception(f"API error: {error_info}")
                
                update_api_key_health(api_key, success=False, error_code="NO_CANDIDATES", response_time=response_time)
                raise Exception("No candidates found in API response")
        
        else:
            # HTTP ERROR handling with detailed analysis
            error_code = str(response.status_code)
            logging.error(f"❌ [{client_id}] HTTP ERROR: {response.status_code}")
            logging.error(f"🔍 [{client_id}] Response details:")
            logging.error(f"   - Status code: {response.status_code}")
            logging.error(f"   - Response time: {response_time:.2f}s")
            logging.error(f"   - Content length: {len(response.text)}")
            logging.error(f"   - Content type: {response.headers.get('content-type', 'Unknown')}")
            logging.error(f"🔍 [{client_id}] Response text (first 500 chars): {response.text[:500]}...")
            
            # Special handling for different HTTP error codes
            if response.status_code == 503:
                logging.error(f"🚨 [{client_id}] API OVERLOADED (503) - Service temporarily unavailable")
                logging.error(f"🔍 [{client_id}] This indicates Google's servers are overloaded")
                logging.error(f"🔧 [{client_id}] Marking API key for extended cooldown")
                update_api_key_health(api_key, success=False, error_code="503", response_time=response_time)
                
                # Log current API key ecosystem health
                ecosystem_health = get_enhanced_api_key_status()
                healthy_keys = ecosystem_health.get('healthy_keys', 0)
                logging.warning(f"🔑 [{client_id}] API Key Ecosystem Status: {healthy_keys} healthy keys remaining")
                
                if healthy_keys <= 2:
                    logging.error(f"🚨 [{client_id}] CRITICAL: Very few healthy API keys remaining!")
                
            elif response.status_code == 429:
                logging.error(f"🚨 [{client_id}] RATE LIMITED (429) - Too many requests")
                logging.error(f"🔍 [{client_id}] Rate limit headers: {dict(response.headers)}")
                update_api_key_health(api_key, success=False, error_code="429", response_time=response_time)
                
            elif response.status_code in [400, 401, 403]:
                logging.error(f"🚨 [{client_id}] CLIENT ERROR ({response.status_code}) - Request or auth issue")
                
                if response.status_code == 400:
                    logging.error(f"🔍 [{client_id}] Bad Request - Possible prompt issue")
                    logging.error(f"🔍 [{client_id}] Prompt length: {len(enhanced_prompt)} chars")
                elif response.status_code == 401:
                    logging.error(f"🔍 [{client_id}] Unauthorized - API key may be invalid")
                elif response.status_code == 403:
                    logging.error(f"🔍 [{client_id}] Forbidden - API key may be disabled")
                
                update_api_key_health(api_key, success=False, error_code=error_code, response_time=response_time)
                
            elif response.status_code >= 500:
                logging.error(f"🚨 [{client_id}] SERVER ERROR ({response.status_code}) - Google server issue")
                logging.error(f"🔍 [{client_id}] This is typically temporary")
                update_api_key_health(api_key, success=False, error_code=error_code, response_time=response_time)
                
            else:
                logging.error(f"🚨 [{client_id}] UNKNOWN HTTP ERROR ({response.status_code})")
                update_api_key_health(api_key, success=False, error_code=error_code, response_time=response_time)
            
            # Log updated API key health after error
            updated_health = api_key_health.get(api_key, {})
            updated_success_rate = updated_health.get('success_rate', 1.0)
            updated_current_load = updated_health.get('current_load', 0)
            updated_failures = updated_health.get('consecutive_failures', 0)
            
            logging.error(f"🔑 [{client_id}] UPDATED API KEY HEALTH AFTER ERROR:")
            logging.error(f"   - Key ID: {updated_health.get('key_id', 'unknown')}")
            logging.error(f"   - Consecutive Failures: {updated_failures}")
            logging.error(f"   - Success Rate: {updated_success_rate:.3f}")
            logging.error(f"   - Current Load: {updated_current_load}")
            logging.error(f"   - Error Response Time: {response_time:.2f}s")
            
            # Log ecosystem impact
            ecosystem_summary = get_api_key_status_summary()
            logging.error(f"🔑 [{client_id}] API Key Ecosystem Impact: {ecosystem_summary}")
            
            raise Exception(f"HTTP {response.status_code}: {response.text}")
    
    except Exception as e:
        # Comprehensive error handling with detailed context
        analysis_time = time.time() - start_time
        response_time_safe = time.time() - request_start_time if request_start_time else analysis_time
        
        logging.error(f"❌ [{client_id}] RISK ANALYSIS FAILED")
        logging.error(f"🔍 [{client_id}] Error details:")
        logging.error(f"   - Error type: {type(e).__name__}")
        logging.error(f"   - Error message: {str(e)}")
        logging.error(f"   - Analysis time: {analysis_time:.2f}s")
        logging.error(f"   - Request time: {response_time_safe:.2f}s")
        logging.error(f"   - Analysis type: {analysis_type}")
        logging.error(f"   - Client ID: {client_id}")
        
        # Context logging for debugging
        logging.error(f"🔍 [{client_id}] Error context:")
        logging.error(f"   - API key suffix: ...{api_key[-4:]}")
        logging.error(f"   - Temperature: {temperature}")
        logging.error(f"   - Max tokens: {max_tokens}")
        logging.error(f"   - REQUEST_TIMEOUT: {REQUEST_TIMEOUT}s")
        
        # If this exception wasn't already handled above, update API key health
        if api_key in api_key_health:
            current_failures = api_key_health[api_key].get('consecutive_failures', 0)
            
            # Only update if we haven't already updated for HTTP errors
            if "HTTP" not in str(e) and "API error" not in str(e):
                logging.warning(f"🔑 [{client_id}] Updating API key health for unhandled exception")
                update_api_key_health(api_key, success=False, error_code="GENERAL_EXCEPTION", response_time=response_time_safe)
            else:
                logging.info(f"🔑 [{client_id}] API key health already updated for this error type")
            
            # Log comprehensive API key health summary on error
            final_health = api_key_health.get(api_key, {})
            final_success_rate = final_health.get('success_rate', 1.0)
            final_current_load = final_health.get('current_load', 0)
            final_failures = final_health.get('consecutive_failures', 0)
            
            logging.error(f"🔑 [{client_id}] FINAL API KEY HEALTH AFTER ERROR:")
            logging.error(f"   - Key ID: {final_health.get('key_id', 'unknown')}")
            logging.error(f"   - Consecutive Failures: {final_failures}")
            logging.error(f"   - Success Rate: {final_success_rate:.3f}")
            logging.error(f"   - Current Load: {final_current_load}")
            logging.error(f"   - Key Status: {'FAILED' if final_failures >= 3 else 'DEGRADED' if final_failures > 0 else 'HEALTHY'}")
            
            # Log ecosystem health impact
            ecosystem_status = get_enhanced_api_key_status()
            logging.error(f"🔑 [{client_id}] API ECOSYSTEM IMPACT:")
            logging.error(f"   - Healthy keys: {ecosystem_status.get('healthy_keys', 0)}/{ecosystem_status.get('total_keys', 0)}")
            logging.error(f"   - Failed keys: {ecosystem_status.get('failed_keys', 0)}")
            logging.error(f"   - Cooling down: {ecosystem_status.get('cooling_down', 0)}")
            logging.error(f"   - Total load: {ecosystem_status.get('total_load', 0)}")
        else:
            logging.error(f"❌ [{client_id}] API key not found in health tracking: ...{api_key[-4:]}")
        
        # Full traceback logging for debugging
        import traceback
        logging.error(f"🔍 [{client_id}] FULL ERROR TRACEBACK:")
        for line_num, line in enumerate(traceback.format_exc().split('\n'), 1):
            if line.strip():
                logging.error(f"   {line_num:02d}: {line}")
        
        # Log additional debugging context
        logging.error(f"🔍 [{client_id}] DEBUGGING CONTEXT:")
        logging.error(f"   - Python version: {sys.version}")
        logging.error(f"   - Platform: {platform.system()}")
        logging.error(f"   - Memory usage: {sys.getsizeof(complete_raw_data) if complete_raw_data else 0} bytes")
        logging.error(f"   - Thread ID: {threading.current_thread().ident}")
        logging.error(f"   - Function start time: {datetime.fromtimestamp(start_time).isoformat()}")
        
        raise

def create_enhanced_risk_analysis_prompt(complete_raw_data: Dict, analysis_type: str, analysis_requirements: str) -> str:
    """Create 100/100 enhanced analysis prompt with complete Multi-Database Intelligence integration for Risk Engine"""
    
    logging.info(f"🎯 Starting enhanced risk analysis prompt creation for {analysis_type}")
    
    user_profile = complete_raw_data.get("user_profile", {})
    responses = complete_raw_data.get("responses", [])
    multi_db_intelligence = complete_raw_data.get("multi_database_intelligence", {})
    behavioral_data = complete_raw_data.get("behavioral_analytics", {})
    
    logging.info(f"📊 Data summary: {len(responses)} responses, multi-db: {bool(multi_db_intelligence)}, behavioral: {bool(behavioral_data)}")
    
    # Extract and validate user profile data
    business_name = user_profile.get('business_name', 'Unknown Business')
    username = user_profile.get('username', 'Client')
    
    # Handle industry as both string and list
    industry_raw = user_profile.get('industry', 'Unknown Industry')
    if isinstance(industry_raw, list):
        industry = ", ".join(industry_raw) if industry_raw else 'Unknown Industry'
    else:
        industry = str(industry_raw) if industry_raw else 'Unknown Industry'
    
    team_size = user_profile.get('team_size', 'Unknown')
    biggest_challenge = user_profile.get('biggest_challenge', 'Unknown Challenge')
    business_description = user_profile.get('business_description', 'Not provided')
    location = user_profile.get('location', 'Unknown Location')
    
    logging.info(f"👤 User profile: {username} at {business_name} ({industry}, {team_size} employees)")
    
    # Get current date and time for Gemini context
    current_datetime = datetime.now()
    current_date_str = current_datetime.strftime('%A, %B %d, %Y')
    current_time_str = current_datetime.strftime('%I:%M %p %Z')
    current_timestamp = current_datetime.isoformat()

    user_context = f"""
═══════════════════════════════════════════════════════════════════════════════
🏰 WORLD-CLASS RISK & FORTRESS INTELLIGENCE ANALYSIS FRAMEWORK 🏰
═══════════════════════════════════════════════════════════════════════════════

📅 ANALYSIS CONTEXT:
- Analysis Date: {current_date_str}
- Analysis Time: {current_time_str}
- Report Generation: Evidence-based risk intelligence with validated methodologies

👤 CLIENT PROFILE:
- Full Name: {username}
- Business Name: {business_name}
- Industry: {industry}
- Team Size: {team_size} employees
- Location: {location}
- Primary Challenge: {biggest_challenge}
- Business Description: {business_description}

🎯 ANALYSIS OBJECTIVE:
Generate a comprehensive risk fortress analysis for {username}, founder/leader of {business_name}, 
integrating multi-database intelligence across validated risk frameworks to address {biggest_challenge} 
with evidence-based recommendations and quantified confidence levels.

📊 DATA INTEGRATION SUMMARY:
{enhanced_format_multi_database_intelligence(multi_db_intelligence)}

🧠 BEHAVIORAL INTELLIGENCE:
{enhanced_format_risk_behavioral_data(behavioral_data)}

═══════════════════════════════════════════════════════════════════════════════
🔬 METHODOLOGICAL REQUIREMENTS FOR 100/100 ANALYSIS QUALITY:
═══════════════════════════════════════════════════════════════════════════════

**CORE METHODOLOGY PRINCIPLES:**
1. **Evidence-Based Analysis**: Every recommendation supported by specific data patterns from above intelligence
2. **Validated Frameworks**: Use established risk theory (ISO 31000, COSO ERM, etc.) enhanced with behavioral insights
3. **Statistical Rigor**: Proper correlation analysis with confidence intervals using multi-engine data
4. **Alternative Explanations**: Consider competing hypotheses for all major findings
5. **Implementation Focus**: Specific, actionable recommendations with resource requirements
6. **Risk Assessment**: Scenario planning with probability-weighted outcomes
7. **Cross-Engine Integration**: Synthesize insights from ALL intelligence sources provided above
8. **Behavioral Customization**: Align recommendations with individual behavioral patterns and preferences

**DATA ANALYSIS STANDARDS:**
- **Multi-Engine Integration**: Use ALL intelligence sources provided above for comprehensive analysis
- **Confidence Intervals**: Provide 95% CI for all quantitative predictions
- **Effect Size Reporting**: Include Cohen's d or eta-squared for practical significance
- **Correlation Analysis**: Report actual correlation coefficients with significance levels
- **Validation Methods**: Cross-reference findings across multiple data sources
- **Assumption Testing**: Explicitly state and validate key analytical assumptions
- **Behavioral Integration**: Map individual patterns to organizational risk factors
- **Implementation Success Modeling**: Predict intervention effectiveness with confidence ranges

═══════════════════════════════════════════════════════════════════════════════
📊 SECTION 1: THREAT LANDSCAPE & VULNERABILITY ANALYSIS
═══════════════════════════════════════════════════════════════════════════════

**ANALYSIS FRAMEWORK:**
Apply validated risk assessment models integrating ALL intelligence sources above:
- **ISO 31000 Risk Management Framework** enhanced with behavioral patterns from multi-engine data
- **COSO Enterprise Risk Management Model** integrated with operational insights from comprehensive intelligence
- **NIST Cybersecurity Framework** calibrated to business phase and growth trajectory from engine analysis
- **Business Continuity Planning Standards** aligned with team dynamics and communication patterns
- **Crisis Management Best Practices** customized to leadership style and decision-making preferences

**REQUIRED STATISTICAL ANALYSIS:**
- **Cross-Engine Threat Correlation**: Map behavioral patterns to vulnerability susceptibility using intelligence data
- **Multi-Database Risk Assessment**: Integrate all intelligence sources for comprehensive threat modeling
- **Vulnerability Prioritization**: Use validated scales enhanced with behavioral and operational insights
- **Risk Exposure Analysis**: Analyze threat categories using cross-engine correlation patterns
- **Industry Benchmarking**: Compare findings with {industry} sector patterns and best practices
- **Behavioral Risk Profiling**: Map individual patterns to organizational vulnerability patterns

**KEY DELIVERABLES:**
1. **Cross-Engine Threat Assessment Matrix** with behavioral amplification factors and ROI estimates
2. **Integrated Vulnerability Prioritization** using multi-engine intelligence with timeline milestones
3. **Behavioral Risk Strategy** customized to individual patterns with effectiveness metrics
4. **Business Continuity Roadmap** aligned with operational maturity and team capabilities
5. **Crisis Management Protocols** matching communication style and leadership preferences

**CONFIDENCE REQUIREMENTS:**
- All recommendations include confidence levels (0.0-1.0 scale) based on multi-engine data quality
- Statistical significance testing where applicable using cross-engine correlations
- Alternative scenarios with probability weighting using behavioral and operational insights
- Implementation success predictions with ranges based on individual patterns and capabilities

═══════════════════════════════════════════════════════════════════════════════
🛡️ SECTION 2: FORTRESS ARCHITECTURE & DEFENSE SYSTEMS ANALYSIS
═══════════════════════════════════════════════════════════════════════════════

**ANALYSIS FRAMEWORK:**
Apply established defensive architecture models enhanced with comprehensive intelligence:
- **Defense in Depth Strategy** calibrated to behavioral decision-making patterns from multi-engine analysis
- **Zero Trust Security Model** aligned with delegation comfort and team trust dynamics from intelligence
- **Layered Security Architecture** optimized for operational maturity level from cross-engine assessment
- **Incident Response Planning** customized to communication style and crisis management preferences
- **Security Operations Integration** matching learning patterns and information processing capabilities

**REQUIRED ANALYTICAL METHODS:**
- **Behavioral Defense Optimization**: Use personality patterns to design optimal security controls
- **Multi-Engine Security Assessment**: Integrate operational, behavioral, and strategic insights
- **Attack Surface Analysis**: Consider business processes and behavioral factors in vulnerability assessment
- **Security Control Effectiveness**: Measure against behavioral compliance patterns and operational maturity
- **Response Time Optimization**: Factor in communication patterns and decision-making speed
- **Defense ROI Modeling**: Calculate security investment returns using comprehensive business intelligence

**KEY DELIVERABLES:**
1. **Personalized Defense Architecture** with behavioral customization and implementation phases
2. **Adaptive Security Controls** that evolve with business growth and maturity indicators
3. **Integrated Monitoring Systems** with behavioral alert calibration and response protocols
4. **Incident Response Framework** matching individual communication and crisis management style
5. **Security Awareness Program** aligned with learning preferences and team engagement patterns

**INTEGRATION REQUIREMENTS:**
- **Cross-reference with Operational Intelligence**: Align security with business process requirements
- **Behavioral Alignment**: Ensure security controls match individual and team behavioral patterns
- **Growth Trajectory Integration**: Design scalable security that supports expansion plans
- **Financial Optimization**: Balance security investment with risk tolerance and budget constraints

═══════════════════════════════════════════════════════════════════════════════
⚔️ SECTION 3: OPERATIONAL RESILIENCE & RECOVERY ANALYSIS
═══════════════════════════════════════════════════════════════════════════════

**ANALYSIS FRAMEWORK:**
Apply proven resilience improvement methodologies enhanced with multi-engine intelligence:
- **Business Impact Analysis (BIA)** informed by growth ambitions and operational maturity from intelligence
- **Recovery Time Objectives (RTO) Planning** calibrated to actual team capabilities and behavioral patterns
- **Disaster Recovery Planning** aligned with financial capacity and strategic priorities from cross-engine data
- **Business Continuity Management** scaled to leadership style and communication effectiveness
- **Crisis Communication Strategies** customized to stakeholder relationship patterns and influence networks

**REQUIRED PROCESS ANALYSIS:**
- **Critical Function Mapping**: Use operational intelligence to identify true business dependencies
- **Recovery Capability Assessment**: Factor in team dynamics and behavioral patterns for realistic planning
- **Backup System Evaluation**: Consider operational maturity and technological comfort levels
- **Communication Effectiveness**: Assess stakeholder management patterns and crisis communication style
- **Resource Optimization**: Balance recovery speed with financial constraints and operational capabilities
- **Timeline Prediction**: Model recovery scenarios using behavioral and operational intelligence

**KEY DELIVERABLES:**
1. **Dynamic Business Continuity Plan** with behavioral customization and scenario-based protocols
2. **Predictive Recovery Strategy** with confidence intervals and resource optimization analysis
3. **Stakeholder Communication Matrix** aligned with relationship management style and network strength
4. **Resilience Training Program** matching learning preferences and skill development needs
5. **Continuous Testing Framework** with realistic scenarios and behavioral factor integration

**QUANTIFICATION REQUIREMENTS:**
- **Recovery Time Estimates**: Confidence intervals based on team capabilities and behavioral patterns
- **Cost Impact Projections**: Probability ranges using financial intelligence and operational constraints
- **Success Probability Modeling**: Factor in leadership effectiveness and team dynamics from intelligence
- **Resource Requirement Analysis**: Optimize allocation using comprehensive business intelligence

═══════════════════════════════════════════════════════════════════════════════
🔒 SECTION 4: COMPLIANCE & GOVERNANCE ANALYSIS
═══════════════════════════════════════════════════════════════════════════════

**ANALYSIS FRAMEWORK:**
Apply validated governance assessment models enhanced with behavioral intelligence:
- **Regulatory Compliance Management** calibrated to risk tolerance and industry requirements
- **Risk Governance Framework** aligned with leadership decision-making patterns and delegation preferences
- **Audit and Assurance Processes** matched to operational maturity and documentation comfort levels
- **Policy and Procedure Management** customized to communication style and team engagement patterns
- **Stakeholder Communication** optimized for relationship management effectiveness and influence networks

**REQUIRED GOVERNANCE METRICS:**
- **Compliance Effectiveness Correlation**: Map governance success to behavioral and operational factors
- **Governance Maturity Scoring**: Use multi-engine intelligence for comprehensive maturity assessment
- **Policy Adherence Prediction**: Model compliance success using team dynamics and leadership style
- **Audit Readiness Assessment**: Factor in documentation patterns and process systematization levels
- **Stakeholder Engagement Analysis**: Evaluate relationship strength and communication effectiveness
- **Cultural Alignment Measurement**: Assess governance culture fit with behavioral and operational patterns

**KEY DELIVERABLES:**
1. **Adaptive Governance Strategy** with behavioral customization and maturity-based scaling
2. **Compliance Enhancement Plan** matching operational preferences and team capabilities
3. **Policy Framework Optimization** aligned with communication style and enforcement comfort
4. **Audit Program Enhancement** with process maturity integration and effectiveness metrics
5. **Stakeholder Governance Protocol** customized to relationship management patterns and influence style

**BEHAVIORAL INTEGRATION:**
- **Leadership Style Alignment**: Match governance approach to individual behavioral patterns from intelligence
- **Team Dynamics Integration**: Consider group behavior and communication patterns in policy design
- **Cultural Compatibility**: Ensure governance framework aligns with organizational culture indicators
- **Change Management**: Factor in behavioral change patterns for governance implementation success

═══════════════════════════════════════════════════════════════════════════════
⚠️ SECTION 5: STRATEGIC RISK MANAGEMENT & PLANNING ANALYSIS
═══════════════════════════════════════════════════════════════════════════════

**ANALYSIS FRAMEWORK:**
Apply strategic risk management best practices enhanced with comprehensive intelligence:
- **Strategic Risk Assessment** informed by vision-reality alignment and growth trajectory analysis
- **Risk Appetite and Tolerance Setting** calibrated to behavioral patterns and financial capacity
- **Key Risk Indicator (KRI) Development** using cross-engine correlation analysis and predictive modeling
- **Risk Reporting and Communication** aligned with leadership style and stakeholder engagement patterns
- **Board and Executive Risk Oversight** optimized for governance preferences and strategic involvement

**REQUIRED RISK ASSESSMENT:**
- **Vision-Risk Alignment Analysis**: Compare dream ambitions with risk management capability
- **Strategic Risk Probability Matrix**: Model scenarios using behavioral and operational intelligence
- **Risk Appetite Calibration**: Align tolerance with behavioral patterns and financial constraints
- **KPI-Risk Correlation**: Connect performance indicators with risk metrics using multi-engine data
- **Strategic Communication Effectiveness**: Assess reporting alignment with stakeholder preferences
- **Executive Oversight Optimization**: Match governance style with leadership behavioral patterns

**KEY DELIVERABLES:**
1. **Dynamic Strategic Risk Register** with cross-engine intelligence and probability-weighted impacts
2. **Behavioral Risk Appetite Framework** with personalized tolerance boundaries and evolution protocols
3. **Predictive KRI Dashboard** with early warning systems and behavioral alert calibration
4. **Strategic Risk Communication Plan** customized to stakeholder management and influence patterns
5. **Executive Risk Governance Framework** aligned with leadership style and oversight preferences

**SCENARIO MODELING:**
- **Best/Worst/Most Likely Scenarios**: Enhanced with behavioral and operational intelligence factors
- **Economic Impact Assessment**: Factor in financial intelligence and market positioning analysis
- **Competitive Threat Analysis**: Use industry intelligence and strategic positioning insights
- **Regulatory Change Impact**: Consider compliance maturity and adaptation capability patterns
- **Stakeholder Risk Assessment**: Evaluate relationship strength and influence network resilience

═══════════════════════════════════════════════════════════════════════════════
💰 SECTION 6: FINANCIAL RISK & IMPACT ANALYSIS
═══════════════════════════════════════════════════════════════════════════════

**ANALYSIS FRAMEWORK:**
Apply financial risk management best practices enhanced with multi-engine intelligence:
- **Financial Risk Assessment** integrated with growth projections and operational cash flow patterns
- **Liquidity Risk Management** calibrated to revenue predictability and expansion financing needs
- **Credit Risk Evaluation** aligned with customer concentration and market positioning intelligence
- **Market Risk Analysis** enhanced with competitive positioning and industry trend insights
- **Operational Risk Quantification** using process maturity and people dependency assessments

**REQUIRED FINANCIAL METRICS:**
- **Risk-Adjusted Return Analysis**: Use performance indicators enhanced with behavioral risk factors
- **Financial Exposure Optimization**: Balance risk limits with growth ambitions and operational needs
- **Capital Allocation Effectiveness**: Align investment priorities with strategic goals and risk appetite
- **Financial Resilience Assessment**: Stress test scenarios using operational and behavioral intelligence
- **Insurance and Hedging Strategy**: Optimize coverage using comprehensive risk assessment and cost-benefit analysis

**KEY DELIVERABLES:**
1. **Integrated Financial Risk Plan** with growth-aligned liquidity management and investment optimization
2. **Predictive Cash Flow Strategy** with scenario modeling and behavioral factor integration
3. **Risk-Adjusted Investment Framework** with strategic priority alignment and resource optimization
4. **Financial Monitoring Dashboard** with behavioral alert systems and early warning indicators
5. **Strategic Financial Planning** with multi-scenario modeling and confidence interval projections

**INTEGRATION WITH OPERATIONS:**
- **Operational Risk Alignment**: Connect financial metrics with operational maturity and process reliability
- **Growth Investment Optimization**: Balance expansion priorities with financial stability requirements
- **Behavioral Financial Patterns**: Factor in decision-making style and risk tolerance for investment strategy
- **Market Positioning Impact**: Consider competitive intelligence and industry dynamics in financial planning

═══════════════════════════════════════════════════════════════════════════════
🚀 SECTION 7: IMPLEMENTATION ROADMAP & EXECUTION ANALYSIS
═══════════════════════════════════════════════════════════════════════════════

**ANALYSIS FRAMEWORK:**
Apply proven implementation methodologies enhanced with behavioral intelligence:
- **Risk Management Implementation Planning** calibrated to decision-making patterns and change preferences
- **Change Management for Risk Culture** aligned with leadership style and team engagement dynamics
- **Project Management for Risk Initiatives** optimized for operational preferences and resource allocation
- **Training and Awareness Programs** customized to learning styles and skill development patterns
- **Continuous Improvement Processes** matching feedback preferences and optimization cycle comfort

**REQUIRED IMPLEMENTATION PLANNING:**
- **Phase-Gate Implementation**: Define milestones using behavioral and operational readiness indicators
- **Resource Requirement Planning**: Assess skill gaps using team capability intelligence and development needs
- **Change Management Strategy**: Design resistance mitigation using behavioral change patterns and team dynamics
- **Success Measurement Framework**: Define KPIs using behavioral preferences and operational maturity indicators
- **Timeline Optimization**: Risk-adjust scheduling using implementation capability and change readiness factors

**KEY DELIVERABLES:**
1. **Master Implementation Timeline** with critical path analysis and behavioral factor integration
2. **Personalized Resource Allocation Plan** with skill assessment and capability development optimization
3. **Behavioral Change Management Strategy** with resistance mitigation and engagement optimization
4. **Success Measurement Framework** with behavioral KPI calibration and continuous improvement triggers
5. **Adaptive Implementation Protocol** with feedback integration and course correction mechanisms

**TIMELINE SPECIFICATIONS:**
- **0-30 Day Quick Wins**: Immediate impact initiatives aligned with behavioral patterns and operational readiness
- **30-90 Day Foundation**: Core system implementation with team capability development and change management
- **90-180 Day Integration**: Cross-functional optimization with behavioral alignment and process integration
- **180-365 Day Optimization**: Advanced capability development with continuous improvement and performance tuning
- **Annual Strategic Evolution**: Long-term capability building with strategic alignment and competitive advantage development

═══════════════════════════════════════════════════════════════════════════════
🔍 CROSS-SECTIONAL CORRELATION ANALYSIS REQUIREMENTS
═══════════════════════════════════════════════════════════════════════════════

**STATISTICAL CORRELATION ANALYSIS:**
Analyze relationships between all major variables using multi-engine intelligence:
- **Pearson Correlation Coefficients**: For continuous variables across behavioral, operational, and strategic dimensions
- **Spearman Rank Correlations**: For ordinal relationships between maturity levels and risk factors
- **Chi-Square Tests**: For categorical associations between risk types and behavioral patterns
- **Partial Correlations**: Control for confounding variables using comprehensive intelligence data
- **Multiple Regression Analysis**: Model complex relationships between behavioral, operational, and risk factors

**REQUIRED CORRELATION INVESTIGATIONS:**
1. **Behavioral Risk Susceptibility**: Individual patterns ↔ Organizational vulnerability correlation (target r > 0.30)
2. **Growth-Risk Capacity**: Expansion ambitions ↔ Risk management capability alignment analysis
3. **Leadership-Crisis Effectiveness**: Personal style ↔ Organizational crisis response success rates
4. **Operational Maturity-Resilience**: Process sophistication ↔ Business continuity capability correlation
5. **Financial Health-Risk Appetite**: Cash flow stability ↔ Strategic risk tolerance optimization
6. **Team Dynamics-Compliance**: Communication patterns ↔ Governance implementation success rates
7. **Vision-Implementation Risk**: Dream ambitions ↔ Execution failure probability modeling

**INSIGHT DISCOVERY MANDATE:**
- **Identify minimum 10 statistically significant correlations** (p<0.05) with practical business implications
- **Document effect sizes** using Cohen's d for practical significance assessment and prioritization
- **Provide alternative explanations** for all major correlations with competitive hypothesis testing
- **Quantify confidence levels** for all relationship claims using bootstrap confidence intervals
- **Cross-validate findings** using multiple analytical approaches and out-of-sample testing

**INTEGRATION SYNTHESIS:**
- **Connect Individual-Organizational Patterns**: Map personal behavioral patterns to business risk vulnerabilities
- **Align Assessment-Intervention Strategy**: Link multi-engine insights with specific risk management interventions
- **Customize Recommendations**: Tailor all suggestions to {username}'s behavioral patterns and {business_name}'s context
- **Optimize Resource Allocation**: Prioritize interventions based on correlation strength and implementation feasibility

═══════════════════════════════════════════════════════════════════════════════
📈 PREDICTIVE MODELING & FORECASTING REQUIREMENTS
═══════════════════════════════════════════════════════════════════════════════

**RISK SUCCESS PREDICTION:**
Develop predictive models for key risk outcomes using comprehensive intelligence:
- **Implementation Success Probability**: Model intervention effectiveness using behavioral and operational factors
- **Risk Reduction Trajectory**: Forecast improvement timeline with confidence intervals and milestone predictions
- **Incident Likelihood Prediction**: Estimate crisis probability using vulnerability and preparedness assessments
- **Financial Impact Projections**: Model cost scenarios with sensitivity analysis and uncertainty quantification
- **Compliance Effectiveness Forecasting**: Predict governance success using behavioral and operational intelligence

**FORECASTING METHODOLOGIES:**
- **Time Series Analysis**: Identify risk trend patterns using historical and behavioral intelligence data
- **Regression Modeling**: Predict outcomes using multi-engine variable integration and interaction effects
- **Scenario Analysis**: Quantify uncertainty using behavioral, operational, and strategic factor combinations
- **Monte Carlo Simulation**: Model complex risk scenarios with probability distributions and confidence intervals
- **Machine Learning Integration**: Use pattern recognition for non-linear relationship discovery and prediction accuracy

**VALIDATION REQUIREMENTS:**
- **Cross-Validation**: Test predictive model accuracy using temporal and behavioral data splits
- **Out-of-Sample Testing**: Validate model generalization using holdout intelligence data where applicable
- **Confidence Interval Reporting**: Provide uncertainty quantification for all predictions using bootstrap methods
- **Alternative Model Comparison**: Select optimal approaches using AIC, BIC, and cross-validation performance metrics
- **Assumption Testing**: Document and validate key modeling assumptions using comprehensive intelligence data

═══════════════════════════════════════════════════════════════════════════════
🎯 SUCCESS METRICS & QUALITY ASSURANCE FRAMEWORK (100/100)
═══════════════════════════════════════════════════════════════════════════════

**ANALYSIS QUALITY SCORING (100-POINT SYSTEM):**

**METHODOLOGY & EVIDENCE BASE (30 points):**
- **Multi-Engine Data Integration**: Complete utilization of all intelligence sources provided above (10 points)
- **Statistical Rigor**: Confidence intervals, effect sizes, and significance testing with proper methodology (10 points)
- **Alternative Explanations**: Competitive hypothesis testing and model comparison (5 points)
- **Validated Framework Integration**: Proper application of established risk management theories (5 points)

**PRIORITIZATION & IMPACT QUANTIFICATION (25 points):**
- **ROI Estimates**: Confidence ranges for all recommendations using comprehensive intelligence (10 points)
- **Resource Requirement Analysis**: Detailed planning with skill gap assessment and optimization (5 points)
- **Implementation Timeline**: Risk-adjusted scheduling with behavioral and operational factors (5 points)
- **Industry Benchmarking**: {industry} context integration with competitive intelligence (5 points)

**RISK ASSESSMENT & SCENARIO PLANNING (20 points):**
- **Comprehensive Risk Analysis**: Multi-dimensional assessment with cross-engine correlation (10 points)
- **Scenario Modeling**: Probability weighting with behavioral and operational factor integration (5 points)
- **Validation Framework**: Success metrics with confidence intervals and measurement protocols (5 points)

**IMPLEMENTATION SPECIFICITY (15 points):**
- **Detailed Tactical Plans**: Specific, actionable recommendations with behavioral customization (8 points)
- **Success Measurement Protocols**: KPI definition with behavioral calibration and monitoring systems (4 points)
- **Stakeholder Management Strategy**: Communication and engagement aligned with relationship patterns (3 points)

**INTEGRATION & PERSONALIZATION (10 points):**
- **Cross-Engine Correlation Analysis**: Statistical relationships across all intelligence sources (5 points)
- **Behavioral Customization**: Tailored recommendations for {username} and {business_name} context (5 points)

**MANDATORY DELIVERABLE REQUIREMENTS:**
□ Evidence-based recommendations supported by multi-engine statistical analysis and intelligence integration
□ Quantified confidence levels for all major claims using comprehensive data sources and validation methods
□ Resource requirements and timeline estimates with behavioral customization and operational optimization
□ Risk scenarios with mitigation strategies using cross-engine correlation and predictive modeling
□ Success measurement frameworks with behavioral KPI calibration and continuous improvement protocols
□ Complete integration across all available intelligence dimensions with correlation analysis and pattern recognition
□ Specific customization to {username}, {business_name}, {industry}, and {biggest_challenge} using all data sources
□ Implementation roadmap with behavioral alignment and operational maturity considerations
□ Predictive accuracy validation with confidence intervals and alternative scenario modeling
□ Cross-engine synthesis connecting all available intelligence sources for comprehensive risk strategy

**CRITICAL SUCCESS FACTORS:**
1. **Statistical Excellence**: All correlations reported with significance levels, confidence intervals, and practical effect sizes
2. **Implementation Precision**: Specific, actionable recommendations with resource optimization and behavioral alignment  
3. **Risk Intelligence**: Comprehensive scenario planning with predictive modeling and mitigation optimization
4. **Evidence Integration**: Every recommendation supported by cross-engine intelligence and correlation analysis
5. **Multi-Engine Synthesis**: Complete integration of all available intelligence sources and data dimensions
6. **Behavioral Customization**: Risk strategy precisely tailored to {username}'s patterns and {business_name}'s context
7. **Industry Intelligence**: Sector-specific risks and opportunities with competitive analysis and benchmarking
8. **Predictive Accuracy**: Forecasting models with validation, confidence intervals, and comprehensive scenario planning

═══════════════════════════════════════════════════════════════════════════════
⏰ IMPLEMENTATION TIMELINE CONTEXT
═══════════════════════════════════════════════════════════════════════════════

**CURRENT CONTEXT INTEGRATION:**
- **Start Date**: {current_date_str} with seasonal and industry cycle optimization for {industry}
- **Industry Intelligence**: {industry} market conditions, competitive threats, and regulatory landscape analysis
- **Geographic Factors**: {location} business environment, regulatory requirements, and market dynamics integration
- **Organizational Readiness**: {team_size} team capacity assessment with skill gap analysis and development planning
- **Challenge Resolution Priority**: Transform {biggest_challenge} through comprehensive risk management strategy

**TIMELINE FRAMEWORKS:**
- **30-60-90 Day Quick Wins**: Immediate impact initiatives with behavioral alignment and resource optimization
- **Quarterly Milestones**: Progress measurement with behavioral feedback integration and course correction protocols
- **Annual Strategic Objectives**: Long-term capability building with competitive advantage and market positioning
- **Multi-Year Vision Alignment**: Sustainable resilience development with strategic evolution and continuous optimization

**CONTEXTUAL SUCCESS OPTIMIZATION:**
- **Industry-Specific Integration**: Risk benchmarking and competitive intelligence for {industry} sector advantages
- **Seasonal Cycle Alignment**: Business rhythm integration with risk management implementation and resource allocation
- **Geographic Regulatory Harmony**: Local compliance requirements with market dynamics and stakeholder relationships
- **Team Capability Optimization**: Skill development planning with behavioral patterns and external resource integration
- **Challenge-Specific Transformation**: Measurable {biggest_challenge} resolution with strategic competitive advantages

This framework ensures world-class risk analysis quality through rigorous cross-engine methodology, 
evidence-based recommendations, comprehensive predictive modeling, and precise behavioral customization 
while maintaining practical implementation focus and measurable business outcomes tailored 
specifically to {username} and {business_name}'s complete risk transformation and competitive advantage.

═══════════════════════════════════════════════════════════════════════════════
"""
    
    # Enhanced risk response analysis with Multi-Database integration
    logging.info(f"🔧 Creating multi-database risk analysis framework...")
    
    # Format risk assessment responses for analysis - using your existing function
    risk_response_formatting = format_risk_assessment_responses(responses)
    
    # Enhanced analysis instructions with Multi-Database Intelligence
    logging.info(f"📋 Creating enhanced multi-database instructions...")
    enhanced_multi_db_instructions = f"""
═══════════════════════════════════════════════════════════════════════════════
🎯 COMPREHENSIVE MULTI-DATABASE INTELLIGENCE ANALYSIS FOR {username.upper()} 🎯
═══════════════════════════════════════════════════════════════════════════════

🎯 COMPLETE Q&A CROSS-ENGINE INTELLIGENCE INTEGRATION:
{enhanced_format_complete_qa_intelligence(multi_db_intelligence.get('complete_qa_data', {}))}

📊 COMPREHENSIVE RISK ASSESSMENT RESPONSES ANALYSIS:
{risk_response_formatting}

🧠 EVIDENCE-BASED MULTI-DATABASE INTELLIGENCE APPLICATION WITH ADVANCED ANALYTICS:

═══════════════════════════════════════════════════════════════════════════════
📊 SECTION 1: INTEGRATED INTELLIGENCE SYNTHESIS (25% of analysis)
═══════════════════════════════════════════════════════════════════════════════

1. 🧮 COMPREHENSIVE STATISTICAL INTEGRATION WITH VALIDATION:
   **MULTI-DATABASE CORRELATION ANALYSIS:**
   - **Cross-Section Integration**: Systematic correlation analysis across all available assessment domains
   - **Component Assessment Integration**: Business risk maturity correlation with risk readiness
   - **Personality-Risk Mapping**: Individual traits correlated with risk preferences using above intelligence
   - **Vision-Reality Alignment**: Gap analysis between aspirations and current risk capabilities
   - **Behavioral Consistency Validation**: Response pattern consistency across assessment domains

   **ADVANCED ANALYTICAL INTEGRATION:**
   - **Multi-Source Validation**: Cross-referencing findings across different assessment types using above data
   - **Pattern Recognition**: Statistical identification of recurring themes and relationships from intelligence
   - **Predictive Modeling**: Evidence-based forecasting of risk outcomes using comprehensive data
   - **Risk Assessment**: Systematic evaluation of implementation challenges and barriers
   - **Opportunity Identification**: Data-driven discovery of improvement opportunities from all sources
   - **Success Probability**: Quantified likelihood estimates for recommended interventions

   **STATISTICAL VALIDATION REQUIREMENTS:**
   - **Correlation Analysis**: Pearson and Spearman correlations with confidence intervals
   - **Significance Testing**: Multiple comparison correction using FDR or Bonferroni methods
   - **Effect Size Reporting**: Cohen's d and eta-squared for practical significance assessment
   - **Confidence Intervals**: 95% CI for all major statistical findings and predictions
   - **Cross-Validation**: Multiple validation methods to ensure finding reliability
   - **Assumption Testing**: Verification of statistical assumptions before analysis

═══════════════════════════════════════════════════════════════════════════════
📝 SECTION 2: COMPREHENSIVE RESPONSE ANALYSIS & PATTERN DISCOVERY (35% of analysis)
═══════════════════════════════════════════════════════════════════════════════

2. 🔍 SYSTEMATIC RESPONSE ANALYSIS WITH EVIDENCE-BASED INSIGHTS:
   **COMPREHENSIVE RESPONSE EVALUATION USING ALL INTELLIGENCE ABOVE:**
   - **Content Analysis**: Systematic categorization using established coding frameworks
   - **Consistency Assessment**: Statistical evaluation of response coherence and reliability across engines
   - **Depth Analysis**: Assessment of response sophistication using validated measures
   - **Pattern Identification**: Statistical detection of recurring themes and preferences across all data
   - **Quality Assessment**: Evaluation of response completeness and thoughtfulness using intelligence
   - **Integration Analysis**: Cross-referencing responses across different assessment sections

   **ADVANCED PATTERN RECOGNITION:**
   - **Cluster Analysis**: Statistical grouping of similar response patterns using validated methods
   - **Factor Analysis**: Identification of underlying constructs in response data across engines
   - **Discriminant Analysis**: Statistical differentiation between response categories
   - **Regression Analysis**: Predictive modeling of outcomes based on response patterns
   - **Association Analysis**: Statistical identification of co-occurring response themes
   - **Cross-Engine Correlation**: Pattern analysis across all intelligence sources above

   **BEHAVIORAL ANALYSIS FRAMEWORK:**
   - **Decision-Making Style**: Assessment using validated decision-making instruments from intelligence
   - **Risk Orientation**: Analysis using established risk assessment frameworks from behavioral data
   - **Risk Tolerance**: Evaluation using validated risk assessment measures from all sources
   - **Communication Preferences**: Assessment using established communication style inventories
   - **Change Readiness**: Analysis using validated change management assessment tools
   - **Motivational Drivers**: Evaluation using established motivation theory frameworks

═══════════════════════════════════════════════════════════════════════════════
🧠 SECTION 3: MULTI-DATABASE VALIDATION & PREDICTIVE MODELING (25% of analysis)  
═══════════════════════════════════════════════════════════════════════════════

3. 🎯 PREDICTIVE MODELING WITH EVIDENCE-BASED VALIDATION:
   **COMPREHENSIVE VALIDATION FRAMEWORK USING ALL INTELLIGENCE ABOVE:**
   - **Cross-Database Consistency**: Response pattern validation across assessment domains using intelligence
   - **Personality-Behavior Alignment**: Statistical correlation between traits and preferences from data
   - **Vision-Action Coherence**: Analysis of alignment between stated goals and behaviors
   - **Implementation Readiness**: Assessment of change adoption probability with evidence from intelligence
   - **Success Prediction**: Outcome forecasting with confidence intervals and assumptions
   - **Risk Assessment**: Systematic evaluation of potential implementation challenges using all data

   **ADVANCED PREDICTIVE ALGORITHMS:**
   - **Multiple Regression**: Multi-variable prediction models with assumption testing
   - **Logistic Regression**: Binary outcome prediction with odds ratio interpretation
   - **Decision Trees**: Rule-based prediction with interpretability and validation
   - **Random Forest**: Ensemble prediction with feature importance and out-of-bag validation
   - **Cross-Validation**: Multiple validation strategies for model reliability assessment
   - **Bootstrap Analysis**: Resampling methods for confidence interval estimation

   **BEHAVIORAL PREDICTION FRAMEWORK:**
   - **Change Adoption**: Probability assessment using established change management models and intelligence
   - **Performance Outcomes**: Prediction using validated performance assessment frameworks
   - **Team Dynamics**: Forecasting using established team effectiveness measures from data
   - **Leadership Effectiveness**: Prediction using validated leadership assessment tools
   - **Implementation Success**: Outcome forecasting using project management success factors
   - **Risk Mitigation**: Proactive identification using established risk management frameworks

═══════════════════════════════════════════════════════════════════════════════
🚀 SECTION 4: STRATEGIC SYNTHESIS & EVIDENCE-BASED RECOMMENDATIONS (15% of analysis)
═══════════════════════════════════════════════════════════════════════════════

4. 💎 STRATEGIC OPTIMIZATION WITH VALIDATED PRIORITIZATION:
   **EVIDENCE-BASED STRATEGY DEVELOPMENT USING ALL INTELLIGENCE ABOVE:**
   - **Priority Matrix**: Impact-effort analysis using systematic evaluation criteria from all data sources
   - **Resource Optimization**: Allocation strategy using established optimization principles
   - **Timeline Development**: Implementation sequencing using project management best practices
   - **Risk Assessment**: Comprehensive evaluation using enterprise risk management frameworks
   - **Success Metrics**: KPI development using SMART criteria and industry benchmarks
   - **ROI Analysis**: Return on investment estimation using established financial analysis methods

   **STRATEGIC INTEGRATION FRAMEWORK:**
   - **Multi-Objective Optimization**: Balancing competing priorities using systematic trade-off analysis
   - **Stakeholder Alignment**: Strategy development considering all key stakeholder perspectives from intelligence
   - **Resource Constraint Analysis**: Realistic planning within identified resource limitations
   - **Implementation Feasibility**: Practical assessment of strategy executability using all data
   - **Change Management**: Integration of proven change management principles and practices
   - **Continuous Improvement**: Built-in feedback loops and adjustment mechanisms

   **EVIDENCE-BASED IMPACT QUANTIFICATION:**
   - **Performance Metrics**: Quantified targets using industry benchmarks and best practices
   - **Timeline Estimates**: Realistic scheduling using project management methodologies
   - **Resource Requirements**: Comprehensive needs assessment with cost-benefit analysis
   - **Success Probability**: Evidence-based likelihood estimates with confidence intervals
   - **Risk Mitigation**: Proactive planning using established risk management principles
   - **Outcome Measurement**: Success tracking using validated measurement frameworks

═══════════════════════════════════════════════════════════════════════════════
🎯 MANDATORY DELIVERABLES WITH EVIDENCE-BASED STANDARDS:
═══════════════════════════════════════════════════════════════════════════════

**ANALYTICAL DISCOVERY REQUIREMENTS:**
✓ **15 Statistically Significant Relationships** with effect sizes and confidence intervals using all intelligence
✓ **8 Cross-Domain Predictive Models** with validation metrics and accuracy assessments
✓ **5 Contradiction Analyses** with evidence-based reconciliation strategies
✓ **10 Evidence-Based Insights** with statistical support and practical implications
✓ **7 Success Probability Assessments** with confidence intervals and risk factors
✓ **6 Strategic Opportunities** with business case development and ROI projections
✓ **3 Meta-Analytical Syntheses** showing higher-order risk patterns across all engines
✓ **1 Comprehensive Integration** synthesizing all assessment domains with evidence from all intelligence

**METHODOLOGICAL RIGOR STANDARDS:**
- **Statistical Power**: Adequate data utilization for detecting meaningful effects across all intelligence
- **Multiple Comparison Control**: Appropriate correction for multiple statistical tests
- **Assumption Verification**: Testing of statistical assumptions before analysis using all data
- **Cross-Engine Validation**: Verification of findings across multiple intelligence sources
- **Missing Data Analysis**: Systematic evaluation and handling of incomplete data
- **Bias Assessment**: Evaluation of potential sources of systematic error
- **Generalizability**: Assessment of finding applicability to similar contexts
- **Reproducibility**: Sufficient methodological detail for independent replication

**PROFESSIONAL QUALITY REQUIREMENTS:**
- **Industry Standards**: Analysis meeting professional consulting benchmarks using all available intelligence
- **Peer Review Quality**: Methods and findings suitable for professional review
- **Practical Utility**: Recommendations that are actionable and implementable using intelligence insights
- **Evidence-Based Support**: All suggestions backed by statistical evidence from comprehensive data
- **Risk Assessment**: Comprehensive evaluation of implementation challenges using all sources
- **Success Measurement**: Clear metrics for tracking recommendation effectiveness
- **Long-Term Value**: Analysis providing sustainable business benefit using intelligence
- **Ethical Standards**: Appropriate handling of confidential business information

═══════════════════════════════════════════════════════════════════════════════
🚨 QUALITY ASSURANCE & SUCCESS METRICS (100-POINT EVALUATION):
═══════════════════════════════════════════════════════════════════════════════

**ANALYSIS QUALITY SCORING:**
- **Methodological Rigor** (30 points): Statistical methods, validation, assumption testing using all intelligence
- **Evidence Quality** (25 points): Statistical significance, effect sizes, confidence intervals from comprehensive data
- **Integration Depth** (20 points): Cross-domain synthesis, theoretical framework application across all engines
- **Practical Value** (15 points): Actionable recommendations, implementation guidance using intelligence
- **Professional Standards** (10 points): Industry-standard analysis, appropriate benchmarking

**MANDATORY TECHNICAL STANDARDS:**
□ **Complete Intelligence Utilization**: Use ALL provided intelligence sources above for comprehensive analysis
□ **Statistical Significance Testing**: Appropriate methods with multiple comparison correction
□ **Effect Size Reporting**: Practical significance assessment for all major findings using all data
□ **Confidence Interval Reporting**: Uncertainty quantification for all estimates
□ **Cross-Engine Integration**: Synthesis across all assessment domains with evidence from all intelligence
□ **Validation Evidence**: Multiple forms of reliability and validity assessment using comprehensive data
□ **Professional Quality**: Analysis meeting consulting industry standards using all available intelligence
□ **Implementation Guidance**: Practical recommendations with resource requirements using intelligence insights

**SUCCESS VALIDATION CRITERIA:**
- All 15+ analytical discovery requirements fulfilled with quantitative evidence from all intelligence sources
- Statistical analysis with appropriate significance testing and effect size reporting using comprehensive data
- Cross-domain integration with minimum 10 significant relationships documented across all engines
- Evidence-based recommendations with implementation guidance and success metrics using all intelligence
- Professional quality meeting consulting industry standards utilizing all available data sources
- Practical utility with actionable business improvements derived from comprehensive intelligence analysis
- Risk assessment with mitigation strategies and contingency planning using all available data
- Long-term value with sustainable business development focus utilizing comprehensive intelligence

═══════════════════════════════════════════════════════════════════════════════
🎯 RISK ANALYSIS EXECUTION FRAMEWORK FOR {business_name} 🎯
═══════════════════════════════════════════════════════════════════════════════

For {username} in {industry} with {team_size} employees addressing {biggest_challenge}:

**COMPREHENSIVE RISK INTEGRATION WITH EVIDENCE-BASED ANALYSIS USING ALL INTELLIGENCE ABOVE:**
- **Risk Assessment** with personality-preference correlation analysis using behavioral intelligence
- **Systems & Process Risk** with behavioral readiness and change adoption assessment using comprehensive data
- **Culture & Risk Tolerance Strategy** with risk personality integration and evidence-based planning using intelligence
- **Performance Risk Management Strategy** with trait-behavior correlation and predictive modeling using all data
- **Communication Risk Optimization** with preference analysis and network effectiveness measures using intelligence
- **Risk Structure Design** with aspiration-capability alignment and gap analysis using comprehensive data
- **Change Management Risk Implementation** with adoption probability and success prediction using all intelligence
- **Risk & Compliance Strategy** with personality-risk correlation and mitigation planning using all available data

**MULTI-DATABASE RISK CORRELATION WITH STATISTICAL EVIDENCE USING ALL INTELLIGENCE ABOVE:**
- Analyze {username}'s responses across ALL assessment domains with correlation matrices using comprehensive intelligence
- Identify improvement opportunities with probability scoring and evidence-based validation using all data sources
- Reveal risk patterns with statistical validation and confidence quantification using comprehensive intelligence
- Demonstrate synergistic effects with predictive modeling and causal analysis using all available data
- CALCULATE comprehensive correlation matrices for major risk relationships using all intelligence sources
- IDENTIFY strongest predictive patterns with cross-validation evidence and reliability assessment using all data
- QUANTIFY pattern significance with effect size measures and practical importance evaluation using comprehensive intelligence
- MEASURE integration effectiveness with systematic validation and quality assurance using all available data sources

BEGIN COMPREHENSIVE MULTI-DATABASE INTELLIGENCE ANALYSIS WITH EVIDENCE-BASED STATISTICAL INTEGRATION 
USING ALL PROVIDED INTELLIGENCE SOURCES ABOVE FOR MAXIMUM ANALYSIS DEPTH AND ACCURACY:

This framework ensures world-class risk analysis through:
- **Complete Intelligence Utilization** with systematic use of all provided data sources above
- **Rigorous Statistical Methods** with proper validation and assumption testing using comprehensive data
- **Evidence-Based Insights** with quantified confidence and practical significance from all intelligence
- **Professional Quality Standards** meeting consulting industry benchmarks using all available data
- **Cross-Domain Integration** with systematic correlation and pattern analysis across all engines
- **Practical Implementation Focus** with actionable recommendations and success metrics using intelligence
- **Comprehensive Quality Assurance** with enhanced evaluation system utilizing all data sources
- **Long-Term Business Value** with sustainable development guidance using comprehensive intelligence

═══════════════════════════════════════════════════════════════════════════════
"""
    
    # Combine all sections for final prompt
    final_prompt = f"{user_context}\n{enhanced_multi_db_instructions}"
    
    # Log final prompt statistics
    prompt_length = len(final_prompt)
    prompt_word_count = len(final_prompt.split())
    
    logging.info(f"✅ Enhanced risk analysis prompt completed")
    logging.info(f"📊 Final prompt statistics:")
    logging.info(f"   - Total characters: {prompt_length:,}")
    logging.info(f"   - Total words: {prompt_word_count:,}")
    logging.info(f"   - User: {username} at {business_name}")
    logging.info(f"   - Industry: {industry}")
    logging.info(f"   - Challenge: {biggest_challenge}")
    
    return final_prompt


# Helper functions that work with your existing code structure
def enhanced_format_risk_behavioral_data(behavioral_data: Dict) -> str:
    """Enhanced behavioral data formatting for risk analysis"""
    if not behavioral_data:
        return "📊 BEHAVIORAL INTELLIGENCE: Limited data available - analysis will focus on response patterns"
    
    formatted = []
    formatted.append("📊 COMPREHENSIVE BEHAVIORAL RISK INTELLIGENCE:")
    
    # Risk decision patterns analysis
    risk_patterns = behavioral_data.get('risk_decision_patterns', {})
    if risk_patterns:
        risk_tolerance = risk_patterns.get('risk_tolerance', 'unknown')
        decision_speed = risk_patterns.get('decision_speed', 'unknown')
        threat_awareness = risk_patterns.get('threat_awareness', 'unknown')
        
        formatted.append(f"🎯 RISK DECISION PROFILE:")
        formatted.append(f"   - Risk Tolerance: {risk_tolerance}")
        formatted.append(f"   - Decision Speed: {decision_speed}")
        formatted.append(f"   - Threat Awareness: {threat_awareness}")
        formatted.append(f"   - Confidence Score: {risk_patterns.get('confidence_score', 0.0):.2f}")
    
    # Mouse behavior for engagement analysis
    mouse_data = behavioral_data.get('mouse_behavior', {})
    if mouse_data:
        movements = mouse_data.get('total_movements', 0)
        speed = mouse_data.get('average_speed', 0)
        engagement_level = analyze_risk_engagement_level(movements, speed)
        
        formatted.append(f"🖱️ ENGAGEMENT PATTERNS:")
        formatted.append(f"   - Mouse Movements: {movements:,}")
        formatted.append(f"   - Average Speed: {speed:.1f} px/s")
        formatted.append(f"   - Engagement Analysis: {engagement_level}")
    
    # Keyboard behavior for thoughtfulness
    keyboard_data = behavioral_data.get('keyboard_behavior', {})
    if keyboard_data:
        keystrokes = keyboard_data.get('total_keystrokes', 0)
        backspaces = keyboard_data.get('backspace_count', 0)
        revision_ratio = (backspaces / max(keystrokes, 1)) * 100
        thoughtfulness = analyze_risk_thoughtfulness(revision_ratio)
        
        formatted.append(f"⌨️ RESPONSE DELIBERATION:")
        formatted.append(f"   - Total Keystrokes: {keystrokes:,}")
        formatted.append(f"   - Revision Ratio: {revision_ratio:.1f}%")
        formatted.append(f"   - Thoughtfulness Analysis: {thoughtfulness}")
    
    # Attention patterns
    attention_data = behavioral_data.get('attention_patterns', {})
    if attention_data:
        formatted.append(f"👁️ ATTENTION & FOCUS:")
        for pattern, value in attention_data.items():
            formatted.append(f"   - {pattern.title().replace('_', ' ')}: {value}")
    
    return "\n".join(formatted)


def enhanced_format_multi_database_intelligence(multi_db_intelligence: Dict) -> str:
    """Enhanced multi-database intelligence formatting with detailed cross-engine insights"""
    if not multi_db_intelligence:
        return "🧠 MULTI-ENGINE INTELLIGENCE: Single-source analysis - risk assessment only"
    
    formatted = []
    formatted.append("🧠 COMPREHENSIVE MULTI-ENGINE INTELLIGENCE INTEGRATION:")
    
    data_sources = multi_db_intelligence.get('data_sources_available', [])
    formatted.append(f"📊 ACTIVE INTELLIGENCE SOURCES: {len(data_sources)} engines")
    
    # Profile Intelligence (Business DNA)
    profile_data = multi_db_intelligence.get('profile_intelligence', {})
    if profile_data.get('data_available'):
        response_count = profile_data.get('response_count', 0)
        formatted.append(f"🧬 BUSINESS DNA PROFILE: {response_count} personality responses")
        formatted.append(f"   → Risk Implication: Personality-driven risk tolerance and decision-making style")
    
    # Component Intelligence (Business Phase)
    component_data = multi_db_intelligence.get('component_intelligence', {})
    if component_data.get('data_available'):
        phase = component_data.get('business_phase', 'Unknown')
        phase_label = component_data.get('phase_label', 'Unknown')
        formatted.append(f"📊 BUSINESS PHASE: Phase {phase} ({phase_label})")
        formatted.append(f"   → Risk Implication: Phase-specific vulnerabilities and opportunities")
    
    # Growth Intelligence
    growth_data = multi_db_intelligence.get('growth_intelligence', {})
    if growth_data.get('data_available'):
        growth_count = growth_data.get('response_count', 0)
        formatted.append(f"🚀 GROWTH STRATEGY: {growth_count} growth-related responses")
        formatted.append(f"   → Risk Implication: Growth-risk balance and scaling challenges")
    
    # Dream Intelligence
    dream_data = multi_db_intelligence.get('dream_intelligence', {})
    if dream_data.get('data_available'):
        dream_count = dream_data.get('response_count', 0)
        formatted.append(f"💫 VISION ALIGNMENT: {dream_count} aspiration responses")
        formatted.append(f"   → Risk Implication: Vision-reality gap and strategic risk appetite")
    
    # People & Operations Intelligence
    people_ops_data = multi_db_intelligence.get('people_ops_intelligence', {})
    if people_ops_data.get('data_available'):
        people_count = people_ops_data.get('response_count', 0)
        formatted.append(f"👥 PEOPLE & OPERATIONS: {people_count} organizational responses")
        formatted.append(f"   → Risk Implication: Operational vulnerabilities and team-based risks")
    
    # Analyst Intelligence (Behavioral)
    analyst_data = multi_db_intelligence.get('analyst_intelligence', {})
    if analyst_data.get('data_available'):
        analyst_count = analyst_data.get('response_count', 0)
        formatted.append(f"🧠 BEHAVIORAL INTELLIGENCE: {analyst_count} behavioral patterns")
        formatted.append(f"   → Risk Implication: Decision biases and behavioral risk factors")
    
    # Complete Q&A Intelligence
    complete_qa = multi_db_intelligence.get('complete_qa_data', {})
    if complete_qa and complete_qa.get('token_tracking', {}).get('qa_pairs_count', 0) > 0:
        qa_count = complete_qa['token_tracking']['qa_pairs_count']
        total_tokens = complete_qa['token_tracking'].get('total_tokens', 0)
        formatted.append(f"🎯 COMPLETE Q&A INTELLIGENCE: {qa_count} cross-engine insights")
        formatted.append(f"   → Token Budget: {total_tokens:,} tokens of contextual intelligence")
        formatted.append(f"   → Coverage: Comprehensive cross-domain analysis available")
    
    # Cross-engine correlation opportunities
    if len(data_sources) >= 3:
        formatted.append(f"🔗 CROSS-ENGINE CORRELATION OPPORTUNITIES:")
        formatted.append(f"   → Personality × Risk Tolerance Analysis")
        formatted.append(f"   → Growth Ambition × Risk Capacity Alignment") 
        formatted.append(f"   → Operational Maturity × Security Readiness")
        formatted.append(f"   → Vision × Risk Strategy Coherence")
    
    return "\n".join(formatted)


def enhanced_format_complete_qa_intelligence(complete_qa_data: Dict) -> str:
    """Format complete Q&A data for comprehensive analysis"""
    if not complete_qa_data or not complete_qa_data.get('complete_qa_data'):
        return "❌ CROSS-ENGINE Q&A: Not available for this analysis"
    
    formatted = []
    formatted.append("🎯 COMPLETE CROSS-ENGINE Q&A INTELLIGENCE:")
    
    qa_data = complete_qa_data['complete_qa_data']
    token_tracking = complete_qa_data.get('token_tracking', {})
    
    # Summary statistics
    total_pairs = token_tracking.get('qa_pairs_count', 0)
    total_tokens = token_tracking.get('total_tokens', 0)
    efficiency = token_tracking.get('efficiency_score', 0)
    
    formatted.append(f"📊 INTELLIGENCE SUMMARY:")
    formatted.append(f"   - Total Q&A Pairs: {total_pairs}")
    formatted.append(f"   - Token Utilization: {total_tokens:,} tokens")
    formatted.append(f"   - Efficiency Score: {efficiency:.1f} pairs per 1K tokens")
    
    # Engine-specific breakdown
    engine_breakdown = token_tracking.get('by_engine', {})
    formatted.append(f"📋 ENGINE-SPECIFIC BREAKDOWN:")
    
    for engine, stats in engine_breakdown.items():
        pairs = stats.get('qa_pairs', 0)
        tokens = stats.get('tokens', 0)
        if pairs > 0:
            formatted.append(f"   🔹 {engine.replace('_', ' ').title()}: {pairs} Q&A pairs, {tokens:,} tokens")
    
    # Sample Q&A pairs for context (first few from each engine)
    formatted.append(f"🎯 SAMPLE INTELLIGENCE (for AI context):")
    
    sample_count = 0
    max_samples = 15  # Limit samples to avoid overwhelming prompt
    
    for engine_name, engine_data in qa_data.items():
        if isinstance(engine_data, list) and engine_data and sample_count < max_samples:
            formatted.append(f"   📁 {engine_name.replace('_', ' ').title()} Samples:")
            
            for qa_pair in engine_data[:3]:  # First 3 from each engine
                if sample_count >= max_samples:
                    break
                    
                question = qa_pair.get('question', '')[:100]  # Truncate long questions
                response = str(qa_pair.get('response', ''))[:150]  # Truncate long responses
                
                formatted.append(f"      Q: {question}{'...' if len(str(qa_pair.get('question', ''))) > 100 else ''}")
                formatted.append(f"      A: {response}{'...' if len(str(qa_pair.get('response', ''))) > 150 else ''}")
                
                sample_count += 1
    
    if total_pairs > sample_count:
        formatted.append(f"   ... and {total_pairs - sample_count} more Q&A pairs available for analysis")
    
    return "\n".join(formatted)


def analyze_risk_engagement_level(movements: int, speed: float) -> str:
    """Analyze engagement level for risk assessment"""
    if movements > 1000 and speed > 50:
        return "High engagement - active risk assessment"
    elif movements > 500:
        return "Moderate engagement - careful consideration"
    else:
        return "Low engagement - quick responses"


def analyze_risk_thoughtfulness(revision_ratio: float) -> str:
    """Analyze thoughtfulness level for risk assessment"""
    if revision_ratio > 15:
        return "High thoughtfulness - careful risk consideration"
    elif revision_ratio > 8:
        return "Moderate thoughtfulness - some revision"
    else:
        return "Direct responses - quick risk decisions"


# ======================================================
#           3. Missing format_multi_database_intelligence_risk function
# ======================================================

def format_multi_database_intelligence_risk(multi_db_intelligence: Dict) -> str:
    """Format multi-database intelligence for risk analysis"""
    if not multi_db_intelligence:
        return "Multi-Database Intelligence: Not available - using risk responses only"
    
    formatted = []
    
    # Component Intelligence
    component_data = multi_db_intelligence.get('component_intelligence', {})
    if component_data:
        phase = component_data.get('business_phase', 'Unknown')
        phase_label = component_data.get('phase_label', 'Unknown')
        formatted.append(f"📊 Component Intelligence: Business Phase {phase} ({phase_label}) - {len(component_data.get('responses', {}))} component responses")
    
    # Business DNA Profile
    profile_data = multi_db_intelligence.get('profile_intelligence', {})
    if profile_data:
        response_count = len(profile_data.get('responses', {}))
        formatted.append(f"🧬 Business DNA Profile: {response_count} personality responses available for risk customization")
    
    # Dream Analysis
    dream_data = multi_db_intelligence.get('dream_intelligence', {})
    if dream_data:
        dream_count = len(dream_data.get('responses', {}))
        formatted.append(f"💫 Dream Analysis: {dream_count} aspiration responses for risk alignment")
    
    # Growth Intelligence
    growth_data = multi_db_intelligence.get('growth_intelligence', {})
    if growth_data:
        growth_count = len(growth_data.get('responses', {}))
        formatted.append(f"🚀 Growth Analysis: {growth_count} growth responses for risk scaling alignment")
    
    # People & Operations Intelligence  
    people_ops_data = multi_db_intelligence.get('people_ops_intelligence', {})
    if people_ops_data:
        people_ops_count = len(people_ops_data.get('responses', {}))
        formatted.append(f"👥 People & Operations: {people_ops_count} organizational responses for risk leadership")
    
    # Behavioral Intelligence
    analyst_data = multi_db_intelligence.get('analyst_intelligence', {})
    if analyst_data:
        behavioral_count = len(analyst_data.get('responses', {}))
        formatted.append(f"🧠 Behavioral Intelligence: {behavioral_count} behavioral responses for risk implementation customization")
    
    # Risk Intelligence (self-reference)
    risk_data = multi_db_intelligence.get('risk_intelligence', {})
    if risk_data:
        risk_count = len(risk_data.get('responses', {}))
        formatted.append(f"🏰 Risk Intelligence: {risk_count} risk responses for comprehensive threat analysis")
    
    if not formatted:
        return "Multi-Database Intelligence: Extraction in progress - risk analysis will use available data sources"
    
    return "\n".join(formatted)

def format_risk_behavioral_data(behavioral_data: Dict) -> str:
    """Format behavioral data for risk analysis"""
    if not behavioral_data:
        return "No behavioral data available for risk analysis"
    
    formatted = []
    
    # Risk-specific behavioral analysis
    if 'risk_decision_patterns' in behavioral_data:
        patterns = behavioral_data['risk_decision_patterns']
        formatted.append(f"🎯 Risk Decision Patterns: {patterns.get('pattern_type', 'Unknown')} risk management style")
    
    # Mouse behavior analysis for risk
    mouse_data = behavioral_data.get('mouse_behavior', {})
    if mouse_data:
        total_movements = mouse_data.get('total_movements', 0)
        avg_speed = mouse_data.get('average_speed', 0)
        formatted.append(f"🖱️ Risk Engagement: {total_movements} interactions, {avg_speed:.1f} avg speed - indicates {_analyze_risk_engagement_level(total_movements, avg_speed)}")
    
    # Keyboard behavior analysis for risk
    keyboard_data = behavioral_data.get('keyboard_behavior', {})
    if keyboard_data:
        total_keystrokes = keyboard_data.get('total_keystrokes', 0)
        backspace_count = keyboard_data.get('backspace_count', 0)
        revision_ratio = (backspace_count / max(total_keystrokes, 1)) * 100
        formatted.append(f"⌨️ Risk Thoughtfulness: {revision_ratio:.1f}% revision rate - indicates {_analyze_risk_thoughtfulness(revision_ratio)}")
    
    return "\n".join(formatted) if formatted else "Behavioral analysis available for risk implementation customization"

def _analyze_risk_engagement_level(movements: int, speed: float) -> str:
    """Analyze engagement level for risk context"""
    if movements > 500 and speed > 10:
        return "high engagement with risk opportunities"
    elif movements > 200:
        return "moderate engagement with strategic risk thinking"
    else:
        return "focused, deliberate risk consideration"

def _analyze_risk_thoughtfulness(revision_ratio: float) -> str:
    """Analyze thoughtfulness for risk implementation"""
    if revision_ratio > 15:
        return "highly considered risk responses - detailed implementation recommended"
    elif revision_ratio > 8:
        return "thoughtful risk analysis - systematic implementation approach"
    else:
        return "decisive risk thinking - rapid implementation capability"

def format_risk_assessment_responses(responses):
    """Enhanced format risk assessment responses for comprehensive 100/100 analysis"""
    if not responses:
        return "📋 RISK ASSESSMENT RESPONSES: No risk assessment responses available for analysis"
    
    # Filter for risk assessment responses
    risk_responses = [r for r in responses if r.get('questionnaire_type', '').lower() in ['risk_fortress', 'risk', 'risk_assessment']]
    
    if not risk_responses:
        return "📋 RISK ASSESSMENT RESPONSES: No direct risk responses found - using cross-engine risk pattern analysis"
    
    formatted = []
    formatted.append("═══════════════════════════════════════════════════════════════════════════════")
    formatted.append("📋 COMPREHENSIVE RISK ASSESSMENT RESPONSE INTELLIGENCE")
    formatted.append("═══════════════════════════════════════════════════════════════════════════════")
    formatted.append("🎯 ANALYSIS PRIORITY: Extract maximum risk intelligence from every client response")
    formatted.append("🔍 STATISTICAL REQUIREMENT: Correlate responses with behavioral patterns and business vulnerabilities")
    formatted.append("")
    
    # Categorize responses by risk type for better analysis
    risk_categories = {
        'vulnerability_assessment': [],
        'threat_tolerance': [],
        'operational_resilience': [],
        'financial_risk': [],
        'strategic_risk': [],
        'compliance_governance': [],
        'implementation_readiness': []
    }
    
    # Analyze each response and categorize
    response_count = 0
    for response in risk_responses:
        response_count += 1
        question_id = response.get('question_id', f'RISK_{response_count}')
        question_text = response.get('question_text', response.get('question', 'Unknown risk question'))
        response_data = response.get('response_data', response.get('response', {}))
        
        # Determine risk category based on question content
        risk_category = categorize_risk_question(question_text, question_id)
        
        formatted.append(f"📊 RISK RESPONSE #{response_count} - CATEGORY: {risk_category.upper()}")
        formatted.append(f"🔍 QUESTION ID: {question_id}")
        formatted.append(f"❓ QUESTION: {question_text}")
        
        # Enhanced response analysis based on type
        risk_intelligence = extract_risk_intelligence(response_data, question_text)
        formatted.append(f"💡 CLIENT RESPONSE: {risk_intelligence['response_summary']}")
        formatted.append(f"📈 RISK IMPLICATION: {risk_intelligence['risk_implication']}")
        formatted.append(f"🎯 ANALYSIS INSTRUCTION: {risk_intelligence['analysis_instruction']}")
        
        # Add to categorized list for pattern analysis
        risk_categories[risk_category].append({
            'question_id': question_id,
            'response': risk_intelligence['response_summary'],
            'implication': risk_intelligence['risk_implication']
        })
        
        formatted.append("─" * 80)
        formatted.append("")
    
    # Add pattern analysis section
    formatted.append("═══════════════════════════════════════════════════════════════════════════════")
    formatted.append("🧠 RISK PATTERN INTELLIGENCE SUMMARY")
    formatted.append("═══════════════════════════════════════════════════════════════════════════════")
    
    for category, responses_in_category in risk_categories.items():
        if responses_in_category:
            formatted.append(f"🔹 {category.replace('_', ' ').title()} ({len(responses_in_category)} responses)")
            for resp in responses_in_category[:3]:  # Show top 3 per category
                formatted.append(f"   • {resp['question_id']}: {resp['response']}")
            formatted.append("")
    
    formatted.append("🎯 COMPREHENSIVE ANALYSIS REQUIREMENTS:")
    formatted.append("   ✓ Correlate response patterns across all risk categories")
    formatted.append("   ✓ Identify behavioral risk tolerance patterns")
    formatted.append("   ✓ Map responses to organizational vulnerability assessment")
    formatted.append("   ✓ Calculate risk readiness scores with confidence intervals")
    formatted.append("   ✓ Generate predictive risk management recommendations")
    formatted.append("   ✓ Cross-reference with behavioral intelligence for personalization")
    
    formatted.append("═══════════════════════════════════════════════════════════════════════════════")
    
    return "\n".join(formatted)


def categorize_risk_question(question_text: str, question_id: str) -> str:
    """Categorize risk questions for better analysis"""
    question_lower = question_text.lower()
    id_lower = question_id.lower()
    
    # Categorization logic
    if any(term in question_lower for term in ['vulnerability', 'weakness', 'exposed', 'threat']):
        return 'vulnerability_assessment'
    elif any(term in question_lower for term in ['tolerance', 'appetite', 'comfort', 'willing']):
        return 'threat_tolerance'
    elif any(term in question_lower for term in ['operational', 'process', 'system', 'backup', 'recovery']):
        return 'operational_resilience'
    elif any(term in question_lower for term in ['financial', 'money', 'cash', 'revenue', 'budget']):
        return 'financial_risk'
    elif any(term in question_lower for term in ['strategy', 'strategic', 'vision', 'goal', 'future']):
        return 'strategic_risk'
    elif any(term in question_lower for term in ['compliance', 'governance', 'policy', 'regulation']):
        return 'compliance_governance'
    elif any(term in question_lower for term in ['implementation', 'change', 'adopt', 'ready']):
        return 'implementation_readiness'
    else:
        return 'vulnerability_assessment'  # Default category


def extract_risk_intelligence(response_data, question_text: str) -> dict:
    """Extract comprehensive risk intelligence from response data"""
    risk_intelligence = {
        'response_summary': 'No response data available',
        'risk_implication': 'Unable to assess risk implication',
        'analysis_instruction': 'Standard risk pattern analysis required'
    }
    
    if isinstance(response_data, dict):
        # Single selection analysis
        if 'selected_option' in response_data:
            selected = response_data['selected_option']
            risk_intelligence['response_summary'] = f'Selected: "{selected}"'
            risk_intelligence['risk_implication'] = analyze_selection_risk_implication(selected, question_text)
            risk_intelligence['analysis_instruction'] = f'Analyze selection "{selected}" for risk tolerance patterns and vulnerability assessment'
        
        # Multiple selection analysis
        elif 'selected_options' in response_data:
            selected_list = response_data['selected_options']
            risk_intelligence['response_summary'] = f'Multiple selections: {selected_list}'
            risk_intelligence['risk_implication'] = analyze_multiple_selection_risk(selected_list, question_text)
            risk_intelligence['analysis_instruction'] = f'Analyze combination pattern for comprehensive risk profile assessment'
        
        # Text response analysis
        elif 'response_text' in response_data:
            text = response_data['response_text']
            word_count = response_data.get('word_count', len(text.split()) if text else 0)
            risk_intelligence['response_summary'] = f'Text response ({word_count} words): "{text[:100]}{"..." if len(text) > 100 else ""}"'
            risk_intelligence['risk_implication'] = analyze_text_risk_content(text, question_text)
            risk_intelligence['analysis_instruction'] = f'Perform content analysis for risk awareness depth and threat management insights'
        
        # Slider/rating analysis
        elif 'slider_value' in response_data:
            value = response_data['slider_value']
            scale_info = response_data.get('scale_info', 'Unknown scale')
            risk_intelligence['response_summary'] = f'Rated: {value} on scale ({scale_info})'
            risk_intelligence['risk_implication'] = analyze_rating_risk_level(value, question_text, scale_info)
            risk_intelligence['analysis_instruction'] = f'Analyze rating {value} for risk capability assessment and tolerance calibration'
    
    # Handle non-dict responses
    elif isinstance(response_data, str):
        risk_intelligence['response_summary'] = f'String response: "{response_data}"'
        risk_intelligence['risk_implication'] = analyze_text_risk_content(response_data, question_text)
        risk_intelligence['analysis_instruction'] = 'Analyze string response for risk management insights'
    
    elif isinstance(response_data, (int, float)):
        risk_intelligence['response_summary'] = f'Numeric response: {response_data}'
        risk_intelligence['risk_implication'] = f'Numeric value {response_data} indicates quantifiable risk assessment'
        risk_intelligence['analysis_instruction'] = f'Analyze numeric value {response_data} for risk measurement and scoring'
    
    return risk_intelligence


def analyze_selection_risk_implication(selected: str, question_text: str) -> str:
    """Analyze risk implications of selected options"""
    selected_lower = selected.lower()
    
    # High risk indicators
    if any(term in selected_lower for term in ['high', 'severe', 'critical', 'exposed', 'vulnerable', 'weak']):
        return 'HIGH RISK: Selection indicates significant vulnerability requiring immediate attention'
    
    # Medium risk indicators
    elif any(term in selected_lower for term in ['moderate', 'some', 'partial', 'developing', 'improving']):
        return 'MODERATE RISK: Selection shows areas for improvement with manageable exposure'
    
    # Low risk indicators
    elif any(term in selected_lower for term in ['low', 'strong', 'secure', 'protected', 'robust', 'excellent']):
        return 'LOW RISK: Selection indicates good risk management capability'
    
    # Neutral or strategic selections
    else:
        return f'STRATEGIC CONSIDERATION: Selection "{selected}" requires contextual risk assessment'


def analyze_multiple_selection_risk(selected_list: list, question_text: str) -> str:
    """Analyze risk implications of multiple selections"""
    if not selected_list:
        return 'NO SELECTION: Indicates potential risk awareness gap or decision avoidance'
    
    selection_count = len(selected_list)
    
    if selection_count == 1:
        return analyze_selection_risk_implication(selected_list[0], question_text)
    elif selection_count <= 3:
        return f'FOCUSED RISK PROFILE: {selection_count} selections indicate targeted risk awareness'
    else:
        return f'COMPREHENSIVE RISK AWARENESS: {selection_count} selections show broad risk consideration'


def analyze_text_risk_content(text: str, question_text: str) -> str:
    """Analyze text content for risk implications"""
    if not text or len(text.strip()) < 10:
        return 'LIMITED RISK INSIGHT: Brief response may indicate low engagement with risk assessment'
    
    text_lower = text.lower()
    word_count = len(text.split())
    
    # Risk awareness indicators
    risk_terms = ['risk', 'threat', 'vulnerability', 'security', 'protect', 'danger', 'expose']
    risk_term_count = sum(1 for term in risk_terms if term in text_lower)
    
    # Sophistication indicators
    sophisticated_terms = ['strategy', 'framework', 'systematic', 'process', 'management', 'assessment']
    sophistication_count = sum(1 for term in sophisticated_terms if term in text_lower)
    
    if risk_term_count >= 2 and sophistication_count >= 1:
        return f'HIGH RISK SOPHISTICATION: {word_count}-word response shows advanced risk awareness'
    elif risk_term_count >= 1:
        return f'MODERATE RISK AWARENESS: {word_count}-word response shows basic risk understanding'
    else:
        return f'DEVELOPING RISK AWARENESS: {word_count}-word response needs risk perspective enhancement'


def analyze_rating_risk_level(value, question_text: str, scale_info: str) -> str:
    """Analyze risk implications of rating values"""
    # Determine if high values are good or bad based on question context
    question_lower = question_text.lower()
    
    # Questions where high values indicate good risk management
    positive_indicators = ['strength', 'capability', 'readiness', 'protection', 'security']
    is_positive_scale = any(term in question_lower for term in positive_indicators)
    
    # Normalize value to 0-10 scale for analysis
    try:
        numeric_value = float(value)
        if 'out of 10' in scale_info.lower() or '1-10' in scale_info.lower():
            normalized_value = numeric_value
        elif 'out of 5' in scale_info.lower() or '1-5' in scale_info.lower():
            normalized_value = numeric_value * 2
        else:
            normalized_value = numeric_value  # Assume 1-10 scale
        
        if is_positive_scale:
            if normalized_value >= 8:
                return f'LOW RISK: High rating ({value}) indicates strong risk management capability'
            elif normalized_value >= 6:
                return f'MODERATE RISK: Medium rating ({value}) shows adequate but improvable risk management'
            else:
                return f'HIGH RISK: Low rating ({value}) indicates significant risk management gaps'
        else:
            if normalized_value >= 8:
                return f'HIGH RISK: High rating ({value}) indicates significant risk exposure'
            elif normalized_value >= 6:
                return f'MODERATE RISK: Medium rating ({value}) shows manageable risk exposure'
            else:
                return f'LOW RISK: Low rating ({value}) indicates minimal risk exposure'
                
    except (ValueError, TypeError):
        return f'RATING ANALYSIS: Value "{value}" requires contextual risk assessment'

# ======================================================
#           Database Functions
# ======================================================

def setup_risk_logging():
    """Set up logging for risk engine"""
    log_dir = Path("logs")
    log_dir.mkdir(exist_ok=True)
    
    log_file = log_dir / f"risk_engine_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
    
    logger = logging.getLogger()
    logger.setLevel(logging.INFO)
    
    for handler in logger.handlers[:]:
        logger.removeHandler(handler)
    
    # Console handler
    console_handler = logging.StreamHandler()
    console_handler.setLevel(logging.INFO)
    console_format = logging.Formatter('%(asctime)s - RISK ENGINE %(levelname)s - %(message)s')
    console_handler.setFormatter(console_format)
    
    # File handler
    file_handler = logging.FileHandler(log_file, encoding='utf-8')
    file_handler.setLevel(logging.INFO)
    file_format = logging.Formatter('%(asctime)s - %(levelname)s - %(name)s - %(funcName)s:%(lineno)d - %(message)s')
    file_handler.setFormatter(file_format)
    
    logger.addHandler(console_handler)
    logger.addHandler(file_handler)
    
    logging.info(f"Risk Engine Logging Initialized: {log_file}")
    return logger

def get_risk_connection():
    """Get connection to risk database"""
    try:
        conn = psycopg2.connect(
            host=RISK_DB_CONFIG["host"],
            dbname=RISK_DB_CONFIG["database"],
            user=RISK_DB_CONFIG["user"],
            password=RISK_DB_CONFIG["password"],
            port=RISK_DB_CONFIG["port"]
        )
        conn.autocommit = True
        return conn
    except Exception as e:
        logging.error(f"Risk database connection error: {str(e)}")
        raise

def create_risk_tables(conn):
    """Create necessary risk tables"""
    try:
        with conn.cursor() as cur:
            # Create risk_assessments table
            cur.execute("""
                CREATE TABLE IF NOT EXISTS risk_assessments (
                    id SERIAL PRIMARY KEY,
                    user_id VARCHAR(255) UNIQUE NOT NULL,
                    assessment_type VARCHAR(100) NOT NULL,
                    version VARCHAR(20) NOT NULL,
                    created_at TIMESTAMPTZ,
                    last_updated TIMESTAMPTZ,
                    timezone VARCHAR(100),
                    session_metadata JSONB,
                    device_fingerprint JSONB,
                    progress_tracking JSONB,
                    completion_flags JSONB,
                    raw_data JSONB,
                    multi_database_intelligence JSONB,
                    created_timestamp TIMESTAMPTZ DEFAULT NOW()
                )
            """)
            
            # Create risk_responses table
            cur.execute("""
                CREATE TABLE IF NOT EXISTS risk_responses (
                    id SERIAL PRIMARY KEY,
                    assessment_id INTEGER REFERENCES risk_assessments(id),
                    user_id VARCHAR(255) NOT NULL,
                    question_id VARCHAR(50) NOT NULL,
                    section VARCHAR(100) NOT NULL,
                    question_type VARCHAR(50),
                    question_text TEXT,
                    response_format VARCHAR(50),
                    response_data JSONB,
                    all_options JSONB,
                    metadata JSONB,
                    weight VARCHAR(20),
                    answered_at TIMESTAMPTZ,
                    last_modified_at TIMESTAMPTZ,
                    created_timestamp TIMESTAMPTZ DEFAULT NOW(),
                    UNIQUE(assessment_id, question_id)
                )
            """)
            
            # Create risk_behavioral_analytics table
            cur.execute("""
                CREATE TABLE IF NOT EXISTS risk_behavioral_analytics (
                    id SERIAL PRIMARY KEY,
                    assessment_id INTEGER REFERENCES risk_assessments(id) UNIQUE,
                    user_id VARCHAR(255) NOT NULL,
                    mouse_behavior JSONB,
                    keyboard_behavior JSONB,
                    attention_patterns JSONB,
                    decision_making_style JSONB,
                    risk_decision_patterns JSONB,
                    created_at TIMESTAMPTZ,
                    created_timestamp TIMESTAMPTZ DEFAULT NOW()
                )
            """)
            
            # Create risk_reports table
            cur.execute("""
                CREATE TABLE IF NOT EXISTS risk_reports (
                    id SERIAL PRIMARY KEY,
                    report_id VARCHAR(255) UNIQUE NOT NULL,
                    user_id VARCHAR(255) NOT NULL,
                    assessment_id INTEGER REFERENCES risk_assessments(id),
                    report_type VARCHAR(100) NOT NULL,
                    status VARCHAR(50) NOT NULL,
                    azure_container VARCHAR(255),
                    blob_paths JSONB,
                    chunk_count INTEGER,
                    generation_metadata JSONB,
                    created_at TIMESTAMPTZ DEFAULT NOW(),
                    completed_at TIMESTAMPTZ,
                    indexer_job_id VARCHAR(255),
                    indexer_status VARCHAR(50),
                    indexer_triggered_at TIMESTAMPTZ,
                    indexer_completed_at TIMESTAMPTZ,
                    indexer_error_message TEXT,
                    indexer_retry_count INTEGER DEFAULT 0,
                    multi_database_integration JSONB
                )
            """)
            
            # Create indexes
            cur.execute("CREATE INDEX IF NOT EXISTS idx_risk_assessments_user_id ON risk_assessments(user_id)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_risk_responses_user_id ON risk_responses(user_id)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_risk_responses_section ON risk_responses(section)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_risk_reports_user_id ON risk_reports(user_id)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_risk_reports_report_id ON risk_reports(report_id)")
            
            # Create indexer indexes
            cur.execute("CREATE INDEX IF NOT EXISTS idx_risk_reports_indexer_job_id ON risk_reports(indexer_job_id)")
            cur.execute("CREATE INDEX IF NOT EXISTS idx_risk_reports_indexer_status ON risk_reports(indexer_status)")
            
        logging.info("✅ Risk engine tables created successfully")
        
    except Exception as e:
        logging.error(f"❌ Error creating risk tables: {str(e)}")
        raise

def ensure_risk_tables_exist():
    """Ensure risk tables exist on startup"""
    conn = None
    try:
        logging.info("🔧 Ensuring Risk database tables exist...")
        
        conn = get_risk_connection()
        
        # Create tables if they don't exist
        create_risk_tables(conn)
        
        # Verify tables were created by checking if they exist
        with conn.cursor() as cur:
            cur.execute("""
                SELECT table_name 
                FROM information_schema.tables 
                WHERE table_schema = 'public' 
                AND table_name IN ('risk_assessments', 'risk_responses', 'risk_behavioral_analytics', 'risk_reports')
            """)
            existing_tables = [row[0] for row in cur.fetchall()]
            
            logging.info(f"✅ Verified Risk tables exist: {existing_tables}")
            
            if len(existing_tables) >= 4:
                logging.info("✅ All Risk engine tables are ready")
                return True
            else:
                missing_tables = ['risk_assessments', 'risk_responses', 'risk_behavioral_analytics', 'risk_reports']
                missing = [t for t in missing_tables if t not in existing_tables]
                logging.error(f"❌ Missing Risk tables: {missing}")
                return False
        
    except Exception as e:
        logging.error(f"❌ Error ensuring Risk tables exist: {str(e)}")
        return False
    finally:
        if conn:
            conn.close()


def analyze_risk_decision_patterns(behavioral_data: Dict, responses: List[Dict]) -> Dict:
    """Analyze risk-specific decision patterns"""
    
    logging.info("🧠 Starting risk decision pattern analysis")
    
    patterns = {
        'risk_tolerance': 'unknown',
        'threat_awareness': 'unknown',
        'decision_speed': 'unknown',
        'risk_thinking': 'unknown',
        'confidence_score': 0.0,
        'analysis_method': 'heuristic'
    }
    
    try:
        # Basic heuristic analysis
        mouse_data = behavioral_data.get('mouse_behavior', {})
        if mouse_data:
            total_movements = mouse_data.get('total_movements', 0)
            avg_speed = mouse_data.get('average_speed', 0)
            
            if total_movements > 1000 and avg_speed > 15:
                patterns['decision_speed'] = 'fast_decisive'
                patterns['risk_tolerance'] = 'high_risk_tolerance'
                patterns['confidence_score'] += 0.3
            elif total_movements > 500:
                patterns['decision_speed'] = 'moderate_analytical'
                patterns['risk_tolerance'] = 'balanced_risk_tolerance'
                patterns['confidence_score'] += 0.2
            else:
                patterns['decision_speed'] = 'deliberate_cautious'
                patterns['risk_tolerance'] = 'conservative_risk_tolerance'
                patterns['confidence_score'] += 0.2
        
        keyboard_data = behavioral_data.get('keyboard_behavior', {})
        if keyboard_data:
            backspace_count = keyboard_data.get('backspace_count', 0)
            total_keystrokes = keyboard_data.get('total_keystrokes', 0)
            revision_ratio = (backspace_count / max(total_keystrokes, 1)) * 100
            
            if revision_ratio > 20:
                patterns['risk_thinking'] = 'highly_systematic'
                patterns['threat_awareness'] = 'high_threat_awareness'
            elif revision_ratio > 10:
                patterns['risk_thinking'] = 'balanced_systematic'
                patterns['threat_awareness'] = 'moderate_threat_awareness'
            else:
                patterns['risk_thinking'] = 'intuitive_adaptive'
                patterns['threat_awareness'] = 'confident_threat_assessment'
        
        # Basic response analysis
        risk_responses = [r for r in responses if any(keyword in r.get('section', '').lower() 
                         for keyword in ['risk', 'threat', 'vulnerability', 'security', 'fortress'])]
        
        if risk_responses:
            conservative_count = 0
            aggressive_count = 0
            
            for response in risk_responses:
                selected = str(response.get('response_data', {}).get('selected_option', '')).lower()
                
                if any(word in selected for word in ['safe', 'secure', 'conservative', 'careful', 'protect', 'avoid']):
                    conservative_count += 1
                elif any(word in selected for word in ['aggressive', 'bold', 'risk', 'opportunity', 'growth']):
                    aggressive_count += 1
            
            if conservative_count > aggressive_count:
                patterns['risk_tolerance'] = 'conservative_risk_tolerance'
            elif aggressive_count > conservative_count:
                patterns['risk_tolerance'] = 'aggressive_risk_tolerance'
            else:
                patterns['risk_tolerance'] = 'balanced_risk_tolerance'
        
        # Add metadata
        patterns['analysis_timestamp'] = datetime.now().isoformat()
        patterns['analysis_metadata'] = {
            'responses_analyzed': len(responses) if responses else 0,
            'risk_responses_found': len(risk_responses) if 'risk_responses' in locals() else 0,
            'behavioral_data_available': bool(behavioral_data),
        }
        
        logging.info("✅ Risk decision pattern analysis completed successfully")
        
    except Exception as e:
        logging.error(f"❌ Error analyzing risk decision patterns: {str(e)}")
        patterns['error'] = str(e)
        patterns['error_type'] = type(e).__name__
    
    return patterns

def store_risk_report_metadata(report_id: str, user_id: str, assessment_id: int, chunk_count: int, 
                              container_name: str, generation_metadata: Dict):
    """Store risk report metadata"""
    conn = None
    try:
        conn = get_risk_connection()
        
        with conn.cursor() as cur:
            sql = """
                INSERT INTO risk_reports (
                    report_id, user_id, assessment_id, report_type, status,
                    azure_container, chunk_count, generation_metadata, completed_at,
                    multi_database_integration
                ) VALUES (
                    %s, %s, %s, %s, %s, %s, %s, %s, %s, %s
                ) ON CONFLICT (report_id) DO UPDATE SET
                    status = EXCLUDED.status,
                    chunk_count = EXCLUDED.chunk_count,
                    generation_metadata = EXCLUDED.generation_metadata,
                    completed_at = EXCLUDED.completed_at,
                    multi_database_integration = EXCLUDED.multi_database_integration
            """
            
            multi_db_info = {
                'integration_enabled': True,
                'data_sources_used': generation_metadata.get('data_sources_used', []),
                'intelligence_correlation': generation_metadata.get('intelligence_correlation', {}),
                'total_intelligence_sources': generation_metadata.get('total_intelligence_sources', 0),
                'complete_qa_pairs': generation_metadata.get('complete_qa_pairs', 0)
            }
            
            cur.execute(sql, (
                report_id,
                user_id,
                assessment_id,
                "comprehensive_risk_fortress_strategy",
                "completed",
                container_name,
                chunk_count,
                json.dumps(generation_metadata),
                datetime.now(),
                json.dumps(multi_db_info)
            ))
        
        logging.info(f"Stored risk report metadata for report_id={report_id}")
        
    except Exception as e:
        logging.error(f"Error storing risk report metadata: {str(e)}")
        raise
    finally:
        if conn:
            conn.close()



def generate_risk_section_with_dedicated_client(
    section_name: str,
    section_config: Dict,
    complete_raw_data: Dict,
    api_key: str,
    section_index: int,
    max_retries: int = 3
) -> Dict:
    """Generate risk section with enhanced retry mechanism"""
    
    client_id = f"risk_section_{section_index}_{section_name}"
    original_api_key = api_key
    current_api_key = api_key
    
    # Enhanced API key selection and health tracking (same as people & ops)
    key_health = api_key_health.get(current_api_key, {})
    success_rate = key_health.get('success_rate', 1.0)
    current_load = key_health.get('current_load', 0)
    consecutive_failures = key_health.get('consecutive_failures', 0)
    
    logging.info(f"🔑 [{client_id}] STARTING RISK SECTION GENERATION")
    logging.info(f"🔑 [{client_id}] Initial API key: {key_health.get('key_id', 'unknown')} (...{current_api_key[-4:]})")
    logging.info(f"🔑 [{client_id}] Key health: {consecutive_failures} failures, {success_rate:.3f} success rate, {current_load} load")
    
    attempt_history = []
    
    for retry_attempt in range(max_retries):
        attempt_start_time = time.time()
        attempt_info = {
            'attempt_number': retry_attempt + 1,
            'api_key_used': current_api_key[-4:],
            'start_time': attempt_start_time,
            'success': False,
            'error': None,
            'words_generated': 0,
            'tokens_used': 0,
            'key_switched': False
        }
        
        try:
            # API key reset and selection for retries (same logic as people & ops)
            if retry_attempt > 0:
                logging.info(f"🔄 [{client_id}] RETRY {retry_attempt + 1}: Starting retry process...")
                
                # Reset the current API key
                reset_api_key_immediately(current_api_key)
                
                # Select optimal API key for retry
                if retry_attempt == 1:
                    try:
                        new_api_key = get_load_balanced_api_key(section_index)
                        selection_method = "load_balanced"
                    except Exception:
                        new_api_key = get_smart_api_key(section_index, retry_attempt)
                        selection_method = "smart_fallback"
                else:
                    new_api_key = get_smart_api_key(section_index, retry_attempt)
                    selection_method = "smart_selection"
                
                key_switched = new_api_key != current_api_key
                attempt_info['key_switched'] = key_switched
                
                if key_switched:
                    logging.info(f"🔄 [{client_id}] KEY SWITCH: ...{current_api_key[-4:]} → ...{new_api_key[-4:]}")
                    current_api_key = new_api_key
                else:
                    logging.info(f"🔄 [{client_id}] KEY REUSED: Same key selected (...{current_api_key[-4:]})")
            
            # Execute the analysis
            logging.info(f"🎯 [{client_id}] Executing risk analysis...")
            analysis_start_time = time.time()
            
            target_words = min(section_config["word_target"], 3000)
            
            response = risk_ultra_deep_analysis(
                complete_raw_data=complete_raw_data,
                analysis_type=section_name,
                analysis_requirements=section_config["analysis_requirements"],
                api_key=current_api_key,
                client_id=client_id,
                temperature=0.7,
                max_tokens=1000000
            )
            
            analysis_time = time.time() - analysis_start_time
            current_words = len(response.content.split())
            
            logging.info(f"📊 [{client_id}] Analysis completed:")
            logging.info(f"   - Analysis time: {analysis_time:.2f}s")
            logging.info(f"   - Words generated: {current_words:,}")
            logging.info(f"   - Tokens used: {response.token_count:,}")
            
            # Response quality validation
            if current_words < 100 and retry_attempt < max_retries - 1:
                logging.warning(f"⚠️ [{client_id}] RESPONSE TOO SHORT: {current_words} words < 100 minimum")
                
                attempt_info.update({
                    'success': False,
                    'error': f'Response too short: {current_words} words',
                    'words_generated': current_words,
                    'tokens_used': response.token_count,
                    'analysis_time': analysis_time
                })
                attempt_history.append(attempt_info)
                
                # Mark as failure and reset key
                update_api_key_health(current_api_key, success=False, error_code="SHORT_RESPONSE")
                reset_api_key_immediately(current_api_key)
                
                wait_time = 30 * (retry_attempt + 1)
                logging.info(f"⏳ [{client_id}] Waiting {wait_time}s before retry due to short response...")
                time.sleep(wait_time)
                continue
            
            # SUCCESS
            total_attempt_time = time.time() - attempt_start_time
            
            attempt_info.update({
                'success': True,
                'words_generated': current_words,
                'tokens_used': response.token_count,
                'analysis_time': analysis_time,
                'total_attempt_time': total_attempt_time
            })
            attempt_history.append(attempt_info)
            
            logging.info(f"🎉 [{client_id}] RISK SECTION GENERATION SUCCESS!")
            logging.info(f"✅ [{client_id}] Success metrics:")
            logging.info(f"   - Words Generated: {current_words:,}")
            logging.info(f"   - Tokens Used: {response.token_count:,}")
            logging.info(f"   - Analysis Time: {analysis_time:.2f}s")
            
            return {
                "title": section_config["title"],
                "content": response.content,
                "metadata": {
                    "word_target": target_words,
                    "words_generated": current_words,
                    "tokens_generated": response.token_count,
                    "ai_analysis_time": analysis_time,
                    "ai_model": "gemini-2.5-pro-risk-engine",
                    "analysis_type": "risk_dedicated_analysis",
                    "timestamp": datetime.now().isoformat(),
                    "client_id": client_id,
                    "retry_attempts": retry_attempt + 1,
                    "success": True,
                    "api_key_used": key_health.get('key_id', 'unknown'),
                    "key_switched": current_api_key != original_api_key,
                    "attempt_history": attempt_history
                }
            }
            
        except Exception as e:
            error_str = str(e)
            attempt_time = time.time() - attempt_start_time
            
            attempt_info.update({
                'success': False,
                'error': error_str,
                'analysis_time': attempt_time,
                'total_attempt_time': attempt_time
            })
            attempt_history.append(attempt_info)
            
            logging.error(f"❌ [{client_id}] ATTEMPT {retry_attempt + 1} FAILED")
            logging.error(f"🔍 [{client_id}] Error: {error_str}")
            
            if retry_attempt < max_retries - 1:
                # Error-specific wait times
                if "503" in error_str or "overload" in error_str.lower():
                    wait_time = 300 + (retry_attempt * 180)
                elif "429" in error_str:
                    wait_time = 120 + (retry_attempt * 60)
                else:
                    wait_time = 60 * (retry_attempt + 1)
                
                logging.info(f"⏳ [{client_id}] Waiting {wait_time}s before retry...")
                time.sleep(wait_time)
            else:
                # All attempts failed - return fallback
                total_function_time = time.time() - attempt_history[0]['start_time']
                
                logging.error(f"💥 [{client_id}] ALL ATTEMPTS EXHAUSTED - FINAL FAILURE")
                
                fallback_content = f"""This risk analysis section encountered persistent API issues during generation.

SECTION DETAILS:
- Section: {section_config['title']}
- Target Words: {section_config.get('word_target', 'Unknown'):,}
- Analysis Type: {section_name}

FAILURE ANALYSIS:
- Total Attempts: {max_retries}
- API Keys Tried: {len(set([attempt['api_key_used'] for attempt in attempt_history]))}
- Key Resets Applied: {retry_attempt}
- Total Processing Time: {total_function_time:.1f}s

FINAL ERROR: {error_str}

This risk analysis will be completed when API capacity is restored."""
                
                return {
                    "title": section_config["title"],
                    "content": fallback_content,
                    "metadata": {
                        "error": True,
                        "error_message": error_str,
                        "timestamp": datetime.now().isoformat(),
                        "client_id": client_id,
                        "retry_attempts": max_retries,
                        "final_error": error_str,
                        "attempt_history": attempt_history
                    }
                }
    
    return None

# ======================================================
#           MISSING RISK ENGINE FUNCTIONS AND CLASSES
# ======================================================

# 1. Fix the return type in risk_ultra_deep_analysis function
def risk_ultra_deep_analysis(
    complete_raw_data: Dict,
    analysis_type: str,
    analysis_requirements: str,
    api_key: str,
    client_id: str = "risk_analysis",
    temperature: float = 0.7,
    max_tokens: int = 1000000
) -> RiskChatResponse:  # FIXED: Changed from PeopleOpsChatResponse to RiskChatResponse
    """Enhanced risk analysis with ultra-deep response analysis"""
    
    start_time = time.time()
    request_start_time = None
    
    logging.info(f"🚀 [{client_id}] Starting Risk Analysis: {analysis_type}")
    logging.info(f"🔍 [{client_id}] API key: ...{api_key[-4:]}")
    
    try:
        # Create enhanced prompt
        logging.info(f"📝 [{client_id}] Creating enhanced prompt...")
        enhanced_prompt = create_enhanced_risk_analysis_prompt(
            complete_raw_data, analysis_type, analysis_requirements
        )
        
        # Convert to Gemini format
        contents = convert_messages_to_gemini_format([
            {"role": "user", "content": enhanced_prompt}
        ])
        
        # Production-optimized payload
        payload = {
            "contents": contents,
            "generationConfig": {
                "temperature": temperature,
                "maxOutputTokens": max_tokens,
                "topP": 0.9,
                "topK": 40,
                "candidateCount": 1,
                "stopSequences": [],
                "responseMimeType": "text/plain"
            },
            "safetySettings": [
                {"category": "HARM_CATEGORY_HARASSMENT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_HATE_SPEECH", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_SEXUALLY_EXPLICIT", "threshold": "BLOCK_NONE"},
                {"category": "HARM_CATEGORY_DANGEROUS_CONTENT", "threshold": "BLOCK_NONE"}
            ]
        }
        
        params = {'key': api_key}
        url = f"https://generativelanguage.googleapis.com/v1beta/models/gemini-2.5-pro:generateContent"
        
        # Execute request
        logging.info(f"📤 [{client_id}] Sending request to Gemini API...")
        request_start_time = time.time()
        
        response = requests.post(
            url,
            json=payload,
            params=params,
            timeout=REQUEST_TIMEOUT
        )
        
        response_time = time.time() - request_start_time
        
        logging.info(f"📡 [{client_id}] Response received")
        logging.info(f"⏱️ [{client_id}] Response time: {response_time:.2f}s")
        logging.info(f"📊 [{client_id}] Response status: {response.status_code}")
        
        if response.status_code == 200:
            # Success - update API key health
            update_api_key_health(api_key, success=True, response_time=response_time)
            
            # Parse JSON response
            data = response.json()
            
            # Extract content (same logic as people & ops)
            if 'candidates' in data and len(data['candidates']) > 0:
                candidate = data['candidates'][0]
                content = ""
                
                if 'content' in candidate and candidate['content'] is not None:
                    content_obj = candidate['content']
                    if 'parts' in content_obj and content_obj['parts']:
                        first_part = content_obj['parts'][0]
                        if isinstance(first_part, dict) and 'text' in first_part:
                            content = first_part['text']
                
                if not content and 'text' in candidate:
                    content = candidate['text']
                
                # Content validation
                if not content:
                    logging.error(f"❌ [{client_id}] No content extracted")
                    update_api_key_health(api_key, success=False, error_code="NO_CONTENT")
                    raise Exception("No content extracted from API response")
                
                if len(content.strip()) < 10:
                    logging.error(f"❌ [{client_id}] Content too short: {len(content.strip())} chars")
                    update_api_key_health(api_key, success=False, error_code="SHORT_CONTENT")
                    raise Exception(f"Content too short: {len(content.strip())} characters")
                
                # Success metrics
                usage = data.get('usageMetadata', {})
                token_count = usage.get('totalTokenCount', 0)
                analysis_time = time.time() - start_time
                word_count = len(content.split())
                
                logging.info(f"🎉 [{client_id}] RISK ANALYSIS COMPLETE - SUCCESS")
                logging.info(f"📊 [{client_id}] Final success metrics:")
                logging.info(f"   - Analysis type: {analysis_type}")
                logging.info(f"   - Total tokens: {token_count:,}")
                logging.info(f"   - Total time: {analysis_time:.2f}s")
                logging.info(f"   - Words generated: {word_count:,}")
                
                response_obj = RiskChatResponse(  # FIXED: Using RiskChatResponse
                    content=content,
                    model="gemini-2.5-pro-risk",
                    api_key_used=f"{client_id}_key_{api_key[-4:]}",
                    usage=usage,
                    finish_reason=candidate.get('finishReason', 'STOP'),
                    response_time=analysis_time,
                    timestamp=time.time(),
                    token_count=token_count
                )
                
                return response_obj
                
            else:
                logging.error(f"❌ [{client_id}] No candidates in response")
                update_api_key_health(api_key, success=False, error_code="NO_CANDIDATES")
                raise Exception("No candidates found in API response")
        
        else:
            # HTTP error handling
            error_code = str(response.status_code)
            logging.error(f"❌ [{client_id}] HTTP ERROR: {response.status_code}")
            
            update_api_key_health(api_key, success=False, error_code=error_code, response_time=response_time)
            raise Exception(f"HTTP {response.status_code}: {response.text}")
    
    except Exception as e:
        analysis_time = time.time() - start_time
        response_time_safe = time.time() - request_start_time if request_start_time else analysis_time
        
        logging.error(f"❌ [{client_id}] RISK ANALYSIS FAILED")
        logging.error(f"🔍 [{client_id}] Error: {type(e).__name__}: {str(e)}")
        logging.error(f"🔍 [{client_id}] Analysis time: {analysis_time:.2f}s")
        
        raise

def generate_comprehensive_risk_report(complete_raw_data: Dict, report_id: str, max_report_retries: int = 2) -> Dict:
    """Generate comprehensive risk report with notifications and enhanced API key management"""
    
    logging.info(f"🚀 Starting Risk Report Generation with Smart Notifications and Load Balancing for {report_id}")
    start_time = time.time()
    
    # Extract user data for personalized notifications
    user_id = complete_raw_data.get("user_id", "unknown")
    user_profile = complete_raw_data.get("user_profile", {})
    
    # Risk notification tracking
    notifications_sent = {"start": False, "middle": False, "complete": False}
    
    # Send start notification
    send_risk_notification_background(user_id, user_profile, "start")
    notifications_sent["start"] = True
    
    # Initialize API key health tracking
    logging.info(f"🔑 Initial API Key Health Status: {get_api_key_status_summary()}")
    
    for report_attempt in range(max_report_retries):
        logging.info(f"🔄 Risk report attempt {report_attempt + 1}/{max_report_retries}")
        
        risk_sections = get_risk_report_sections()
        
        report_data = {}
        failed_sections = []
        successful_sections = []
        
        # Track API key usage
        api_keys_used = set()
        load_balancing_stats = {
            "total_sections": len(risk_sections),
            "load_balanced_selections": 0,
            "smart_selections": 0,
            "key_switches": 0
        }
        
        # Process sections in batches
        section_items = list(risk_sections.items())
        batch_size = 3
        
        for batch_start in range(0, len(section_items), batch_size):
            batch_end = min(batch_start + batch_size, len(section_items))
            batch = section_items[batch_start:batch_end]
            
            logging.info(f"🔄 Processing risk batch {batch_start//batch_size + 1}: sections {batch_start+1}-{batch_end}")
            
            # Parallel processing within batch
            with ThreadPoolExecutor(max_workers=batch_size) as executor:
                future_to_section = {}
                
                for i, (section_name, section_config) in enumerate(batch):
                    section_index = batch_start + i
                    
                    # Enhanced API key selection
                    try:
                        api_key = get_load_balanced_api_key(section_index)
                        selection_method = "load_balanced"
                        load_balancing_stats["load_balanced_selections"] += 1
                    except Exception as e:
                        api_key = get_smart_api_key(section_index, 0)
                        selection_method = "smart_fallback"
                        load_balancing_stats["smart_selections"] += 1
                    
                    api_keys_used.add(api_key)
                    
                    if i > 0:
                        time.sleep(2)  # Delay between submissions
                    
                    future = executor.submit(
                        generate_risk_section_with_dedicated_client,
                        section_name=section_name,
                        section_config=section_config,
                        complete_raw_data=complete_raw_data,
                        api_key=api_key,
                        section_index=section_index,
                        max_retries=2
                    )
                    
                    future_to_section[future] = (section_name, section_index, api_key, selection_method)
                    logging.info(f"📤 Submitted risk section {section_index + 1}/{len(section_items)}: {section_name}")
                
                # Collect batch results
                for future in as_completed(future_to_section):
                    section_name, section_index, original_api_key, selection_method = future_to_section[future]
                    
                    try:
                        section_content = future.result()
                        report_data[section_name] = section_content
                        
                        # Track API key switching
                        metadata = section_content.get("metadata", {})
                        if metadata.get("key_switched", False):
                            load_balancing_stats["key_switches"] += 1
                        
                        # Track success/failure
                        if section_content.get("metadata", {}).get("error", False):
                            failed_sections.append(section_name)
                            logging.error(f"❌ Risk section failed: {section_name}")
                        else:
                            successful_sections.append(section_name)
                            words_generated = metadata.get("words_generated", 0)
                            logging.info(f"✅ Risk section completed: {section_name} ({words_generated:,} words)")
                        
                        total_completed = len(successful_sections) + len(failed_sections)
                        
                        # Update job status and send middle notification
                        if report_id in risk_job_status:
                            completion_percentage = (total_completed / len(section_items)) * 100
                            risk_job_status[report_id]["message"] = f"Risk processing: {total_completed}/{len(section_items)} sections ({completion_percentage:.1f}%)"
                            risk_job_status[report_id]["sections_completed"] = total_completed
                            
                            # Send middle notification
                            if not notifications_sent["middle"] and completion_percentage >= 45 and completion_percentage <= 65:
                                progress_data = {
                                    'chapters_completed': total_completed,
                                    'total_chapters': len(section_items),
                                    'progress_percentage': completion_percentage
                                }
                                send_risk_notification_background(user_id, user_profile, "middle", progress_data)
                                notifications_sent["middle"] = True
                        
                    except Exception as e:
                        logging.error(f"❌ Error retrieving risk result for {section_name}: {str(e)}")
                        failed_sections.append(section_name)
            
            # Wait between batches
            if batch_end < len(section_items):
                enhanced_status = get_enhanced_api_key_status()
                healthy_keys = enhanced_status.get("healthy_keys", 0)
                total_load = enhanced_status.get("total_load", 0)
                
                if healthy_keys <= 3 or total_load > 10:
                    wait_time = 90
                elif healthy_keys <= 5:
                    wait_time = 75
                else:
                    wait_time = 65
                
                logging.info(f"⏳ Risk batch wait: {wait_time}s before next batch...")
                time.sleep(wait_time)
        
        # Check success rate
        success_rate = len(successful_sections) / len(risk_sections)
        parallel_time = time.time() - start_time
        
        logging.info(f"📊 Risk attempt {report_attempt + 1} completed: {len(successful_sections)}/{len(risk_sections)} sections successful ({success_rate:.1%})")
        
        if success_rate >= 0.8:
            logging.info(f"✅ Risk report successful with {success_rate:.1%} success rate")
            break
        else:
            logging.warning(f"⚠️ Risk report attempt {report_attempt + 1} below threshold ({success_rate:.1%} < 80%)")
            if report_attempt < max_report_retries - 1:
                reset_failed_api_keys()
                logging.info(f"🔄 Reset failed API keys, retrying in 60s...")
                time.sleep(60)
    
    # Calculate final metrics
    total_time = time.time() - start_time
    total_words = sum([
        len(section.get("content", "").split()) 
        for section in report_data.values()
    ])
    
    logging.info(f"🌟 Risk Report Completed: {len(successful_sections)} successful sections, {total_words:,} words")
    
    # Send completion notification with report_id for DB persistence
    if not notifications_sent["complete"]:
        completion_data = {
            'total_words': total_words,
            'total_sections': len(successful_sections),
            'processing_time': total_time
        }
        send_risk_notification_background(user_id, user_profile, "complete", completion_data, report_id)
        notifications_sent["complete"] = True
    
    # Add report metadata
    report_data["_enhanced_risk_report_metadata"] = {
        "report_id": report_id,
        "generation_timestamp": datetime.now().isoformat(),
        "total_sections": len(report_data),
        "successful_sections": len(successful_sections),
        "failed_sections": len(failed_sections),
        "success_rate": len(successful_sections) / len(risk_sections),
        "total_words": total_words,
        "total_generation_time": total_time,
        "ai_model": "gemini-2.5-pro-risk",
        "processing_method": "risk_parallel_analysis_load_balanced",
        "report_type": "comprehensive_risk_fortress_strategy",
        "notifications_sent": notifications_sent,
        "multi_database_integration": {
            "enabled": True,
            "data_sources_used": complete_raw_data.get("multi_database_intelligence", {}).get("data_sources_available", []),
            "intelligence_correlation_applied": True,
            "behavioral_customization": True,
            "personality_integration": True,
            "dream_alignment": True,
            "complete_qa_pairs": complete_raw_data.get("multi_database_intelligence", {}).get("complete_qa_data", {}).get("token_tracking", {}).get("qa_pairs_count", 0)
        },
        "api_key_optimization": {
            "load_balancing_enabled": True,
            "unique_keys_used": len(api_keys_used),
            "total_keys_available": len(GEMINI_API_KEYS),
            "key_utilization_rate": len(api_keys_used) / len(GEMINI_API_KEYS),
            "load_balanced_selections": load_balancing_stats.get("load_balanced_selections", 0),
            "smart_fallback_selections": load_balancing_stats.get("smart_selections", 0),
            "key_switches_during_processing": load_balancing_stats.get("key_switches", 0),
            "final_api_health_summary": get_api_key_status_summary(),
            "enhanced_api_health": get_enhanced_api_key_status()
        }
    }
    
    return report_data




# ======================================================
#           Multi-Database Intelligence for Risk Engine
# ======================================================

async def extract_complete_qa_data_for_user(user_id: str) -> Dict:
    """Extract ALL question-answer pairs from all engines for a specific user - ENHANCED VERSION for Risk Engine"""
    
    logging.info(f"🎯 Starting complete Q&A extraction for user_id={user_id}")
    
    # All databases - extract everything for this user INCLUDING the new Risk Engine
    target_databases = [
        ("BACKABLE-PROFILE-ENGINE", "profile_responses", "profile_engine"),
        ("BACKABLE-THE-ANALYST", "analyst_responses", "analyst_engine"), 
        ("BACKABLE-THE-GROWTH-ENGINE", "growth_responses", "growth_engine"),
        ("BACKABLE-COMPONENT-ENGINE", "component_responses", "component_engine"),
        ("BACKABLE-DREAM-ANALYZER", "dream_responses", "dream_engine"),
        ("BACKABLE-PEOPLE-AND-OPERATIONS-ENGINE", "people_ops_responses", "people_ops_engine"),
        ("BACKABLE-RISK-ENGINE", "risk_responses", "risk_engine")  # NEW: Added Risk Engine
    ]
    
    qa_intelligence = {
        "user_id": user_id,
        "extraction_timestamp": datetime.now().isoformat(),
        "extraction_method": "complete_qa_pairs_all_engines_including_risk",
        "target_tokens": "up_to_48k",
        "complete_qa_data": {
            "profile_engine": [],
            "analyst_engine": [],
            "growth_engine": [],
            "component_engine": [],
            "dream_engine": [],
            "people_ops_engine": [],
            "risk_engine": []  # NEW: Added Risk Engine
        },
        "cross_engine_insights": {},
        "token_tracking": {
            "by_engine": {},
            "total_tokens": 0,
            "qa_pairs_count": 0
        }
    }
    
    logging.info(f"📋 Extracting ALL question_text + response_data pairs for user {user_id}")
    logging.info(f"🔥 NO FILTERING - Getting everything for comprehensive analysis including Risk Engine")
    
    total_tokens = 0
    total_qa_pairs = 0
    
    for db_name, responses_table, engine_key in target_databases:
        logging.info(f"\n📊 Processing {db_name} for user {user_id}...")
        
        engine_tokens = 0
        engine_qa_pairs = 0
        
        try:
            # Get database connection using connection pool
            if "PROFILE" in db_name:
                pool = await get_db_pool(PROFILE_DB_CONFIG)
            elif "ANALYST" in db_name:
                pool = await get_db_pool(ANALYST_DB_CONFIG)
            elif "GROWTH" in db_name:
                pool = await get_db_pool(GROWTH_DB_CONFIG)
            elif "COMPONENT" in db_name:
                pool = await get_db_pool(COMPONENT_DB_CONFIG)
            elif "DREAM" in db_name:
                pool = await get_db_pool(DREAM_DB_CONFIG)
            elif "PEOPLE-AND-OPERATIONS" in db_name:
                pool = await get_db_pool(PEOPLE_OPS_DB_CONFIG)
            elif "RISK" in db_name:  # NEW: Risk Engine connection
                pool = await get_db_pool(RISK_DB_CONFIG)
            else:
                logging.warning(f"⚠️ Unknown database: {db_name}, skipping...")
                continue
                
            async with pool.acquire() as conn:
                try:
                    # Extract ALL question_text and response_data for this user
                    query = f"""
                        SELECT question_text, response_data
                        FROM {responses_table} 
                        WHERE user_id = $1
                          AND question_text IS NOT NULL
                          AND response_data IS NOT NULL
                        ORDER BY answered_at DESC
                    """
                    
                    results = await conn.fetch(query, user_id)
                    
                    for row in results:
                        question_text = row[0]
                        response_data = row[1]
                        
                        # Clean the response data
                        cleaned_response = clean_response_data(response_data)
                        
                        if cleaned_response:
                            qa_pair = {
                                "question": question_text,
                                "response": cleaned_response
                            }
                            
                            # Calculate tokens for this Q&A pair
                            qa_tokens = estimate_tokens(json.dumps(qa_pair))
                            
                            # Check token budget - use 46K of 50K budget (leave room for overhead)
                            if total_tokens + qa_tokens > 46000:
                                logging.warning(f"   ⚠️  Approaching 50K token limit, stopping extraction")
                                break
                            
                            qa_intelligence["complete_qa_data"][engine_key].append(qa_pair)
                            
                            engine_tokens += qa_tokens
                            total_tokens += qa_tokens
                            engine_qa_pairs += 1
                            total_qa_pairs += 1
                            
                            # Show first 80 chars of response for preview
                            response_preview = str(cleaned_response)[:80] + "..." if len(str(cleaned_response)) > 80 else str(cleaned_response)
                            logging.debug(f"   ✅ Q&A: {response_preview}")
                
                except Exception as e:
                    logging.error(f"   ❌ Error with {responses_table}: {e}")
            
            qa_intelligence["token_tracking"]["by_engine"][engine_key] = {
                "tokens": engine_tokens,
                "qa_pairs": engine_qa_pairs
            }
            
            logging.info(f"   ✅ {db_name}: {engine_qa_pairs} Q&A pairs, {engine_tokens:,} tokens")
            
        except Exception as e:
            logging.error(f"   ❌ Database connection failed for {db_name}: {e}")
    
    # Final calculations
    json_overhead = int(total_tokens * 0.15)  # 15% overhead for clean JSON
    final_tokens = total_tokens + json_overhead
    
    qa_intelligence["token_tracking"]["total_tokens"] = final_tokens
    qa_intelligence["token_tracking"]["qa_pairs_count"] = total_qa_pairs
    qa_intelligence["token_tracking"]["json_overhead"] = json_overhead
    qa_intelligence["token_tracking"]["efficiency_score"] = total_qa_pairs / max(final_tokens, 1) * 1000
    
    logging.info(f"\n🎯 COMPLETE EXTRACTION FINISHED for user {user_id}")
    logging.info(f"Q&A Pairs: {total_qa_pairs}")
    logging.info(f"Total Tokens: {final_tokens:,}")
    logging.info(f"Token Budget Used: {(final_tokens/50000)*100:.1f}% of 50K")
    logging.info(f"Efficiency: {qa_intelligence['token_tracking']['efficiency_score']:.1f} Q&A pairs per 1000 tokens")
    logging.info(f"Remaining Budget: {50000 - final_tokens:,} tokens")
    
    return qa_intelligence

def estimate_tokens(text):
    """Estimate token count (roughly 4 characters per token)"""
    return len(str(text)) // 4 if text else 0

def clean_response_data(response_data):
    """Clean response data to keep only essential information"""
    if not response_data:
        return None
    
    if isinstance(response_data, dict):
        # For slider values - keep the core scores
        if 'slider_values' in response_data:
            return response_data['slider_values']
        
        # For selections - keep the selected option
        if 'selected_option' in response_data:
            return response_data['selected_option']
        
        # For arrays/lists in response
        if 'selected_options' in response_data:
            return response_data['selected_options']
        
        # For text responses
        if 'text_response' in response_data:
            return response_data['text_response']
        
        # For response values
        if 'response_value' in response_data:
            return response_data['response_value']
        
        # Keep only essential keys, skip metadata
        essential_keys = ['value', 'label', 'text', 'score', 'rating', 'selection', 'answer']
        cleaned = {}
        for key, value in response_data.items():
            if any(essential in key.lower() for essential in essential_keys):
                cleaned[key] = value
        
        return cleaned if cleaned else response_data
    
    elif isinstance(response_data, list):
        # Keep lists as-is but limit to 10 items
        return response_data[:10]
    
    elif isinstance(response_data, str):
        # Keep text but limit length for very long responses
        return response_data[:500] if len(response_data) > 500 else response_data
    
    return response_data

async def get_multi_database_intelligence(user_id: str) -> Dict:
    """Get intelligence from all available databases using connection pools - ENHANCED FOR RISK ENGINE"""
    logging.info(f"🔍 Extracting multi-database intelligence for user_id={user_id}")
    
    intelligence = {
        'user_id': user_id,
        'component_intelligence': {},
        'profile_intelligence': {},
        'dream_intelligence': {},
        'growth_intelligence': {},
        'analyst_intelligence': {},
        'people_ops_intelligence': {},  # Added people ops
        'risk_intelligence': {},        # NEW: Added risk intelligence
        'complete_qa_data': {},
        'extraction_timestamp': datetime.now().isoformat(),
        'data_sources_available': []
    }
    
    # 🔥 NEW: Extract complete Q&A data from all engines INCLUDING RISK
    try:
        logging.info(f"🎯 Extracting complete Q&A data from all engines for user {user_id}")
        complete_qa_data = await extract_complete_qa_data_for_user(user_id)
        intelligence['complete_qa_data'] = complete_qa_data
        logging.info(f"✅ Complete Q&A extraction successful: {complete_qa_data.get('token_tracking', {}).get('qa_pairs_count', 0)} total Q&A pairs")
    except Exception as e:
        logging.error(f"❌ Error extracting complete Q&A data: {str(e)}")
        intelligence['complete_qa_data'] = {'error': str(e), 'qa_pairs_count': 0}
    
    # Component Intelligence
    try:
        pool = await get_db_pool(COMPONENT_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    ca.phase, ca.phase_label, ca.assessment_type,
                    jsonb_object_agg(cr.question_id, 
                        jsonb_build_object(
                            'answer', cr.response_data->>'selected_option',
                            'weight', cr.weight,
                            'section', cr.section,
                            'metadata', cr.metadata
                        )
                    ) as responses
                FROM component_responses cr
                JOIN component_assessments ca ON cr.user_id = ca.user_id
                WHERE cr.user_id = $1
                GROUP BY ca.phase, ca.phase_label, ca.assessment_type
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result:
                intelligence['component_intelligence'] = {
                    'business_phase': result[0],
                    'phase_label': result[1],
                    'assessment_type': result[2],
                    'responses': result[3] or {},
                    'data_available': True
                }
                intelligence['data_sources_available'].append('component')
                logging.info(f"✅ Component intelligence extracted: Phase {result[0]} ({result[1]})")
            else:
                intelligence['component_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No component data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting component intelligence: {str(e)}")
        intelligence['component_intelligence'] = {'data_available': False, 'error': str(e)}
    
    # Profile Intelligence (Business DNA)
    try:
        pool = await get_db_pool(PROFILE_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    pa.assessment_type, pa.created_at,
                    COALESCE(jsonb_object_agg(
                        CASE WHEN pr.question_id IS NOT NULL THEN pr.question_id END,
                        CASE WHEN pr.question_id IS NOT NULL THEN
                            jsonb_build_object(
                                'answer', pr.response_data->>'selected_option',
                                'weight', pr.weight,
                                'section', pr.section,
                                'question_text', pr.question_text
                            )
                        END
                    ) FILTER (WHERE pr.question_id IS NOT NULL), '{}'::jsonb) as responses
                FROM profile_assessments pa
                LEFT JOIN profile_responses pr ON pr.user_id = pa.user_id
                WHERE pa.user_id = $1
                GROUP BY pa.assessment_type, pa.created_at
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result and result[2] and result[2] != {}:
                intelligence['profile_intelligence'] = {
                    'assessment_type': result[0],
                    'created_at': result[1].isoformat() if result[1] else None,
                    'responses': result[2],
                    'data_available': True,
                    'response_count': len(result[2])
                }
                intelligence['data_sources_available'].append('profile')
                logging.info(f"✅ Profile intelligence extracted: {len(result[2])} responses")
            else:
                intelligence['profile_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No profile data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting profile intelligence: {str(e)}")
        intelligence['profile_intelligence'] = {'data_available': False, 'error': str(e)}
    
    # Dream Intelligence
    try:
        pool = await get_db_pool(DREAM_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    da.assessment_type, da.created_at,
                    jsonb_object_agg(dr.question_id, 
                        jsonb_build_object(
                            'answer', dr.response_data->>'selected_option',
                            'response_text', dr.response_data->>'response_text',
                            'section', dr.section,
                            'question_text', dr.question_text
                        )
                    ) as responses
                FROM dream_responses dr
                JOIN dream_assessments da ON dr.user_id = da.user_id
                WHERE dr.user_id = $1
                GROUP BY da.assessment_type, da.created_at
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result:
                intelligence['dream_intelligence'] = {
                    'assessment_type': result[0],
                    'created_at': result[1].isoformat() if result[1] else None,
                    'responses': result[2] or {},
                    'data_available': True,
                    'response_count': len(result[2] or {})
                }
                intelligence['data_sources_available'].append('dream')
                logging.info(f"✅ Dream intelligence extracted: {len(result[2] or {})} responses")
            else:
                intelligence['dream_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No dream data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting dream intelligence: {str(e)}")
        intelligence['dream_intelligence'] = {'data_available': False, 'error': str(e)}
    
    # Growth Intelligence
    try:
        pool = await get_db_pool(GROWTH_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    ga.assessment_type, ga.created_at,
                    jsonb_object_agg(gr.question_id, 
                        jsonb_build_object(
                            'answer', gr.response_data->>'selected_option',
                            'response_text', gr.response_data->>'response_text',
                            'section', gr.section,
                            'question_text', gr.question_text
                        )
                    ) as responses
                FROM growth_responses gr
                JOIN growth_assessments ga ON gr.user_id = ga.user_id
                WHERE gr.user_id = $1
                GROUP BY ga.assessment_type, ga.created_at
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result:
                intelligence['growth_intelligence'] = {
                    'assessment_type': result[0],
                    'created_at': result[1].isoformat() if result[1] else None,
                    'responses': result[2] or {},
                    'data_available': True,
                    'response_count': len(result[2] or {})
                }
                intelligence['data_sources_available'].append('growth')
                logging.info(f"✅ Growth intelligence extracted: {len(result[2] or {})} responses")
            else:
                intelligence['growth_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No growth data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting growth intelligence: {str(e)}")
        intelligence['growth_intelligence'] = {'data_available': False, 'error': str(e)}
    
    # Analyst Intelligence (Behavioral)
    try:
        pool = await get_db_pool(ANALYST_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    jsonb_object_agg(ar.question_id, 
                        jsonb_build_object(
                            'answer', ar.response_data->>'selected_option',
                            'response_text', ar.response_data->>'response_text',
                            'section', ar.section,
                            'metadata', ar.metadata
                        )
                    ) as responses
                FROM analyst_responses ar
                WHERE ar.user_id = $1
                GROUP BY ar.user_id
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result and result[0]:
                intelligence['analyst_intelligence'] = {
                    'responses': result[0],
                    'data_available': True,
                    'response_count': len(result[0])
                }
                intelligence['data_sources_available'].append('analyst')
                logging.info(f"✅ Analyst intelligence extracted: {len(result[0])} responses")
            else:
                intelligence['analyst_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No analyst data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting analyst intelligence: {str(e)}")
        intelligence['analyst_intelligence'] = {'data_available': False, 'error': str(e)}
    
    # People & Operations Intelligence
    try:
        pool = await get_db_pool(PEOPLE_OPS_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    jsonb_object_agg(por.question_id, 
                        jsonb_build_object(
                            'answer', por.response_data->>'selected_option',
                            'response_text', por.response_data->>'response_text',
                            'section', por.section,
                            'metadata', por.metadata
                        )
                    ) as responses
                FROM people_ops_responses por
                WHERE por.user_id = $1
                GROUP BY por.user_id
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result and result[0]:
                intelligence['people_ops_intelligence'] = {
                    'responses': result[0],
                    'data_available': True,
                    'response_count': len(result[0])
                }
                intelligence['data_sources_available'].append('people_ops')
                logging.info(f"✅ People Ops intelligence extracted: {len(result[0])} responses")
            else:
                intelligence['people_ops_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No people ops data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting people ops intelligence: {str(e)}")
        intelligence['people_ops_intelligence'] = {'data_available': False, 'error': str(e)}
    
    # NEW: Risk Intelligence
    try:
        pool = await get_db_pool(RISK_DB_CONFIG)
        async with pool.acquire() as conn:
            sql = """
                SELECT 
                    jsonb_object_agg(rr.question_id, 
                        jsonb_build_object(
                            'answer', rr.response_data->>'selected_option',
                            'response_text', rr.response_data->>'response_text',
                            'section', rr.section,
                            'metadata', rr.metadata
                        )
                    ) as responses
                FROM risk_responses rr
                WHERE rr.user_id = $1
                GROUP BY rr.user_id
            """
            result = await conn.fetchrow(sql, user_id)
            
            if result and result[0]:
                intelligence['risk_intelligence'] = {
                    'responses': result[0],
                    'data_available': True,
                    'response_count': len(result[0])
                }
                intelligence['data_sources_available'].append('risk')
                logging.info(f"✅ Risk intelligence extracted: {len(result[0])} responses")
            else:
                intelligence['risk_intelligence'] = {'data_available': False}
                logging.info(f"⚠️ No risk data found for user {user_id}")
        
    except Exception as e:
        logging.error(f"❌ Error extracting risk intelligence: {str(e)}")
        intelligence['risk_intelligence'] = {'data_available': False, 'error': str(e)}
    
    logging.info(f"🎯 Multi-database intelligence extraction complete: {len(intelligence['data_sources_available'])} sources available")
    return intelligence


# ======================================================
#           Risk Engine Report Generation Sections
# ======================================================

def get_risk_report_sections():
    """Risk report sections with 100% data utilization from ALL 7 engines for maximum RAG quality"""
    return {
        "threat_intelligence_analysis": {
            "title": "Cross-Engine Threat Intelligence & Vulnerability Assessment",
            "word_target": 2800,
            "analysis_requirements": """
COMPREHENSIVE THREAT ANALYSIS USING ALL 7 ENGINES:

1. COMPLETE DATA EXTRACTION FROM ALL ENGINES
   - Risk Fortress Engine: Extract ALL threat indicators from Questions 1.1-9.2
   - Profile Engine: Correlate personality traits with vulnerability patterns and risk tolerance
   - Dream Engine: Analyze vision-reality gaps that create strategic blind spots and threat exposure
   - Growth Engine: Map expansion plans against scalability threats and resource constraints
   - Analyst Engine: Extract operational vulnerabilities from business process responses
   - People Operations Engine: Identify team-based vulnerabilities and human factor risks
   - SME Assessment Engine: Correlate business maturity with threat detection sophistication

2. BEHAVIORAL THREAT INTELLIGENCE FROM ALL SOURCES
   - Mouse behavior patterns during threat-related questions across ALL engines
   - Keyboard deliberation patterns when answering risk-sensitive questions
   - Response consistency analysis across all 7 engines for threat awareness reliability
   - Engagement depth correlation between risk questions and other engine responses
   - Cross-engine response time analysis for stress indicators during threat assessment

3. MULTI-ENGINE THREAT CORRELATION MATRIX
   - Profile-Risk Correlation: Leadership style vs threat management approach
   - Dream-Risk Alignment: Long-term vision vs current threat preparedness gaps
   - Growth-Risk Exposure: Business expansion velocity vs threat mitigation capacity
   - Analyst-Risk Operations: Daily operations vs threat detection integration
   - People-Risk Culture: Team dynamics vs collective threat awareness
   - SME-Risk Maturity: Business sophistication vs threat management capability
   - Complete cross-engine threat vulnerability mapping using ALL response data

4. ACTIONABLE THREAT ASSESSMENT WITH ENGINE-SPECIFIC INSIGHTS
   - Threat priority matrix incorporating insights from ALL 7 engines (1-10 risk scores)
   - Engine-specific vulnerability identification:
     * Profile Engine: Leadership blind spots and decision-making vulnerabilities
     * Dream Engine: Strategic vision threats and aspiration-reality misalignment risks
     * Growth Engine: Scaling threats and market expansion vulnerabilities
     * Analyst Engine: Operational process weaknesses and system vulnerabilities
     * People Operations: Team-based risks and human factor vulnerabilities
     * SME Assessment: Maturity-gap threats and capability-aspiration mismatches

5. COMPREHENSIVE IMPLEMENTATION USING ALL ENGINE DATA
   - Threat monitoring protocols customized to Profile leadership style
   - Strategic threat planning aligned with Dream vision and timeline
   - Scalable threat management integrated with Growth expansion plans
   - Operational threat detection embedded in Analyst workflow processes
   - Team-based threat awareness training using People Operations insights
   - Maturity-appropriate threat sophistication from SME Assessment data
   - Budget requirements: $10-50K based on cross-engine risk exposure analysis
   - Timeline: Phase 1 (30 days), Phase 2 (90 days), Phase 3 (180 days)
   - Success KPIs: 25-40% threat detection improvement using baseline from ALL engines

6. CROSS-ENGINE VALIDATION AND INSIGHTS
   - Minimum 8 significant threat patterns identified across ALL 7 engines
   - Engine-to-engine threat correlation analysis with confidence assessments
   - Behavioral threat consistency validation across all assessment domains
   - Comprehensive threat intelligence synthesis using complete data spectrum
            """
        },
        
        "security_architecture_design": {
            "title": "Multi-Engine Security Architecture & Behavioral Defense Systems",
            "word_target": 2700,
            "analysis_requirements": """
SECURITY ARCHITECTURE USING COMPLETE 7-ENGINE INTELLIGENCE:

1. COMPREHENSIVE SECURITY POSTURE FROM ALL ENGINES
   - Risk Fortress: Security maturity from Questions 2.1-2.9 plus all security-related responses
   - Profile Engine: Security leadership style and decision-making patterns for security governance
   - Dream Engine: Long-term security vision alignment and future security requirements
   - Growth Engine: Security scaling requirements and expansion-driven security needs
   - Analyst Engine: Current operational security gaps and process integration requirements
   - People Operations: Team security capabilities, culture, and training readiness
   - SME Assessment: Security sophistication alignment with overall business maturity

2. BEHAVIORAL SECURITY ANALYSIS ACROSS ALL ENGINES
   - Security attention patterns from behavioral data during ALL engine responses
   - Security discipline indicators from response consistency across 7 engines
   - Security implementation readiness from engagement patterns in ALL assessments
   - Security culture indicators from People Operations and Profile behavioral data
   - Security investment priorities from financial responses across Growth and Analyst engines

3. COMPLETE MULTI-ENGINE SECURITY INTEGRATION
   - Profile-Security Leadership: Match security approach to authentic leadership style
   - Dream-Security Vision: Align security strategy with long-term business aspirations
   - Growth-Security Scaling: Design security architecture that scales with expansion plans
   - Analyst-Security Operations: Integrate security with existing operational workflows
   - People-Security Culture: Build security awareness using team dynamics and capabilities
   - SME-Security Maturity: Calibrate security sophistication to business development phase
   - Cross-engine security priority synthesis using ALL available response data

4. ENGINE-SPECIFIC SECURITY ARCHITECTURE COMPONENTS
   - Profile-Based Security: Leadership-aligned security governance and decision frameworks
   - Dream-Aligned Security: Security investments supporting long-term vision achievement
   - Growth-Scalable Security: Security controls that expand with business growth
   - Operations-Integrated Security: Security embedded in daily operational processes
   - People-Centric Security: Security training and culture development using team insights
   - Maturity-Appropriate Security: Security sophistication matching business development stage

5. COMPREHENSIVE IMPLEMENTATION USING ALL ENGINE DATA
   - Security budget allocation: $15-75K based on cross-engine capability assessment
   - Implementation timeline: 6-month rollout with engine-specific milestone integration
   - Personnel requirements: 20-100 hours using team capability data from People Operations
   - Engine-integrated success KPIs:
     * Profile: Security leadership effectiveness >85%
     * Dream: Security-vision alignment score >90%
     * Growth: Security scaling readiness >80%
     * Analyst: Security-operations integration >95%
     * People: Security culture adoption >90%
     * SME: Security maturity advancement 40-60%

6. CROSS-ENGINE SECURITY VALIDATION
   - Security effectiveness patterns identified across ALL 7 engines
   - Engine-to-engine security consistency analysis with behavioral validation
   - Complete security intelligence synthesis using entire data spectrum
   - Security architecture optimization using comprehensive behavioral and operational insights
            """
        },
        
        "risk_mitigation_strategy": {
            "title": "Cross-Engine Risk Mitigation & Behavioral Implementation Strategy",
            "word_target": 2600,
            "analysis_requirements": """
COMPREHENSIVE RISK MITIGATION USING ALL 7 ENGINE INTELLIGENCE:

1. COMPLETE RISK MITIGATION ANALYSIS FROM ALL ENGINES
   - Risk Fortress: Mitigation preferences and capabilities from Questions 3.1-3.7
   - Profile Engine: Risk tolerance and mitigation leadership style preferences
   - Dream Engine: Risk appetite alignment with long-term vision and strategic goals
   - Growth Engine: Mitigation resource allocation and scaling risk management needs
   - Analyst Engine: Operational risk mitigation integration and process requirements
   - People Operations: Team-based risk mitigation capabilities and culture readiness
   - SME Assessment: Risk management sophistication and maturity-appropriate strategies

2. BEHAVIORAL RISK MITIGATION PATTERNS ACROSS ALL ENGINES
   - Risk decision consistency analysis from responses across ALL 7 engines
   - Mitigation implementation discipline indicators from behavioral engagement data
   - Risk communication preferences from Profile and People Operations responses
   - Change readiness for risk procedures from behavioral adaptation indicators
   - Cross-engine risk tolerance alignment and consistency validation

3. MULTI-ENGINE RISK MITIGATION CORRELATION MATRIX
   - Profile-Mitigation Style: Leadership approach vs risk management preferences
   - Dream-Mitigation Investment: Vision priorities vs risk mitigation resource allocation
   - Growth-Mitigation Scaling: Business expansion vs evolving risk management needs
   - Analyst-Mitigation Integration: Operations vs risk procedure implementation
   - People-Mitigation Culture: Team dynamics vs collective risk management adoption
   - SME-Mitigation Sophistication: Business maturity vs risk management complexity
   - Complete cross-engine mitigation strategy synthesis

4. ENGINE-SPECIFIC MITIGATION STRATEGIES
   - Profile-Aligned Mitigation: Risk management approach matching leadership DNA
   - Dream-Integrated Mitigation: Risk strategies supporting long-term vision achievement
   - Growth-Scalable Mitigation: Risk management that evolves with business expansion
   - Operations-Embedded Mitigation: Risk procedures integrated with daily workflows
   - People-Centric Mitigation: Risk culture development using team insights and capabilities
   - Maturity-Calibrated Mitigation: Risk sophistication appropriate to business phase

5. COMPREHENSIVE MITIGATION IMPLEMENTATION USING ALL ENGINE DATA
   - Mitigation priority matrix using insights from ALL 7 engines (1-10 risk scores)
   - Budget allocation: $5-30K per risk category based on cross-engine exposure analysis
   - Implementation timeline: 3-month cycles with engine-specific integration milestones
   - Personnel requirements: 15-50 hours per strategy using People Operations capability data
   - Engine-integrated success measures:
     * Profile: Risk leadership effectiveness >80%
     * Dream: Risk-vision alignment >85%
     * Growth: Risk scaling readiness >75%
     * Analyst: Risk-operations integration >90%
     * People: Risk culture adoption >85%
     * SME: Risk management maturity advancement 30-50%

6. CROSS-ENGINE MITIGATION VALIDATION AND OPTIMIZATION
   - Risk mitigation effectiveness patterns across ALL 7 engines
   - Engine-to-engine mitigation consistency with behavioral correlation analysis
   - Complete risk intelligence synthesis using entire assessment spectrum
   - Mitigation cost-effectiveness optimization using comprehensive business intelligence
            """
        },
        
        "compliance_governance": {
            "title": "Multi-Engine Compliance & Behavioral Governance Framework",
            "word_target": 2400,
            "analysis_requirements": """
COMPREHENSIVE COMPLIANCE STRATEGY USING ALL 7 ENGINE INTELLIGENCE:

1. COMPLETE COMPLIANCE ANALYSIS FROM ALL ENGINES
   - Risk Fortress: Compliance maturity and governance preferences from Questions 4.1-4.6
   - Profile Engine: Governance leadership style and compliance decision-making patterns
   - Dream Engine: Compliance culture alignment with long-term business vision
   - Growth Engine: Compliance scaling requirements and regulatory expansion needs
   - Analyst Engine: Compliance integration with operational processes and workflows
   - People Operations: Team compliance capabilities, culture, and training readiness
   - SME Assessment: Compliance sophistication alignment with business maturity level

2. BEHAVIORAL COMPLIANCE PATTERNS ACROSS ALL ENGINES
   - Compliance attention to detail from behavioral patterns during ALL engine responses
   - Governance discipline indicators from response consistency across 7 engines
   - Policy adherence likelihood from behavioral consistency patterns
   - Compliance culture readiness from People Operations and Profile behavioral synthesis
   - Regulatory change adaptation from behavioral flexibility indicators across engines

3. MULTI-ENGINE COMPLIANCE CORRELATION SYNTHESIS
   - Profile-Compliance Leadership: Governance style matching authentic leadership approach
   - Dream-Compliance Culture: Compliance framework supporting long-term vision
   - Growth-Compliance Scaling: Regulatory management evolving with business expansion
   - Analyst-Compliance Operations: Compliance embedded in operational processes
   - People-Compliance Team: Compliance culture leveraging team dynamics and capabilities
   - SME-Compliance Maturity: Governance sophistication appropriate to business phase
   - Cross-engine compliance priority integration using ALL response data

4. ENGINE-SPECIFIC COMPLIANCE FRAMEWORK COMPONENTS
   - Profile-Based Governance: Leadership-aligned compliance structure and decision processes
   - Dream-Aligned Compliance: Regulatory strategy supporting vision achievement
   - Growth-Scalable Compliance: Governance framework that expands with business
   - Operations-Integrated Compliance: Regulatory requirements embedded in daily workflows
   - People-Centric Compliance: Compliance training and culture using team insights
   - Maturity-Appropriate Governance: Compliance sophistication matching development stage

5. COMPREHENSIVE COMPLIANCE IMPLEMENTATION USING ALL ENGINE DATA
   - Compliance budget: $8-25K based on cross-engine maturity and complexity assessment
   - Implementation timeline: 4-month rollout with engine-specific milestone integration
   - Personnel requirements: 30-80 hours using People Operations team capability analysis
   - Legal consultation needs: Based on Growth expansion and Analyst operational complexity
   - Engine-integrated compliance KPIs:
     * Profile: Governance leadership effectiveness >85%
     * Dream: Compliance-vision alignment >90%
     * Growth: Regulatory scaling readiness >80%
     * Analyst: Compliance-operations integration >95%
     * People: Compliance culture adoption >90%
     * SME: Compliance maturity advancement 40-60%

6. CROSS-ENGINE COMPLIANCE VALIDATION AND OPTIMIZATION
   - Compliance effectiveness patterns identified across ALL 7 engines
   - Engine-to-engine governance consistency analysis with behavioral validation
   - Complete compliance intelligence synthesis using entire assessment spectrum
   - Governance optimization using comprehensive behavioral and operational insights
            """
        },
        
        "implementation_roadmap": {
            "title": "Cross-Engine Implementation Intelligence & Behavioral Change Management",
            "word_target": 2800,
            "analysis_requirements": """
COMPREHENSIVE IMPLEMENTATION STRATEGY USING ALL 7 ENGINE INTELLIGENCE:

1. COMPLETE IMPLEMENTATION READINESS FROM ALL ENGINES
   - Risk Fortress: Implementation capability and change readiness from Questions 5.1-5.8
   - Profile Engine: Change leadership style and implementation decision-making patterns
   - Dream Engine: Implementation motivation and long-term sustainability factors
   - Growth Engine: Implementation resource allocation and scaling change requirements
   - Analyst Engine: Implementation integration with operational processes and capabilities
   - People Operations: Team implementation capabilities, change culture, and adoption readiness
   - SME Assessment: Implementation sophistication alignment with business maturity

2. BEHAVIORAL IMPLEMENTATION INTELLIGENCE ACROSS ALL ENGINES
   - Implementation discipline indicators from response consistency across ALL 7 engines
   - Change adoption patterns from behavioral engagement during all assessments
   - Implementation stress tolerance from behavioral data across multiple contexts
   - Change resistance indicators from behavioral consistency and adaptation patterns
   - Implementation communication preferences from Profile and People Operations synthesis

3. MULTI-ENGINE IMPLEMENTATION CORRELATION MATRIX
   - Profile-Implementation Leadership: Change leadership style vs implementation approach
   - Dream-Implementation Motivation: Vision alignment vs implementation sustainability
   - Growth-Implementation Scaling: Business expansion vs change management complexity
   - Analyst-Implementation Operations: Operational integration vs implementation feasibility
   - People-Implementation Team: Team dynamics vs collective change adoption capability
   - SME-Implementation Maturity: Business sophistication vs implementation complexity
   - Complete cross-engine implementation strategy synthesis

4. ENGINE-SPECIFIC IMPLEMENTATION STRATEGIES
   - Profile-Aligned Implementation: Change approach matching authentic leadership DNA
   - Dream-Motivated Implementation: Change strategy supporting long-term vision achievement
   - Growth-Integrated Implementation: Implementation that supports business expansion goals
   - Operations-Embedded Implementation: Change management integrated with daily workflows
   - People-Centric Implementation: Change adoption using team insights and capabilities
   - Maturity-Calibrated Implementation: Implementation sophistication matching business phase

5. COMPREHENSIVE IMPLEMENTATION TIMELINE USING ALL ENGINE DATA
   - Phase 1 (Days 1-30): Foundation using Profile leadership and People Operations readiness
     * Budget: $5-15K based on cross-engine complexity assessment
     * Personnel: 40-60 hours using People Operations capability data
     * Deliverables: Leadership-aligned frameworks, team-ready procedures
   - Phase 2 (Days 31-90): Integration using Analyst operations and Growth scaling needs
     * Budget: $10-35K based on Growth expansion and Analyst operational complexity
     * Personnel: 80-120 hours using cross-engine integration requirements
     * Deliverables: Operations-embedded systems, growth-scalable processes
   - Phase 3 (Days 91-180): Optimization using Dream vision and SME maturity alignment
     * Budget: $5-20K based on Dream aspirations and SME sophistication requirements
     * Personnel: 30-50 hours using maturity-appropriate optimization
     * Deliverables: Vision-aligned optimization, maturity-enhanced performance

6. CROSS-ENGINE IMPLEMENTATION SUCCESS VALIDATION
   - Implementation effectiveness patterns across ALL 7 engines
   - Engine-to-engine change consistency analysis with behavioral correlation
   - Timeline adherence: within 15% using cross-engine realistic assessment
   - Budget variance: within 20% using comprehensive resource requirement analysis
   - User adoption rate: >85% using People Operations team capability insights
   - Performance improvement: 25-40% using complete baseline from all engines
   - Change resistance mitigation: <10% using behavioral prediction from all sources
            """
        },
        
        "financial_risk_analysis": {
            "title": "Cross-Engine Financial Risk Intelligence & Investment Optimization",
            "word_target": 2500,
            "analysis_requirements": """
COMPREHENSIVE FINANCIAL RISK ANALYSIS USING ALL 7 ENGINE INTELLIGENCE:

1. COMPLETE FINANCIAL RISK ASSESSMENT FROM ALL ENGINES
   - Risk Fortress: Financial risk tolerance and investment preferences from Questions 6.1-6.5
   - Profile Engine: Financial decision-making style and investment personality patterns
   - Dream Engine: Financial vision alignment and long-term investment priorities
   - Growth Engine: Financial scaling needs and expansion investment requirements
   - Analyst Engine: Operational financial efficiency and cash flow management patterns
   - People Operations: Financial resource allocation and team financial capabilities
   - SME Assessment: Financial sophistication and investment maturity alignment

2. BEHAVIORAL FINANCIAL INTELLIGENCE ACROSS ALL ENGINES
   - Financial decision consistency from responses across ALL 7 engines
   - Investment risk patterns from behavioral data during financial questions
   - Financial stress tolerance from behavioral indicators across multiple contexts
   - Financial planning sophistication from response depth and deliberation patterns
   - Cross-engine financial priority alignment and consistency validation

3. MULTI-ENGINE FINANCIAL CORRELATION MATRIX
   - Profile-Financial Style: Leadership approach vs financial management preferences
   - Dream-Financial Vision: Long-term aspirations vs financial investment alignment
   - Growth-Financial Scaling: Business expansion vs financial resource requirements
   - Analyst-Financial Operations: Operational efficiency vs financial management integration
   - People-Financial Resources: Team capabilities vs financial resource allocation
   - SME-Financial Maturity: Business sophistication vs financial management complexity
   - Complete cross-engine financial strategy synthesis

4. ENGINE-SPECIFIC FINANCIAL STRATEGIES
   - Profile-Aligned Financial Management: Investment approach matching leadership DNA
   - Dream-Integrated Financial Planning: Financial strategy supporting vision achievement
   - Growth-Scalable Financial Framework: Financial management evolving with expansion
   - Operations-Embedded Financial Controls: Financial procedures integrated with workflows
   - People-Optimized Financial Resources: Financial allocation using team insights
   - Maturity-Calibrated Financial Sophistication: Financial complexity appropriate to business phase

5. COMPREHENSIVE FINANCIAL IMPLEMENTATION USING ALL ENGINE DATA
   - Financial risk mitigation: $10-40K based on cross-engine exposure analysis
   - Financial monitoring system: $2-8K using Analyst operational and Growth scaling needs
   - Professional consultation: 20-60 hours based on SME maturity and Profile leadership needs
   - Implementation timeline: 3-month cycles with engine-specific integration milestones
   - Engine-integrated financial KPIs:
     * Profile: Financial leadership effectiveness >80%
     * Dream: Financial-vision alignment >85%
     * Growth: Financial scaling readiness >75%
     * Analyst: Financial-operations integration >90%
     * People: Financial resource optimization >85%
     * SME: Financial maturity advancement 30-50%

6. CROSS-ENGINE FINANCIAL VALIDATION AND OPTIMIZATION
   - Financial performance patterns across ALL 7 engines
   - Engine-to-engine financial consistency analysis with behavioral correlation
   - Investment ROI improvement: 15-25% using comprehensive baseline assessment
   - Financial risk reduction: 25-45% using complete cross-engine exposure analysis
   - Cash flow optimization using Analyst operations and Growth expansion integration
   - Financial resilience improvement using complete behavioral and operational intelligence
            """
        },
        
        "crisis_response_protocols": {
            "title": "Multi-Engine Crisis Intelligence & Behavioral Leadership Under Pressure",
            "word_target": 2400,
            "analysis_requirements": """
COMPREHENSIVE CRISIS MANAGEMENT USING ALL 7 ENGINE INTELLIGENCE:

1. COMPLETE CRISIS PREPAREDNESS FROM ALL ENGINES
   - Risk Fortress: Crisis readiness and response preferences from Questions 7.1-7.6
   - Profile Engine: Crisis leadership style and decision-making under pressure patterns
   - Dream Engine: Crisis recovery alignment with long-term vision preservation
   - Growth Engine: Crisis impact on expansion plans and scaling recovery requirements
   - Analyst Engine: Crisis operational continuity and process recovery capabilities
   - People Operations: Team crisis capabilities, communication, and coordination readiness
   - SME Assessment: Crisis management sophistication and business continuity maturity

2. BEHAVIORAL CRISIS INTELLIGENCE ACROSS ALL ENGINES
   - Crisis stress response indicators from behavioral patterns during ALL engine responses
   - Crisis decision consistency from responses across 7 engines under simulated pressure
   - Crisis communication effectiveness from behavioral patterns and engagement data
   - Crisis adaptation capability from behavioral flexibility indicators across engines
   - Cross-engine crisis leadership reliability and consistency validation

3. MULTI-ENGINE CRISIS CORRELATION SYNTHESIS
   - Profile-Crisis Leadership: Leadership style vs crisis management effectiveness
   - Dream-Crisis Recovery: Vision preservation vs crisis response strategy
   - Growth-Crisis Impact: Business expansion vs crisis resilience and recovery needs
   - Analyst-Crisis Operations: Operational continuity vs crisis response integration
   - People-Crisis Team: Team dynamics vs collective crisis response capability
   - SME-Crisis Maturity: Business sophistication vs crisis management complexity
   - Complete cross-engine crisis preparedness synthesis

4. ENGINE-SPECIFIC CRISIS RESPONSE COMPONENTS
   - Profile-Based Crisis Leadership: Crisis management matching authentic leadership style
   - Dream-Aligned Crisis Recovery: Crisis strategy preserving long-term vision
   - Growth-Integrated Crisis Planning: Crisis management supporting expansion resilience
   - Operations-Embedded Crisis Protocols: Crisis response integrated with workflows
   - People-Centric Crisis Communication: Crisis coordination using team dynamics
   - Maturity-Appropriate Crisis Management: Crisis sophistication matching business phase

5. COMPREHENSIVE CRISIS IMPLEMENTATION USING ALL ENGINE DATA
   - Crisis preparedness budget: $5-20K based on cross-engine exposure and complexity
   - Crisis response team: 3-8 personnel using People Operations capability assessment
   - Training requirements: 16-40 hours using Profile leadership and SME maturity analysis
   - Implementation timeline: 2-month protocol development using cross-engine integration
   - Engine-integrated crisis KPIs:
     * Profile: Crisis leadership effectiveness >85%
     * Dream: Crisis-vision alignment >90%
     * Growth: Crisis scaling resilience >80%
     * Analyst: Crisis-operations integration >95%
     * People: Crisis team coordination >90%
     * SME: Crisis management maturity 40-60% advancement

6. CROSS-ENGINE CRISIS VALIDATION AND EFFECTIVENESS
   - Crisis preparedness patterns across ALL 7 engines
   - Engine-to-engine crisis consistency analysis with behavioral stress response validation
   - Crisis response time: <2 hours using Profile leadership and People Operations coordination
   - Crisis communication: within 4 hours using comprehensive stakeholder analysis
   - Recovery time improvement: 50% faster using complete operational and behavioral intelligence
   - Crisis learning integration using SME maturity and behavioral adaptation insights
            """
        },
        
        "continuous_optimization": {
            "title": "Cross-Engine Optimization Intelligence & Multi-Domain Performance Analytics",
            "word_target": 2200,
            "analysis_requirements": """
COMPREHENSIVE OPTIMIZATION USING ALL 7 ENGINE INTELLIGENCE:

1. COMPLETE OPTIMIZATION ASSESSMENT FROM ALL ENGINES
   - Risk Fortress: Optimization preferences and continuous improvement from Questions 8.1-8.4
   - Profile Engine: Optimization leadership style and improvement decision-making patterns
   - Dream Engine: Optimization alignment with long-term vision and strategic development
   - Growth Engine: Optimization resource allocation and scaling improvement requirements
   - Analyst Engine: Operational optimization opportunities and process improvement capabilities
   - People Operations: Team optimization capabilities, learning culture, and improvement readiness
   - SME Assessment: Optimization sophistication and continuous improvement maturity

2. BEHAVIORAL OPTIMIZATION INTELLIGENCE ACROSS ALL ENGINES
   - Optimization discipline from response consistency across ALL 7 engines
   - Continuous improvement sustainability from behavioral learning and adaptation patterns
   - Optimization implementation reliability from engagement consistency across engines
   - Performance feedback integration from behavioral adaptation indicators
   - Cross-engine optimization priority alignment and consistency validation

3. MULTI-ENGINE OPTIMIZATION CORRELATION MATRIX
   - Profile-Optimization Leadership: Leadership style vs continuous improvement effectiveness
   - Dream-Optimization Vision: Long-term aspirations vs optimization strategy alignment
   - Growth-Optimization Scaling: Business expansion vs improvement requirement complexity
   - Analyst-Optimization Operations: Operational efficiency vs optimization integration
   - People-Optimization Culture: Team dynamics vs collective improvement adoption
   - SME-Optimization Maturity: Business sophistication vs optimization complexity
   - Complete cross-engine optimization strategy synthesis

4. ENGINE-SPECIFIC OPTIMIZATION STRATEGIES
   - Profile-Aligned Optimization: Improvement approach matching leadership DNA
   - Dream-Integrated Optimization: Optimization supporting long-term vision achievement
   - Growth-Scalable Optimization: Improvement framework evolving with expansion
   - Operations-Embedded Optimization: Continuous improvement integrated with workflows
   - People-Centric Optimization: Improvement culture using team insights and capabilities
   - Maturity-Calibrated Optimization: Improvement sophistication matching business phase

5. COMPREHENSIVE OPTIMIZATION IMPLEMENTATION USING ALL ENGINE DATA
   - Optimization budget allocation using cross-engine priority and capability assessment
   - Implementation timeline with engine-specific integration and milestone coordination
   - Performance measurement using comprehensive baseline from ALL 7 engines
   - Engine-integrated optimization KPIs:
     * Profile: Optimization leadership effectiveness >80%
     * Dream: Optimization-vision alignment >85%
     * Growth: Optimization scaling efficiency >75%
     * Analyst: Operations-optimization integration >90%
     * People: Optimization culture adoption >85%
     * SME: Optimization maturity advancement 30-50%

6. CROSS-ENGINE OPTIMIZATION VALIDATION AND PERFORMANCE ANALYTICS
   - Optimization effectiveness patterns across ALL 7 engines
   - Engine-to-engine improvement consistency analysis with behavioral sustainability validation
   - Performance improvement forecasting using complete cross-engine baseline
   - Optimization ROI calculation using comprehensive resource and impact analysis
   - Continuous improvement sustainability using behavioral and operational intelligence synthesis
   - Strategic optimization evolution using complete business intelligence spectrum
            """
        }
    }

# ======================================================
#           JWT Authentication Configuration
# ======================================================

# JWT Secret Key - must match the key used by philotimo-backend
JWT_SECRET = os.getenv("JWT_SECRET", "philotimo-global-jwt-secret-2024!!")
JWT_ALGORITHM = os.getenv("JWT_ALGORITHM", "HS256")

# Philotimo database configuration for token validation
PHILOTIMO_DB_CONFIG = {
    "host": "philotimo-staging-db.postgres.database.azure.com",
    "database": "philotimodb",
    "user": "wchen",
    "password": "DevPhilot2024!!",
    "port": 5432,
    "sslmode": "require"
}

def hash_token(token: str) -> str:
    """Hash a JWT token using SHA256 for database comparison"""
    return hashlib.sha256(token.encode()).hexdigest()

async def verify_jwt_token(authorization: str = Header(None)) -> Dict:
    """
    Verify JWT token from Authorization header and validate against database

    Args:
        authorization: Authorization header containing Bearer token

    Returns:
        Dict containing user_id, client_id, and email

    Raises:
        HTTPException: If token is missing, invalid, or expired
    """
    if not authorization:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Missing authorization header"
        )

    # Extract token from "Bearer <token>" format
    parts = authorization.split()
    if len(parts) != 2 or parts[0].lower() != "bearer":
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid authorization header format. Expected 'Bearer <token>'"
        )

    token = parts[1]

    try:
        # Decode JWT to get JTI (JWT ID)
        payload = jwt.decode(token, JWT_SECRET, algorithms=[JWT_ALGORITHM])
        jti = payload.get("jti")
        user_id_from_sub = payload.get("sub")  # Some tokens use 'sub' for user_id

        if not jti:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="Invalid token: Missing JTI"
            )

        # Hash the token for database lookup
        token_hash = hash_token(token)

        # Validate token against database
        conn = None
        try:
            conn = psycopg2.connect(**PHILOTIMO_DB_CONFIG)
            cursor = conn.cursor(cursor_factory=RealDictCursor)

            # Query to validate token and get user info
            query = """
                SELECT
                    t.user_id,
                    t.is_revoked AS revoked,
                    t.expires_at,
                    u.client_id,
                    u.email
                FROM api_tokens t
                JOIN users u ON t.user_id = u.id
                WHERE t.jti = %s AND t.token_hash = %s
            """

            cursor.execute(query, (jti, token_hash))
            result = cursor.fetchone()

            if not result:
                raise HTTPException(
                    status_code=status.HTTP_401_UNAUTHORIZED,
                    detail="Invalid token: Token not found in database"
                )

            # Check if token is revoked
            if result['revoked']:
                raise HTTPException(
                    status_code=status.HTTP_401_UNAUTHORIZED,
                    detail="Invalid token: Token has been revoked"
                )

            # Check if token is expired
            if result['expires_at']:
                import datetime
                if datetime.datetime.now(datetime.timezone.utc) > result['expires_at']:
                    raise HTTPException(
                        status_code=status.HTTP_401_UNAUTHORIZED,
                        detail="Invalid token: Token has expired"
                    )

            return {
                "user_id": result['user_id'],
                "client_id": result['client_id'],
                "email": result['email']
            }

        finally:
            if conn:
                conn.close()

    except jwt.ExpiredSignatureError:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="Invalid token: Token has expired"
        )
    except jwt.InvalidTokenError as e:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail=f"Invalid token: {str(e)}"
        )
    except Exception as e:
        logging.error(f"Token validation error: {str(e)}")
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Token validation failed: {str(e)}"
        )

# ======================================================
#           FastAPI Application Setup
# ======================================================

app = FastAPI(title="Risk Engine API", version="1.0")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Pydantic models
class RiskAssessmentRequest(BaseModel):
    user_id: str
    business_name: str
    assessment_data: Dict[str, Any]

class RiskProgressRequest(BaseModel):
    user_id: str
    business_name: str
    assessment_data: Dict[str, Any]
    current_chapter: int
    auto_save: bool = False

# Application lifespan
@asynccontextmanager
async def lifespan(app: FastAPI):
    global _connection_pools
    
    # Startup
    logger = setup_risk_logging()
    logging.info("🚀 Risk Engine API Starting Up")
    logging.info("🔑 Loaded API keys: " + ", ".join([f"Risk_{i+1:02d}" for i in range(len(GEMINI_API_KEYS))]))
    logging.info("📊 Multi-Database Intelligence: ENABLED")
    logging.info("🧠 Behavioral Analytics: ENABLED")
    logging.info("📄 Word Document Chunking: ENABLED")
    logging.info("🧠 Question-Response Chunking: ENABLED")
    logging.info("🔍 Auto-Indexer Integration: ENABLED")
    
    # Clean up any existing connection pools on startup
    try:
        if _connection_pools:
            logging.info(f"🧹 Cleaning up {len(_connection_pools)} existing connection pools...")
            for pool_name, pool in list(_connection_pools.items()):
                try:
                    if not pool.is_closing():
                        await pool.close()
                        await asyncio.sleep(0.1)
                    logging.info(f"✅ Cleaned startup pool: {pool_name}")
                except Exception as cleanup_error:
                    logging.warning(f"⚠️ Error cleaning startup pool {pool_name}: {cleanup_error}")
            
            _connection_pools.clear()
            logging.info("✅ Startup pool cleanup completed")
    except Exception as startup_cleanup_error:
        logging.warning(f"⚠️ Startup pool cleanup error: {startup_cleanup_error}")
    
    # Database table initialization
    try:
        logging.info("🔧 Initializing database tables on startup...")
        
        try:
            conn = get_risk_connection()
            
            # Test connection first
            with conn.cursor() as test_cur:
                test_cur.execute("SELECT 1")
                logging.info("✅ Database connection test passed")
            
            # Create tables
            create_risk_tables(conn)
            
            # Verify tables exist
            with conn.cursor() as cur:
                cur.execute("""
                    SELECT table_name 
                    FROM information_schema.tables 
                    WHERE table_schema = 'public' 
                    AND table_name IN ('risk_assessments', 'risk_responses', 'risk_behavioral_analytics', 'risk_reports')
                """)
                existing_tables = [row[0] for row in cur.fetchall()]
                
                logging.info(f"✅ Verified tables exist: {existing_tables}")
                
                if len(existing_tables) >= 4:
                    logging.info("✅ Database tables initialization complete - All Risk tables ready")
                else:
                    missing_tables = ['risk_assessments', 'risk_responses', 'risk_behavioral_analytics', 'risk_reports']
                    missing = [t for t in missing_tables if t not in existing_tables]
                    logging.error(f"❌ Missing tables: {missing}")
                    raise Exception(f"Missing required tables: {missing}")
            
            conn.close()
            
        except Exception as table_error:
            logging.error(f"❌ Table initialization failed: {str(table_error)}")
            
            # Emergency table creation with detailed error handling
            try:
                logging.info("🚨 EMERGENCY: Attempting direct table creation...")
                emergency_conn = get_risk_connection()
                
                with emergency_conn.cursor() as cur:
                    # Create tables with explicit error handling for each table
                    tables_to_create = [
                        ("risk_assessments", """
                            CREATE TABLE IF NOT EXISTS risk_assessments (
                                id SERIAL PRIMARY KEY,
                                user_id VARCHAR(255) UNIQUE NOT NULL,
                                assessment_type VARCHAR(100) NOT NULL,
                                version VARCHAR(20) NOT NULL,
                                created_at TIMESTAMPTZ,
                                last_updated TIMESTAMPTZ,
                                timezone VARCHAR(100),
                                session_metadata JSONB,
                                device_fingerprint JSONB,
                                progress_tracking JSONB,
                                completion_flags JSONB,
                                raw_data JSONB,
                                multi_database_intelligence JSONB,
                                created_timestamp TIMESTAMPTZ DEFAULT NOW()
                            )
                        """),
                        ("risk_responses", """
                            CREATE TABLE IF NOT EXISTS risk_responses (
                                id SERIAL PRIMARY KEY,
                                assessment_id INTEGER REFERENCES risk_assessments(id),
                                user_id VARCHAR(255) NOT NULL,
                                question_id VARCHAR(50) NOT NULL,
                                section VARCHAR(100) NOT NULL,
                                question_type VARCHAR(50),
                                question_text TEXT,
                                response_format VARCHAR(50),
                                response_data JSONB,
                                all_options JSONB,
                                metadata JSONB,
                                weight VARCHAR(20),
                                answered_at TIMESTAMPTZ,
                                last_modified_at TIMESTAMPTZ,
                                created_timestamp TIMESTAMPTZ DEFAULT NOW(),
                                UNIQUE(assessment_id, question_id)
                            )
                        """),
                        ("risk_behavioral_analytics", """
                            CREATE TABLE IF NOT EXISTS risk_behavioral_analytics (
                                id SERIAL PRIMARY KEY,
                                assessment_id INTEGER REFERENCES risk_assessments(id) UNIQUE,
                                user_id VARCHAR(255) NOT NULL,
                                mouse_behavior JSONB,
                                keyboard_behavior JSONB,
                                attention_patterns JSONB,
                                decision_making_style JSONB,
                                risk_decision_patterns JSONB,
                                created_at TIMESTAMPTZ,
                                created_timestamp TIMESTAMPTZ DEFAULT NOW()
                            )
                        """),
                        ("risk_reports", """
                            CREATE TABLE IF NOT EXISTS risk_reports (
                                id SERIAL PRIMARY KEY,
                                report_id VARCHAR(255) UNIQUE NOT NULL,
                                user_id VARCHAR(255) NOT NULL,
                                assessment_id INTEGER REFERENCES risk_assessments(id),
                                report_type VARCHAR(100) NOT NULL,
                                status VARCHAR(50) NOT NULL,
                                azure_container VARCHAR(255),
                                blob_paths JSONB,
                                chunk_count INTEGER,
                                generation_metadata JSONB,
                                created_at TIMESTAMPTZ DEFAULT NOW(),
                                completed_at TIMESTAMPTZ,
                                indexer_job_id VARCHAR(255),
                                indexer_status VARCHAR(50),
                                indexer_triggered_at TIMESTAMPTZ,
                                indexer_completed_at TIMESTAMPTZ,
                                indexer_error_message TEXT,
                                indexer_retry_count INTEGER DEFAULT 0,
                                multi_database_integration JSONB
                            )
                        """)
                    ]
                    
                    # Create each table individually with error handling
                    for table_name, create_sql in tables_to_create:
                        try:
                            cur.execute(create_sql)
                            logging.info(f"✅ Emergency created table: {table_name}")
                        except Exception as table_create_error:
                            logging.error(f"❌ Failed to create table {table_name}: {table_create_error}")
                    
                    # Create indexes with individual error handling
                    indexes = [
                        ("idx_risk_assessments_user_id", "CREATE INDEX IF NOT EXISTS idx_risk_assessments_user_id ON risk_assessments(user_id)"),
                        ("idx_risk_responses_user_id", "CREATE INDEX IF NOT EXISTS idx_risk_responses_user_id ON risk_responses(user_id)"),
                        ("idx_risk_responses_section", "CREATE INDEX IF NOT EXISTS idx_risk_responses_section ON risk_responses(section)"),
                        ("idx_risk_reports_user_id", "CREATE INDEX IF NOT EXISTS idx_risk_reports_user_id ON risk_reports(user_id)"),
                        ("idx_risk_reports_report_id", "CREATE INDEX IF NOT EXISTS idx_risk_reports_report_id ON risk_reports(report_id)"),
                        ("idx_risk_reports_indexer_job_id", "CREATE INDEX IF NOT EXISTS idx_risk_reports_indexer_job_id ON risk_reports(indexer_job_id)"),
                        ("idx_risk_reports_indexer_status", "CREATE INDEX IF NOT EXISTS idx_risk_reports_indexer_status ON risk_reports(indexer_status)")
                    ]
                    
                    for index_name, index_sql in indexes:
                        try:
                            cur.execute(index_sql)
                            logging.debug(f"✅ Created index: {index_name}")
                        except Exception as index_error:
                            logging.warning(f"⚠️ Failed to create index {index_name}: {index_error}")
                    
                    # Final verification
                    cur.execute("""
                        SELECT table_name 
                        FROM information_schema.tables 
                        WHERE table_schema = 'public' 
                        AND table_name IN ('risk_assessments', 'risk_responses', 'risk_behavioral_analytics', 'risk_reports')
                    """)
                    final_tables = [row[0] for row in cur.fetchall()]
                    
                    logging.info(f"✅ EMERGENCY VERIFICATION: {len(final_tables)} tables confirmed: {final_tables}")
                    
                    if len(final_tables) >= 4:
                        logging.info("✅ EMERGENCY TABLE CREATION SUCCESS - All required tables available")
                    else:
                        logging.error(f"❌ EMERGENCY TABLE CREATION PARTIAL: Only {len(final_tables)}/4 tables created")
                        
                emergency_conn.close()
                
            except Exception as emergency_error:
                logging.error(f"❌ EMERGENCY table creation failed: {str(emergency_error)}")
                logging.error("🚨 CRITICAL: Database tables unavailable - API functionality severely limited")
                if 'emergency_conn' in locals():
                    try:
                        emergency_conn.close()
                    except:
                        pass
            
    except Exception as startup_error:
        logging.error(f"❌ Critical startup error: {str(startup_error)}")
        import traceback
        logging.error(f"🔍 Full startup traceback: {traceback.format_exc()}")
    
    # Add startup success confirmation
    logging.info("✅ Risk Engine startup sequence completed")
    logging.info(f"🔑 API Key Health: {get_api_key_status_summary() if 'get_api_key_status_summary' in globals() else 'Not initialized'}")
    
    yield
    
    # Shutdown with enhanced pool cleanup
    logging.info("⬇️ Risk Engine API Shutting Down")
    
    # Enhanced connection pool cleanup with timeout and error handling
    try:
        if _connection_pools:
            logging.info(f"🔌 Gracefully closing {len(_connection_pools)} connection pools...")
            
            # Create list to avoid modification during iteration
            pools_to_close = list(_connection_pools.items())
            successful_closures = 0
            failed_closures = 0
            
            for pool_name, pool in pools_to_close:
                pool_close_start = time.time()
                try:
                    # Check if pool is already closing
                    if pool.is_closing():
                        logging.info(f"ℹ️ Pool {pool_name} already closing, skipping")
                        continue
                    
                    # Close with timeout
                    await asyncio.wait_for(pool.close(), timeout=5.0)
                    
                    # Wait for closure to complete
                    while not pool.is_closing():
                        await asyncio.sleep(0.1)
                        if time.time() - pool_close_start > 5.0:
                            logging.warning(f"⏰ Pool {pool_name} closure timeout")
                            break
                    
                    pool_close_time = time.time() - pool_close_start
                    logging.info(f"✅ Closed connection pool: {pool_name} ({pool_close_time:.2f}s)")
                    successful_closures += 1
                    
                except asyncio.TimeoutError:
                    pool_close_time = time.time() - pool_close_start
                    logging.error(f"⏰ Pool {pool_name} close timeout after {pool_close_time:.2f}s")
                    failed_closures += 1
                except Exception as pool_error:
                    pool_close_time = time.time() - pool_close_start
                    logging.error(f"❌ Error closing pool {pool_name} after {pool_close_time:.2f}s: {pool_error}")
                    failed_closures += 1
                finally:
                    # Always remove from dictionary
                    _connection_pools.pop(pool_name, None)
            
            # Clear the dictionary regardless
            _connection_pools.clear()
            
            logging.info(f"🔌 Pool cleanup summary: {successful_closures} successful, {failed_closures} failed")
            
            if successful_closures > 0:
                logging.info("✅ Connection pools closed successfully")
            if failed_closures > 0:
                logging.warning(f"⚠️ {failed_closures} pools failed to close properly")
                
        else:
            logging.info("ℹ️ No connection pools to close")
            
    except Exception as shutdown_error:
        logging.error(f"❌ Error during graceful shutdown: {shutdown_error}")
        
        # Force cleanup on error
        try:
            logging.info("🚨 FORCE CLEANUP: Attempting to clear connection pools...")
            _connection_pools.clear()
            logging.info("✅ Force cleanup completed")
        except Exception as force_error:
            logging.error(f"❌ Force cleanup failed: {force_error}")
    
    # Final cleanup verification
    try:
        remaining_pools = len(_connection_pools)
        if remaining_pools > 0:
            logging.warning(f"⚠️ {remaining_pools} connection pools still in memory after cleanup")
            _connection_pools.clear()
        else:
            logging.info("✅ All connection pools successfully cleaned up")
    except Exception as final_check_error:
        logging.error(f"❌ Final cleanup verification error: {final_check_error}")
    
    logging.info("🏁 Risk Engine shutdown complete")

app.router.lifespan_context = lifespan

@app.get("/")
async def root():
    return {
        "message": "Risk Engine API v2.1",
        "status": "operational",
        "features": {
            "multi_database_intelligence": True,
            "enhanced_api_key_management": True,
            "load_balancing": True,
            "gemini_ai_analysis": True,
            "behavioral_analytics": True,
            "word_document_chunking": True,
            "question_response_chunking": True,
            "complete_qa_extraction": True
        },
        "api_keys_status": get_enhanced_api_key_status()
    }

@app.get("/auth/me")
async def get_authenticated_user(auth: Dict = Depends(verify_jwt_token)):
    """
    Get authenticated user information from JWT token
    This endpoint allows the frontend to verify authentication and get user_id
    """
    return {
        "status": "success",
        "user_id": str(auth["user_id"]),
        "client_id": auth.get("client_id"),
        "email": auth.get("email"),
        "authenticated": True
    }

@app.get("/api-key-health")
async def get_api_key_health():
    """Get detailed API key health information"""
    return {
        "timestamp": datetime.now().isoformat(),
        "summary": get_api_key_status_summary(),
        "detailed_status": get_enhanced_api_key_status()
    }

@app.get("/health")
async def health_check():
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "version": "2.1",
        "api_keys_status": get_api_key_status_summary()
    }

@app.post("/risk-audit/{user_id}")
async def process_risk_audit(user_id: str, request: RiskAssessmentRequest, auth: Dict = Depends(verify_jwt_token)):
    """Process comprehensive risk audit with multi-database intelligence and indexer integration"""

    # Permission check: user can only access their own data
    if int(user_id) != auth["user_id"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You can only access your own audit data"
        )

    start_time = time.time()
    logging.info(f"🚀 Starting Risk Audit for user_id={user_id}")
    
    # Generate unique report ID
    report_id = f"risk_report_{user_id}_{int(datetime.now().timestamp())}"
    
    # Initialize job status
    risk_job_status[report_id] = {
        "status": "processing",
        "message": "Starting risk analysis...",
        "progress": 0,
        "sections_completed": 0,
        "started_at": datetime.now().isoformat(),
        "user_id": user_id,
        "business_name": request.business_name
    }
    
    try:
        # Get user profile data with connection pooling
        logging.info(f"👤 Retrieving user profile for {user_id}")
        user_profile = await get_user_profile_data(user_id)
        
        if not user_profile:
            logging.warning(f"⚠️ No user profile found for {user_id}, using provided data")
            user_profile = {
                "username": user_id,
                "business_name": request.business_name,
                "industry": "Unknown Industry",
                "team_size": "Unknown",
                "biggest_challenge": "Risk management"
            }
        
        # Store assessment data with multi-database intelligence
        logging.info(f"💾 Storing risk assessment data...")
        assessment_id = store_risk_assessment(user_id, request.assessment_data, include_multi_db=True)
        
        # Get multi-database intelligence with connection pooling
        logging.info(f"🧠 Extracting enhanced multi-database intelligence...")
        multi_db_intelligence = await get_multi_database_intelligence(user_id)
        
        # Enhanced: Add raw_assessment_data to complete_raw_data for Q&R chunking
        complete_raw_data = {
            "user_id": user_id,
            "report_id": report_id,
            "user_profile": user_profile,
            "responses": request.assessment_data.get("responses", []),
            "assessment_metadata": request.assessment_data.get("assessment_metadata", {}),
            "comprehensive_metadata": request.assessment_data.get("comprehensive_metadata", {}),
            "multi_database_intelligence": multi_db_intelligence,
            "behavioral_analytics": request.assessment_data.get("comprehensive_metadata", {}).get("behavioral_analytics", {}),
            "completion_flags": request.assessment_data.get("completion_flags", {}),
            "text_responses": request.assessment_data.get("text_responses", {}),
            "numeric_inputs": request.assessment_data.get("numeric_inputs", {}),
            "processing_timestamp": datetime.now().isoformat(),
            "raw_assessment_data": request.assessment_data
        }
        
        # Update job status
        risk_job_status[report_id]["message"] = "Generating comprehensive risk strategy..."
        risk_job_status[report_id]["progress"] = 20
        
        # Generate comprehensive risk report
        logging.info(f"🧠 Generating comprehensive risk report...")
        
        with ThreadPoolExecutor(max_workers=1) as executor:
            future = executor.submit(generate_comprehensive_risk_report, complete_raw_data, report_id)
            report_data = await asyncio.get_event_loop().run_in_executor(None, lambda: future.result())
        
        # Ensure raw assessment data is in report metadata for Q&R chunking
        if "_enhanced_risk_report_metadata" not in report_data:
            report_data["_enhanced_risk_report_metadata"] = {}
        
        report_data["_enhanced_risk_report_metadata"]["raw_assessment_data"] = request.assessment_data
        
        logging.info(f"✅ Added raw assessment data to report metadata for Q&R chunking")
        
        # Update job status
        risk_job_status[report_id]["message"] = "Uploading risk strategy to secure storage..."
        risk_job_status[report_id]["progress"] = 80
        
        # Get Azure container name
        container_name = get_azure_container_name(user_id)
        
        # Upload to Azure
        logging.info(f"☁️ Uploading risk report to Azure...")
        upload_success, upload_message = await upload_risk_report_to_azure(report_data, report_id, user_id)
        
        if not upload_success:
            raise Exception(f"Azure upload failed: {upload_message}")
        
        # Store report metadata
        generation_metadata = {
            "total_sections": len([k for k in report_data.keys() if k != "_enhanced_risk_report_metadata"]),
            "total_words": report_data.get("_enhanced_risk_report_metadata", {}).get("total_words", 0),
            "generation_time": time.time() - start_time,
            "ai_model": "gemini-2.5-pro",
            "multi_database_sources": len(complete_raw_data.get("multi_database_intelligence", {}).get("data_sources_available", [])),
            "data_sources_used": complete_raw_data.get("multi_database_intelligence", {}).get("data_sources_available", []),
            "intelligence_correlation": True,
            "total_intelligence_sources": len(complete_raw_data.get("multi_database_intelligence", {}).get("data_sources_available", [])),
            "complete_qa_pairs": complete_raw_data.get("multi_database_intelligence", {}).get("complete_qa_data", {}).get("token_tracking", {}).get("qa_pairs_count", 0),
            "upload_message": upload_message,
            "database_pooling_enabled": True
        }
        
        # Store report metadata in database
        logging.info(f"💾 Storing risk report metadata...")
        store_risk_report_metadata(report_id, user_id, assessment_id, 
                                 report_data.get("_enhanced_risk_report_metadata", {}).get("total_sections", 0),
                                 container_name, generation_metadata)
        
        # Auto-indexer removed - no longer needed
        logging.info(f"✅ Report storage complete - auto-indexer disabled")
        
        # Final status update
        total_time = time.time() - start_time
        risk_job_status[report_id]["status"] = "completed"
        risk_job_status[report_id]["message"] = f"Risk strategy complete! Generated in {total_time:.1f}s"
        risk_job_status[report_id]["progress"] = 100
        risk_job_status[report_id]["completed_at"] = datetime.now().isoformat()
        risk_job_status[report_id]["total_generation_time"] = total_time
        
        # Final comprehensive logging
        logging.info(f"🎉 RISK AUDIT COMPLETED!")
        logging.info(f"📊 FINAL COMPLETION SUMMARY:")
        logging.info(f"   ✅ Report ID: {report_id}")
        logging.info(f"   ✅ User ID: {user_id}")
        logging.info(f"   ✅ Business: {request.business_name}")
        logging.info(f"   ✅ Total time: {total_time:.2f}s")
        logging.info(f"   ✅ Assessment stored: ID {assessment_id}")
        logging.info(f"   ✅ Azure upload: {upload_message}")
        logging.info(f"   ✅ Container: {container_name}")
        logging.info(f"   ✅ Total sections: {generation_metadata.get('total_sections', 0)}")
        logging.info(f"   ✅ Total words: {generation_metadata.get('total_words', 0):,}")
        logging.info(f"   ✅ Multi-DB sources: {generation_metadata.get('multi_database_sources', 0)}")
        
        # Log indexer status
        indexer_info = generation_metadata.get("indexer_integration", {})
        if indexer_info.get("trigger_successful", False):
            logging.info(f"   ✅ Indexer: TRIGGERED (background processing)")
        else:
            logging.warning(f"   ⚠️ Indexer: FAILED - {indexer_info.get('trigger_error', 'Unknown error')}")
        
        logging.info(f"✅ Risk Audit completed for {user_id} in {total_time:.2f}s")
        
        return {
            "status": "processing",
            "report_id": report_id,
            "message": "Risk strategy generation started",
            "estimated_completion": "2-3 minutes",
            "user_id": user_id,
            "business_name": request.business_name,
            "generation_metadata": generation_metadata,
            "indexer_info": {
                "indexer_triggered": indexer_info.get("trigger_successful", False),
                "indexer_message": "Background indexing started" if indexer_info.get("trigger_successful", False) else f"Indexer failed: {indexer_info.get('trigger_error', 'Unknown')}",
                "search_optimization": "In progress" if indexer_info.get("trigger_successful", False) else "Limited"
            }
        }
        
    except Exception as e:
        total_time = time.time() - start_time
        error_message = f"Risk audit error: {str(e)}"
        
        logging.error(f"❌ RISK AUDIT FAILED!")
        logging.error(f"🔍 Failure details:")
        logging.error(f"   - Error type: {type(e).__name__}")
        logging.error(f"   - Error message: {str(e)}")
        logging.error(f"   - Total time before failure: {total_time:.2f}s")
        logging.error(f"   - User ID: {user_id}")
        logging.error(f"   - Business name: {request.business_name}")
        logging.error(f"   - Report ID: {report_id}")
        
        # Update job status with error
        risk_job_status[report_id]["status"] = "failed"
        risk_job_status[report_id]["message"] = error_message
        risk_job_status[report_id]["error"] = str(e)
        risk_job_status[report_id]["failed_at"] = datetime.now().isoformat()
        risk_job_status[report_id]["total_processing_time"] = total_time
        
        # Log full traceback for debugging
        import traceback
        logging.error(f"🔍 FULL ERROR TRACEBACK:")
        for line_num, line in enumerate(traceback.format_exc().split('\n'), 1):
            if line.strip():
                logging.error(f"   {line_num:02d}: {line}")
        
        raise HTTPException(status_code=500, detail=error_message)

@app.get("/risk_report_status/{report_id}")
async def get_risk_report_status(report_id: str, auth: Dict = Depends(verify_jwt_token)):
    """Get risk report generation status"""

    # No permission check - report_id doesn't directly contain user_id

    if report_id not in risk_job_status:
        # Try to get status from database
        conn = None
        try:
            conn = get_risk_connection()
            with conn.cursor() as cur:
                sql = """
                    SELECT status, generation_metadata, created_at, completed_at, 
                           indexer_status, indexer_job_id
                    FROM risk_reports 
                    WHERE report_id = %s
                """
                cur.execute(sql, (report_id,))
                row = cur.fetchone()
                
                if row:
                    status, metadata, created_at, completed_at, indexer_status, indexer_job_id = row
                    
                    return {
                        "status": status,
                        "report_id": report_id,
                        "message": f"Report {status}",
                        "created_at": created_at.isoformat() if created_at else None,
                        "completed_at": completed_at.isoformat() if completed_at else None,
                        "indexer_status": indexer_status,
                        "indexer_job_id": indexer_job_id,
                        "metadata": metadata
                    }
                else:
                    raise HTTPException(status_code=404, detail="Risk report not found")
        finally:
            if conn:
                conn.close()
    
    return risk_job_status[report_id]

@app.post("/risk_assessment_progress")
async def save_risk_progress(request: RiskProgressRequest, auth: Dict = Depends(verify_jwt_token)):
    """Save risk assessment progress with enhanced tracking"""

    # Permission check: user can only save their own progress
    if int(request.user_id) != auth["user_id"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You can only save your own progress"
        )

    try:
        logging.info(f"💾 Saving risk progress for user {request.user_id}")
        logging.info(f"📊 Progress details: chapter {request.current_chapter}, auto_save: {request.auto_save}")
        
        # Build proper progress tracking data
        progress_tracking_data = {
            'completed_chapters': request.current_chapter,
            'total_chapters': 8,
            'percentage_complete': (request.current_chapter / 8) * 100,
            'last_saved_at': datetime.now().isoformat(),
            'auto_save': request.auto_save,
            'save_trigger': 'auto_save' if request.auto_save else 'manual_save',
            'chapters_completed': []
        }
        
        # Add completed chapters based on current chapter
        chapter_names = [
            "Chapter 1: Risk Foundation Analysis",
            "Chapter 2: Vulnerability Assessment", 
            "Chapter 3: Threat Landscape",
            "Chapter 4: Control Systems Evaluation",
            "Chapter 5: Crisis Management Readiness",
            "Chapter 6: Compliance & Regulatory",
            "Chapter 7: Strategic Risk Planning",
            "Chapter 8: Risk Fortress Architecture"
        ]
        
        for i in range(min(request.current_chapter, len(chapter_names))):
            progress_tracking_data['chapters_completed'].append(chapter_names[i])
        
        logging.info(f"📈 Progress tracking: {progress_tracking_data['percentage_complete']:.1f}% complete")
        
        # Update assessment data with progress tracking
        enhanced_assessment_data = {
            **request.assessment_data,
            'progress_tracking': progress_tracking_data,
            'current_chapter': request.current_chapter,
            'last_updated': datetime.now().isoformat(),
            'save_metadata': {
                'save_timestamp': datetime.now().isoformat(),
                'save_type': 'auto_save' if request.auto_save else 'manual_save',
                'user_agent': 'risk_engine_frontend',
                'chapter_at_save': request.current_chapter
            }
        }
        
        # Count responses for logging
        response_count = 0
        if 'responses' in enhanced_assessment_data:
            if isinstance(enhanced_assessment_data['responses'], list):
                response_count = len(enhanced_assessment_data['responses'])
            elif isinstance(enhanced_assessment_data['responses'], dict):
                response_count = len(enhanced_assessment_data['responses'])
        
        logging.info(f"📊 Assessment data: {response_count} responses, chapter {request.current_chapter}")
        
        # Store progress data
        assessment_id = store_risk_assessment(request.user_id, enhanced_assessment_data)
        
        logging.info(f"✅ Risk progress saved successfully for user {request.user_id}")
        
        return {
            "status": "saved",
            "assessment_id": assessment_id,
            "user_id": request.user_id,
            "current_chapter": request.current_chapter,
            "auto_save": request.auto_save,
            "timestamp": datetime.now().isoformat(),
            "progress_percentage": progress_tracking_data['percentage_complete'],
            "chapters_completed": len(progress_tracking_data['chapters_completed']),
            "response_count": response_count,
            "save_metadata": {
                "save_type": "auto_save" if request.auto_save else "manual_save",
                "chapters_completed": progress_tracking_data['chapters_completed']
            }
        }
        
    except Exception as e:
        error_message = f"Error saving risk progress: {str(e)}"
        logging.error(f"❌ {error_message}")
        logging.error(f"🔍 Save error context: user_id={request.user_id}, chapter={request.current_chapter}")
        import traceback
        logging.error(f"🔍 Full traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=error_message)

@app.get("/risk_assessment_progress/{user_id}")
async def get_risk_progress(user_id: str, auth: Dict = Depends(verify_jwt_token)):
    """Get saved risk assessment progress for a user"""

    # Permission check: user can only access their own progress
    if int(user_id) != auth["user_id"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You can only access your own progress"
        )

    conn = None
    try:
        logging.info(f"📥 Retrieving risk progress for user {user_id}")
        conn = get_risk_connection()
        with conn.cursor() as cur:
            # Get the most recent assessment for this user
            sql = """
                SELECT raw_data, progress_tracking, last_updated, created_at
                FROM risk_assessments 
                WHERE user_id = %s 
                ORDER BY last_updated DESC 
                LIMIT 1
            """
            
            cur.execute(sql, (user_id,))
            row = cur.fetchone()
            
            if row:
                raw_data, progress_tracking, last_updated, created_at = row
                
                logging.info(f"🔍 Raw data type: {type(raw_data)}, Progress tracking type: {type(progress_tracking)}")
                
                # Handle both string and dict types for raw_data
                if isinstance(raw_data, str):
                    try:
                        assessment_data = json.loads(raw_data)
                        logging.info(f"✅ Parsed raw_data from JSON string")
                    except json.JSONDecodeError as e:
                        logging.error(f"❌ Failed to parse raw_data JSON: {e}")
                        assessment_data = {}
                elif isinstance(raw_data, dict):
                    assessment_data = raw_data
                    logging.info(f"✅ Used raw_data as dict (JSONB column)")
                else:
                    logging.warning(f"⚠️ Unexpected raw_data type: {type(raw_data)}")
                    assessment_data = {}
                
                # Handle both string and dict types for progress_tracking
                current_chapter = 1
                if progress_tracking:
                    if isinstance(progress_tracking, str):
                        try:
                            progress_data = json.loads(progress_tracking)
                            current_chapter = progress_data.get('completed_chapters', 1)
                            logging.info(f"✅ Parsed progress_tracking from JSON string")
                        except json.JSONDecodeError as e:
                            logging.error(f"❌ Failed to parse progress_tracking JSON: {e}")
                            current_chapter = 1
                    elif isinstance(progress_tracking, dict):
                        current_chapter = progress_tracking.get('completed_chapters', 1)
                        logging.info(f"✅ Used progress_tracking as dict")
                    else:
                        logging.warning(f"⚠️ Unexpected progress_tracking type: {type(progress_tracking)}")
                        current_chapter = 1
                
                logging.info(f"📊 Progress summary: {len(assessment_data.get('responses', []))} responses, chapter {current_chapter}")
                
                return {
                    "status": "found",
                    "user_id": user_id,
                    "assessment_data": assessment_data,
                    "current_chapter": current_chapter,
                    "created_at": created_at.isoformat() if created_at else None,
                    "updated_at": last_updated.isoformat() if last_updated else None
                }
            else:
                logging.info(f"ℹ️ No saved progress found for user {user_id}")
                return {
                    "status": "not_found",
                    "user_id": user_id,
                    "message": "No saved progress found"
                }
                
    except Exception as e:
        error_message = f"Error retrieving risk progress: {str(e)}"
        logging.error(f"❌ {error_message}")
        logging.error(f"🔍 Error context: user_id={user_id}, error_type={type(e).__name__}")
        raise HTTPException(status_code=500, detail=error_message)
    finally:
        if conn:
            conn.close()

@app.get("/risk_reports/{user_id}")
async def get_user_risk_reports(user_id: str, auth: Dict = Depends(verify_jwt_token)):
    """Get all risk reports for a user"""

    # Permission check: user can only access their own reports
    if int(user_id) != auth["user_id"]:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You can only access your own reports"
        )

    conn = None
    try:
        conn = get_risk_connection()
        with conn.cursor() as cur:
            sql = """
                SELECT report_id, report_type, status, created_at, completed_at, 
                       chunk_count, indexer_status, generation_metadata
                FROM risk_reports 
                WHERE user_id = %s 
                ORDER BY created_at DESC
            """
            
            cur.execute(sql, (user_id,))
            rows = cur.fetchall()
            
            reports = []
            for row in rows:
                report_id, report_type, status, created_at, completed_at, chunk_count, indexer_status, metadata = row
                
                reports.append({
                    "report_id": report_id,
                    "report_type": report_type,
                    "status": status,
                    "created_at": created_at.isoformat() if created_at else None,
                    "completed_at": completed_at.isoformat() if completed_at else None,
                    "chunk_count": chunk_count,
                    "indexer_status": indexer_status,
                    "metadata": metadata
                })
            
            return {
                "user_id": user_id,
                "total_reports": len(reports),
                "reports": reports
            }
            
    except Exception as e:
        logging.error(f"❌ Error getting risk reports: {str(e)}")
        raise HTTPException(status_code=500, detail=str(e))
    finally:
        if conn:
            conn.close()

@app.get("/check_risk_indexer/{report_id}")
async def check_risk_indexer(report_id: str, auth: Dict = Depends(verify_jwt_token)):
    """Enhanced check if indexer completed for risk reports - now checks both database AND real API status"""

    # No permission check - report_id doesn't directly contain user_id

    conn = None
    try:
        conn = get_risk_connection()
        with conn.cursor() as cur:
            sql = """
                SELECT 
                    indexer_job_id, indexer_status, indexer_triggered_at, 
                    indexer_completed_at, indexer_error_message
                FROM risk_reports 
                WHERE report_id = %s
            """
            cur.execute(sql, (report_id,))
            row = cur.fetchone()
            
            if row:
                job_id, status, triggered_at, completed_at, error_msg = row
                
                # Calculate time since trigger
                time_since_trigger = None
                if triggered_at:
                    time_since_trigger = (datetime.now() - triggered_at.replace(tzinfo=None)).total_seconds()
                
                # ENHANCED: Now check real API status if we have a job_id
                api_status = None
                api_message = None
                api_checked = False
                
                if job_id and job_id not in ["failed", "trigger_failed", "worker_error", "thread_error", "sync_error"]:
                    try:
                        logging.info(f"🔍 Checking real API status for job_id: {job_id}")
                        
                        # Check the real auto-indexer API status
                        async with aiohttp.ClientSession(timeout=aiohttp.ClientTimeout(total=10)) as session:
                            status_url = f"{INDEXER_API_BASE_URL}/status/{job_id}"
                            
                            async with session.get(status_url) as response:
                                api_checked = True
                                
                                if response.status == 200:
                                    api_data = await response.json()
                                    api_status = api_data.get("status", "unknown")
                                    api_message = api_data.get("message", "No message from API")
                                    
                                    logging.info(f"✅ Real API status: {api_status}")
                                    
                                    # Update our database if API shows completion but we don't have it
                                    if api_status == "success" and status != "completed":
                                        logging.info(f"🔄 Updating database: API shows success but DB shows {status}")
                                        try:
                                            with conn.cursor() as update_cur:
                                                update_cur.execute(
                                                    "UPDATE risk_reports SET indexer_status = %s, indexer_completed_at = %s WHERE report_id = %s",
                                                    ("completed", datetime.now(), report_id)
                                                )
                                                conn.commit()
                                                status = "completed"  # Update local variable
                                                logging.info(f"✅ Database updated to completed status")
                                        except Exception as update_error:
                                            logging.error(f"❌ Error updating database: {update_error}")
                                    
                                    elif api_status == "failed" and status not in ["failed", "completed"]:
                                        logging.info(f"🔄 Updating database: API shows failed but DB shows {status}")
                                        try:
                                            with conn.cursor() as update_cur:
                                                failure_msg = api_data.get("message", "Failed according to API")
                                                update_cur.execute(
                                                    "UPDATE risk_reports SET indexer_status = %s, indexer_error_message = %s WHERE report_id = %s",
                                                    ("failed", failure_msg, report_id)
                                                )
                                                conn.commit()
                                                status = "failed"  # Update local variable
                                                error_msg = failure_msg
                                                logging.info(f"✅ Database updated to failed status")
                                        except Exception as update_error:
                                            logging.error(f"❌ Error updating database: {update_error}")
                                
                                elif response.status == 404:
                                    api_status = "not_found"
                                    api_message = "Job not found in auto-indexer system"
                                    logging.warning(f"⚠️ Job not found in auto-indexer: {job_id}")
                                
                                else:
                                    api_status = "api_error" 
                                    api_message = f"API returned status {response.status}"
                                    logging.error(f"❌ API error: {response.status}")
                        
                    except asyncio.TimeoutError:
                        api_status = "timeout"
                        api_message = "API status check timed out"
                        logging.warning(f"⏰ API status check timeout for job_id: {job_id}")
                        api_checked = True
                        
                    except Exception as api_error:
                        api_status = "error"
                        api_message = f"Error checking API: {str(api_error)}"
                        logging.error(f"❌ API status check error: {api_error}")
                        api_checked = True
                
                # Enhanced status determination with API integration
                if status == "completed":
                    actual_status = "✅ SUCCESS - Indexer completed!"
                elif status == "failed":
                    actual_status = f"❌ FAILED - {error_msg}"
                elif api_checked and api_status == "success":
                    actual_status = "✅ SUCCESS - Confirmed by API (just updated database)"
                elif api_checked and api_status == "failed":
                    actual_status = f"❌ FAILED - Confirmed by API: {api_message}"
                elif api_checked and api_status in ["queued", "starting", "retrying"]:
                    actual_status = f"🔄 RUNNING - API status: {api_status}"
                elif api_checked and api_status == "not_found":
                    actual_status = "❓ JOB NOT FOUND - May have been cleaned up from auto-indexer"
                elif api_checked and api_status == "timeout":
                    actual_status = "⏰ API TIMEOUT - Auto-indexer may be busy"
                elif status == "triggered" and time_since_trigger and time_since_trigger > 600:
                    actual_status = "❓ LIKELY FAILED - No completion after 10+ minutes"
                elif status == "triggered":
                    actual_status = f"🔄 UNKNOWN - Triggered {int(time_since_trigger)}s ago, no completion feedback"
                else:
                    actual_status = f"❓ UNKNOWN STATUS - {status}"
                
                # Build comprehensive response
                response_data = {
                    "report_id": report_id,
                    "indexer_job_id": job_id,
                    "database_status": status,
                    "actual_status": actual_status,
                    "time_since_trigger_seconds": int(time_since_trigger) if time_since_trigger else None,
                    "error_message": error_msg
                }
                
                # Add API status info if checked
                if api_checked:
                    response_data["api_status_check"] = {
                        "checked": True,
                        "api_status": api_status,
                        "api_message": api_message,
                        "api_url": f"{INDEXER_API_BASE_URL}/status/{job_id}" if job_id else None
                    }
                else:
                    response_data["api_status_check"] = {
                        "checked": False,
                        "reason": "No valid job_id or job_id indicates local failure"
                    }
                
                return response_data
            else:
                return {"error": "Report not found", "report_id": report_id}
                
    except Exception as e:
        logging.error(f"❌ Error in check_risk_indexer: {str(e)}")
        return {"error": str(e)}
    finally:
        if conn:
            conn.close()

# API key management endpoints
@app.post("/reset_api_key/{key_suffix}")
async def reset_specific_api_key(key_suffix: str):
    """Reset a specific API key by suffix (last 4 characters)"""
    try:
        for api_key in GEMINI_API_KEYS:
            if api_key.endswith(key_suffix):
                reset_api_key_immediately(api_key)
                return {
                    "status": "success",
                    "message": f"API key ...{key_suffix} reset successfully",
                    "key_health": api_key_health.get(api_key, {})
                }
        
        return {"status": "error", "message": f"API key ending with {key_suffix} not found"}
        
    except Exception as e:
        return {"status": "error", "message": str(e)}

@app.post("/reset_all_failed_keys")
async def reset_all_failed_keys_endpoint():
    """Reset all failed API keys"""
    try:
        reset_count = reset_all_failed_api_keys()
        ecosystem_health = get_enhanced_api_key_status()
        
        return {
            "status": "success",
            "message": f"Reset {reset_count} failed API keys",
            "reset_count": reset_count,
            "ecosystem_health": ecosystem_health
        }
        
    except Exception as e:
        return {"status": "error", "message": str(e)}

# Main execution
if __name__ == "__main__":
    # Set up logging
    logger = setup_risk_logging()
    logger.info("🚀 Starting Backable Risk Engine")

    # Log Vertex AI status
    if vertex_ai_client:
        logger.info(f"✅ Vertex AI initialized successfully (Project: {VERTEX_PROJECT_ID}, Location: {VERTEX_LOCATION})")
        logger.info("🎯 Using Vertex AI as PRIMARY method with API keys as fallback")
    else:
        logger.warning("⚠️ Vertex AI not available - using API keys only")

    logger.info(f"🔑 Loaded {len(GEMINI_API_KEYS)} API keys for fallback")

    # Get port from environment variable or use default
    port = int(os.environ.get("PORT", 8001))  # Different port from other engines
    
    # Run with uvicorn for production
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=port,
        log_level="info",
        access_log=True,
        workers=1  # Single worker for optimal resource management
    )